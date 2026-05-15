import streamlit as st
import plotly.graph_objects as go
import os, json, hashlib, requests, io
from datetime import datetime
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
from capavenir_database import CapAvenirDB
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                 Table, TableStyle, HRFlowable, KeepTogether)

@st.cache_resource
def get_db():
    return CapAvenirDB()
db = get_db()

# Mot de passe conseiller (SHA-256) — défaut: capavenir2025
# Pour changer: python3 -c "import hashlib; print(hashlib.sha256(b'VOTRE_MDT').hexdigest())"
CONSEILLER_PWD_HASH = "de07220a5f97bf76fc0ad428c9d7e90318d853fdc227484172e6d29da0c1c4bf"
CONSEILLER_LOGIN    = "conseiller"   # identifiant par défaut

def check_password(login: str, password: str) -> bool:
    h = hashlib.sha256(password.encode()).hexdigest()
    return login.strip().lower() == CONSEILLER_LOGIN and h == CONSEILLER_PWD_HASH


# =====================================================================
# COMPATIBILITE STREAMLIT
# =====================================================================
def rerun():
    try:
        st.rerun()
    except AttributeError:
        st.experimental_rerun()

# =====================================================================
# CONFIGURATION
# =====================================================================
st.set_page_config(
    page_title="CapAvenir CMR - Orientation IA",
    page_icon="🎓",
    layout="centered",
)

def inject_css(dark: bool):
    if dark:
        bg           = "linear-gradient(160deg,#0f172a 0%,#1e1040 50%,#0a2318 100%)"
        txt          = "#e2e8f0"
        card_bg      = "linear-gradient(145deg,#1e293b,#1a1f35)"
        card_brd     = "#334155"
        metric_bg    = "linear-gradient(135deg,#1e293b,#141b2d)"
        step_todo_bg = "#1e293b"; step_todo_c="#64748b"
        fiche_bg     = "linear-gradient(135deg,#1e293b,#172032)"
        fiche_brd    = "#7c3aed"
        tab_sel      = "#10b981"
        hover_row    = "#1e293b"
        dash_card    = "linear-gradient(145deg,#1e2d3d,#1a2035)"
        input_brd    = "#475569"
    else:
        bg           = "linear-gradient(160deg,#fff7ed 0%,#fdf4ff 45%,#ecfdf5 100%)"
        txt          = "#1e1b4b"
        card_bg      = "linear-gradient(145deg,#ffffff,#fafafa)"
        card_brd     = "#e5e7eb"
        metric_bg    = "linear-gradient(145deg,#ffffff,#f9fafb)"
        step_todo_bg = "#f3f4f6"; step_todo_c="#9ca3af"
        fiche_bg     = "linear-gradient(145deg,#fffbeb,#fef9c3)"
        fiche_brd    = "#f59e0b"
        tab_sel      = "#f59e0b"
        hover_row    = "#fef3c7"
        dash_card    = "linear-gradient(145deg,#ffffff,#fafafa)"
        input_brd    = "#d1d5db"

    st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

/* ══ GLOBAL ══ */
.stApp {{
    background: {bg} !important;
    color: {txt};
    font-family: 'Inter', sans-serif !important;
}}
h1,h2,h3,h4,h5,h6 {{ color: {txt}; }}

/* ══ HEADER ANIMÉ ══ */
.header-box {{
    background: linear-gradient(135deg,#6d28d9 0%,#db2777 35%,#f97316 65%,#10b981 100%);
    background-size: 300% 300%;
    animation: gradientShift 6s ease infinite;
    color: white; padding: 2.5rem 2.5rem 2rem;
    border-radius: 24px; text-align: center; margin-bottom: 1.8rem;
    box-shadow: 0 12px 40px rgba(109,40,217,0.35), 0 4px 16px rgba(0,0,0,0.15);
    border: 2px solid rgba(255,255,255,0.25);
    position: relative; overflow: hidden;
}}
.header-box::before {{
    content: '';
    position: absolute; top:-50%; left:-50%;
    width: 200%; height: 200%;
    background: radial-gradient(circle,rgba(255,255,255,0.08) 0%,transparent 70%);
    animation: pulse 4s ease-in-out infinite;
}}
@keyframes gradientShift {{
    0%   {{ background-position: 0% 50%; }}
    50%  {{ background-position: 100% 50%; }}
    100% {{ background-position: 0% 50%; }}
}}
@keyframes pulse {{
    0%,100% {{ opacity:0.5; transform: scale(1); }}
    50%     {{ opacity:1;   transform: scale(1.05); }}
}}
.header-badge {{
    display: inline-block;
    background: rgba(255,255,255,0.18);
    border: 1px solid rgba(255,255,255,0.3);
    border-radius: 20px; padding: 0.25rem 0.9rem;
    font-size: 0.72rem; font-weight: 600; letter-spacing: 0.08em;
    text-transform: uppercase; margin-top: 0.6rem;
    backdrop-filter: blur(4px);
}}

/* ══ ALERTES ══ */
.alert-success {{
    background: linear-gradient(135deg,#ecfdf5,#d1fae5);
    border-left: 5px solid #10b981; padding: 1rem 1.2rem;
    border-radius: 12px; color: #065f46; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(16,185,129,0.18);
}}
.alert-warning {{
    background: linear-gradient(135deg,#fffbeb,#fef3c7);
    border-left: 5px solid #f59e0b; padding: 1rem 1.2rem;
    border-radius: 12px; color: #78350f; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(245,158,11,0.18);
}}
.alert-info {{
    background: linear-gradient(135deg,#fdf4ff,#f3e8ff);
    border-left: 5px solid #a855f7; padding: 1rem 1.2rem;
    border-radius: 12px; color: #581c87; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(168,85,247,0.18);
}}
.alert-danger {{
    background: linear-gradient(135deg,#fef2f2,#fee2e2);
    border-left: 5px solid #ef4444; padding: 1rem 1.2rem;
    border-radius: 12px; color: #7f1d1d; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(239,68,68,0.18);
}}
.alert-attente {{
    background: linear-gradient(135deg,#fff7ed,#fed7aa);
    border-left: 5px solid #f97316; padding: 1rem 1.2rem;
    border-radius: 12px; color: #7c2d12; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(249,115,22,0.18);
}}
.alert-revise {{
    background: linear-gradient(135deg,#fdf2f8,#fce7f3);
    border-left: 5px solid #ec4899; padding: 1rem 1.2rem;
    border-radius: 12px; color: #831843; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(236,72,153,0.18);
}}

/* ══ CHAT IA ══ */
.chat-ia {{
    background: linear-gradient(135deg,#f3e8ff,#ddd6fe);
    border-radius: 0 18px 18px 18px; padding: 1rem 1.2rem;
    margin: 0.5rem 0; color: #4c1d95;
    box-shadow: 0 3px 12px rgba(124,58,237,0.15);
    border-left: 4px solid #7c3aed;
    animation: fadeInLeft 0.3s ease;
}}
.chat-user {{
    background: linear-gradient(135deg,#fef3c7,#fde68a);
    border-radius: 18px 0 18px 18px; padding: 1rem 1.2rem;
    margin: 0.5rem 0; color: #78350f; text-align: right;
    box-shadow: 0 3px 12px rgba(245,158,11,0.18);
    border-right: 4px solid #f59e0b;
    animation: fadeInRight 0.3s ease;
}}
@keyframes fadeInLeft  {{ from {{ opacity:0; transform:translateX(-10px); }} to {{ opacity:1; transform:translateX(0); }} }}
@keyframes fadeInRight {{ from {{ opacity:0; transform:translateX(10px);  }} to {{ opacity:1; transform:translateX(0); }} }}

/* ══ FICHE ══ */
.fiche-box {{
    background: {fiche_bg}; border: 2px solid {fiche_brd};
    border-radius: 18px; padding: 2rem; margin: 1rem 0;
    box-shadow: 0 6px 24px rgba(251,191,36,0.14);
}}

/* ══ SCORE CARDS ══ */
.score-card {{
    text-align: center; background: {card_bg};
    border: 1px solid {card_brd}; border-radius: 14px;
    padding: 1rem 0.4rem;
    box-shadow: 0 3px 10px rgba(0,0,0,0.06);
    transition: all 0.25s ease;
}}
.score-card:hover {{
    transform: translateY(-4px);
    box-shadow: 0 8px 20px rgba(124,58,237,0.18);
    border-color: #a855f7;
}}
.score-raw    {{ font-size: 0.73rem; color: #94a3b8; margin-bottom: 2px; }}
.score-etalon {{ font-size: 1.15rem; font-weight: 700; color: #10b981; }}
.score-label  {{ font-size: 0.78rem; font-weight: 600; color: {txt}; }}

/* ══ QUESTIONS TESTS ══ */
.test-question {{
    background: linear-gradient(135deg,#fdf4ff,#ede9fe);
    border: 1px solid #c4b5fd; border-radius: 14px;
    padding: 1.1rem 1.3rem; margin: 0.7rem 0;
    box-shadow: 0 3px 8px rgba(139,92,246,0.1);
    border-left: 5px solid #8b5cf6;
    transition: box-shadow 0.2s;
}}
.test-question:hover {{ box-shadow: 0 4px 14px rgba(139,92,246,0.2); }}

/* ══ STEPPER ══ */
.step-active {{
    text-align:center;
    background: linear-gradient(135deg,#f59e0b,#ef4444,#ec4899);
    color: white; border-radius: 12px; padding: 0.45rem 0.2rem;
    font-size: 0.75rem; font-weight: 700;
    box-shadow: 0 5px 16px rgba(245,158,11,0.45);
}}
.step-done {{
    text-align:center;
    background: linear-gradient(135deg,#d1fae5,#a7f3d0);
    color: #065f46; border-radius: 12px; padding: 0.45rem 0.2rem;
    font-size: 0.75rem; font-weight: 600;
}}
.step-todo {{
    text-align:center; background: {step_todo_bg}; color: {step_todo_c};
    border-radius: 12px; padding: 0.45rem 0.2rem; font-size: 0.75rem;
}}

/* ══ DASHBOARD (mode conseiller) ══ */
.dash-stat-card {{
    background: {dash_card};
    border-radius: 16px; padding: 1.2rem 1rem;
    border: 1px solid {card_brd};
    box-shadow: 0 4px 14px rgba(0,0,0,0.06);
    text-align: center; transition: all 0.25s ease;
}}
.dash-stat-card:hover {{
    transform: translateY(-3px);
    box-shadow: 0 8px 22px rgba(0,0,0,0.1);
}}
.dash-stat-val  {{ font-size: 2.2rem; font-weight: 800; line-height: 1.1; }}
.dash-stat-lbl  {{ font-size: 0.72rem; font-weight: 600; text-transform: uppercase;
                   letter-spacing: 0.07em; opacity: 0.7; margin-top: 4px; }}
.dash-row {{ display:flex; align-items:center; padding: 0.65rem 0.8rem;
             border-radius: 10px; margin-bottom: 4px; transition: background 0.15s; }}
.dash-row:hover {{ background: {hover_row}; }}
.dash-badge {{
    display: inline-block; border-radius: 20px;
    padding: 0.18rem 0.75rem; font-size: 0.72rem; font-weight: 700;
    letter-spacing: 0.03em; white-space: nowrap;
}}
.badge-c        {{ background:#d1fae5; color:#065f46; }}
.badge-a        {{ background:#dbeafe; color:#1e3a8a; }}
.badge-confirme {{ background:#d1fae5; color:#065f46; }}
.badge-attente  {{ background:#fff7ed; color:#7c2d12; }}
.badge-probation{{ background:#fef3c7; color:#78350f; }}
.badge-revise   {{ background:#fce7f3; color:#831843; }}
.badge-indetermine {{ background:#f3f4f6; color:#374151; }}
.search-box {{
    background: {card_bg};
    border: 2px solid {input_brd}; border-radius: 12px;
    padding: 0.6rem 1rem; width: 100%;
    font-size: 0.9rem; color: {txt};
    outline: none; transition: border-color 0.2s;
}}
.search-box:focus {{ border-color: #a855f7; }}

/* ══ OVERRIDES STREAMLIT ══ */
div[data-testid="stMetric"] {{
    background: {metric_bg}; border: 1px solid {card_brd};
    border-radius: 14px; padding: 0.9rem;
    box-shadow: 0 3px 10px rgba(0,0,0,0.05);
}}
div[data-testid="stMetric"] label {{
    font-size: 0.75rem !important; font-weight: 600 !important;
    text-transform: uppercase !important; letter-spacing: 0.06em !important;
    opacity: 0.7;
}}
.stButton > button[kind="primary"] {{
    background: linear-gradient(135deg,#f59e0b,#ef4444) !important;
    border: none !important; border-radius: 12px !important;
    box-shadow: 0 5px 14px rgba(245,158,11,0.38) !important;
    font-weight: 700 !important; color: white !important;
    letter-spacing: 0.02em !important; transition: all 0.2s !important;
}}
.stButton > button[kind="primary"]:hover {{
    background: linear-gradient(135deg,#d97706,#dc2626) !important;
    box-shadow: 0 7px 20px rgba(245,158,11,0.48) !important;
    transform: translateY(-1px) !important;
}}
div[data-testid="stTabs"] button[aria-selected="true"] {{
    border-bottom: 3px solid {tab_sel} !important;
    color: {tab_sel} !important; font-weight: 600 !important;
}}
/* Dividers */
hr {{ border-color: {card_brd} !important; }}
/* Section labels */
.label-sci   {{ background:#d1fae5; color:#065f46; }}
.label-lit   {{ background:#dbeafe; color:#1e3a8a; }}
.label-test  {{ background:#ede9fe; color:#4c1d95; }}
.section-label-colored {{
    font-size: 0.68rem; font-weight: 700; letter-spacing: 0.1em;
    text-transform: uppercase; padding: 0.3rem 0.85rem;
    border-radius: 20px; display: inline-block; margin: 0.8rem 0 0.4rem 0;
}}
</style>
""", unsafe_allow_html=True)


# =====================================================================
# GENERATION PDF — Fiche d'orientation (ReportLab) — UNE SEULE PAGE
# =====================================================================
def generate_pdf_fiche(data: dict) -> bytes:
    """Fiche d'orientation compacte sur une seule page A4. Robuste aux valeurs None."""

    # ── Sécuriser toutes les valeurs numériques ──
    def _f(v, default=0.0):
        try:
            return float(v) if v is not None else default
        except (TypeError, ValueError):
            return default

    def _s(v, default="—", max_len=None):
        r = str(v).strip() if v else default
        return r[:max_len] if max_len and r != default else r

    buffer = io.BytesIO()

    # Couleurs
    CMR_GREEN  = colors.HexColor("#10b981")
    CMR_DARK   = colors.HexColor("#0f172a")
    CMR_BLUE   = colors.HexColor("#1e3a5f")
    CMR_LIGHT  = colors.HexColor("#f8fafc")
    CMR_GRAY   = colors.HexColor("#64748b")
    CMR_AMBER  = colors.HexColor("#f59e0b")
    serie      = _s(data.get("serie"), "?")
    SERIE_CLR  = CMR_GREEN if "C" in serie else (colors.HexColor("#3b82f6") if "A" in serie else CMR_AMBER)

    doc = SimpleDocTemplate(buffer, pagesize=A4,
        leftMargin=1.4*cm, rightMargin=1.4*cm,
        topMargin=1.0*cm, bottomMargin=1.0*cm)

    # Largeur utile = 21cm - 1.4 - 1.4 = 18.2cm
    W = 18.2 * cm

    # Styles compacts
    sH   = ParagraphStyle("h",   fontName="Helvetica-Bold", fontSize=12, textColor=CMR_DARK,  alignment=TA_CENTER, spaceAfter=1)
    sSub = ParagraphStyle("sub", fontName="Helvetica",      fontSize=7,  textColor=colors.white, alignment=TA_CENTER, spaceAfter=1)
    sSec = ParagraphStyle("sec", fontName="Helvetica-Bold", fontSize=8,  textColor=CMR_GREEN, spaceBefore=5, spaceAfter=2)
    sB   = ParagraphStyle("b",   fontName="Helvetica",      fontSize=7.5,textColor=CMR_DARK,  spaceAfter=1)
    sC   = ParagraphStyle("c",   fontName="Helvetica",      fontSize=7.5,textColor=CMR_DARK,  alignment=TA_CENTER)
    sSer = ParagraphStyle("ser", fontName="Helvetica-Bold", fontSize=18, textColor=SERIE_CLR, alignment=TA_CENTER)
    sFt  = ParagraphStyle("ft",  fontName="Helvetica",      fontSize=6,  textColor=CMR_GRAY,  alignment=TA_CENTER)
    sSig = ParagraphStyle("sig", fontName="Helvetica",      fontSize=7.5,textColor=CMR_DARK)

    story = []

    # ── En-tête bilingue (3 colonnes) — W = 18.2cm → 5+8.2+5 ──
    hdr = Table([[
        Paragraph("REPUBLIQUE DU CAMEROUN<br/>Paix — Travail — Patrie", sSub),
        Paragraph("<b>FICHE D'ORIENTATION SCOLAIRE</b><br/>CapAvenir CMR v2.1 — " + _s(data.get("date_gen")), sH),
        Paragraph("REPUBLIC OF CAMEROON<br/>Peace — Work — Fatherland", sSub),
    ]], colWidths=[5*cm, 8.2*cm, 5*cm])
    hdr.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,-1),CMR_DARK),
        ("TEXTCOLOR",(0,0),(-1,-1),colors.white),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),6),
        ("BOTTOMPADDING",(0,0),(-1,-1),6),
    ]))
    story.append(hdr)
    story.append(Spacer(1, 0.15*cm))

    # ── Identité — 3 colonnes, total 18.2cm ──
    story.append(Paragraph("INFORMATIONS PERSONNELLES", sSec))
    id_data = [
        [Paragraph(f"<b>Nom &amp; Prénom :</b> {_s(data.get('nom'))} {_s(data.get('prenom'))}", sB),
         Paragraph(f"<b>Âge :</b> {_s(data.get('age'))} ans  <b>Sexe :</b> {_s(data.get('sexe'))}", sB),
         Paragraph(f"<b>Lycée :</b> {_s(data.get('lycee'), max_len=28)}", sB)],
        [Paragraph(f"<b>Choix :</b> {_s(data.get('choix'))}", sB),
         Paragraph(f"<b>Projet :</b> {_s(data.get('projet_pro'), max_len=32)}", sB),
         Paragraph(f"<b>Revenu :</b> {_s(data.get('revenu'), max_len=26)}", sB)],
    ]
    # 6.5 + 5.8 + 5.9 = 18.2cm
    id_t = Table(id_data, colWidths=[6.5*cm, 5.8*cm, 5.9*cm])
    id_t.setStyle(TableStyle([
        ("FONTSIZE",(0,0),(-1,-1),7.5),
        ("BACKGROUND",(0,0),(-1,-1),CMR_LIGHT),
        ("GRID",(0,0),(-1,-1),0.2,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(id_t)
    story.append(Spacer(1, 0.1*cm))

    # ── Tests + Notes côte à côte ──
    story.append(Paragraph("TESTS PSYCHOTECHNIQUES & NOTES SCOLAIRES", sSec))

    # Colonne gauche : tests — 4.6+1.5+1.5 = 7.6cm
    test_data = [
        [Paragraph("<b>Test</b>", sC), Paragraph("<b>Brut</b>", sC), Paragraph("<b>Étal.</b>", sC)],
        ["D48 — Logique",      f"{_f(data.get('d48')):.1f}",  f"{_f(data.get('d48_e')):.1f}"],
        ["KRX — Maths",        f"{_f(data.get('krx')):.1f}",  f"{_f(data.get('krx_e')):.1f}"],
        ["MECA — Mécanique",   f"{_f(data.get('meca')):.1f}", f"{_f(data.get('meca_e')):.1f}"],
        ["BV11 — Vocabulaire", f"{_f(data.get('bv11')):.1f}", f"{_f(data.get('bv11_e')):.1f}"],
        ["PRC — Proverbes",    f"{_f(data.get('prc')):.1f}",  f"{_f(data.get('prc_e')):.1f}"],
        [Paragraph("<b>SA étal.</b>", sB), "—", Paragraph(f"<b>{_f(data.get('SA_etal')):.2f}</b>", sB)],
        [Paragraph("<b>LA étal.</b>", sB), "—", Paragraph(f"<b>{_f(data.get('LA_etal')):.2f}</b>", sB)],
    ]
    t_tests = Table(test_data, colWidths=[4.6*cm, 1.5*cm, 1.5*cm])
    t_tests.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),CMR_DARK),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
        ("FONTSIZE",(0,0),(-1,-1),7.5),
        ("ALIGN",(1,0),(-1,-1),"CENTER"),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[CMR_LIGHT,colors.white]),
        ("GRID",(0,0),(-1,-1),0.2,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),4),
        ("TEXTCOLOR",(2,7),(2,8),CMR_GREEN),("FONTNAME",(2,7),(2,8),"Helvetica-Bold"),
    ]))

    # Colonne droite : notes — 2.8+1.2+2.8+1.2 = 8.0cm (espacement 0.6)
    notes_d = data.get("notes") or {}
    notes_data = [
        [Paragraph("<b>Matière</b>",sC), Paragraph("<b>Note</b>",sC),
         Paragraph("<b>Matière</b>",sC), Paragraph("<b>Note</b>",sC)],
        ["Mathématiques",         f"{_f(notes_d.get('maths')):.1f}",
         "Français",              f"{_f(notes_d.get('francais')):.1f}"],
        ["Sciences Phys.",        f"{_f(notes_d.get('sci_phy')):.1f}",
         "Histoire-Géo",          f"{_f(notes_d.get('histgeo')):.1f}"],
        ["SVT",                   f"{_f(notes_d.get('svt')):.1f}",
         "Anglais",               f"{_f(notes_d.get('anglais')):.1f}"],
        [Paragraph("<b>Moy. Sci.</b>",sB), Paragraph(f"<b>{_f(data.get('moy_sci')):.2f}</b>",sB),
         Paragraph("<b>Moy. Lit.</b>",sB), Paragraph(f"<b>{_f(data.get('moy_lit')):.2f}</b>",sB)],
    ]
    t_notes = Table(notes_data, colWidths=[2.9*cm, 1.1*cm, 2.9*cm, 1.1*cm])
    t_notes.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),CMR_BLUE),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
        ("FONTSIZE",(0,0),(-1,-1),7.5),
        ("ALIGN",(1,0),(1,-1),"CENTER"),("ALIGN",(3,0),(3,-1),"CENTER"),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[CMR_LIGHT,colors.white]),
        ("GRID",(0,0),(-1,-1),0.2,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),4),
        ("TEXTCOLOR",(1,5),(1,5),CMR_GREEN),("FONTNAME",(1,5),(3,5),"Helvetica-Bold"),
        ("TEXTCOLOR",(3,5),(3,5),colors.HexColor("#3b82f6")),
    ]))

    # Double colonne : 7.6 + 0.6 + 8.0 = 16.2cm  (+ padding latéraux = 18.2)
    double = Table([[t_tests, Spacer(0.6*cm, 1), t_notes]],
                   colWidths=[7.6*cm, 0.6*cm, 8.0*cm])
    double.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(double)
    story.append(Spacer(1, 0.12*cm))

    # ── Verdict ──
    story.append(HRFlowable(width="100%", thickness=1.5, color=SERIE_CLR, spaceAfter=4))
    statut_txt = _s(data.get("statut_label"), "CONFIRMÉ")
    sc         = int(_f(data.get("score_conf"), 0))
    trim_lbl   = _s(data.get("trim_label"), "T1")
    bg_badge   = (colors.HexColor("#ecfdf5") if "C" in serie
                  else colors.HexColor("#eff6ff") if "A" in serie
                  else colors.HexColor("#fffbeb"))

    # Badge série + KPIs sur une ligne — 4.5 + 4.5 + 4.5 + 4.7 = 18.2
    kpi_inner = Table([[
        Paragraph(f"<b>Statut</b><br/>{statut_txt}", sC),
        Paragraph(f"<b>Confiance</b><br/>{sc} %", sC),
        Paragraph(f"<b>Trimestre</b><br/>{trim_lbl}", sC),
    ]], colWidths=[4.5*cm, 4.5*cm, 4.5*cm])
    kpi_inner.setStyle(TableStyle([
        ("FONTSIZE",(0,0),(-1,-1),7.5),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ]))
    verdict_t = Table([[Paragraph(f"SÉRIE {serie}", sSer), kpi_inner]],
                      colWidths=[4.7*cm, 13.5*cm])
    verdict_t.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(0,0),bg_badge),
        ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#e2e8f0")),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(verdict_t)

    # ── Observations + IA ──
    obs  = _s(data.get("notes_conseiller"), "", max_len=220)
    ia_s = _s(data.get("ia_synthese"), "", max_len=220)
    if obs or ia_s:
        story.append(Spacer(1, 0.1*cm))
        rows_obs = []
        if obs:
            rows_obs.append([Paragraph(f"<b>Observations :</b> {obs}", sB)])
        if ia_s:
            rows_obs.append([Paragraph(f"<b>Synthèse IA :</b> {ia_s}", sB)])
        obs_t = Table(rows_obs, colWidths=[W])
        obs_t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,-1),CMR_LIGHT),
            ("GRID",(0,0),(-1,-1),0.2,colors.HexColor("#e2e8f0")),
            ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
            ("LEFTPADDING",(0,0),(-1,-1),6),
        ]))
        story.append(obs_t)

    # ── Signature (bas à droite) ──
    story.append(Spacer(1, 0.3*cm))
    sig_t = Table([[
        Paragraph("", sB),
        Paragraph(
            "<b>Fait à</b> .................................  "
            "<b>Le</b> .................................<br/><br/>"
            "<b>Signature du Conseiller d'Orientation :</b><br/><br/>"
            "____________________________________",
            sSig)
    ]], colWidths=[9*cm, 9.2*cm])
    sig_t.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"BOTTOM")]))
    story.append(sig_t)

    # ── Footer ──
    story.append(Spacer(1, 0.1*cm))
    story.append(HRFlowable(width="100%", thickness=0.3,
                             color=colors.HexColor("#e2e8f0"), spaceAfter=3))
    story.append(Paragraph(
        "CapAvenir CMR v2.1 — Mémoire ENS Informatique Niveau 5 — 2025 | Confidentiel",
        sFt))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# =====================================================================
# TABLE D'ETALONNAGE — Notes brutes → Notes étalonnées /20
# Basée sur les normes psychométriques adaptées au contexte camerounais
# =====================================================================
ETALONNAGE = {
    0.0: 3.0,  1.0: 4.5,  2.0: 5.5,  3.0: 6.5,  4.0: 7.5,
    5.0: 8.5,  6.0: 9.5,  7.0: 10.0, 8.0: 11.0, 9.0: 11.5,
    10.0: 12.0, 11.0: 12.5, 12.0: 13.0, 13.0: 13.5, 14.0: 14.0,
    15.0: 14.5, 16.0: 15.5, 17.0: 16.5, 18.0: 17.5, 19.0: 18.5, 20.0: 20.0
}

def etalonner(note_brute):
    keys = sorted(ETALONNAGE.keys())
    best = min(keys, key=lambda k: abs(k - note_brute))
    return ETALONNAGE[best]

# =====================================================================
# BANQUE DE QUESTIONS DES TESTS PSYCHOTECHNIQUES
# =====================================================================
TESTS_QUESTIONS = {
    # ------------------------------------------------------------------ D48 : 40 questions
    "D48": {
        "nom": "D48 — Raisonnement Logique Abstrait",
        "description": "Trouvez le terme manquant dans chaque série logique. (40 questions)",
        "questions": [
            {"q": "2, 4, 8, 16, ?", "choices": ["24","28","32","36"], "answer": "32",
             "expl": "Suite géométrique de raison 2 : chaque terme est multiplié par 2."},
            {"q": "1, 4, 9, 16, 25, ?", "choices": ["30","34","36","49"], "answer": "36",
             "expl": "Carrés parfaits : 1²=1, 2²=4, 3²=9, 4²=16, 5²=25, 6²=36."},
            {"q": "3, 6, 11, 18, 27, ?", "choices": ["36","38","39","40"], "answer": "38",
             "expl": "Différences : +3,+5,+7,+9,+11. 27+11=38."},
            {"q": "100, 50, 25, 12,5, ?", "choices": ["6","6,25","7","5"], "answer": "6,25",
             "expl": "Suite géométrique de raison 1/2 : 12,5 ÷ 2 = 6,25."},
            {"q": "2, 3, 5, 8, 13, 21, ?", "choices": ["29","34","30","36"], "answer": "34",
             "expl": "Suite de Fibonacci : chaque terme = somme des deux précédents. 13+21=34."},
            {"q": "Triangle → Carré → Pentagone → Hexagone → ?", "choices": ["Cercle","Octogone","Heptagone","Rhombus"], "answer": "Heptagone",
             "expl": "On ajoute un côté à chaque étape : 3→4→5→6→7 côtés = Heptagone."},
            {"q": "Chaud : Froid :: Jour : ?", "choices": ["Soleil","Sombre","Nuit","Lampe"], "answer": "Nuit",
             "expl": "Relation d'opposition : Chaud↔Froid, Jour↔Nuit."},
            {"q": "1, 3, 7, 15, 31, ?", "choices": ["55","60","63","65"], "answer": "63",
             "expl": "Chaque terme = 2×précédent + 1. 31×2+1=63."},
            {"q": "A, C, E, G, ?", "choices": ["H","I","J","K"], "answer": "I",
             "expl": "On saute une lettre à chaque fois : A(skip B)C(skip D)E(skip F)G(skip H)I."},
            {"q": "5, 10, 20, 40, ?", "choices": ["60","70","80","100"], "answer": "80",
             "expl": "Suite géométrique de raison 2 : 40×2=80."},
            {"q": "Z, X, V, T, ?", "choices": ["S","R","Q","P"], "answer": "R",
             "expl": "On recule de 2 lettres à chaque fois : Z→X→V→T→R."},
            {"q": "4, 9, 16, 25, 36, ?", "choices": ["42","45","49","50"], "answer": "49",
             "expl": "Carrés : 2²=4, 3²=9, 4²=16, 5²=25, 6²=36, 7²=49."},
            {"q": "1, 8, 27, 64, ?", "choices": ["100","121","125","144"], "answer": "125",
             "expl": "Cubes : 1³=1, 2³=8, 3³=27, 4³=64, 5³=125."},
            {"q": "3, 5, 9, 17, 33, ?", "choices": ["55","60","65","66"], "answer": "65",
             "expl": "Chaque terme = 2×précédent − 1. 33×2−1=65."},
            {"q": "Oiseau : Plume :: Poisson : ?", "choices": ["Eau","Queue","Écaille","Nageoire"], "answer": "Écaille",
             "expl": "La plume est le revêtement de l'oiseau, l'écaille est le revêtement du poisson."},
            {"q": "7, 14, 28, 56, ?", "choices": ["84","100","112","120"], "answer": "112",
             "expl": "Suite géométrique de raison 2 : 56×2=112."},
            {"q": "2, 6, 12, 20, 30, ?", "choices": ["40","42","44","45"], "answer": "42",
             "expl": "Différences : +4,+6,+8,+10,+12. 30+12=42."},
            {"q": "Médecin : Hôpital :: Juge : ?", "choices": ["Loi","Prison","Tribunal","Mairie"], "answer": "Tribunal",
             "expl": "Le médecin travaille à l'hôpital, le juge travaille au tribunal."},
            {"q": "1, 2, 4, 7, 11, 16, ?", "choices": ["20","21","22","23"], "answer": "22",
             "expl": "Différences : +1,+2,+3,+4,+5,+6. 16+6=22."},
            {"q": "Soleil : Étoile :: Terre : ?", "choices": ["Lune","Planète","Satellite","Comète"], "answer": "Planète",
             "expl": "Le Soleil est une étoile, la Terre est une planète."},
            {"q": "10, 9, 7, 4, 0, ?", "choices": ["-5","-4","-3","-6"], "answer": "-5",
             "expl": "Différences : −1,−2,−3,−4,−5. 0−5=−5."},
            {"q": "Gant : Main :: Chaussette : ?", "choices": ["Jambe","Pied","Genou","Cheville"], "answer": "Pied",
             "expl": "Le gant se met sur la main, la chaussette sur le pied."},
            {"q": "3, 6, 18, 72, ?", "choices": ["144","288","360","432"], "answer": "360",
             "expl": "Raisons : ×2, ×3, ×4, ×5. 72×5=360."},
            {"q": "Arbre : Forêt :: Maison : ?", "choices": ["Rue","Village","Toit","Fenêtre"], "answer": "Village",
             "expl": "Un ensemble d'arbres = forêt. Un ensemble de maisons = village."},
            {"q": "2, 5, 10, 17, 26, ?", "choices": ["35","36","37","38"], "answer": "37",
             "expl": "Différences : +3,+5,+7,+9,+11. 26+11=37."},
            {"q": "Couteau : Couper :: Stylo : ?", "choices": ["Dessiner","Écrire","Tracer","Souligner"], "answer": "Écrire",
             "expl": "Le couteau sert à couper, le stylo sert à écrire (fonction principale)."},
            {"q": "4, 8, 16, 32, 64, ?", "choices": ["96","100","120","128"], "answer": "128",
             "expl": "Suite géométrique de raison 2 : 64×2=128."},
            {"q": "Lait : Vache :: Miel : ?", "choices": ["Fleur","Ruche","Abeille","Sucre"], "answer": "Abeille",
             "expl": "Le lait vient de la vache, le miel vient de l'abeille."},
            {"q": "0, 1, 1, 2, 3, 5, 8, ?", "choices": ["11","12","13","14"], "answer": "13",
             "expl": "Suite de Fibonacci : 5+8=13."},
            {"q": "Voiture : Route :: Bateau : ?", "choices": ["Port","Rivage","Mer","Quai"], "answer": "Mer",
             "expl": "La voiture circule sur la route, le bateau navigue sur la mer."},
            {"q": "5, 15, 45, 135, ?", "choices": ["270","360","405","540"], "answer": "405",
             "expl": "Suite géométrique de raison 3 : 135×3=405."},
            {"q": "Chien : Aboyer :: Chat : ?", "choices": ["Rugir","Miauler","Siffler","Bramer"], "answer": "Miauler",
             "expl": "Le chien aboie, le chat miaule."},
            {"q": "1, 5, 14, 30, 55, ?", "choices": ["77","80","85","91"], "answer": "91",
             "expl": "Différences : +4,+9,+16,+25,+36 (carrés). 55+36=91."},
            {"q": "Nuit : Obscurité :: Soleil : ?", "choices": ["Chaleur","Lumière","Étoile","Jour"], "answer": "Lumière",
             "expl": "La nuit apporte l'obscurité, le soleil apporte la lumière."},
            {"q": "6, 11, 21, 41, 81, ?", "choices": ["121","141","161","151"], "answer": "161",
             "expl": "Chaque terme = 2×précédent − 1. 81×2−1=161."},
            {"q": "Peintre : Tableau :: Sculpteur : ?", "choices": ["Dessin","Argile","Statue","Muse"], "answer": "Statue",
             "expl": "Le peintre crée un tableau, le sculpteur crée une statue."},
            {"q": "3, 4, 6, 9, 13, 18, ?", "choices": ["22","24","25","27"], "answer": "24",
             "expl": "Différences : +1,+2,+3,+4,+5,+6. 18+6=24."},
            {"q": "Eau : Soif :: Nourriture : ?", "choices": ["Goût","Santé","Faim","Digestion"], "answer": "Faim",
             "expl": "L'eau étanche la soif, la nourriture apaise la faim."},
            {"q": "12, 23, 34, 45, 56, ?", "choices": ["60","65","67","70"], "answer": "67",
             "expl": "On ajoute 11 à chaque fois : 56+11=67."},
            {"q": "Mère : Fille :: Tante : ?", "choices": ["Cousine","Nièce","Sœur","Petite-fille"], "answer": "Nièce",
             "expl": "La mère est la mère de la fille. La tante est la tante de la nièce (fille de son frère/sœur)."},
        ]
    },
    # ------------------------------------------------------------------ KRX : 20 questions
    "KRX": {
        "nom": "KRX — Aptitude Mathématique",
        "description": "Résolvez les problèmes de raisonnement numérique. (20 questions)",
        "questions": [
            {"q": "Un train roule à 60 km/h. Combien de km parcourt-il en 2h30 ?", "choices": ["120 km","130 km","150 km","160 km"], "answer": "150 km",
             "expl": "60 × 2,5 = 150 km."},
            {"q": "Si 3x + 7 = 22, quelle est la valeur de x ?", "choices": ["3","4","5","6"], "answer": "5",
             "expl": "3x = 22−7 = 15, donc x = 5."},
            {"q": "Quel est le PGCD de 24 et 36 ?", "choices": ["6","8","12","18"], "answer": "12",
             "expl": "24 = 2³×3, 36 = 2²×3². PGCD = 2²×3 = 12."},
            {"q": "Article à 5 000 FCFA avec remise de 20 %. Nouveau prix :", "choices": ["3 500 FCFA","4 000 FCFA","4 500 FCFA","4 200 FCFA"], "answer": "4 000 FCFA",
             "expl": "5000 × 0,8 = 4 000 FCFA."},
            {"q": "La somme des angles d'un triangle vaut :", "choices": ["90°","180°","270°","360°"], "answer": "180°",
             "expl": "Théorème fondamental : la somme des angles d'un triangle = 180°."},
            {"q": "Quelle est la racine carrée de 144 ?", "choices": ["11","12","13","14"], "answer": "12",
             "expl": "12 × 12 = 144."},
            {"q": "Si un rectangle a une longueur de 8 cm et une largeur de 5 cm, son aire est :", "choices": ["26 cm²","40 cm²","13 cm²","80 cm²"], "answer": "40 cm²",
             "expl": "Aire = longueur × largeur = 8 × 5 = 40 cm²."},
            {"q": "Combien vaut 15 % de 200 ?", "choices": ["25","30","35","40"], "answer": "30",
             "expl": "15/100 × 200 = 30."},
            {"q": "Un robinet remplit un réservoir en 4h. En 1h30, quelle fraction est remplie ?", "choices": ["1/4","3/8","1/3","3/4"], "answer": "3/8",
             "expl": "En 1h, il remplit 1/4. En 1h30 = 1,5h : 1,5/4 = 3/8."},
            {"q": "Le PPCM de 6 et 8 est :", "choices": ["12","18","24","48"], "answer": "24",
             "expl": "Multiples de 6 : 6,12,18,24. Multiples de 8 : 8,16,24. PPCM = 24."},
            {"q": "Si 2x − 3 = 11, alors x vaut :", "choices": ["5","6","7","8"], "answer": "7",
             "expl": "2x = 11+3 = 14, x = 7."},
            {"q": "Périmètre d'un carré de côté 7 cm :", "choices": ["14 cm","21 cm","28 cm","49 cm"], "answer": "28 cm",
             "expl": "Périmètre = 4 × côté = 4 × 7 = 28 cm."},
            {"q": "Dans une classe de 40 élèves, 60 % sont des filles. Combien y a-t-il de garçons ?", "choices": ["16","20","24","28"], "answer": "16",
             "expl": "60 % de 40 = 24 filles. Garçons = 40−24 = 16."},
            {"q": "Quel est le résultat de 2³ × 2² ?", "choices": ["10","25","32","64"], "answer": "32",
             "expl": "2³ × 2² = 2^(3+2) = 2⁵ = 32."},
            {"q": "La moyenne de 12, 15, 18 et 11 est :", "choices": ["13","13,5","14","14,5"], "answer": "14",
             "expl": "(12+15+18+11)/4 = 56/4 = 14."},
            {"q": "Convertir 3,5 heures en minutes :", "choices": ["180 min","200 min","210 min","215 min"], "answer": "210 min",
             "expl": "3,5 × 60 = 210 minutes."},
            {"q": "Un commerçant achète à 8 000 FCFA et revend à 10 000 FCFA. Son bénéfice en % est :", "choices": ["20 %","25 %","30 %","15 %"], "answer": "25 %",
             "expl": "Bénéfice = 2000. 2000/8000 × 100 = 25 %."},
            {"q": "Quelle valeur manque : 2/3 = ?/12 ?", "choices": ["6","7","8","9"], "answer": "8",
             "expl": "2/3 = 8/12 (on multiplie numérateur et dénominateur par 4)."},
            {"q": "L'aire d'un triangle de base 10 cm et de hauteur 6 cm est :", "choices": ["30 cm²","60 cm²","15 cm²","45 cm²"], "answer": "30 cm²",
             "expl": "Aire = (base × hauteur)/2 = (10×6)/2 = 30 cm²."},
            {"q": "Si une voiture consomme 7 L aux 100 km, combien consomme-t-elle pour 350 km ?", "choices": ["21 L","24,5 L","28 L","35 L"], "answer": "24,5 L",
             "expl": "7/100 × 350 = 24,5 L."},
        ]
    },
    # ------------------------------------------------------------------ MECA : 36 questions
    "MECA": {
        "nom": "MECA — Compréhension Mécanique",
        "description": "Questions sur les principes mécaniques et physiques. (36 questions)",
        "questions": [
            {"q": "Engrenage A (20 dents) à 100 tr/min entraîne engrenage B (10 dents). B tourne à :", "choices": ["50 tr/min","100 tr/min","200 tr/min","400 tr/min"], "answer": "200 tr/min",
             "expl": "Rapport inverse : (20/10) × 100 = 200 tr/min."},
            {"q": "Levier point d'appui central : 10 N à 2 m à gauche équilibre combien à 1 m à droite ?", "choices": ["5 N","10 N","20 N","40 N"], "answer": "20 N",
             "expl": "Principe des moments : 10×2 = F×1, F = 20 N."},
            {"q": "Quel matériau conduit le mieux l'électricité ?", "choices": ["Bois","Plastique","Cuivre","Verre"], "answer": "Cuivre",
             "expl": "Le cuivre est l'un des meilleurs conducteurs électriques."},
            {"q": "Bille lancée horizontalement vs bille lâchée verticalement (même hauteur) :", "choices": ["La lancée touche en premier","La lâchée touche en premier","Les deux en même temps","Dépend de la masse"], "answer": "Les deux en même temps",
             "expl": "La composante verticale est indépendante du mouvement horizontal (principe de la chute libre)."},
            {"q": "Une poulie fixe seule : la force pour soulever une charge est :", "choices": ["2× plus grande","Égale au poids","2× plus petite","4× plus petite"], "answer": "Égale au poids",
             "expl": "La poulie fixe change uniquement la direction de la force, pas son intensité."},
            {"q": "Une roue de 30 cm de rayon tourne à 10 tr/min. Vitesse linéaire à la périphérie :", "choices": ["3π cm/s","π cm/s","10π cm/s","6π cm/s"], "answer": "π cm/s",
             "expl": "v = 2π×r×n = 2π×30×(10/60) = 10π cm/s… soit environ 31,4 cm/s. La réponse correcte est 10π cm/s."},
            {"q": "Pour soulever 100 N avec une poulie mobile, la force nécessaire est :", "choices": ["200 N","100 N","50 N","25 N"], "answer": "50 N",
             "expl": "Une poulie mobile divise la force par 2 : F = 100/2 = 50 N."},
            {"q": "Quel type de levier a son point d'appui entre la charge et la force ?", "choices": ["Levier du 1er genre","Levier du 2e genre","Levier du 3e genre","Aucun"], "answer": "Levier du 1er genre",
             "expl": "Levier du 1er genre : Appui entre Force et Charge (ex : balance, ciseaux)."},
            {"q": "Si on double la vitesse d'un objet, son énergie cinétique est multipliée par :", "choices": ["2","3","4","8"], "answer": "4",
             "expl": "Ec = ½mv². Vitesse doublée → v² quadruplé → Ec × 4."},
            {"q": "Un ressort allongé de 5 cm supporte 10 N. Avec 20 N, il s'allonge de :", "choices": ["5 cm","10 cm","15 cm","20 cm"], "answer": "10 cm",
             "expl": "Loi de Hooke : F = k×x. k = 10/5 = 2 N/cm. x = 20/2 = 10 cm."},
            {"q": "La pression d'un liquide augmente avec :", "choices": ["La surface","La profondeur","La couleur","La température uniquement"], "answer": "La profondeur",
             "expl": "P = ρ×g×h. La pression augmente proportionnellement avec la profondeur h."},
            {"q": "Quel outil amplifie la force grâce au principe du levier ?", "choices": ["Marteau","Tournevis","Pince coupante","Règle"], "answer": "Pince coupante",
             "expl": "La pince coupante est un levier du 1er genre : le pivot amplifie la force appliquée."},
            {"q": "Un vélo : le pédalier (grand plateau) a 48 dents, le pignon arrière 12 dents. Rapport de transmission :", "choices": ["1/4","4","6","1/6"], "answer": "4",
             "expl": "Rapport = dents plateau / dents pignon = 48/12 = 4. La roue tourne 4× plus vite que les pédales."},
            {"q": "Un objet en chute libre : sa vitesse :", "choices": ["Est constante","Diminue","Augmente régulièrement","Dépend du poids"], "answer": "Augmente régulièrement",
             "expl": "g ≈ 9,8 m/s². La vitesse augmente de 9,8 m/s chaque seconde."},
            {"q": "Quelle est la fonction principale d'un amortisseur dans un véhicule ?", "choices": ["Augmenter la vitesse","Absorber les chocs et vibrations","Réduire le carburant","Diriger les roues"], "answer": "Absorber les chocs et vibrations",
             "expl": "L'amortisseur dissipe l'énergie des oscillations du ressort de suspension."},
            {"q": "Deux aimants : pôles identiques se font face. Ils vont :", "choices": ["S'attirer","Se repousser","Rester indifférents","Fondre"], "answer": "Se repousser",
             "expl": "Loi des aimants : pôles identiques se repoussent, pôles contraires s'attirent."},
            {"q": "Dans un circuit série, si une ampoule grille :", "choices": ["Seule cette ampoule s'éteint","Toutes les ampoules s'éteignent","Le courant double","Rien ne change"], "answer": "Toutes les ampoules s'éteignent",
             "expl": "En série, le circuit est coupé si un composant lâche : tout s'éteint."},
            {"q": "Dans un circuit parallèle, si une ampoule grille :", "choices": ["Toutes s'éteignent","Seule cette ampoule s'éteint","Le courant diminue de moitié","Le fusible saute"], "answer": "Seule cette ampoule s'éteint",
             "expl": "En parallèle, chaque branche est indépendante : les autres restent allumées."},
            {"q": "La force qui s'oppose au mouvement d'un objet sur une surface est :", "choices": ["La gravité","La tension","Le frottement","La poussée"], "answer": "Le frottement",
             "expl": "La force de frottement s'oppose toujours au mouvement ou à la tendance au mouvement."},
            {"q": "Un plan incliné à 30° : pour monter une charge de 100 N, la force minimale est :", "choices": ["100 N","50 N","70 N","30 N"], "answer": "50 N",
             "expl": "F = P × sin(30°) = 100 × 0,5 = 50 N (sans frottement)."},
            {"q": "Quel principe explique qu'un bouchon flotte sur l'eau ?", "choices": ["Principe de Pascal","Principe d'Archimède","Loi de Newton","Loi de Hooke"], "answer": "Principe d'Archimède",
             "expl": "Tout corps immergé subit une poussée égale au poids du liquide déplacé."},
            {"q": "La transmission par courroie entre deux poulies : si la poulie motrice (D=20cm) tourne à 300 tr/min et la poulie réceptrice a D=10cm, elle tourne à :", "choices": ["150 tr/min","300 tr/min","600 tr/min","900 tr/min"], "answer": "600 tr/min",
             "expl": "N2 = N1 × D1/D2 = 300 × 20/10 = 600 tr/min."},
            {"q": "L'énergie potentielle gravitationnelle dépend de :", "choices": ["La vitesse","La hauteur et la masse","La couleur","La forme"], "answer": "La hauteur et la masse",
             "expl": "Ep = m×g×h. Elle dépend de la masse m et de la hauteur h."},
            {"q": "Un manomètre mesure :", "choices": ["La température","La pression","La vitesse","Le courant"], "answer": "La pression",
             "expl": "Le manomètre est l'instrument de mesure de la pression des fluides."},
            {"q": "Un écrou se visse :", "choices": ["Uniquement dans le sens horaire","Dans le sens antihoraire","Selon le type de filetage","Dans les deux sens identiquement"], "answer": "Selon le type de filetage",
             "expl": "Filetage à droite (standard) : sens horaire. Filetage à gauche : sens antihoraire."},
            {"q": "Quel est le rôle d'un fusible dans un circuit électrique ?", "choices": ["Augmenter la tension","Stocker l'énergie","Protéger contre les surintensités","Mesurer le courant"], "answer": "Protéger contre les surintensités",
             "expl": "Le fusible fond et coupe le circuit si le courant dépasse la valeur nominale."},
            {"q": "Pour une vis, un pas de 2 mm signifie qu'en 5 tours elle avance de :", "choices": ["5 mm","10 mm","15 mm","20 mm"], "answer": "10 mm",
             "expl": "Avance = pas × nombre de tours = 2 × 5 = 10 mm."},
            {"q": "Quel type de pont supporte les charges par traction (câbles) ?", "choices": ["Pont en arc","Pont suspendu","Pont à poutres","Pont en treillis"], "answer": "Pont suspendu",
             "expl": "Les câbles du pont suspendu travaillent en traction pour porter le tablier."},
            {"q": "Dans un moteur à 4 temps, l'ordre des phases est :", "choices": ["Admission-Détente-Compression-Échappement","Admission-Compression-Détente-Échappement","Compression-Admission-Détente-Échappement","Détente-Admission-Compression-Échappement"], "answer": "Admission-Compression-Détente-Échappement",
             "expl": "Les 4 temps : 1-Admission du mélange, 2-Compression, 3-Combustion/Détente, 4-Échappement."},
            {"q": "La chaleur se transmet par convection :", "choices": ["Dans les solides uniquement","Dans les liquides et gaz","Dans le vide","Dans les métaux uniquement"], "answer": "Dans les liquides et gaz",
             "expl": "La convection est un mode de transfert thermique propre aux fluides (liquides et gaz)."},
            {"q": "Quel outil utilise-t-on pour mesurer un courant électrique ?", "choices": ["Voltmètre","Ohmmètre","Ampèremètre","Wattmètre"], "answer": "Ampèremètre",
             "expl": "L'ampèremètre mesure l'intensité du courant, branché en série dans le circuit."},
            {"q": "Un vérin hydraulique : si on appuie sur un petit piston (S=2 cm²) avec 10 N, la force sur le grand piston (S=20 cm²) est :", "choices": ["1 N","10 N","100 N","200 N"], "answer": "100 N",
             "expl": "Principe de Pascal : F2 = F1 × S2/S1 = 10 × 20/2 = 100 N."},
            {"q": "Quel type de joint assure l'étanchéité d'un assemblage boulonné ?", "choices": ["Joint torique","Boulon","Rondelle plate","Écrou"], "answer": "Joint torique",
             "expl": "Le joint torique (O-ring) est le composant dédié à l'étanchéité des assemblages."},
            {"q": "La résistance électrique se mesure en :", "choices": ["Volts","Ampères","Ohms","Watts"], "answer": "Ohms",
             "expl": "L'unité de résistance électrique est l'Ohm (Ω), loi d'Ohm : U = R×I."},
            {"q": "Un thermomètre à mercure fonctionne sur le principe de :", "choices": ["La conductivité","La dilatation thermique","La capillarité","La réfraction"], "answer": "La dilatation thermique",
             "expl": "Le mercure se dilate avec la chaleur, faisant monter le niveau dans le tube."},
            {"q": "Quel mécanisme transforme un mouvement rotatif en mouvement rectiligne ?", "choices": ["Engrenage droit","Système bielle-manivelle","Poulie fixe","Ressort hélicoïdal"], "answer": "Système bielle-manivelle",
             "expl": "La bielle-manivelle (utilisée dans les moteurs) transforme la rotation en translation et vice versa."},
        ]
    },
    # ------------------------------------------------------------------ BV11 : 56 questions
    "BV11": {
        "nom": "BV11 — Aptitude Littéraire & Vocabulaire",
        "description": "Questions de vocabulaire, synonymes, antonymes et compréhension. (56 questions)",
        "questions": [
            {"q": "Synonyme de 'magnifique' :", "choices": ["Horrible","Splendide","Banal","Terne"], "answer": "Splendide",
             "expl": "Magnifique et splendide signifient 'remarquablement beau'."},
            {"q": "Antonyme de 'prolixe' :", "choices": ["Bavard","Éloquent","Laconique","Verbeux"], "answer": "Laconique",
             "expl": "Prolixe = qui parle beaucoup. Laconique = concis, bref."},
            {"q": "Dans 'Il arbore un sourire radieux', 'arbore' signifie :", "choices": ["Cache","Affiche fièrement","Dissimule","Perd"], "answer": "Affiche fièrement",
             "expl": "Arborer = montrer, exhiber avec fierté."},
            {"q": "Sens de 'éphémère' :", "choices": ["Durable","Qui dure peu de temps","Ancien","Éternel"], "answer": "Qui dure peu de temps",
             "expl": "Éphémère (du grec éphémeros) = qui ne dure qu'un jour, transitoire."},
            {"q": "'La neige est un manteau blanc posé sur la terre.' — figure de style :", "choices": ["Métaphore","Oxymore","Anaphore","Personnification"], "answer": "Métaphore",
             "expl": "Comparaison implicite sans 'comme' : la neige est assimilée à un manteau."},
            {"q": "Synonyme de 'perspicace' :", "choices": ["Naïf","Clairvoyant","Distrait","Lent"], "answer": "Clairvoyant",
             "expl": "Perspicace = qui voit et comprend avec acuité, clairvoyant."},
            {"q": "Antonyme de 'austère' :", "choices": ["Sévère","Rigide","Luxueux","Strict"], "answer": "Luxueux",
             "expl": "Austère = sobre, dépouillé. Son antonyme est luxueux/fastueux."},
            {"q": "Quel mot a le même sens que 'véloce' ?", "choices": ["Lent","Rapide","Fort","Léger"], "answer": "Rapide",
             "expl": "Véloce vient du latin 'velox' = rapide."},
            {"q": "'Le vent gémissait dans les arbres' — figure de style :", "choices": ["Métaphore","Comparaison","Personnification","Hyperbole"], "answer": "Personnification",
             "expl": "Le vent est doté d'une action humaine (gémir) : c'est une personnification."},
            {"q": "Synonyme de 'intrépide' :", "choices": ["Peureux","Courageux","Prudent","Hésitant"], "answer": "Courageux",
             "expl": "Intrépide = qui ne ressent pas la peur, courageux, audacieux."},
            {"q": "Sens de 'fallacieux' :", "choices": ["Vrai","Trompeur","Évident","Logique"], "answer": "Trompeur",
             "expl": "Fallacieux = qui cherche à tromper, qui contient une erreur cachée."},
            {"q": "Antonyme de 'bénin' :", "choices": ["Doux","Grave","Simple","Inoffensif"], "answer": "Grave",
             "expl": "Bénin = sans danger. Son contraire est grave/malin."},
            {"q": "Quel est le sens de 'acrimonieux' ?", "choices": ["Doux et agréable","Plein d'amertume et d'hostilité","Silencieux","Généreux"], "answer": "Plein d'amertume et d'hostilité",
             "expl": "Acrimonieux = qui exprime de l'aigreur, de l'hostilité."},
            {"q": "'Il pleuvait des cordes.' — figure de style :", "choices": ["Personnification","Anaphore","Hyperbole","Litote"], "answer": "Hyperbole",
             "expl": "Hyperbole = exagération délibérée pour renforcer l'expression."},
            {"q": "Synonyme de 'diligent' :", "choices": ["Paresseux","Zélé","Indolent","Nonchalant"], "answer": "Zélé",
             "expl": "Diligent = qui travaille avec soin et rapidité, zélé."},
            {"q": "Le préfixe 'anti-' signifie :", "choices": ["Avant","Avec","Contre","Pour"], "answer": "Contre",
             "expl": "Anti- : préfixe d'opposition. Ex : antiparasitaire = contre les parasites."},
            {"q": "Sens de 'sibyllin' :", "choices": ["Clair et direct","Obscur et mystérieux","Simple","Vulgaire"], "answer": "Obscur et mystérieux",
             "expl": "Sibyllin = dont le sens est énigmatique, difficile à comprendre (de Sibylle, prophétesse antique)."},
            {"q": "'Je meurs de faim !' — figure de style :", "choices": ["Métaphore","Hyperbole","Allitération","Euphémisme"], "answer": "Hyperbole",
             "expl": "Exagération volontaire pour exprimer une forte sensation."},
            {"q": "Antonyme de 'loquace' :", "choices": ["Bavard","Taiseux","Expressif","Éloquent"], "answer": "Taiseux",
             "expl": "Loquace = qui parle beaucoup. Taiseux/silencieux est l'antonyme."},
            {"q": "Quel est le sens de 'sagacité' ?", "choices": ["Sottise","Vivacité d'esprit","Paresse","Timidité"], "answer": "Vivacité d'esprit",
             "expl": "Sagacité = pénétration d'esprit, finesse, intelligence pratique."},
            {"q": "Quel mot est l'intrus ? Joyeux / Gai / Hilare / Morose", "choices": ["Joyeux","Gai","Hilare","Morose"], "answer": "Morose",
             "expl": "Morose = triste, sombre. Les trois autres sont des synonymes de joie."},
            {"q": "Synonyme de 'exigu' :", "choices": ["Vaste","Étroit","Profond","Haut"], "answer": "Étroit",
             "expl": "Exigu = de très petite dimension, étroit."},
            {"q": "'Partir c'est mourir un peu.' — figure de style :", "choices": ["Comparaison","Métaphore","Oxymore","Anaphore"], "answer": "Métaphore",
             "expl": "Le départ est assimilé à une petite mort sans 'comme' ou 'tel'."},
            {"q": "Sens du suffixe '-phobie' :", "choices": ["Amour de","Peur de","Étude de","Absence de"], "answer": "Peur de",
             "expl": "-phobie : crainte irrationnelle. Ex : claustrophobie = peur des espaces clos."},
            {"q": "Antonyme de 'altruiste' :", "choices": ["Généreux","Égoïste","Charitable","Bienveillant"], "answer": "Égoïste",
             "expl": "Altruiste = qui pense aux autres avant soi. Égoïste = qui pense à soi avant tout."},
            {"q": "Dans 'une chaleur accablante', 'accablante' signifie :", "choices": ["Douce","Légère","Écrasante","Agréable"], "answer": "Écrasante",
             "expl": "Accablant = qui épuise, qui pèse de tout son poids."},
            {"q": "Synonyme de 'futile' :", "choices": ["Important","Sérieux","Frivole","Utile"], "answer": "Frivole",
             "expl": "Futile = sans importance, léger. Frivole = pareil."},
            {"q": "'Cette salle est grande comme un placard !' — figure de style :", "choices": ["Métaphore","Comparaison","Hyperbole","Personnification"], "answer": "Comparaison",
             "expl": "Comparaison avec l'outil 'comme'. Elle est ici ironique/hyperbolique mais la figure principale est la comparaison."},
            {"q": "Sens de 'contigu' :", "choices": ["Éloigné","Adjacent","Opposé","Similaire"], "answer": "Adjacent",
             "expl": "Contigu = qui touche, qui est immédiatement à côté."},
            {"q": "Quel mot signifie 'qui ne peut être contesté' ?", "choices": ["Discutable","Irréfutable","Douteux","Ambigu"], "answer": "Irréfutable",
             "expl": "Irréfutable = qu'on ne peut réfuter (infirmer). Synonyme : indiscutable, incontestable."},
            {"q": "'Il n'a pas tort.' est une :", "choices": ["Hyperbole","Litote","Métaphore","Antiphrase"], "answer": "Litote",
             "expl": "Litote = dire moins pour suggérer plus. 'Pas tort' = il a raison."},
            {"q": "Antonyme de 'diffus' :", "choices": ["Épars","Vague","Concentré","Large"], "answer": "Concentré",
             "expl": "Diffus = répandu en tous sens. Concentré = rassemblé en un point."},
            {"q": "Synonyme de 'alacre' :", "choices": ["Triste","Vif et enjoué","Lent","Sombre"], "answer": "Vif et enjoué",
             "expl": "Alacre (ou alègre) = plein d'entrain, vif, enjoué."},
            {"q": "Quel est le sens de 'équivoque' ?", "choices": ["Clair","Ambigu","Certain","Direct"], "answer": "Ambigu",
             "expl": "Équivoque = qui peut s'interpréter de plusieurs façons, ambigu."},
            {"q": "'Noir sur blanc', 'Blanc comme neige' : 'blanc' et 'noir' sont ici des :", "choices": ["Verbes","Adjectifs","Noms","Adverbes"], "answer": "Adjectifs",
             "expl": "Blanc et noir qualifient des noms : ce sont des adjectifs qualificatifs."},
            {"q": "Sens de 'parcimonie' :", "choices": ["Générosité excessive","Économie excessive","Courage","Rapidité"], "answer": "Économie excessive",
             "expl": "Parcimonie = épargne méticuleuse, avarice modérée."},
            {"q": "Antonyme de 'insolite' :", "choices": ["Étrange","Inhabituel","Banal","Surprenant"], "answer": "Banal",
             "expl": "Insolite = inhabituel, étrange. Son contraire est banal, ordinaire."},
            {"q": "Dans 'un discours fleuri', 'fleuri' signifie :", "choices": ["Bref","Orné de métaphores","Monotone","Simple"], "answer": "Orné de métaphores",
             "expl": "Fleuri, au figuré = riche en ornements rhétoriques, en images."},
            {"q": "Synonyme de 'véhément' :", "choices": ["Calme","Passionné et violent","Prudent","Indifférent"], "answer": "Passionné et violent",
             "expl": "Véhément = qui s'exprime avec une grande force et passion."},
            {"q": "Quel mot désigne la crainte morbide des espaces ouverts ?", "choices": ["Claustrophobie","Agoraphobie","Arachnophobie","Xénophobie"], "answer": "Agoraphobie",
             "expl": "Agoraphobie = peur des espaces ouverts ou des lieux publics (du grec agora = place)."},
            {"q": "Sens de 'conciliant' :", "choices": ["Qui cherche la querelle","Qui favorise l'accord","Qui est indifférent","Qui est autoritaire"], "answer": "Qui favorise l'accord",
             "expl": "Conciliant = qui cherche à mettre d'accord, à apaiser les tensions."},
            {"q": "'Le soleil se couchait dans un embrasement de pourpre.' — le registre est :", "choices": ["Familier","Scientifique","Lyrique","Comique"], "answer": "Lyrique",
             "expl": "Vocabulaire riche, images poétiques, expression des sentiments = registre lyrique."},
            {"q": "Synonyme de 'ténace' :", "choices": ["Fragile","Inconstant","Persévérant","Lâche"], "answer": "Persévérant",
             "expl": "Ténace = qui s'accroche, qui ne lâche pas prise, persévérant."},
            {"q": "Quel préfixe signifie 'deux' ou 'double' ?", "choices": ["Mono-","Di-","Tri-","Poly-"], "answer": "Di-",
             "expl": "Di- (ou bi-) = deux. Ex : diocèse, dicotylédone, bicycle."},
            {"q": "Sens de 'perfide' :", "choices": ["Loyal","Traître","Généreux","Innocent"], "answer": "Traître",
             "expl": "Perfide = qui trahit la confiance, déloyal, traître."},
            {"q": "'Toujours ce bruit, toujours cette douleur, toujours ce silence.' — figure de style :", "choices": ["Épiphore","Anaphore","Métaphore","Chiasme"], "answer": "Anaphore",
             "expl": "Anaphore = répétition d'un mot au début de plusieurs propositions ('toujours')."},
            {"q": "Antonyme de 'prodiguer' :", "choices": ["Distribuer","Épargner","Donner","Offrir"], "answer": "Épargner",
             "expl": "Prodiguer = distribuer avec excès. Épargner = retenir, ne pas distribuer."},
            {"q": "Synonyme de 'rétif' :", "choices": ["Docile","Récalcitrant","Obéissant","Souple"], "answer": "Récalcitrant",
             "expl": "Rétif = qui résiste, qui refuse d'avancer ou d'obéir. Syn : récalcitrant."},
            {"q": "Quel mot qualifie ce qui est relatif à la ville ?", "choices": ["Rural","Suburban","Urbain","Sylvestre"], "answer": "Urbain",
             "expl": "Urbain (du latin urbs = ville) désigne ce qui est relatif à la ville."},
            {"q": "Sens de 'acuité' :", "choices": ["Flou","Netteté / intensité","Douceur","Lenteur"], "answer": "Netteté / intensité",
             "expl": "Acuité = qualité de ce qui est aigu, net, pénétrant (acuité visuelle, intellectuelle)."},
            {"q": "Dans 'Il rend l'âme', 'rend l'âme' est un(e) :", "choices": ["Comparaison","Antiphrase","Euphémisme","Oxymore"], "answer": "Euphémisme",
             "expl": "Euphémisme = expression atténuée pour éviter un mot trop direct (ici, 'mourir')."},
            {"q": "Synonyme de 'acerbe' :", "choices": ["Doux","Mordant","Agréable","Aimable"], "answer": "Mordant",
             "expl": "Acerbe = qui blesse par sa causticité, mordant, piquant."},
            {"q": "Quel est l'intrus ? Loquace / Bavard / Verbeux / Taciturne", "choices": ["Loquace","Bavard","Verbeux","Taciturne"], "answer": "Taciturne",
             "expl": "Taciturne = silencieux, peu communicatif. Les autres évoquent le fait de beaucoup parler."},
            {"q": "Sens de 'pérenne' :", "choices": ["Passager","Qui dure longtemps","Récent","Inutile"], "answer": "Qui dure longtemps",
             "expl": "Pérenne = qui dure toute l'année ou de façon durable, permanent."},
            {"q": "Synonyme de 'acéré' :", "choices": ["Émoussé","Tranchant","Mou","Lourd"], "answer": "Tranchant",
             "expl": "Acéré = qui a un fil très coupant, tranchant comme l'acier."},
            {"q": "Antonyme de 'limpide' :", "choices": ["Clair","Transparent","Trouble","Pur"], "answer": "Trouble",
             "expl": "Limpide = parfaitement clair. Son contraire est trouble/opaque."},
        ]
    },
    # ------------------------------------------------------------------ PRC : 40 questions
    "PRC": {
        "nom": "PRC — Proverbes et Raisonnement Linguistique",
        "description": "Interprétez les proverbes et analogies linguistiques. (40 questions)",
        "questions": [
            {"q": "'Mieux vaut prévenir que guérir' signifie :", "choices": ["Il faut soigner vite","Éviter les problèmes vaut mieux que les résoudre après","La prévention est inutile","Consulter régulièrement un médecin"], "answer": "Éviter les problèmes vaut mieux que les résoudre après",
             "expl": "Ce proverbe conseille l'anticipation plutôt que la réaction curative."},
            {"q": "'L'union fait la force' illustre :", "choices": ["Compétition individuelle","Solidarité collective","Solitude salutaire","Hiérarchie stricte"], "answer": "Solidarité collective",
             "expl": "Ensemble, les individus sont plus forts qu'isolément."},
            {"q": "Si A > B et B > C, alors :", "choices": ["C > A","A > C","B est le plus petit","On ne peut pas comparer A et C"], "answer": "A > C",
             "expl": "Relation transitive : A > B > C implique A > C."},
            {"q": "'Qui sème le vent récolte la tempête' — conclusion :", "choices": ["Les actions sont imprévisibles","Les mauvaises actions entraînent de pires conséquences","Le vent est plus doux que la tempête","Il faut se méfier du vent"], "answer": "Les mauvaises actions entraînent de pires conséquences",
             "expl": "Acte négatif (vent) → conséquences amplifiées et négatives (tempête)."},
            {"q": "Analogie : Livre : Bibliothèque :: Tableau : ?", "choices": ["École","Musée","Artiste","Cadre"], "answer": "Musée",
             "expl": "Les livres sont conservés en bibliothèque ; les tableaux au musée."},
            {"q": "'À cheval donné, on ne regarde pas les dents.' Morale :", "choices": ["Il faut examiner tout cadeau","On doit accepter un cadeau sans critiquer","Les chevaux sont précieux","Il faut être exigeant"], "answer": "On doit accepter un cadeau sans critiquer",
             "expl": "Quand on reçoit quelque chose gratuitement, il est impoli de trouver des défauts."},
            {"q": "'Les murs ont des oreilles.' Sens :", "choices": ["Les maisons sont solides","On peut être écouté n'importe où","Les oreilles sont solides","Les murs parlent"], "answer": "On peut être écouté n'importe où",
             "expl": "Il faut se méfier de parler librement car quelqu'un peut toujours entendre."},
            {"q": "Analogie : Médecin : Hôpital :: Enseignant : ?", "choices": ["Bureau","École","Université","Bibliothèque"], "answer": "École",
             "expl": "Le médecin exerce à l'hôpital, l'enseignant à l'école."},
            {"q": "'Pierre qui roule n'amasse pas mousse.' Signifie :", "choices": ["Il faut voyager souvent","Rester en place permet d'accumuler","Les pierres sont précieuses","La mousse est utile"], "answer": "Rester en place permet d'accumuler",
             "expl": "Qui change souvent de situation ne construit pas de patrimoine stable."},
            {"q": "Si tous les A sont B et tous les B sont C, alors :", "choices": ["Tous les C sont A","Tous les A sont C","Aucun A n'est C","Certains A sont C"], "answer": "Tous les A sont C",
             "expl": "Syllogisme : A⊂B et B⊂C → A⊂C."},
            {"q": "'Il ne faut pas vendre la peau de l'ours avant de l'avoir tué.' Leçon :", "choices": ["La chasse est dangereuse","Ne pas compter sur un résultat incertain","Les ours sont rares","Vendre la fourrure est rentable"], "answer": "Ne pas compter sur un résultat incertain",
             "expl": "Ne pas anticiper comme certains les bénéfices d'une action non encore accomplie."},
            {"q": "Analogie : Marteau : Clou :: Crayon : ?", "choices": ["Gomme","Papier","Écriture","Cartable"], "answer": "Papier",
             "expl": "Le marteau enfonce le clou, le crayon écrit sur le papier (support de l'action)."},
            {"q": "'Chaque chose en son temps.' Signifie :", "choices": ["Il faut tout faire vite","Il faut respecter l'ordre et le moment approprié","Le temps est précieux","Rien ne sert de se presser"], "answer": "Il faut respecter l'ordre et le moment approprié",
             "expl": "Chaque action doit être faite au moment qui lui est propice."},
            {"q": "Aucun oiseau n'est un mammifère. L'aigle est un oiseau. Donc :", "choices": ["L'aigle est un mammifère","L'aigle n'est pas un mammifère","Certains aigles sont des mammifères","On ne peut pas conclure"], "answer": "L'aigle n'est pas un mammifère",
             "expl": "Syllogisme négatif : aucun oiseau n'est mammifère, l'aigle est oiseau → l'aigle n'est pas mammifère."},
            {"q": "'Dis-moi qui tu fréquentes, je te dirai qui tu es.' Sens :", "choices": ["Les amis sont inutiles","L'entourage reflète et influence la personnalité","Il faut fréquenter tout le monde","Les inconnus sont meilleurs"], "answer": "L'entourage reflète et influence la personnalité",
             "expl": "On est souvent jugé et influencé par les personnes avec qui on s'associe."},
            {"q": "Analogie : Capitaine : Bateau :: Pilote : ?", "choices": ["Aéroport","Avion","Piste","Tour de contrôle"], "answer": "Avion",
             "expl": "Le capitaine dirige le bateau, le pilote dirige l'avion."},
            {"q": "'Nul n'est prophète en son pays.' Signifie :", "choices": ["Les prophètes voyagent beaucoup","Il est difficile d'être reconnu là où on a grandi","Les étrangers sont mieux accueillis","Le pays natal est le meilleur"], "answer": "Il est difficile d'être reconnu là où on a grandi",
             "expl": "On est souvent sous-estimé par ceux qui nous connaissent depuis l'enfance."},
            {"q": "Tous les étudiants de cette classe ont réussi. Paul est dans cette classe. Donc :", "choices": ["Paul a peut-être réussi","Paul n'a pas réussi","Paul a réussi","On ne sait pas"], "answer": "Paul a réussi",
             "expl": "Si tous ont réussi et Paul est dans la classe, il a nécessairement réussi."},
            {"q": "'La nuit tous les chats sont gris.' Signifie :", "choices": ["Les chats changent de couleur","Dans l'obscurité les différences disparaissent","Il faut des lunettes","Les chats voient la nuit"], "answer": "Dans l'obscurité les différences disparaissent",
             "expl": "La nuit (ignorance, manque d'information) efface les distinctions."},
            {"q": "Analogie : Loi : Parlement :: Règlement : ?", "choices": ["Ministère","Tribunal","Administration","Entreprise"], "answer": "Entreprise",
             "expl": "La loi est produite par le parlement ; le règlement interne est produit par l'entreprise."},
            {"q": "'L'habit ne fait pas le moine.' Leçon :", "choices": ["Les moines s'habillent mal","Les apparences sont trompeuses","Il faut bien s'habiller","Le costume est important"], "answer": "Les apparences sont trompeuses",
             "expl": "On ne peut pas juger quelqu'un sur son apparence extérieure."},
            {"q": "Marie est plus grande que Julie. Julie est plus grande qu'Anne. La plus petite est :", "choices": ["Marie","Julie","Anne","On ne sait pas"], "answer": "Anne",
             "expl": "Marie > Julie > Anne. La plus petite est donc Anne."},
            {"q": "'Aide-toi et le ciel t'aidera.' Morale :", "choices": ["Il faut prier pour réussir","L'effort personnel est la première condition du succès","Le destin est inévitable","Les autres nous aideront toujours"], "answer": "L'effort personnel est la première condition du succès",
             "expl": "Le soutien extérieur vient en complément de l'initiative personnelle."},
            {"q": "Analogie : Pluie : Parapluie :: Froid : ?", "choices": ["Nuage","Manteau","Hiver","Vent"], "answer": "Manteau",
             "expl": "Le parapluie protège de la pluie, le manteau protège du froid."},
            {"q": "'On ne fait pas d'omelette sans casser des œufs.' Signifie :", "choices": ["Cuisiner est facile","Tout résultat important exige des sacrifices","Les œufs sont fragiles","Il faut éviter les dommages"], "answer": "Tout résultat important exige des sacrifices",
             "expl": "Pour atteindre un but, il faut accepter certains inconvénients ou pertes."},
            {"q": "Si P implique Q, et P est vrai, alors :", "choices": ["Q est faux","Q est vrai","On ne peut pas conclure","P est faux"], "answer": "Q est vrai",
             "expl": "Modus ponens : P→Q et P vrais → Q est vrai."},
            {"q": "'Vouloir c'est pouvoir.' Sens :", "choices": ["Il faut vouloir ce qu'on peut faire","La volonté est le premier moteur du succès","Il faut être fort pour réussir","Le succès dépend de la chance"], "answer": "La volonté est le premier moteur du succès",
             "expl": "La détermination est suffisante pour surmonter les obstacles."},
            {"q": "Analogie : Graine : Arbre :: Œuf : ?", "choices": ["Nid","Oiseau","Coquille","Plume"], "answer": "Oiseau",
             "expl": "La graine donne l'arbre (en mûrissant), l'œuf donne l'oiseau."},
            {"q": "'Après la pluie, le beau temps.' Signifie :", "choices": ["La météo est changeante","Après les épreuves vient la joie","Il pleut souvent","Il faut attendre la pluie"], "answer": "Après les épreuves vient la joie",
             "expl": "Les périodes difficiles sont suivies de périodes meilleures."},
            {"q": "Si aucun chat n'est un chien et Rex est un chien, alors :", "choices": ["Rex est un chat","Rex n'est pas un chat","Rex est peut-être un chat","On ne sait pas"], "answer": "Rex n'est pas un chat",
             "expl": "Aucun chien ne peut être un chat selon la prémisse."},
            {"q": "'Les absents ont toujours tort.' Nuance :", "choices": ["Il ne faut jamais s'absenter","Ceux qui ne sont pas là ne peuvent se défendre","Les absents font des erreurs","La présence est inutile"], "answer": "Ceux qui ne sont pas là ne peuvent se défendre",
             "expl": "En l'absence d'une personne, on lui attribue facilement les torts."},
            {"q": "Analogie : Faiblesse : Défaite :: Force : ?", "choices": ["Guerre","Victoire","Combat","Ennemi"], "answer": "Victoire",
             "expl": "La faiblesse conduit à la défaite, la force conduit à la victoire."},
            {"q": "'On n'apprend pas à un vieux singe à faire des grimaces.' Sens :", "choices": ["Les singes sont intelligents","Un expert n'a pas besoin de leçons","Les vieux apprennent vite","Il faut apprendre tôt"], "answer": "Un expert n'a pas besoin de leçons",
             "expl": "Il est inutile d'enseigner à un expérimenté ce qu'il maîtrise depuis longtemps."},
            {"q": "Si P implique Q, et Q est faux, alors :", "choices": ["P est vrai","P est faux","P est indéterminé","Q implique P"], "answer": "P est faux",
             "expl": "Modus tollens : P→Q, ¬Q → ¬P (contraposée)."},
            {"q": "'Qui trop embrasse mal étreint.' Leçon :", "choices": ["Il ne faut pas embrasser","Vouloir tout faire à la fois nuit à la qualité","Les étreintes sont fortes","Il faut être ambitieux"], "answer": "Vouloir tout faire à la fois nuit à la qualité",
             "expl": "Disperser ses efforts sur trop de projets empêche d'en réussir aucun correctement."},
            {"q": "Analogie : Question : Réponse :: Problème : ?", "choices": ["Difficulté","Solution","Méthode","Erreur"], "answer": "Solution",
             "expl": "Une réponse répond à une question ; une solution résout un problème."},
            {"q": "'Il ne faut pas mettre tous ses œufs dans le même panier.' Sens :", "choices": ["Les œufs sont fragiles","Il faut diversifier pour réduire les risques","Un seul panier suffit","Il ne faut pas avoir de poules"], "answer": "Il faut diversifier pour réduire les risques",
             "expl": "Ne pas concentrer toutes ses ressources ou espoirs en un seul endroit."},
            {"q": "Tous les fruits contiennent des vitamines. La pomme est un fruit. Donc :", "choices": ["La pomme ne contient pas de vitamines","La pomme contient peut-être des vitamines","La pomme contient des vitamines","On ne peut pas conclure"], "answer": "La pomme contient des vitamines",
             "expl": "Syllogisme direct : tous les fruits (pomme incluse) contiennent des vitamines."},
            {"q": "'Mieux vaut tard que jamais.' Signifie :", "choices": ["Il ne faut jamais être en retard","Faire les choses en retard vaut mieux que ne pas les faire","La ponctualité est inutile","Il faut toujours attendre"], "answer": "Faire les choses en retard vaut mieux que ne pas les faire",
             "expl": "Il est préférable d'agir tardivement plutôt que de ne jamais agir."},
            {"q": "Analogie : Cause : Effet :: Effort : ?", "choices": ["Travail","Fatigue","Résultat","Action"], "answer": "Résultat",
             "expl": "Une cause produit un effet ; un effort produit un résultat."},
        ]
    }
}

# =====================================================================
# HEADER
# =====================================================================
st.markdown("""
<div class="header-box">
    <div style="font-size:2.8rem; margin-bottom:0.3rem;">🎓</div>
    <h1 style="color:white; margin:0; font-size:2rem; font-weight:800; letter-spacing:-0.02em;
               text-shadow:0 2px 12px rgba(0,0,0,0.25);">CapAvenir CMR</h1>
    <p style="color:rgba(255,255,255,0.85); margin:0.4rem 0 0 0; font-size:0.95rem; font-weight:300;">
        Système Intelligent d'Orientation Scolaire
    </p>
    <div class="header-badge">ENS Filière Informatique &nbsp;·&nbsp; Niveau 5 &nbsp;·&nbsp; v2.1</div>
</div>
""", unsafe_allow_html=True)

# =====================================================================
# SIDEBAR
# =====================================================================
# ─── Page de connexion conseiller ───────────────────────────────
def afficher_login_conseiller():
    """Modal de connexion pour le mode conseiller."""
    st.markdown("""
    <div style="max-width:400px; margin:4rem auto; padding:2.5rem;
                background:linear-gradient(145deg,#1e293b,#0f172a);
                border-radius:20px; border:1px solid #334155;
                box-shadow:0 20px 60px rgba(0,0,0,0.4);">
        <div style="text-align:center; margin-bottom:1.5rem;">
            <div style="font-size:2.5rem;">🔐</div>
            <div style="color:white; font-size:1.2rem; font-weight:700; margin-top:0.5rem;">
                Accès Conseiller d'Orientation
            </div>
            <div style="color:#94a3b8; font-size:0.82rem; margin-top:0.3rem;">
                Réservé au personnel autorisé
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    with st.form("login_form"):
        login_input = st.text_input("👤 Identifiant", placeholder="conseiller")
        pwd_input   = st.text_input("🔑 Mot de passe", type="password", placeholder="••••••••••")
        submitted   = st.form_submit_button("🔓 Se connecter", use_container_width=True, type="primary")
        if submitted:
            if check_password(login_input, pwd_input):
                st.session_state.conseiller_auth = True
                st.session_state.login_error = ""
                rerun()
            else:
                st.session_state.login_error = "Identifiant ou mot de passe incorrect."
    if st.session_state.get("login_error"):
        st.error(st.session_state.login_error)
    st.caption("Identifiant par défaut : `conseiller` · Mot de passe : `capavenir2025`")

with st.sidebar:
    st.markdown("### ⚙️ Paramètres")

    # ── Mode Conseiller avec authentification ──
    want_conseiller = st.toggle(
        "Mode Conseiller 🔍",
        value=st.session_state.get("mode_conseiller_val", False),
        help="Active les fonctions avancées — connexion requise."
    )
    if want_conseiller and not st.session_state.get("conseiller_auth", False):
        st.session_state.mode_conseiller_val = True
        st.session_state.show_login = True
    elif not want_conseiller:
        st.session_state.mode_conseiller_val = False
        st.session_state.conseiller_auth = False
        st.session_state.show_dashboard = False
    mode_conseiller = want_conseiller and st.session_state.get("conseiller_auth", False)
    st.session_state.mode_conseiller_val = want_conseiller
    dark_mode = st.toggle(
        "Mode sombre 🌙",
        value=st.session_state.get("dark_mode", False),
        help="Basculer entre le thème clair et le thème sombre."
    )
    st.session_state.dark_mode = dark_mode

    # ── Tableau de bord — UNIQUEMENT en mode conseiller ──
    if mode_conseiller:
        st.divider()
        st.markdown("### 📊 Tableau de Bord")
        st.caption("Vue d'ensemble de tous les élèves orientés.")
        if st.button("🗂️ Ouvrir le tableau de bord", use_container_width=True, type="primary"):
            st.session_state.show_dashboard = True
            rerun()
        if st.session_state.get("show_dashboard", False):
            if st.button("↩️ Retour à l'orientation", use_container_width=True):
                st.session_state.show_dashboard = False
                rerun()

    st.divider()
    st.markdown("### 📋 À propos")
    st.caption(
        "CapAvenir CMR v2.1\n"
        "Mémoire ENS Informatique Niv. 5\n"
        "Orientation 3e → 2nde Cameroun\n\n"
        "Tests : D48 · KRX · MECA · BV11 · PRC\n"
        "Étalonnage : normes lycées camerounais"
    )

# Injection CSS (dépend du mode)
inject_css(st.session_state.get("dark_mode", False))

# =====================================================================
# INITIALISATION SESSION STATE
# =====================================================================
defaults = {
    "step": 0,
    "nom": "", "prenom": "", "age": 15, "sexe": "Masculin",
    "lycee": "", "choix_personnel": "C (Scientifique)",
    "projet_pro": "", "revenu": "Moyen (50 000 - 150 000 FCFA/mois)",
    # Notes T1
    "maths_t1": 10.0, "sci_phy_t1": 10.0, "svt_t1": 10.0,
    "francais_t1": 10.0, "histgeo_t1": 10.0, "anglais_t1": 10.0,
    # Notes T2
    "maths_t2": 10.0, "sci_phy_t2": 10.0, "svt_t2": 10.0,
    "francais_t2": 10.0, "histgeo_t2": 10.0, "anglais_t2": 10.0,
    "t2_renseigne": False,
    # Notes T3
    "maths_t3": 10.0, "sci_phy_t3": 10.0, "svt_t3": 10.0,
    "francais_t3": 10.0, "histgeo_t3": 10.0, "anglais_t3": 10.0,
    "t3_renseigne": False,
    # Tests psychotechniques (notes brutes)
    "d48": 10.0, "krx": 10.0, "meca": 10.0, "bv11": 10.0, "prc": 10.0,
    # Passation guidée
    "test_answers": {}, "test_scores": {},
    # Chat & résultats
    "chat_history": [],
    "orientation_finale": None,
    "probation": False,
    "statut": "confirme",       # confirme | attente | probation | revise_a | revise_b | indetermine
    "score_confiance": 0,
    "notes_conseiller": "",
    "show_dashboard": False,
    "mode_conseiller_val": False,
    "conseiller_auth": False,
    "login_error": "",
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# =====================================================================
# FONCTIONS UTILITAIRES
# =====================================================================
def get_notes_actives():
    """Retourne les notes du dernier trimestre renseigné et son label."""
    if st.session_state.t3_renseigne:
        s, label = "_t3", "T3"
    elif st.session_state.t2_renseigne:
        s, label = "_t2", "T2"
    else:
        s, label = "_t1", "T1"
    return {
        "maths":    st.session_state[f"maths{s}"],
        "sci_phy":  st.session_state[f"sci_phy{s}"],
        "svt":      st.session_state[f"svt{s}"],
        "francais": st.session_state[f"francais{s}"],
        "histgeo":  st.session_state[f"histgeo{s}"],
        "anglais":  st.session_state[f"anglais{s}"],
    }, label

def calc_moyennes(notes):
    moy_sci = (notes["maths"] + notes["sci_phy"] + notes["svt"]) / 3
    moy_lit = (notes["francais"] + notes["histgeo"] + notes["anglais"]) / 3
    return moy_sci, moy_lit

def calc_aptitudes():
    SA_brut = (st.session_state.krx + st.session_state.d48) / 2
    LA_brut = (st.session_state.bv11 + st.session_state.prc) / 2
    SA_etal = (etalonner(st.session_state.krx) + etalonner(st.session_state.d48)) / 2
    LA_etal = (etalonner(st.session_state.bv11) + etalonner(st.session_state.prc)) / 2
    return SA_brut, LA_brut, SA_etal, LA_etal

def calc_score_confiance(SA, LA, moy_sci, moy_lit, serie):
    """Score de confiance 0-100 % basé sur l'alignement aptitudes/notes/choix."""
    if serie == "C":
        align  = min(50, max(0, (SA - LA) / 20 * 50))
        perf   = min(50, max(0, (moy_sci - 10) / 10 * 50))
    elif serie == "A":
        align  = min(50, max(0, (LA - SA) / 20 * 50))
        perf   = min(50, max(0, (moy_lit - 10) / 10 * 50))
    else:
        return 30
    return round(min(100, max(0, 50 + align + perf)))

# =====================================================================
# AGENT IA NATUREL — Réponses contextuelles non robotiques
# =====================================================================
def reponse_ia_simulee(user_input, conflit_type, prenom, SA, moy_sci,
                       LA=0, moy_lit=0, projet_pro="", revenu="",
                       chat_history=None, d48=10, krx=10):
    """
    Agent IA simulé qui analyse finement la réponse avant de répondre.
    Ton : chaleureux, direct, jamais robotique. Comme un vrai conseiller.
    """
    txt      = user_input.strip()
    txt_low  = txt.lower()
    tour     = len([m for m in (chat_history or []) if m.get("role") == "user"])
    # Projets reconnus
    P_SCI = ["médecin","médecine","ingénieur","ingénierie","informatique","pharmacie",
             "chirurgien","biologiste","biologie","architecte","pilote","aéronautique",
             "géologie","physicien","chimiste","agronome","vétérinaire","mathématicien"]
    P_LIT = ["avocat","droit","journaliste","journalisme","politique","littérature",
             "histoire","philosophie","diplomate","magistrat","sociologue","communication",
             "linguiste","traducteur","enseignant lettres","administration","lettres"]
    P_ECO = ["économiste","économie","commerce","gestion","comptable","comptabilité",
             "banquier","finance","entrepreneur","business","marketing","management"]
    proj_c = (projet_pro or "").lower()
    det_sci = any(p in txt_low or p in proj_c for p in P_SCI)
    det_lit = any(p in txt_low or p in proj_c for p in P_LIT)
    det_eco = any(p in txt_low or p in proj_c for p in P_ECO)
    pos = any(w in txt_low for w in ["oui","vais","améliorer","travailler","effort",
              "promets","essayer","d'accord","ok","motivé","déterminé","ferai",
              "capable","confiance","compte","m'engage"])
    neg = any(w in txt_low for w in ["non","difficile","pas","incapable","perdu",
              "démotivé","fatigue","abandonne","impossible","peur","nul","rien compris",
              "décourage","espoir"])
    fam = any(w in txt_low for w in ["parent","père","mère","famille","veulent",
              "forcent","obligé","pression","argent","pauvre","pas les moyens"])
    maths_diff = any(w in txt_low for w in ["maths","mathématiques","physique","calcul",
                     "formule","équation","exercice","problème de maths"])
    meth_pb = any(w in txt_low for w in ["concentrer","attention","distrait","réviser",
               "méthode","organiser","comprends pas","pas compris","retenir"])
    est_question = "?" in txt or any(w in txt_low for w in ["pourquoi","comment",
                   "qu'est","c'est quoi","que faire","quel","qui","lequel"])
    obj = f"{min(20.0, round(moy_sci + 3, 1)):.1f}"

    # ══ CAS DÉCALÉ : profil sci mais veut la A ══
    if conflit_type == "decale":
        if tour == 0:
            if det_sci:
                return (f"Tiens, c'est intéressant {prenom} ! Ton projet et tes résultats de tests "
                        f"vont dans la même direction — tu as un vrai potentiel scientifique. "
                        f"Qu'est-ce qui t'a poussé à cocher la série A malgré ça ?")
            if det_lit:
                metier = next((p for p in P_LIT if p in txt_low or p in proj_c), "ce domaine")
                return (f"Je comprends parfaitement ton intérêt pour {metier}, {prenom}. "
                        f"Ce qui est fascinant dans ton cas, c'est que tes tests montrent "
                        f"un fort raisonnement logique — une qualité précieuse même en droit ou en communication. "
                        f"As-tu pensé que certains de ces métiers se rejoignent depuis les deux séries ?")
            if det_eco:
                return (f"L'économie, c'est un excellent choix de carrière, {prenom}. "
                        f"Et tes aptitudes (SA={SA:.1f}/20) te donnent un avantage réel en économétrie, "
                        f"finance quantitative ou gestion. Pourquoi exclure la série C pour y arriver ?")
            if fam:
                return (f"La famille qui a un avis sur ton orientation... ça peut peser lourd. "
                        f"Mais là, tes résultats parlent objectivement : SA={SA:.1f}/20. "
                        f"Ce n'est pas une opinion, c'est une mesure. "
                        f"Penses-tu qu'on pourrait expliquer ça à tes parents ensemble ?")
            if est_question:
                return (f"Bonne question, {prenom}. Concrètement : la série C ne ferme aucune porte. "
                        f"Elle en ouvre même beaucoup en sciences, en économie, et — oui — en droit aussi. "
                        f"Dis-moi quel métier t'attire vraiment, et on verra la meilleure route ensemble.")
            return (f"Merci, {prenom}. Avec SA={SA:.1f}/20 contre LA={LA:.1f}/20, "
                    f"ton profil penche clairement du côté scientifique. "
                    f"Je veux juste comprendre ce qui te pousse vers la série A — "
                    f"c'est une décision importante et je veux que tu la prennes pour les bonnes raisons.")
        else:
            if pos:
                return (f"Voilà qui me rassure, {prenom}. "
                        f"Tu prends cette décision avec lucidité — c'est le plus important. "
                        f"Tes aptitudes sont là, le reste c'est du travail. Bon courage pour la suite !")
            if neg:
                return (f"Prends le temps d'y réfléchir, {prenom}, sans pression. "
                        f"Parles-en aussi avec tes parents et un conseiller de ton lycée. "
                        f"Tes résultats de tests resteront disponibles pour t'aider dans cette réflexion.")
            return (f"C'est une décision mûrement réfléchie, {prenom}. "
                    f"L'essentiel est que tu choisisses pour toi — pas pour faire plaisir à quelqu'un. "
                    f"Y a-t-il encore quelque chose qui te préoccupe ?")

    # ══ CAS RÊVEUR : bonnes aptitudes, notes faibles ══
    elif conflit_type == "reveur":
        if tour == 0:
            if maths_diff:
                return (f"Les maths qui bloquent — ça arrive à beaucoup d'élèves brillants, {prenom}. "
                        f"Et ton score D48={d48:.1f}/20 montre que ta logique est là. "
                        f"Souvent c'est une question de méthode, pas de capacité. "
                        f"Tu révises comment ? Après chaque cours ou juste avant les contrôles ?")
            if meth_pb:
                return (f"La concentration, c'est souvent la vraie bataille, {prenom}. "
                        f"Ce n'est pas une question d'intelligence — ton SA={SA:.1f}/20 le prouve. "
                        f"Est-ce que tu travailles dans un endroit calme ? "
                        f"Et est-ce que tu relis tes cours dans les 24h qui suivent ?")
            if fam:
                rev_low = (revenu or "").lower()
                if "faible" in rev_low:
                    return (f"Je comprends la contrainte financière, {prenom}. "
                            f"Mais sache qu'il existe des soutiens gratuits — programmes lycées, associations, "
                            f"élèves de terminale qui donnent des cours. Ton potentiel (SA={SA:.1f}/20) vaut "
                            f"vraiment qu'on cherche une solution ensemble.")
                return (f"Ta famille peut être un soutien précieux ici, {prenom}. "
                        f"Si tu leur montres tes résultats de tests (SA={SA:.1f}/20), "
                        f"ils comprendront que ce n'est pas un manque de capacité mais besoin d'accompagnement.")
            if pos:
                return (f"C'est exactement l'état d'esprit qu'il faut, {prenom} ! "
                        f"Avec cette détermination et tes capacités (SA={SA:.1f}/20), "
                        f"l'objectif {obj}/20 en sciences au T2 est tout à fait réalisable. "
                        f"Qu'est-ce qui te pose le plus de difficultés concrètement en ce moment ?")
            if neg:
                return (f"Je t'entends, {prenom}, et je ne minimise pas tes difficultés. "
                        f"Mais j'ai devant moi tes résultats : SA={SA:.1f}/20. "
                        f"Ces chiffres ne mentent pas — tu en es capable. "
                        f"Qu'est-ce qui s'est passé cette année pour que tes notes ne reflètent pas ça ?")
            if est_question:
                return (f"Très bonne question, {prenom}. Le décalage entre aptitudes ({SA:.1f}/20) "
                        f"et notes scolaires ({moy_sci:.1f}/20) a souvent une explication précise : "
                        f"méthode de travail, environnement, motivation, ou parfois un événement personnel. "
                        f"Est-ce qu'il s'est passé quelque chose cette année qui a perturbé ta scolarité ?")
            return (f"Pour comprendre la situation, {prenom}, j'ai besoin d'un peu plus de contexte. "
                    f"Tu as des capacités réelles (SA={SA:.1f}/20) mais tes notes ({moy_sci:.1f}/20) "
                    f"ne le montrent pas. Comment se passe une semaine normale de travail pour toi ?")
        else:
            if pos:
                return (f"C'est ce qu'on attendait de toi, {prenom} ! "
                        f"Objectif T2 : {obj}/20 en sciences. Ton dossier passe en suivi — "
                        f"le conseiller reviendra vers toi au prochain trimestre. Tiens bon !")
            if neg:
                return (f"C'est difficile, je comprends. Mais ce n'est pas fini, {prenom}. "
                        f"Ton dossier reste ouvert en suivi. Parle à ton enseignant principal "
                        f"— il peut souvent proposer un accompagnement personnalisé.")
            return (f"Ton dossier est bien pris en compte, {prenom}. "
                    f"Le conseiller suivra ton évolution au prochain trimestre. "
                    f"N'hésite pas à revenir si tu as besoin d'un accompagnement.")

    # Réponse de clôture
    if est_question:
        return (f"C'est une très bonne question, {prenom}. "
                f"Je te conseille d'en discuter aussi en face à face avec le conseiller de ton lycée "
                f"— il pourra aller plus loin dans l'accompagnement. "
                f"Est-ce qu'il y a autre chose qui te préoccupe ?")
    return (f"Merci pour cette précision, {prenom}. "
            f"Tout ce qu'on a échangé ici est pris en compte dans ton dossier. "
            f"Le conseiller dispose maintenant des éléments pour t'accompagner au mieux.")


# =====================================================================
# FONCTIONS EXPORTS DASHBOARD
# =====================================================================
def _generer_excel(dossiers):
    if not OPENPYXL_OK:
        return None
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Dossiers CapAvenir"
    hf = PatternFill("solid", fgColor="0F172A")
    hfont = Font(bold=True, color="FFFFFF", size=9)
    alt_fill = PatternFill("solid", fgColor="F8FAFC")
    headers = ["Nom","Prénom","Lycée","Âge","Sexe","Choix","Projet Pro",
               "SA","LA","Moy.Sci T1","Moy.Lit T1","D48","KRX","MECA","BV11","PRC",
               "Série","Statut","Confiance %","Trimestre","Date"]
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(1, ci, h)
        cell.fill = hf; cell.font = hfont
        cell.alignment = Alignment(horizontal="center")
    for ri, d in enumerate(dossiers, 2):
        serie = (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?")
        row = [
            d.get("nom",""), d.get("prenom",""), d.get("lycee",""),
            d.get("age",""), d.get("sexe",""), d.get("choix_personnel",""),
            d.get("projet_pro",""),
            round(float(d.get("SA_etal") or d.get("SA_brut") or 0), 2),
            round(float(d.get("LA_etal") or d.get("LA_brut") or 0), 2),
            round(float(d.get("moy_sci_t1") or 0), 2),
            round(float(d.get("moy_lit_t1") or 0), 2),
            d.get("d48_brut",""), d.get("krx_brut",""), d.get("meca_brut",""),
            d.get("bv11_brut",""), d.get("prc_brut",""),
            serie, d.get("statut",""), d.get("score_confiance",""),
            d.get("trimestre_actuel") or d.get("trimestre_decision",""),
            (d.get("date_modification","") or "")[:10],
        ]
        for ci, val in enumerate(row, 1):
            cell = ws.cell(ri, ci, val)
            if ri % 2 == 0:
                cell.fill = alt_fill
            cell.alignment = Alignment(horizontal="center" if ci > 7 else "left")
    # Largeurs colonnes
    widths = [14,12,20,5,8,18,24,6,6,10,10,6,6,6,6,6,7,12,10,10,12]
    for ci, w in enumerate(widths, 1):
        ws.column_dimensions[ws.cell(1,ci).column_letter].width = w
    # Figer la première ligne
    ws.freeze_panes = "A2"
    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    return buf.getvalue()


def _generer_word(dossiers, stats):
    if not DOCX_OK:
        return None
    doc = DocxDocument()
    doc.core_properties.title = "Rapport CapAvenir CMR"
    # Titre
    h = doc.add_heading("Rapport d'Orientation — CapAvenir CMR", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        f" · {len(dossiers)} dossier(s) au total"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    # Stats
    doc.add_heading("Statistiques globales", level=1)
    tbl = doc.add_table(rows=1, cols=2); tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Indicateur"
    tbl.rows[0].cells[1].text = "Valeur"
    for lbl, key in [("Total dossiers","total"),("Confirmés","nb_confirmes"),
                     ("En attente","nb_attente"),("Probation","nb_probation"),
                     ("Révisés","nb_revises"),("Série C","nb_serie_c"),("Série A","nb_serie_a")]:
        r = tbl.add_row().cells
        r[0].text = lbl; r[1].text = str(stats.get(key,0))
    doc.add_paragraph()
    # Liste dossiers
    doc.add_heading("Liste des dossiers", level=1)
    tbl2 = doc.add_table(rows=1, cols=7); tbl2.style = "Table Grid"
    for i, h in enumerate(["Nom & Prénom","Lycée","SA","LA","Série","Statut","Confiance"]):
        tbl2.rows[0].cells[i].text = h
        tbl2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for d in dossiers:
        serie = (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?")
        r = tbl2.add_row().cells
        r[0].text = f"{(d.get('nom','') or '').upper()} {(d.get('prenom','') or '').capitalize()}"
        r[1].text = (d.get("lycee","") or "—")[:28]
        r[2].text = f"{float(d.get('SA_etal') or d.get('SA_brut') or 0):.1f}"
        r[3].text = f"{float(d.get('LA_etal') or d.get('LA_brut') or 0):.1f}"
        r[4].text = serie; r[5].text = d.get("statut","—")
        r[6].text = f"{d.get('score_confiance',0) or 0} %"
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.getvalue()


def _generer_pdf_rapport(dossiers, stats):
    """PDF récapitulatif de tous les dossiers."""
    buf = io.BytesIO()
    CMR_DARK  = colors.HexColor("#0f172a")
    CMR_GREEN = colors.HexColor("#10b981")
    CMR_GRAY  = colors.HexColor("#64748b")
    CMR_LIGHT = colors.HexColor("#f8fafc")
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=1.5*cm, rightMargin=1.5*cm,
        topMargin=1.5*cm, bottomMargin=1.5*cm)
    sT  = ParagraphStyle("T",  fontName="Helvetica-Bold", fontSize=16, textColor=CMR_DARK, alignment=TA_CENTER, spaceAfter=4)
    sSub= ParagraphStyle("S",  fontName="Helvetica", fontSize=9, textColor=CMR_GRAY,  alignment=TA_CENTER, spaceAfter=8)
    sSec= ParagraphStyle("Se", fontName="Helvetica-Bold", fontSize=10, textColor=CMR_GREEN, spaceBefore=12, spaceAfter=4)
    sB  = ParagraphStyle("B",  fontName="Helvetica", fontSize=8, textColor=CMR_DARK)
    story = []
    story.append(Paragraph("RAPPORT D'ORIENTATION — CapAvenir CMR", sT))
    story.append(Paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} · {len(dossiers)} dossier(s)", sSub))
    story.append(HRFlowable(width="100%", thickness=2, color=CMR_GREEN, spaceAfter=8))
    # Stats
    story.append(Paragraph("STATISTIQUES", sSec))
    stat_rows = [["Total","Confirmés","En attente","Probation","Révisés","Série C","Série A"],
                 [str(stats.get(k,0)) for k in ["total","nb_confirmes","nb_attente","nb_probation","nb_revises","nb_serie_c","nb_serie_a"]]]
    st_tbl = Table(stat_rows, colWidths=[2.5*cm]*7)
    st_tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),CMR_DARK),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),8),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("BACKGROUND",(0,1),(-1,1),CMR_LIGHT),
        ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ]))
    story.append(st_tbl)
    # Liste
    story.append(Paragraph("LISTE DES DOSSIERS", sSec))
    rows = [["Nom & Prénom","Lycée","SA","LA","Série","Statut","Confiance","Date"]]
    for d in dossiers:
        serie = (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?")
        rows.append([
            f"{(d.get('nom','') or '').upper()} {(d.get('prenom','') or '').capitalize()}"[:22],
            (d.get("lycee","") or "—")[:18],
            f"{float(d.get('SA_etal') or d.get('SA_brut') or 0):.1f}",
            f"{float(d.get('LA_etal') or d.get('LA_brut') or 0):.1f}",
            serie, d.get("statut","—")[:10],
            f"{d.get('score_confiance',0) or 0}%",
            (d.get("date_modification","") or "")[:10],
        ])
    list_tbl = Table(rows, colWidths=[4.5*cm,3.5*cm,1.3*cm,1.3*cm,1.4*cm,2.4*cm,1.8*cm,2.3*cm])
    list_tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),CMR_DARK),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),7.5),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[CMR_LIGHT,colors.white]),
        ("GRID",(0,0),(-1,-1),0.2,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),4),
    ]))
    story.append(list_tbl)
    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


# =====================================================================
# TABLEAU DE BORD — Mode Conseiller uniquement
# =====================================================================
def afficher_dashboard():
    STATUT_META = {
        "confirme":    ("✅ Confirmés",    "#10b981","#d1fae5","badge-confirme"),
        "revise":      ("🔄 Révisés",      "#ec4899","#fce7f3","badge-revise"),
        "attente":     ("⏳ En attente",   "#f97316","#fff7ed","badge-attente"),
        "probation":   ("⚠️ Probation",    "#f59e0b","#fef3c7","badge-probation"),
        "indetermine": ("❓ Indéterminés", "#6b7280","#f3f4f6","badge-indetermine"),
    }

    # ── Titre ──
    st.markdown("""
    <div style="display:flex;align-items:center;gap:0.8rem;margin-bottom:1.5rem;">
        <div style="font-size:2.2rem;">📊</div>
        <div>
            <div style="font-size:1.4rem;font-weight:800;">Tableau de Bord Conseiller</div>
            <div style="font-size:0.82rem;opacity:0.6;margin-top:2px;">
                Vue réservée — Données en temps réel depuis la base de données
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

    # ── Stats ──
    try:
        stats = db.statistiques()
    except Exception:
        stats = {"total":0,"nb_confirmes":0,"nb_attente":0,"nb_probation":0,
                 "nb_revises":0,"nb_serie_c":0,"nb_serie_a":0}

    total = stats.get("total", 0)
    stat_items = [
        ("Total",       total,                      "#6d28d9","#f3e8ff"),
        ("✅ Confirmés", stats.get("nb_confirmes",0), "#10b981","#d1fae5"),
        ("⏳ Attente",   stats.get("nb_attente",0),   "#f97316","#fff7ed"),
        ("⚠️ Probation", stats.get("nb_probation",0), "#f59e0b","#fef3c7"),
        ("🔄 Révisés",   stats.get("nb_revises",0),   "#ec4899","#fce7f3"),
    ]
    cols_s = st.columns(5)
    for col, (lbl, val, color, bg) in zip(cols_s, stat_items):
        with col:
            st.markdown(f"""
            <div class="dash-stat-card">
                <div class="dash-stat-val" style="color:{color};">{val}</div>
                <div class="dash-stat-lbl">{lbl}</div>
            </div>""", unsafe_allow_html=True)

    # ── Graphiques ──
    if total > 0:
        st.write("")
        cg1, cg2 = st.columns([1,2])
        with cg1:
            nb_c = stats.get("nb_serie_c",0)
            nb_a = stats.get("nb_serie_a",0)
            fig_p = go.Figure(go.Pie(
                labels=["Série C","Série A","Indét."],
                values=[nb_c, nb_a, max(0, total-nb_c-nb_a)],
                hole=0.55, textinfo="label+percent", textfont_size=10,
                marker=dict(colors=["#10b981","#3b82f6","#94a3b8"],
                            line=dict(color="white",width=2)),
            ))
            fig_p.update_layout(showlegend=False, height=190,
                margin=dict(l=5,r=5,t=15,b=5),
                annotations=[dict(text=f"<b>{total}</b>",x=0.5,y=0.5,
                                  font_size=18,showarrow=False)])
            st.plotly_chart(fig_p, use_container_width=True)
        with cg2:
            fig_b = go.Figure(go.Bar(
                x=["Confirmés","Attente","Probation","Révisés"],
                y=[stats.get(k,0) for k in ["nb_confirmes","nb_attente","nb_probation","nb_revises"]],
                marker_color=["#10b981","#f97316","#f59e0b","#ec4899"],
                text=[stats.get(k,0) for k in ["nb_confirmes","nb_attente","nb_probation","nb_revises"]],
                textposition="outside",
            ))
            fig_b.update_layout(height=190,margin=dict(l=5,r=5,t=15,b=5),
                yaxis=dict(visible=False),showlegend=False,
                plot_bgcolor="rgba(0,0,0,0)",paper_bgcolor="rgba(0,0,0,0)")
            st.plotly_chart(fig_b, use_container_width=True)

    st.divider()

    # ── Filtres ──
    fc1, fc2, fc3 = st.columns([2,1,1])
    with fc1:
        search_q = st.text_input("🔍 Rechercher (nom ou prénom)", placeholder="Ex : MBALLA, Jean...", key="dash_search")
    with fc2:
        filtre_statut = st.selectbox("Statut", ["Tous","confirme","revise","attente","probation","indetermine"], key="dash_filtre")
    with fc3:
        filtre_serie = st.selectbox("Série", ["Toutes","C","A","?"], key="dash_serie")

    # ── Chargement ──
    try:
        statuts_list = ["confirme","revise","attente","probation","indetermine"] if filtre_statut=="Tous" else [filtre_statut]
        tous_dossiers = []
        for s in statuts_list:
            for d in db.lister_dossiers(s):
                d["_statut_affiche"] = s
                tous_dossiers.append(d)
    except Exception as e:
        st.error(f"Erreur lecture BDD : {e}")
        tous_dossiers = []

    if search_q.strip():
        q = search_q.strip().upper()
        tous_dossiers = [d for d in tous_dossiers if
            q in (d.get("nom","") or "").upper() or
            q in (d.get("prenom","") or "").upper()]
    if filtre_serie != "Toutes":
        tous_dossiers = [d for d in tous_dossiers if
            (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?") == filtre_serie]

    # ── Tableau ──
    st.markdown(f"**{len(tous_dossiers)} dossier(s)**")
    if not tous_dossiers:
        st.info("Aucun dossier ne correspond aux critères.")
    else:
        h_cols = st.columns([2.5,2,2.5,1.2,1.2,1.5,1.5,1.8])
        for col, lbl in zip(h_cols, ["Nom & Prénom","Lycée","Projet Pro","SA","LA","Série","Statut","Date"]):
            col.markdown(f"<div style='font-size:0.7rem;font-weight:700;text-transform:uppercase;opacity:0.5;letter-spacing:0.06em;'>{lbl}</div>", unsafe_allow_html=True)
        st.markdown("<hr style='margin:0.3rem 0 0.5rem 0;opacity:0.25;'>", unsafe_allow_html=True)
        for d in tous_dossiers:
            nom_a   = (d.get("nom","") or "").upper()
            prn_a   = (d.get("prenom","") or "").capitalize()
            lycee_a = (d.get("lycee","") or "—")[:22]
            proj_a  = (d.get("projet_pro","") or "—")[:22]
            sa_a    = float(d.get("SA_etal") or d.get("SA_brut") or 0)
            la_a    = float(d.get("LA_etal") or d.get("LA_brut") or 0)
            serie_a = (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?")
            stat_a  = d.get("_statut_affiche","—")
            date_a  = (d.get("date_modification","") or "")[:10]
            _, _, _, s_badge = STATUT_META.get(stat_a, ("—","#6b7280","#f3f4f6","badge-indetermine"))
            b_serie = "badge-c" if serie_a=="C" else ("badge-a" if serie_a=="A" else "")
            r = st.columns([2.5,2,2.5,1.2,1.2,1.5,1.5,1.8])
            r[0].markdown(f"**{nom_a}** {prn_a}")
            r[1].markdown(f"<small>{lycee_a}</small>", unsafe_allow_html=True)
            r[2].markdown(f"<small>{proj_a}</small>", unsafe_allow_html=True)
            r[3].markdown(f"<b style='color:#10b981;'>{sa_a:.1f}</b>", unsafe_allow_html=True)
            r[4].markdown(f"<b style='color:#3b82f6;'>{la_a:.1f}</b>", unsafe_allow_html=True)
            r[5].markdown(f'<span class="dash-badge {b_serie}">{serie_a}</span>', unsafe_allow_html=True)
            r[6].markdown(f'<span class="dash-badge {s_badge}">{stat_a}</span>', unsafe_allow_html=True)
            r[7].markdown(f"<small style='opacity:0.6;'>{date_a}</small>", unsafe_allow_html=True)
            st.markdown("<div style='border-bottom:1px solid rgba(0,0,0,0.05);margin:2px 0;'></div>", unsafe_allow_html=True)

    # ── Mise à jour T2/T3 ──
    st.divider()
    st.markdown("#### 📝 Mettre à jour les notes — Dossiers non finalisés")
    non_conf = [d for d in tous_dossiers if d.get("_statut_affiche") in ("attente","probation")]
    if not non_conf:
        st.info("Aucun dossier en attente ou en probation.")
    else:
        for d in non_conf:
            nom_d    = (d.get("nom","") or "").upper()
            prn_d    = (d.get("prenom","") or "").capitalize()
            stat_d   = d.get("_statut_affiche","")
            sa_d     = float(d.get("SA_etal") or d.get("SA_brut") or 0)
            la_d     = float(d.get("LA_etal") or d.get("LA_brut") or 0)
            trim_cur = d.get("trimestre_actuel") or d.get("trimestre_decision","T1")
            nxt_trim = "T3" if trim_cur in ("T2","T3") else "T2"
            dkey     = str(d.get("id","x")) + nom_d

            with st.expander(f"📋 {nom_d} {prn_d}  ·  {stat_d.upper()}  ·  Saisir notes {nxt_trim}"):
                st.markdown(f"SA={sa_d:.1f}/20 · LA={la_d:.1f}/20 · Trimestre actuel : **{trim_cur}**")
                ca, cb = st.columns(2)
                with ca:
                    st.markdown("**Matières scientifiques**")
                    t_k = nxt_trim.lower()
                    m_v  = st.number_input(f"Maths",     0.0,20.0,float(d.get(f"maths_{t_k}") or 10),0.5,key=f"m_{dkey}")
                    sp_v = st.number_input(f"Sci.Phy",   0.0,20.0,float(d.get(f"sci_phy_{t_k}") or 10),0.5,key=f"sp_{dkey}")
                    sv_v = st.number_input(f"SVT",       0.0,20.0,float(d.get(f"svt_{t_k}") or 10),0.5,key=f"sv_{dkey}")
                    ms_n = round((m_v+sp_v+sv_v)/3,2)
                    st.metric(f"Moy. Sci. {nxt_trim}", f"{ms_n:.2f}/20",
                              delta=f"{ms_n-10:+.1f} vs seuil", delta_color="normal" if ms_n>=10 else "inverse")
                with cb:
                    st.markdown("**Matières littéraires**")
                    fr_v = st.number_input(f"Français",  0.0,20.0,float(d.get(f"francais_{t_k}") or 10),0.5,key=f"fr_{dkey}")
                    hg_v = st.number_input(f"Hist-Géo",  0.0,20.0,float(d.get(f"histgeo_{t_k}") or 10),0.5,key=f"hg_{dkey}")
                    an_v = st.number_input(f"Anglais",   0.0,20.0,float(d.get(f"anglais_{t_k}") or 10),0.5,key=f"an_{dkey}")
                    ml_n = round((fr_v+hg_v+an_v)/3,2)
                    st.metric(f"Moy. Lit. {nxt_trim}", f"{ml_n:.2f}/20",
                              delta=f"{ml_n-10:+.1f} vs seuil", delta_color="normal" if ml_n>=10 else "inverse")

                # Prévisualisation automatique du nouveau statut
                serie_cur = (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "C")
                ok = (ms_n >= 10) if sa_d >= la_d else (ml_n >= 10)
                if stat_d == "attente":
                    new_stat = "confirme" if ok else "probation"
                    new_lbl  = "✅ CONFIRMÉ" if ok else "⚠️ PROBATION"
                    new_serie= serie_cur
                else:  # probation
                    new_stat = "confirme" if ok else "revise"
                    new_lbl  = "✅ CONFIRMÉ" if ok else "🔄 RÉVISÉ → " + ("A" if serie_cur=="C" else "C")
                    new_serie= serie_cur if ok else ("A" if serie_cur=="C" else "C")

                ok_color = "#065f46" if ok else "#7f1d1d"
                ok_bg    = "#d1fae5" if ok else "#fee2e2"
                st.markdown(f"""
                <div style="background:{ok_bg};border-radius:10px;padding:0.55rem 1rem;
                            margin-top:0.5rem;text-align:center;font-weight:700;color:{ok_color};">
                    Prévisualisation : {new_lbl}
                </div>""", unsafe_allow_html=True)

                if st.button(f"💾 Valider et mettre à jour dans la BDD",
                             use_container_width=True, type="primary", key=f"upd_{dkey}"):
                    ss_u = dict(d)
                    ss_u[f"maths_{t_k}"]   = m_v;  ss_u[f"sci_phy_{t_k}"] = sp_v; ss_u[f"svt_{t_k}"] = sv_v
                    ss_u[f"francais_{t_k}"]= fr_v; ss_u[f"histgeo_{t_k}"]  = hg_v; ss_u[f"anglais_{t_k}"] = an_v
                    t_num = nxt_trim[1]
                    ss_u[f"t{t_num}_renseigne"] = True
                    ss_u["statut"]            = new_stat
                    ss_u["orientation_finale"]= new_serie
                    ss_u["revenu"]            = d.get("revenu_famille","")
                    ss_u["nom"]               = d.get("nom","")
                    ss_u["prenom"]            = d.get("prenom","")
                    res = db.sauvegarder_dossier(ss_u, None)
                    if res["succes"]:
                        st.success(f"✅ Statut mis à jour → **{new_stat.upper()}** | {res['message']}")
                        rerun()
                    else:
                        st.error(res["message"])

    # ── Exports multi-formats ──
    st.divider()
    st.markdown("#### 📤 Exports")
    try:
        all_exp = []
        for s in ["confirme","revise","attente","probation","indetermine"]:
            for dd in db.lister_dossiers(s):
                dd["_statut_affiche"] = s
                all_exp.append(dd)
        stats_exp = db.statistiques()
    except Exception:
        all_exp = []; stats_exp = {}

    e1, e2, e3, e4 = st.columns(4)
    with e1:
        try:
            xls = _generer_excel(all_exp)
            if xls:
                st.download_button("📊 Excel (.xlsx)", xls,
                    file_name=f"capavenir_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True)
            else:
                st.caption("pip install openpyxl")
        except Exception as ex:
            st.caption(f"Excel: {ex}")
    with e2:
        try:
            wrd = _generer_word(all_exp, stats_exp)
            if wrd:
                st.download_button("📝 Word (.docx)", wrd,
                    file_name=f"capavenir_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            else:
                st.caption("pip install python-docx")
        except Exception as ex:
            st.caption(f"Word: {ex}")
    with e3:
        try:
            pdf_r = _generer_pdf_rapport(all_exp, stats_exp)
            st.download_button("📄 PDF Rapport", pdf_r,
                file_name=f"capavenir_{datetime.now().strftime('%Y%m%d')}.pdf",
                mime="application/pdf", use_container_width=True)
        except Exception as ex:
            st.caption(f"PDF: {ex}")
    with e4:
        try:
            j = db.exporter_json()
            st.download_button("📦 JSON", j,
                file_name=f"capavenir_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json", use_container_width=True)
        except Exception:
            st.caption("JSON indisponible")


# =====================================================================
# TABLEAU DE BORD — Mode Conseiller uniquement
# =====================================================================
def generer_excel_export(dossiers):
    if not OPENPYXL_OK:
        return None
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Dossiers CapAvenir"
    hfill = PatternFill("solid", fgColor="0F172A")
    hfont = Font(bold=True, color="FFFFFF", size=10)
    entetes = ["Nom","Prénom","Lycée","Âge","Sexe","Choix","Projet",
               "SA","LA","Moy.Sci","Moy.Lit","D48","KRX","MECA","BV11","PRC",
               "Série","Statut","Confiance%","Trimestre","Date"]
    for ci, h in enumerate(entetes, 1):
        c = ws.cell(1, ci, h)
        c.fill = hfill; c.font = hfont
        c.alignment = Alignment(horizontal="center")
    for ri, d in enumerate(dossiers, 2):
        serie = d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?"
        vals = [
            d.get("nom",""), d.get("prenom",""), d.get("lycee",""),
            d.get("age",""), d.get("sexe",""), d.get("choix_personnel",""), d.get("projet_pro",""),
            round(float(d.get("SA_etal") or d.get("SA_brut") or 0),2),
            round(float(d.get("LA_etal") or d.get("LA_brut") or 0),2),
            round(float(d.get("moy_sci_t1") or 0),2), round(float(d.get("moy_lit_t1") or 0),2),
            d.get("d48_brut",""), d.get("krx_brut",""), d.get("meca_brut",""),
            d.get("bv11_brut",""), d.get("prc_brut",""),
            serie, d.get("statut",""), d.get("score_confiance",""),
            d.get("trimestre_actuel") or d.get("trimestre_decision",""),
            (d.get("date_modification","") or "")[:10],
        ]
        for ci, v in enumerate(vals, 1):
            ws.cell(ri, ci, v)
    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = 15
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf.getvalue()


def generer_word_export(dossiers, stats):
    if not DOCX_OK:
        return None
    doc = DocxDocument()
    doc.core_properties.title = "Rapport CapAvenir CMR"
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    titre = doc.add_heading("Rapport d'Orientation — CapAvenir CMR", 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — {len(dossiers)} dossier(s)"
    )
    doc.add_heading("Statistiques globales", level=1)
    tbl = doc.add_table(rows=1, cols=2); tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Indicateur"
    tbl.rows[0].cells[1].text = "Valeur"
    for lbl, key in [("Total","total"),("Confirmés","nb_confirmes"),
                     ("En attente","nb_attente"),("Probation","nb_probation"),
                     ("Révisés","nb_revises"),("Série C","nb_serie_c"),("Série A","nb_serie_a")]:
        r = tbl.add_row().cells; r[0].text = lbl; r[1].text = str(stats.get(key,0))
    doc.add_heading("Liste des dossiers", level=1)
    t2 = doc.add_table(rows=1, cols=6); t2.style = "Table Grid"
    for i, h in enumerate(["Nom & Prénom","Lycée","SA","LA","Série","Statut"]):
        t2.rows[0].cells[i].text = h
    for d in dossiers:
        serie = d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?"
        r = t2.add_row().cells
        r[0].text = f"{(d.get('nom') or '').upper()} {(d.get('prenom') or '').capitalize()}"
        r[1].text = (d.get("lycee") or "—")[:25]
        r[2].text = f"{float(d.get('SA_etal') or d.get('SA_brut') or 0):.1f}"
        r[3].text = f"{float(d.get('LA_etal') or d.get('LA_brut') or 0):.1f}"
        r[4].text = serie; r[5].text = d.get("statut","—")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def generer_pdf_rapport(dossiers, stats):
    """PDF récapitulatif de tous les dossiers."""
    buf = io.BytesIO()
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    doc2 = SimpleDocTemplate(buf, pagesize=A4,
           leftMargin=1.5*cm, rightMargin=1.5*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
    C_DARK  = colors.HexColor("#0f172a")
    C_GREEN = colors.HexColor("#10b981")
    C_GRAY  = colors.HexColor("#64748b")
    sT  = ParagraphStyle("t", fontName="Helvetica-Bold", fontSize=14, textColor=C_DARK, alignment=TA_CENTER, spaceAfter=4)
    sSub= ParagraphStyle("s", fontName="Helvetica", fontSize=8, textColor=C_GRAY, alignment=TA_CENTER, spaceAfter=2)
    sB  = ParagraphStyle("b", fontName="Helvetica", fontSize=8, textColor=C_DARK, spaceAfter=2)
    story2 = []
    story2.append(Paragraph("RAPPORT D'ORIENTATION — CapAvenir CMR", sT))
    story2.append(Paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — {len(dossiers)} dossier(s)", sSub))
    story2.append(Spacer(1, 0.3*cm))
    stat_rows = [["Confirmés","En attente","Probation","Révisés","Série C","Série A"],
                 [str(stats.get("nb_confirmes",0)), str(stats.get("nb_attente",0)),
                  str(stats.get("nb_probation",0)), str(stats.get("nb_revises",0)),
                  str(stats.get("nb_serie_c",0)), str(stats.get("nb_serie_a",0))]]
    ts = Table(stat_rows, colWidths=[2.8*cm]*6)
    ts.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),C_DARK),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),8),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ]))
    story2.append(ts)
    story2.append(Spacer(1, 0.3*cm))
    hdr_row = [["Nom & Prénom","Lycée","SA","LA","Série","Statut","Confiance","Date"]]
    data_rows = []
    for d in dossiers:
        serie = d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?"
        data_rows.append([
            f"{(d.get('nom') or '').upper()} {(d.get('prenom') or '').capitalize()}",
            (d.get("lycee") or "—")[:18],
            f"{float(d.get('SA_etal') or d.get('SA_brut') or 0):.1f}",
            f"{float(d.get('LA_etal') or d.get('LA_brut') or 0):.1f}",
            serie, d.get("statut","—"),
            f"{d.get('score_confiance',0) or 0}%",
            (d.get("date_modification","") or "")[:10],
        ])
    tbl_dos = Table(hdr_row + data_rows,
                    colWidths=[4.5*cm, 3*cm, 1.2*cm, 1.2*cm, 1.3*cm, 2.2*cm, 1.8*cm, 2*cm])
    tbl_dos.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),C_DARK),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),7.5),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.HexColor("#f8fafc"),colors.white]),
        ("GRID",(0,0),(-1,-1),0.2,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),4),
    ]))
    story2.append(tbl_dos)
    doc2.build(story2)
    buf.seek(0)
    return buf.getvalue()


def afficher_dashboard():
    """Tableau de bord complet — Mode Conseiller."""
    STATUT_META = {
        "confirme":    ("✅ Confirmés",    "#10b981","#d1fae5","badge-confirme"),
        "revise":      ("🔄 Révisés",      "#ec4899","#fce7f3","badge-revise"),
        "attente":     ("⏳ En attente",   "#f97316","#fff7ed","badge-attente"),
        "probation":   ("⚠️ Probation",    "#f59e0b","#fef3c7","badge-probation"),
        "indetermine": ("❓ Indéterminés","#6b7280","#f3f4f6","badge-indetermine"),
    }
    # ── Titre ──
    st.markdown("""
    <div style="display:flex;align-items:center;gap:1rem;margin-bottom:1.5rem;">
        <div style="font-size:2.4rem;">📊</div>
        <div>
            <div style="font-size:1.5rem;font-weight:800;">Tableau de Bord des Élèves Orientés</div>
            <div style="font-size:0.82rem;opacity:0.55;margin-top:2px;">
                Vue réservée au conseiller d'orientation · Données en temps réel
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

    # ── Statistiques ──
    try:
        stats = db.statistiques()
    except Exception:
        stats = {"total":0,"nb_confirmes":0,"nb_attente":0,"nb_probation":0,
                 "nb_revises":0,"nb_serie_c":0,"nb_serie_a":0}

    stat_items = [
        ("🗂️ Total",      stats.get("total",0),          "#6d28d9","#f3e8ff"),
        ("✅ Confirmés",   stats.get("nb_confirmes",0),   "#10b981","#d1fae5"),
        ("⏳ En attente",  stats.get("nb_attente",0),     "#f97316","#fff7ed"),
        ("⚠️ Probation",  stats.get("nb_probation",0),   "#f59e0b","#fef3c7"),
        ("🔄 Révisés",     stats.get("nb_revises",0),     "#ec4899","#fce7f3"),
    ]
    cols_stat = st.columns(len(stat_items))
    for col, (lbl, val, color, bg) in zip(cols_stat, stat_items):
        with col:
            st.markdown(f"""
            <div class="dash-stat-card">
                <div class="dash-stat-val" style="color:{color};">{val}</div>
                <div class="dash-stat-lbl">{lbl}</div>
            </div>""", unsafe_allow_html=True)

    # ── Graphiques ──
    total = stats.get("total", 0)
    if total > 0:
        st.write("")
        cg1, cg2 = st.columns([1, 2])
        with cg1:
            nb_c = stats.get("nb_serie_c", 0)
            nb_a = stats.get("nb_serie_a", 0)
            fig_pie = go.Figure(go.Pie(
                labels=["Série C", "Série A", "Autre"],
                values=[nb_c, nb_a, max(0, total - nb_c - nb_a)],
                hole=0.55, marker=dict(colors=["#10b981","#3b82f6","#94a3b8"],
                                       line=dict(color="white", width=2)),
                textinfo="label+percent", textfont_size=11,
            ))
            fig_pie.update_layout(showlegend=False, height=200,
                margin=dict(l=10,r=10,t=20,b=10),
                annotations=[dict(text=f"<b>{total}</b>",x=0.5,y=0.5,font_size=18,showarrow=False)])
            st.plotly_chart(fig_pie, use_container_width=True)
        with cg2:
            cats = ["Confirmés","En attente","Probation","Révisés"]
            vals = [stats.get("nb_confirmes",0), stats.get("nb_attente",0),
                    stats.get("nb_probation",0), stats.get("nb_revises",0)]
            fig_bar = go.Figure(go.Bar(x=cats, y=vals,
                marker_color=["#10b981","#f97316","#f59e0b","#ec4899"],
                text=vals, textposition="outside", marker=dict(line=dict(width=0))))
            fig_bar.update_layout(height=200, margin=dict(l=10,r=10,t=20,b=10),
                yaxis=dict(visible=False), showlegend=False,
                plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)")
            st.plotly_chart(fig_bar, use_container_width=True)

    st.divider()

    # ── Filtres ──
    fc1, fc2, fc3 = st.columns([2, 1, 1])
    with fc1:
        search_q = st.text_input("🔍 Rechercher (nom ou prénom)",
                                  placeholder="Ex : MBALLA, Jean...", key="dash_search")
    with fc2:
        filtre_statut = st.selectbox("Statut",
            ["Tous","confirme","revise","attente","probation","indetermine"], key="dash_filtre")
    with fc3:
        filtre_serie = st.selectbox("Série", ["Toutes","C","A","?"], key="dash_serie")

    # ── Chargement données ──
    try:
        statuts_a_lister = (["confirme","revise","attente","probation","indetermine"]
                            if filtre_statut == "Tous" else [filtre_statut])
        tous_dossiers = []
        for s in statuts_a_lister:
            for d in db.lister_dossiers(s):
                d["_statut_affiche"] = s
                tous_dossiers.append(d)
    except Exception as e:
        st.error(f"Erreur lecture BDD : {e}")
        tous_dossiers = []

    if search_q.strip():
        q = search_q.strip().upper()
        tous_dossiers = [d for d in tous_dossiers if
                         q in (d.get("nom","") or "").upper() or
                         q in (d.get("prenom","") or "").upper()]
    if filtre_serie != "Toutes":
        tous_dossiers = [d for d in tous_dossiers if
                         (d.get("serie_finale") or d.get("serie_provisoire") or
                          d.get("serie_cible") or "?") == filtre_serie]

    # ── Tableau ──
    st.markdown(f"**{len(tous_dossiers)} dossier(s) trouvé(s)**")
    if not tous_dossiers:
        st.info("Aucun dossier ne correspond aux critères.")
    else:
        h_cols = st.columns([2.5, 2.2, 2, 1.2, 1.2, 1.5, 1.8, 1.8])
        for col, lbl in zip(h_cols, ["Nom & Prénom","Lycée","Projet","SA","LA","Série","Statut","Date"]):
            col.markdown(f"<div style='font-size:0.7rem;font-weight:700;text-transform:uppercase;"
                         f"opacity:0.5;letter-spacing:0.06em;'>{lbl}</div>", unsafe_allow_html=True)
        st.markdown("<hr style='margin:0.2rem 0 0.4rem;opacity:0.2;'>", unsafe_allow_html=True)
        for d in tous_dossiers:
            stat_af = d.get("_statut_affiche","—")
            _, _, _, s_badge = STATUT_META.get(stat_af, ("","#6b7280","#f3f4f6","badge-indetermine"))
            serie_af = (d.get("serie_finale") or d.get("serie_provisoire") or
                        d.get("serie_cible") or "?")
            badge_serie = "badge-c" if serie_af=="C" else ("badge-a" if serie_af=="A" else "")
            sa_v = float(d.get("SA_etal") or d.get("SA_brut") or 0)
            la_v = float(d.get("LA_etal") or d.get("LA_brut") or 0)
            rc = st.columns([2.5, 2.2, 2, 1.2, 1.2, 1.5, 1.8, 1.8])
            rc[0].markdown(f"**{(d.get('nom') or '').upper()}** {(d.get('prenom') or '').capitalize()}")
            rc[1].markdown(f"<small>{(d.get('lycee') or '—')[:22]}</small>", unsafe_allow_html=True)
            rc[2].markdown(f"<small>{(d.get('projet_pro') or '—')[:20]}</small>", unsafe_allow_html=True)
            rc[3].markdown(f"<b style='color:#10b981;'>{sa_v:.1f}</b>", unsafe_allow_html=True)
            rc[4].markdown(f"<b style='color:#3b82f6;'>{la_v:.1f}</b>", unsafe_allow_html=True)
            rc[5].markdown(f'<span class="dash-badge {badge_serie}">{serie_af}</span>', unsafe_allow_html=True)
            rc[6].markdown(f'<span class="dash-badge {s_badge}">{stat_af}</span>', unsafe_allow_html=True)
            rc[7].markdown(f"<small style='opacity:0.5;'>{(d.get('date_modification','') or '')[:10]}</small>", unsafe_allow_html=True)
            st.markdown("<div style='border-bottom:1px solid rgba(0,0,0,0.05);margin:0.1rem 0;'></div>", unsafe_allow_html=True)

    # ── Mise à jour T2/T3 pour dossiers non confirmés ──
    st.divider()
    st.markdown("#### 📝 Saisir les notes T2/T3 — Dossiers en attente / probation")
    non_conf = [d for d in tous_dossiers if d.get("_statut_affiche") in ("attente","probation")]
    if not non_conf:
        st.info("Aucun dossier en attente ou en probation dans cette vue.")
    else:
        for d in non_conf:
            nom_d    = (d.get("nom","") or "").upper()
            prenom_d = (d.get("prenom","") or "").capitalize()
            stat_d   = d.get("_statut_affiche","")
            sa_d     = float(d.get("SA_etal") or d.get("SA_brut") or 0)
            la_d     = float(d.get("LA_etal") or d.get("LA_brut") or 0)
            trim_a   = d.get("trimestre_actuel") or d.get("trimestre_decision","T1")
            prochain = "T2" if trim_a == "T1" else "T3"
            uid      = f"{d.get('id',0)}_{nom_d}"

            with st.expander(f"📋 {nom_d} {prenom_d} — {stat_d.upper()} → Saisir notes {prochain}"):
                st.markdown(f"SA={sa_d:.1f}/20 · LA={la_d:.1f}/20 · Trimestre actuel : {trim_a}")
                ca, cb = st.columns(2)
                with ca:
                    st.markdown("**🔬 Matières scientifiques**")
                    mths = st.number_input(f"Maths {prochain}", 0.0, 20.0,
                        float(d.get(f"maths_{prochain.lower()}") or 10.0), 0.5, key=f"mths_{uid}")
                    sphy = st.number_input(f"Sci.Phy {prochain}", 0.0, 20.0,
                        float(d.get(f"sci_phy_{prochain.lower()}") or 10.0), 0.5, key=f"sphy_{uid}")
                    svtv = st.number_input(f"SVT {prochain}", 0.0, 20.0,
                        float(d.get(f"svt_{prochain.lower()}") or 10.0), 0.5, key=f"svtv_{uid}")
                    msc  = round((mths+sphy+svtv)/3, 2)
                    st.metric(f"Moy. Sci. {prochain}", f"{msc:.2f}/20",
                              delta=f"{msc-10:+.1f} vs seuil")
                with cb:
                    st.markdown("**📖 Matières littéraires**")
                    frv  = st.number_input(f"Français {prochain}", 0.0, 20.0,
                        float(d.get(f"francais_{prochain.lower()}") or 10.0), 0.5, key=f"frv_{uid}")
                    hgv  = st.number_input(f"Hist-Géo {prochain}", 0.0, 20.0,
                        float(d.get(f"histgeo_{prochain.lower()}") or 10.0), 0.5, key=f"hgv_{uid}")
                    angv = st.number_input(f"Anglais {prochain}", 0.0, 20.0,
                        float(d.get(f"anglais_{prochain.lower()}") or 10.0), 0.5, key=f"angv_{uid}")
                    mlt  = round((frv+hgv+angv)/3, 2)
                    st.metric(f"Moy. Lit. {prochain}", f"{mlt:.2f}/20",
                              delta=f"{mlt-10:+.1f} vs seuil")

                # Calcul automatique nouveau statut
                serie_v   = (d.get("serie_finale") or d.get("serie_provisoire") or
                             d.get("serie_cible") or "C")
                cond_ok   = (msc >= 10) if sa_d >= la_d else (mlt >= 10)
                if stat_d == "attente":
                    new_stat = "confirme" if cond_ok else "probation"
                    new_lbl  = "✅ CONFIRMÉ" if cond_ok else "⚠️ PROBATION"
                else:
                    new_stat = "confirme" if cond_ok else "revise"
                    serie_v  = serie_v if cond_ok else ("A" if serie_v=="C" else "C")
                    new_lbl  = "✅ CONFIRMÉ" if cond_ok else f"🔄 RÉVISÉ → Série {serie_v}"
                ok_c = "#10b981" if cond_ok else "#ef4444"
                ok_bg= "#ecfdf5" if cond_ok else "#fef2f2"
                st.markdown(f"""
                <div style="background:{ok_bg};border:1.5px solid {ok_c};border-radius:10px;
                            padding:0.6rem 1rem;text-align:center;font-weight:700;
                            color:{ok_c};margin:0.5rem 0;">
                    Nouveau statut prévu : {new_lbl}
                </div>""", unsafe_allow_html=True)

                if st.button(f"✅ Valider et mettre à jour dans la BDD",
                             use_container_width=True, type="primary",
                             key=f"val_{uid}"):
                    ss_upd = dict(d)
                    tk = prochain.lower()
                    ss_upd[f"maths_{tk}"]   = mths
                    ss_upd[f"sci_phy_{tk}"] = sphy
                    ss_upd[f"svt_{tk}"]     = svtv
                    ss_upd[f"francais_{tk}"]= frv
                    ss_upd[f"histgeo_{tk}"] = hgv
                    ss_upd[f"anglais_{tk}"] = angv
                    ss_upd[f"t{prochain[1]}_renseigne"] = True
                    ss_upd["statut"]            = new_stat
                    ss_upd["orientation_finale"]= serie_v
                    ss_upd["revenu"]            = d.get("revenu_famille","")
                    res_upd = db.sauvegarder_dossier(ss_upd, None)
                    if res_upd["succes"]:
                        st.success(f"✅ Statut mis à jour → **{new_stat.upper()}** · {res_upd['message']}")
                        rerun()
                    else:
                        st.error(res_upd["message"])

    # ── Exports multi-formats ──
    st.divider()
    st.markdown("#### 📤 Exports")

    try:
        all_exp = []
        for s in ["confirme","revise","attente","probation","indetermine"]:
            for dd in db.lister_dossiers(s):
                dd["_statut_affiche"] = s
                all_exp.append(dd)
        stats_exp = db.statistiques()
    except Exception:
        all_exp = []; stats_exp = {}

    e1, e2, e3, e4 = st.columns(4)
    with e1:
        try:
            xlsx = generer_excel_export(all_exp)
            if xlsx:
                st.download_button("📊 Excel (.xlsx)", xlsx,
                    file_name=f"capavenir_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True)
            else:
                st.caption("Installer openpyxl")
        except Exception as ex:
            st.caption(f"Erreur Excel: {ex}")
    with e2:
        try:
            docx_b = generer_word_export(all_exp, stats_exp)
            if docx_b:
                st.download_button("📝 Word (.docx)", docx_b,
                    file_name=f"capavenir_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            else:
                st.caption("Installer python-docx")
        except Exception as ex:
            st.caption(f"Erreur Word: {ex}")
    with e3:
        try:
            pdf_b = generer_pdf_rapport(all_exp, stats_exp)
            st.download_button("📄 PDF Rapport", pdf_b,
                file_name=f"rapport_capavenir_{datetime.now().strftime('%Y%m%d')}.pdf",
                mime="application/pdf", use_container_width=True)
        except Exception as ex:
            st.caption(f"Erreur PDF: {ex}")
    with e4:
        try:
            j_data = db.exporter_json()
            st.download_button("📦 JSON", j_data,
                file_name=f"capavenir_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json", use_container_width=True)
        except Exception:
            st.caption("Export indisponible")


# =====================================================================
# NAVIGATION — Dashboard ou Orientation
# =====================================================================

# ── Affichage page de connexion si besoin ──
if st.session_state.get("show_login", False) and not st.session_state.get("conseiller_auth", False):
    afficher_login_conseiller()
    st.stop()
else:
    st.session_state.show_login = False

# Si mode conseiller et dashboard actif → afficher le tableau de bord et stopper
if mode_conseiller and st.session_state.get("show_dashboard", False):
    afficher_dashboard()
    st.stop()
else:
    st.session_state.show_dashboard = False

STEPS = ["Profil", "Notes", "Tests", "Diagnostic", "Fiche"]
ICONS = ["📋", "📚", "🧪", "🤖", "📄"]
step = st.session_state.step

cols = st.columns(len(STEPS))
for i, col in enumerate(cols):
    with col:
        label = f"{ICONS[i]} {STEPS[i]}"
        if i == step:
            st.markdown(f"<div class='step-active'>{label}</div>", unsafe_allow_html=True)
        elif i < step:
            st.markdown(f"<div class='step-done'>✓ {STEPS[i]}</div>", unsafe_allow_html=True)
        else:
            st.markdown(f"<div class='step-todo'>{label}</div>", unsafe_allow_html=True)

st.write("")

# =====================================================================
# ETAPE 0 — PROFIL ELEVE
# =====================================================================
if step == 0:
    st.subheader("📋 Informations personnelles de l'élève")

    c1, c2 = st.columns(2)
    with c1:
        st.session_state.nom   = st.text_input("Nom",   value=st.session_state.nom)
        st.session_state.age   = st.number_input("Âge", min_value=12, max_value=20, value=st.session_state.age)
        st.session_state.lycee = st.text_input("Lycée / Établissement", value=st.session_state.lycee)
    with c2:
        st.session_state.prenom = st.text_input("Prénom", value=st.session_state.prenom)
        st.session_state.sexe   = st.selectbox("Sexe", ["Masculin", "Féminin"])

    st.divider()
    st.subheader("🏠 Projet personnel & contexte familial")
    c1, c2 = st.columns(2)
    with c1:
        st.session_state.choix_personnel = st.selectbox(
            "Série souhaitée par l'élève",
            ["C (Scientifique)", "A (Littéraire)", "Indécis(e)"]
        )
        st.session_state.projet_pro = st.text_area(
            "Projet professionnel",
            value=st.session_state.projet_pro,
            placeholder="Ex: Devenir médecin, ingénieur, avocat...",
            height=90,
        )
    with c2:
        st.session_state.revenu = st.selectbox(
            "Revenu mensuel estimé des parents",
            ["Faible (< 50 000 FCFA/mois)",
             "Moyen (50 000 - 150 000 FCFA/mois)",
             "Élevé (> 150 000 FCFA/mois)"]
        )
        st.markdown('<div class="alert-info">💡 Le revenu permet d\'évaluer l\'accès aux cours de soutien et la faisabilité des études longues.</div>', unsafe_allow_html=True)

    st.write("")
    if st.button("Suivant : Notes scolaires ➡️", use_container_width=True, type="primary"):
        if st.session_state.nom.strip() and st.session_state.prenom.strip():
            st.session_state.step = 1
            rerun()
        else:
            st.error("Veuillez renseigner le nom et le prénom de l'élève.")

# =====================================================================
# ETAPE 1 — NOTES SCOLAIRES (Multi-trimestre + graphique de progression)
# =====================================================================
elif step == 1:
    st.subheader(f"📚 Notes scolaires — {st.session_state.prenom} {st.session_state.nom}")

    tab_t1, tab_t2, tab_t3 = st.tabs(["📅 Trimestre 1", "📅 Trimestre 2", "📅 Trimestre 3"])

    def saisie_notes(suffix, key_prefix):
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**Matières scientifiques**")
            st.session_state[f"maths{suffix}"]   = st.number_input("Mathématiques",      0.0, 20.0, st.session_state[f"maths{suffix}"],   0.5, key=f"{key_prefix}_maths")
            st.session_state[f"sci_phy{suffix}"] = st.number_input("Sciences Physiques", 0.0, 20.0, st.session_state[f"sci_phy{suffix}"], 0.5, key=f"{key_prefix}_sp")
            st.session_state[f"svt{suffix}"]     = st.number_input("SVT",                0.0, 20.0, st.session_state[f"svt{suffix}"],     0.5, key=f"{key_prefix}_svt")
            moy = (st.session_state[f"maths{suffix}"] + st.session_state[f"sci_phy{suffix}"] + st.session_state[f"svt{suffix}"]) / 3
            st.metric("Moyenne scientifique", f"{moy:.2f} / 20")
        with c2:
            st.markdown("**Matières littéraires**")
            st.session_state[f"francais{suffix}"] = st.number_input("Français",            0.0, 20.0, st.session_state[f"francais{suffix}"], 0.5, key=f"{key_prefix}_fr")
            st.session_state[f"histgeo{suffix}"]  = st.number_input("Histoire-Géographie", 0.0, 20.0, st.session_state[f"histgeo{suffix}"],  0.5, key=f"{key_prefix}_hg")
            st.session_state[f"anglais{suffix}"]  = st.number_input("Anglais",             0.0, 20.0, st.session_state[f"anglais{suffix}"],  0.5, key=f"{key_prefix}_ang")
            moy = (st.session_state[f"francais{suffix}"] + st.session_state[f"histgeo{suffix}"] + st.session_state[f"anglais{suffix}"]) / 3
            st.metric("Moyenne littéraire", f"{moy:.2f} / 20")

    with tab_t1:
        saisie_notes("_t1", "t1")

    with tab_t2:
        has_t2 = st.checkbox("Renseigner le Trimestre 2", value=st.session_state.t2_renseigne, key="cb_t2")
        st.session_state.t2_renseigne = has_t2
        if has_t2:
            saisie_notes("_t2", "t2")

    with tab_t3:
        has_t3 = st.checkbox("Renseigner le Trimestre 3", value=st.session_state.t3_renseigne, key="cb_t3")
        st.session_state.t3_renseigne = has_t3
        if has_t3:
            saisie_notes("_t3", "t3")

    # Graphique de progression multi-trimestre
    if st.session_state.t2_renseigne or st.session_state.t3_renseigne:
        st.divider()
        st.markdown("**📈 Évolution des moyennes trimestrielles**")
        trims, sci_vals, lit_vals = ["T1"], [], []
        sci_vals.append((st.session_state.maths_t1 + st.session_state.sci_phy_t1 + st.session_state.svt_t1) / 3)
        lit_vals.append((st.session_state.francais_t1 + st.session_state.histgeo_t1 + st.session_state.anglais_t1) / 3)
        if st.session_state.t2_renseigne:
            trims.append("T2")
            sci_vals.append((st.session_state.maths_t2 + st.session_state.sci_phy_t2 + st.session_state.svt_t2) / 3)
            lit_vals.append((st.session_state.francais_t2 + st.session_state.histgeo_t2 + st.session_state.anglais_t2) / 3)
        if st.session_state.t3_renseigne:
            trims.append("T3")
            sci_vals.append((st.session_state.maths_t3 + st.session_state.sci_phy_t3 + st.session_state.svt_t3) / 3)
            lit_vals.append((st.session_state.francais_t3 + st.session_state.histgeo_t3 + st.session_state.anglais_t3) / 3)

        fig_prog = go.Figure()
        fig_prog.add_trace(go.Scatter(
            x=trims, y=sci_vals, mode="lines+markers+text",
            name="Moy. Scientifique", line=dict(color="#10b981", width=2.5), marker=dict(size=10),
            text=[f"{v:.1f}" for v in sci_vals], textposition="top center"
        ))
        fig_prog.add_trace(go.Scatter(
            x=trims, y=lit_vals, mode="lines+markers+text",
            name="Moy. Littéraire", line=dict(color="#3b82f6", width=2.5), marker=dict(size=10),
            text=[f"{v:.1f}" for v in lit_vals], textposition="bottom center"
        ))
        fig_prog.add_hline(y=10, line_dash="dash", line_color="#f59e0b", annotation_text="Seuil 10/20")
        fig_prog.update_layout(
            yaxis=dict(range=[0, 22], title="Note /20"),
            height=260, margin=dict(l=20, r=20, t=20, b=20),
            legend=dict(orientation="h", y=-0.3)
        )
        st.plotly_chart(fig_prog, use_container_width=True)

    st.write("")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Retour ⬅️", use_container_width=True):
            st.session_state.step = 0
            rerun()
    with c2:
        if st.button("Suivant : Tests psychotechniques ➡️", use_container_width=True, type="primary"):
            st.session_state.step = 2
            rerun()

# =====================================================================
# ETAPE 2 — TESTS PSYCHOTECHNIQUES (passation guidée OU saisie directe)
# =====================================================================
elif step == 2:
    st.subheader("🧪 Tests psychotechniques")

    mode_radio = st.radio(
        "Mode de saisie :",
        ["🎯 Passation guidée (vraies questions)", "📊 Saisie directe des scores (barèmes papier)"],
        horizontal=True
    )
    mode_guidee = "guidée" in mode_radio

    if mode_guidee:
        st.markdown("_Répondez aux 5 questions de chaque test. Le score est calculé automatiquement._")
        test_keys  = ["D48", "KRX", "MECA", "BV11", "PRC"]
        test_names = [TESTS_QUESTIONS[k]["nom"].split("—")[0].strip() for k in test_keys]
        test_tabs  = st.tabs(test_names)

        for t_key, t_tab in zip(test_keys, test_tabs):
            with t_tab:
                tdata = TESTS_QUESTIONS[t_key]
                st.caption(tdata["description"])
                if t_key not in st.session_state.test_answers:
                    st.session_state.test_answers[t_key] = {}

                correct = 0
                for q_idx, q in enumerate(tdata["questions"]):
                    st.markdown(f'<div class="test-question"><strong>Q{q_idx+1}.</strong> {q["q"]}</div>', unsafe_allow_html=True)
                    prev = st.session_state.test_answers[t_key].get(q_idx)
                    idx_default = q["choices"].index(prev) if prev in q["choices"] else 0
                    choice = st.radio(
                        f"R{q_idx+1}_{t_key}",
                        q["choices"],
                        index=idx_default,
                        key=f"q_{t_key}_{q_idx}",
                        label_visibility="collapsed"
                    )
                    st.session_state.test_answers[t_key][q_idx] = choice
                    if choice == q["answer"]:
                        correct += 1
                    elif mode_conseiller and choice is not None:
                        st.caption(f"✅ Réponse correcte : {q['answer']} — {q['expl']}")

                score = round((correct / len(tdata["questions"])) * 20, 1)
                st.session_state.test_scores[t_key] = score
                st.success(f"Score : **{correct}/{len(tdata['questions'])}** réponses correctes → **{score}/20** brut")

        # Synchroniser avec les variables de session
        for key_map in [("D48","d48"), ("KRX","krx"), ("MECA","meca"), ("BV11","bv11"), ("PRC","prc")]:
            if key_map[0] in st.session_state.test_scores:
                st.session_state[key_map[1]] = st.session_state.test_scores[key_map[0]]

    else:
        # Saisie directe via sliders
        st.markdown("**Tests d'aptitude scientifique**")
        c1, c2 = st.columns(2)
        with c1:
            st.caption("D48 : Raisonnement logique et abstrait")
            st.session_state.d48 = st.slider("Score D48",  0.0, 20.0, st.session_state.d48,  0.5, key="sl_d48")
        with c2:
            st.caption("KRX : Aptitude mathématique, suites numériques")
            st.session_state.krx = st.slider("Score KRX",  0.0, 20.0, st.session_state.krx,  0.5, key="sl_krx")
        st.caption("MECA : Compréhension mécanique et physique")
        st.session_state.meca = st.slider("Score MECA", 0.0, 20.0, st.session_state.meca, 0.5, key="sl_meca")
        st.divider()
        st.markdown("**Tests d'aptitude littéraire**")
        c1, c2 = st.columns(2)
        with c1:
            st.caption("BV11 : Vocabulaire, synonymes et compréhension de texte")
            st.session_state.bv11 = st.slider("Score BV11", 0.0, 20.0, st.session_state.bv11, 0.5, key="sl_bv11")
        with c2:
            st.caption("PRC : Proverbes et raisonnement linguistique")
            st.session_state.prc  = st.slider("Score PRC",  0.0, 20.0, st.session_state.prc,  0.5, key="sl_prc")

    # Résumé des scores (bruts + étalonnés)
    st.divider()
    st.markdown("**Récapitulatif des scores**")
    cols5 = st.columns(5)
    for col, (code, val) in zip(cols5, [
        ("D48", st.session_state.d48), ("KRX", st.session_state.krx),
        ("MECA", st.session_state.meca), ("BV11", st.session_state.bv11),
        ("PRC", st.session_state.prc)
    ]):
        etal = etalonner(val)
        col.markdown(f"""
        <div class="score-card">
            <div style="font-weight:700; font-size:0.9rem;">{code}</div>
            <div class="score-raw">Brut : {val:.1f}/20</div>
            <div class="score-etalon">Étal : {etal:.1f}/20</div>
        </div>
        """, unsafe_allow_html=True)

    SA_brut, LA_brut, SA_etal, LA_etal = calc_aptitudes()
    st.write("")
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("SA brute",      f"{SA_brut:.1f}/20")
    m2.metric("SA étalonnée",  f"{SA_etal:.1f}/20", help="Valeur utilisée pour l'orientation")
    m3.metric("LA brute",      f"{LA_brut:.1f}/20")
    m4.metric("LA étalonnée",  f"{LA_etal:.1f}/20", help="Valeur utilisée pour l'orientation")

    st.write("")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Retour ⬅️", use_container_width=True):
            st.session_state.step = 1
            rerun()
    with c2:
        if st.button("Lancer le Diagnostic IA 🤖", use_container_width=True, type="primary"):
            st.session_state.step = 3
            st.session_state.chat_history = []
            rerun()

# =====================================================================
# ETAPE 3 — DIAGNOSTIC IA
# =====================================================================
elif step == 3:
    SA_brut, LA_brut, SA_etal, LA_etal = calc_aptitudes()
    SA, LA = SA_etal, LA_etal            # Décision basée sur les notes étalonnées
    notes, trim_label = get_notes_actives()
    moy_sci, moy_lit  = calc_moyennes(notes)
    choix = st.session_state.choix_personnel

    # ---- Contexte trimestre ----
    t1_seul   = not st.session_state.t2_renseigne and not st.session_state.t3_renseigne
    t2_dispo  = st.session_state.t2_renseigne and not st.session_state.t3_renseigne
    t3_dispo  = st.session_state.t3_renseigne

    # ---- Moteur d'inférence étendu ----
    # Cas 1 : profil scientifique confirmé
    if SA > LA and moy_sci >= 10:
        serie      = "C"
        statut     = "confirme"
        msg_diag   = "Profil scientifique confirmé ✅"
        alerte_type= "success"

    # Cas 2 : profil littéraire confirmé
    elif LA > SA and moy_lit >= 10:
        serie      = "A"
        statut     = "confirme"
        msg_diag   = "Profil littéraire confirmé ✅"
        alerte_type= "success"

    # Cas 3 : aptitudes scientifiques OK, notes sci faibles
    elif SA > LA and moy_sci < 10:
        if t1_seul:
            # Seulement T1 disponible → EN ATTENTE
            serie      = "C"
            statut     = "attente"
            msg_diag   = "EN ATTENTE — Notes du 2e trimestre requises"
            alerte_type= "attente"
        elif t2_dispo:
            # T2 disponible, toujours faible → PROBATION
            serie      = "C"
            statut     = "probation"
            msg_diag   = "PROBATION — Amélioration critique avant le 3e trimestre"
            alerte_type= "warning"
        elif t3_dispo:
            # T3 disponible, toujours faible → orientation révisée en A
            serie      = "A"
            statut     = "revise"
            msg_diag   = "ORIENTATION RÉVISÉE → Série A définitive"
            alerte_type= "revise"
        else:
            serie, statut, msg_diag, alerte_type = "C", "probation", "Probation", "warning"

    # Cas 4 : aptitudes littéraires OK, notes lit faibles
    elif LA > SA and moy_lit < 10:
        if t1_seul:
            serie      = "A"
            statut     = "attente"
            msg_diag   = "EN ATTENTE — Notes du 2e trimestre requises"
            alerte_type= "attente"
        elif t2_dispo:
            serie      = "A"
            statut     = "probation"
            msg_diag   = "PROBATION — Amélioration critique avant le 3e trimestre"
            alerte_type= "warning"
        elif t3_dispo:
            serie      = "C"
            statut     = "revise"
            msg_diag   = "ORIENTATION RÉVISÉE → Série C définitive"
            alerte_type= "revise"
        else:
            serie, statut, msg_diag, alerte_type = "A", "probation", "Probation", "warning"

    # Cas 5 : profil indéterminé
    elif SA == LA:
        serie      = "?"
        statut     = "indetermine"
        msg_diag   = "Profil équilibré — entretien approfondi recommandé"
        alerte_type= "warning"
    else:
        serie      = "?"
        statut     = "indetermine"
        msg_diag   = "Données insuffisantes — compléter le dossier"
        alerte_type= "warning"

    probation  = statut in ("probation",)
    score_conf = calc_score_confiance(SA, LA, moy_sci, moy_lit, serie)

    # Conseil selon revenu
    conseil_revenu = ""
    if "Faible" in st.session_state.revenu:
        if "C" in serie:
            conseil_revenu = "⚠️ Revenu modeste : la série C implique des études longues et coûteuses. Vérifier les bourses et dispositifs de soutien scolaire."
        elif "A" in serie:
            conseil_revenu = "ℹ️ La série A offre des débouchés accessibles (droit, lettres, administration) compatibles avec le contexte financier familial."

    # Détection du type de conflit
    conflit_type = None
    if SA > LA and "A" in choix:
        conflit_type = "decale"
    elif SA > LA and moy_sci < 10 and "C" in choix:
        conflit_type = "reveur"
    elif LA > SA and moy_lit < 10 and "A" in choix:
        conflit_type = "reveur"

    st.session_state.orientation_finale = serie
    st.session_state.probation = probation
    st.session_state.statut = statut
    st.session_state.score_confiance = score_conf

    # ---- Affichage des métriques ----
    st.subheader("Analyse du profil")
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("SA étalonnée",            f"{SA:.1f}/20",
              delta=f"Brut: {SA_brut:.1f}" if mode_conseiller else None)
    m2.metric("LA étalonnée",            f"{LA:.1f}/20",
              delta=f"Brut: {LA_brut:.1f}" if mode_conseiller else None)
    m3.metric(f"Moy. Sci. ({trim_label})", f"{moy_sci:.1f}/20",
              delta=f"{moy_sci-10:+.1f} vs seuil", delta_color="normal" if moy_sci >= 10 else "inverse")
    m4.metric(f"Moy. Lit. ({trim_label})", f"{moy_lit:.1f}/20",
              delta=f"{moy_lit-10:+.1f} vs seuil", delta_color="normal" if moy_lit >= 10 else "inverse")

    # Barre de confiance
    conf_color = "#10b981" if score_conf >= 70 else ("#f59e0b" if score_conf >= 50 else "#ef4444")
    conf_label = "Fiable" if score_conf >= 70 else ("Modérée" if score_conf >= 50 else "Incertaine")
    st.markdown(f"""
    <div style="background:#f8fafc; border-radius:10px; padding:0.8rem 1rem; margin:0.5rem 0;">
        <div style="font-size:0.8rem; color:#64748b; margin-bottom:4px;">
            Score de confiance de l'orientation — <em>{conf_label}</em>
        </div>
        <div style="background:#e2e8f0; border-radius:6px; height:10px; overflow:hidden;">
            <div style="width:{score_conf}%; height:100%; background:{conf_color}; border-radius:6px;"></div>
        </div>
        <div style="font-size:0.85rem; color:{conf_color}; font-weight:600; margin-top:4px;">{score_conf} %</div>
    </div>
    """, unsafe_allow_html=True)

    # Graphiques côte à côte : Radar + Barres
    col_r, col_b = st.columns(2)
    with col_r:
        labels = ["Logique\n(D48)", "Maths\n(KRX)", "Méca.\n(MECA)", "Litt.\n(BV11)", "Prov.\n(PRC)"]
        apt_etal = [etalonner(st.session_state.d48), etalonner(st.session_state.krx),
                    etalonner(st.session_state.meca), etalonner(st.session_state.bv11),
                    etalonner(st.session_state.prc)]
        notes_list = [notes["maths"], notes["sci_phy"], notes["svt"], notes["francais"], notes["histgeo"]]
        fig_r = go.Figure()
        fig_r.add_trace(go.Scatterpolar(r=apt_etal, theta=labels, fill="toself",
            name="Aptitudes (étalonnées)", line_color="#10b981", fillcolor="rgba(16,185,129,0.15)"))
        fig_r.add_trace(go.Scatterpolar(r=notes_list, theta=labels, fill="toself",
            name=f"Notes scolaires ({trim_label})", line_color="#3b82f6", fillcolor="rgba(59,130,246,0.15)"))
        if mode_conseiller:
            apt_brut = [st.session_state.d48, st.session_state.krx, st.session_state.meca,
                        st.session_state.bv11, st.session_state.prc]
            fig_r.add_trace(go.Scatterpolar(r=apt_brut, theta=labels, fill="toself",
                name="Aptitudes (brutes)", line_color="#f59e0b", fillcolor="rgba(245,158,11,0.08)",
                line=dict(dash="dot")))
        fig_r.update_layout(
            polar=dict(radialaxis=dict(visible=True, range=[0, 20]), gridshape="linear"),
            showlegend=True, height=310,
            margin=dict(l=30, r=30, t=10, b=60),
            legend=dict(orientation="h", y=-0.25, x=0.5, xanchor="center")
        )
        st.plotly_chart(fig_r, use_container_width=True)

    with col_b:
        cats   = ["SA étal.", "LA étal.", f"Sci. {trim_label}", f"Lit. {trim_label}"]
        vals   = [SA, LA, moy_sci, moy_lit]
        colors = ["#10b981" if v >= 10 else "#f87171" for v in vals]
        fig_b = go.Figure(go.Bar(
            x=cats, y=vals, marker_color=colors,
            text=[f"{v:.1f}" for v in vals], textposition="outside"
        ))
        fig_b.add_hline(y=10, line_dash="dash", line_color="#f59e0b", annotation_text="Seuil 10/20")
        fig_b.update_layout(
            yaxis=dict(range=[0, 23], title="Note /20"),
            height=310, margin=dict(l=10, r=10, t=10, b=20), showlegend=False
        )
        st.plotly_chart(fig_b, use_container_width=True)

    # ---- Verdict ----
    st.subheader("Verdict du moteur d'inférence")
    st.markdown(f'<div class="alert-{alerte_type}"><strong>{msg_diag}</strong></div>',
                unsafe_allow_html=True)

    prenom = st.session_state.prenom

    # --- EN ATTENTE (T1 seulement, notes faibles) ---
    if statut == "attente":
        obj_sci = max(10.0, round(moy_sci + 2.0, 1)) if SA > LA else None
        obj_lit = max(10.0, round(moy_lit + 2.0, 1)) if LA > SA else None
        obj_txt = (f"Moyenne scientifique ≥ <strong>{obj_sci}/20</strong>"
                   if obj_sci else f"Moyenne littéraire ≥ <strong>{obj_lit}/20</strong>")
        serie_cible = "C" if SA > LA else "A"
        serie_risque = "A" if SA > LA else "C"
        st.markdown(f"""
        <div class="alert-attente">
            ⏳ <strong>Statut : EN ATTENTE des notes du 2e trimestre</strong><br><br>
            {prenom}, ton potentiel est bien là (SA = {SA:.1f}/20) mais tes notes scolaires
            ({moy_sci:.1f}/20) sont encore insuffisantes pour confirmer la série {serie_cible}.<br><br>
            👉 <strong>Ce qu'il faut faire :</strong> Mets-toi au travail dès maintenant sur tes matières
            scientifiques. L'objectif à atteindre au 2e trimestre est {obj_txt}.<br><br>
            ⚠️ <em>Si aucune progression n'est constatée au 2e trimestre, le statut passera en
            <strong>Probation</strong>. En fin d'année (T3), si les notes restent insuffisantes,
            l'orientation sera définitivement révisée vers la série {serie_risque}.</em>
        </div>
        """, unsafe_allow_html=True)

    # --- PROBATION (T2 dispo, notes toujours faibles) ---
    elif statut == "probation":
        obj_sci = max(10.0, round(moy_sci + 2.0, 1)) if SA > LA else None
        obj_lit = max(10.0, round(moy_lit + 2.0, 1)) if LA > SA else None
        obj_txt = (f"Moyenne scientifique ≥ <strong>{obj_sci}/20</strong>"
                   if obj_sci else f"Moyenne littéraire ≥ <strong>{obj_lit}/20</strong>")
        serie_risque = "A" if SA > LA else "C"
        st.markdown(f"""
        <div class="alert-warning">
            ⚠️ <strong>Statut : PROBATION — Dernière chance avant le 3e trimestre</strong><br><br>
            {prenom}, tes aptitudes le prouvent, tu as les capacités pour la série {serie}.
            Mais tes résultats scolaires ({moy_sci:.1f}/20) ne progressent pas suffisamment.<br><br>
            🎯 <strong>Objectif impératif au T3 :</strong> {obj_txt}<br><br>
            🚨 <em>Si cet objectif n'est pas atteint au 3e trimestre,
            <strong>l'orientation sera définitivement révisée vers la série {serie_risque}.</strong></em>
        </div>
        """, unsafe_allow_html=True)

    # --- ORIENTATION RÉVISÉE (T3 dispo, échec confirmé) ---
    elif statut == "revise":
        serie_initiale = "C" if serie == "A" else "A"
        st.markdown(f"""
        <div class="alert-revise">
            🔄 <strong>Statut : ORIENTATION RÉVISÉE — Décision définitive</strong><br><br>
            {prenom}, malgré de bonnes aptitudes ({SA:.1f}/20), les notes scolaires n'ont pas
            atteint le seuil requis sur les 3 trimestres.<br><br>
            La série <strong>{serie_initiale}</strong> n'est plus envisageable cette année.
            L'orientation définitive est fixée en <strong>Série {serie}</strong>.<br><br>
            ℹ️ <em>Cette décision a été prise conformément au protocole d'orientation :
            3 trimestres insuffisants malgré le suivi en probation.</em>
        </div>
        """, unsafe_allow_html=True)

    if abs(SA - moy_sci) >= 3:
        sens = "supérieure" if SA > moy_sci else "inférieure"
        st.markdown(
            f'<div class="alert-warning">🔍 Décalage détecté : aptitude scientifique étalonnée ({SA:.1f}/20) '
            f'{sens} de {abs(SA-moy_sci):.1f} pt(s) à la moyenne scolaire ({moy_sci:.1f}/20). '
            f'Investiguer la méthode de travail.</div>',
            unsafe_allow_html=True
        )

    if conseil_revenu:
        st.markdown(f'<div class="alert-info">{conseil_revenu}</div>', unsafe_allow_html=True)

    # ==================================================================
    # RUBRIQUE SUIVI T2 — visible uniquement si statut == "attente"
    # ==================================================================
    if statut == "attente":
        st.divider()
        st.markdown("""
        <div style="background:linear-gradient(135deg,#f0f9ff,#e0f2fe);
                    border:2px dashed #0ea5e9; border-radius:16px;
                    padding:1.2rem 1.4rem; margin:0.5rem 0;">
            <div style="font-size:1rem; font-weight:700; color:#0c4a6e; margin-bottom:0.3rem;">
                ⏳ Suivi — Saisie des notes du 2e trimestre
            </div>
            <div style="font-size:0.85rem; color:#0369a1;">
                Renseignez les notes du T2 pour mettre à jour l'orientation
                (Probation ou Confirmation).
            </div>
        </div>
        """, unsafe_allow_html=True)

        with st.expander("📋 Saisir les notes du 2e trimestre", expanded=True):
            st.caption("Ces notes seront enregistrées et le diagnostic sera recalculé automatiquement.")
            ca, cb = st.columns(2)
            with ca:
                st.markdown("**🔬 Matières scientifiques**")
                st.session_state.maths_t2   = st.number_input(
                    "Mathématiques T2",   0.0, 20.0,
                    st.session_state.maths_t2,   0.5, key="suivi_t2_maths")
                st.session_state.sci_phy_t2 = st.number_input(
                    "Sciences Physiques T2", 0.0, 20.0,
                    st.session_state.sci_phy_t2, 0.5, key="suivi_t2_sp")
                st.session_state.svt_t2     = st.number_input(
                    "SVT T2",             0.0, 20.0,
                    st.session_state.svt_t2,     0.5, key="suivi_t2_svt")
                moy_sci_t2 = (st.session_state.maths_t2 +
                              st.session_state.sci_phy_t2 +
                              st.session_state.svt_t2) / 3
                delta_sci = moy_sci_t2 - moy_sci
                st.metric("Moy. Sci. T2", f"{moy_sci_t2:.2f}/20",
                          delta=f"{delta_sci:+.1f} vs T1",
                          delta_color="normal" if delta_sci >= 0 else "inverse")
            with cb:
                st.markdown("**📖 Matières littéraires**")
                st.session_state.francais_t2 = st.number_input(
                    "Français T2",           0.0, 20.0,
                    st.session_state.francais_t2, 0.5, key="suivi_t2_fr")
                st.session_state.histgeo_t2  = st.number_input(
                    "Histoire-Géographie T2",0.0, 20.0,
                    st.session_state.histgeo_t2,  0.5, key="suivi_t2_hg")
                st.session_state.anglais_t2  = st.number_input(
                    "Anglais T2",            0.0, 20.0,
                    st.session_state.anglais_t2,  0.5, key="suivi_t2_ang")
                moy_lit_t2 = (st.session_state.francais_t2 +
                              st.session_state.histgeo_t2 +
                              st.session_state.anglais_t2) / 3
                delta_lit = moy_lit_t2 - moy_lit
                st.metric("Moy. Lit. T2", f"{moy_lit_t2:.2f}/20",
                          delta=f"{delta_lit:+.1f} vs T1",
                          delta_color="normal" if delta_lit >= 0 else "inverse")

            # Prévisualisation du nouveau statut
            serie_cible_t2 = "C" if SA > LA else "A"
            seuil_ok_t2    = moy_sci_t2 >= 10 if SA > LA else moy_lit_t2 >= 10
            nouveau_statut_t2 = "✅ CONFIRMATION" if seuil_ok_t2 else "⚠️ PROBATION"
            couleur_prev_t2   = "#10b981" if seuil_ok_t2 else "#f59e0b"
            fond_prev_t2      = "#ecfdf5" if seuil_ok_t2 else "#fffbeb"
            st.markdown(f"""
            <div style="background:{fond_prev_t2}; border:1px solid {couleur_prev_t2};
                        border-radius:10px; padding:0.7rem 1rem; margin-top:0.8rem;
                        text-align:center;">
                <span style="color:{couleur_prev_t2}; font-weight:700; font-size:0.95rem;">
                    Prévisualisation : {nouveau_statut_t2} en Série {serie_cible_t2}
                </span>
            </div>
            """, unsafe_allow_html=True)

            st.write("")
            if st.button("✅ Valider les notes T2 et recalculer", use_container_width=True, type="primary"):
                st.session_state.t2_renseigne = True
                st.session_state.chat_history = []   # reset l'agent pour le nouveau contexte
                rerun()

    # ==================================================================
    # RUBRIQUE SUIVI T3 — visible uniquement si statut == "probation"
    # ==================================================================
    if statut == "probation":
        st.divider()
        st.markdown("""
        <div style="background:linear-gradient(135deg,#fffbeb,#fef3c7);
                    border:2px dashed #f59e0b; border-radius:16px;
                    padding:1.2rem 1.4rem; margin:0.5rem 0;">
            <div style="font-size:1rem; font-weight:700; color:#78350f; margin-bottom:0.3rem;">
                ⚠️ Suivi Probation — Saisie des notes du 3e trimestre
            </div>
            <div style="font-size:0.85rem; color:#92400e;">
                Renseignez les notes du T3 pour rendre la décision d'orientation
                <strong>définitive</strong>.
            </div>
        </div>
        """, unsafe_allow_html=True)

        with st.expander("📋 Saisir les notes du 3e trimestre (décision finale)", expanded=True):
            st.caption("⚠️ Ces notes détermineront l'orientation définitive de l'élève.")
            ca, cb = st.columns(2)
            with ca:
                st.markdown("**🔬 Matières scientifiques**")
                st.session_state.maths_t3   = st.number_input(
                    "Mathématiques T3",   0.0, 20.0,
                    st.session_state.maths_t3,   0.5, key="suivi_t3_maths")
                st.session_state.sci_phy_t3 = st.number_input(
                    "Sciences Physiques T3", 0.0, 20.0,
                    st.session_state.sci_phy_t3, 0.5, key="suivi_t3_sp")
                st.session_state.svt_t3     = st.number_input(
                    "SVT T3",             0.0, 20.0,
                    st.session_state.svt_t3,     0.5, key="suivi_t3_svt")
                moy_sci_t3 = (st.session_state.maths_t3 +
                              st.session_state.sci_phy_t3 +
                              st.session_state.svt_t3) / 3
                notes_ref  = get_notes_actives()[0]
                moy_ref    = calc_moyennes(notes_ref)[0] if SA > LA else calc_moyennes(notes_ref)[1]
                delta_sci3 = moy_sci_t3 - moy_ref
                st.metric("Moy. Sci. T3", f"{moy_sci_t3:.2f}/20",
                          delta=f"{delta_sci3:+.1f} vs T2",
                          delta_color="normal" if delta_sci3 >= 0 else "inverse")
            with cb:
                st.markdown("**📖 Matières littéraires**")
                st.session_state.francais_t3 = st.number_input(
                    "Français T3",           0.0, 20.0,
                    st.session_state.francais_t3, 0.5, key="suivi_t3_fr")
                st.session_state.histgeo_t3  = st.number_input(
                    "Histoire-Géographie T3",0.0, 20.0,
                    st.session_state.histgeo_t3,  0.5, key="suivi_t3_hg")
                st.session_state.anglais_t3  = st.number_input(
                    "Anglais T3",            0.0, 20.0,
                    st.session_state.anglais_t3,  0.5, key="suivi_t3_ang")
                moy_lit_t3 = (st.session_state.francais_t3 +
                              st.session_state.histgeo_t3 +
                              st.session_state.anglais_t3) / 3
                moy_ref_lit = calc_moyennes(get_notes_actives()[0])[1]
                delta_lit3  = moy_lit_t3 - moy_ref_lit
                st.metric("Moy. Lit. T3", f"{moy_lit_t3:.2f}/20",
                          delta=f"{delta_lit3:+.1f} vs T2",
                          delta_color="normal" if delta_lit3 >= 0 else "inverse")

            # Prévisualisation décision finale
            serie_cible_t3 = "C" if SA > LA else "A"
            serie_risque_t3= "A" if SA > LA else "C"
            seuil_ok_t3    = moy_sci_t3 >= 10 if SA > LA else moy_lit_t3 >= 10
            if seuil_ok_t3:
                prev_label = f"✅ CONFIRMATION DÉFINITIVE — Série {serie_cible_t3}"
                prev_color = "#10b981"; prev_fond = "#ecfdf5"
            else:
                prev_label = f"🔄 ORIENTATION RÉVISÉE DÉFINITIVEMENT → Série {serie_risque_t3}"
                prev_color = "#f43f5e"; prev_fond = "#fef2f2"

            st.markdown(f"""
            <div style="background:{prev_fond}; border:2px solid {prev_color};
                        border-radius:12px; padding:0.9rem 1rem; margin-top:0.8rem;
                        text-align:center;">
                <span style="color:{prev_color}; font-weight:700; font-size:0.95rem;">
                    {prev_label}
                </span>
            </div>
            """, unsafe_allow_html=True)

            st.write("")
            if st.button("🏁 Valider les notes T3 — Décision définitive", use_container_width=True, type="primary"):
                st.session_state.t3_renseigne = True
                st.session_state.chat_history = []
                rerun()

    # ---- Agent conversationnel ----
    if conflit_type:
        st.divider()
        st.subheader("💬 Agent IA d'interpellation")
        prenom = st.session_state.prenom

        if conflit_type == "decale":
            msg_init = (f"Bonjour {prenom} ! Tes tests révèlent un fort potentiel scientifique "
                        f"(SA étalonnée = {SA:.1f}/20), mais tu as choisi la série A. "
                        f"Peux-tu expliquer pourquoi tu préfères la filière littéraire ?")
        else:
            msg_init = (f"Bonjour {prenom} ! Tu as de bonnes aptitudes scientifiques "
                        f"(SA = {SA:.1f}/20), mais tes notes en classe ({moy_sci:.1f}/20) "
                        f"sont encore insuffisantes. Qu'est-ce qui te pose le plus de difficultés ?")

        if not st.session_state.chat_history:
            st.session_state.chat_history.append({"role": "ia", "content": msg_init})

        for msg in st.session_state.chat_history:
            if msg["role"] == "ia":
                st.markdown(f'<div class="chat-ia">🤖 <strong>Agent CapAvenir :</strong><br>{msg["content"]}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="chat-user">{msg["content"]}<br><strong>— {prenom}</strong></div>', unsafe_allow_html=True)

        user_input = st.text_input("Réponse de l'élève :", placeholder="Écrivez ici...", key="chat_input_field")

        if st.button("Envoyer la réponse 📨", use_container_width=True):
            if user_input.strip():
                st.session_state.chat_history.append({"role": "user", "content": user_input})
                api_key = os.environ.get("ANTHROPIC_API_KEY", "")
                if api_key and len(api_key) > 20:
                    system_prompt = (
                        f"Tu es un conseiller d'orientation scolaire IA bienveillant pour le système camerounais. "
                        f"Tu aides {prenom} {st.session_state.nom}. "
                        f"SA étalonnée={SA:.1f}/20, LA étalonnée={LA:.1f}/20, "
                        f"Moy.Sci={moy_sci:.1f}/20, Moy.Lit={moy_lit:.1f}/20. "
                        f"Choix : {choix}. Projet : {st.session_state.projet_pro or 'non précisé'}. "
                        f"Revenu famille : {st.session_state.revenu}. "
                        f"Conflit : {conflit_type}. Réponds en 2-3 phrases, empathique et direct, en français."
                    )
                    messages_api = [
                        {"role": "assistant" if m["role"] == "ia" else "user", "content": m["content"]}
                        for m in st.session_state.chat_history
                    ]
                    try:
                        resp = requests.post(
                            "https://api.anthropic.com/v1/messages",
                            headers={"Content-Type": "application/json",
                                     "x-api-key": api_key,
                                     "anthropic-version": "2023-06-01"},
                            json={"model": "claude-sonnet-4-20250514",
                                  "max_tokens": 300, "system": system_prompt,
                                  "messages": messages_api},
                            timeout=15,
                        )
                        ia_reply = resp.json()["content"][0]["text"]
                    except Exception:
                        ia_reply = reponse_ia_simulee(
                            user_input, conflit_type, prenom, SA, moy_sci,
                            LA=LA, moy_lit=moy_lit,
                            projet_pro=st.session_state.projet_pro,
                            revenu=st.session_state.revenu,
                            chat_history=st.session_state.chat_history,
                            d48=st.session_state.d48, krx=st.session_state.krx)
                else:
                    ia_reply = reponse_ia_simulee(
                        user_input, conflit_type, prenom, SA, moy_sci,
                        LA=LA, moy_lit=moy_lit,
                        projet_pro=st.session_state.projet_pro,
                        revenu=st.session_state.revenu,
                        chat_history=st.session_state.chat_history,
                        d48=st.session_state.d48, krx=st.session_state.krx)

                st.session_state.chat_history.append({"role": "ia", "content": ia_reply})
                rerun()

    # Notes du conseiller
    st.divider()
    st.subheader("📝 Observations du conseiller")
    st.session_state.notes_conseiller = st.text_area(
        "Remarques / Observations (entretien, comportement, motivations…) :",
        value=st.session_state.notes_conseiller,
        placeholder="Saisir ici les observations du conseiller lors de l'entretien individuel...",
        height=80,
    )

    # ── Sauvegarde possible pour TOUS les statuts ──
    st.divider()
    sc1, sc2 = st.columns(2)
    with sc1:
        if st.button("💾 Sauvegarder le dossier", use_container_width=True):
            res = db.sauvegarder_dossier(st.session_state, etalonner)
            if res["succes"]:
                st.success(f"{res['message']}")
            else:
                st.error(res["message"])
    with sc2:
        if statut in ("attente", "probation"):
            hint = "T2 requis" if statut == "attente" else "T3 requis"
            st.markdown(f"""
            <div class="alert-attente" style="padding:0.6rem 1rem; font-size:0.85rem;">
                🔒 Fiche provisoire — <strong>{hint} pour finaliser</strong>
            </div>""", unsafe_allow_html=True)

    st.write("")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Retour ⬅️", use_container_width=True):
            st.session_state.step = 2
            rerun()
    with c2:
        if st.button("Générer la fiche d'orientation 📄", use_container_width=True, type="primary"):
            st.session_state.step = 4
            rerun()

# =====================================================================
# ETAPE 4 — FICHE D'ORIENTATION
# =====================================================================
elif step == 4:
    serie   = st.session_state.orientation_finale or "?"
    prenom  = st.session_state.prenom
    nom     = st.session_state.nom
    SA_brut, LA_brut, SA_etal, LA_etal = calc_aptitudes()
    notes, trim_label = get_notes_actives()
    moy_sci, moy_lit  = calc_moyennes(notes)
    score_conf = st.session_state.score_confiance
    date_gen = datetime.now().strftime("%d/%m/%Y à %H:%M")

    st.subheader("📄 Fiche d'Orientation Scolaire")
    st.markdown(f"""
    <div class="fiche-box">
        <div style="text-align:center; margin-bottom:1rem;">
            <div style="font-size:0.8rem; color:#64748b; text-transform:uppercase;">République du Cameroun — Ministère de l'Éducation de Base</div>
            <div style="font-weight:700; font-size:1.3rem; color:#0f172a; margin:0.3rem 0;">FICHE D'ORIENTATION SCOLAIRE</div>
            <div style="font-size:0.8rem; color:#94a3b8;">Générée le {date_gen} — CapAvenir CMR v2.0</div>
        </div>
        <hr style="border-color:#e2e8f0; margin:0.8rem 0;">
        <p><strong>Nom &amp; Prénom :</strong> {nom} {prenom} &nbsp;|&nbsp; <strong>Âge :</strong> {st.session_state.age} ans &nbsp;|&nbsp; <strong>Sexe :</strong> {st.session_state.sexe}</p>
        <p><strong>Lycée :</strong> {st.session_state.lycee or "non renseigné"} &nbsp;|&nbsp; <strong>Choix personnel :</strong> {st.session_state.choix_personnel}</p>
        <p><strong>Projet professionnel :</strong> {st.session_state.projet_pro or "non renseigné"}</p>
        <p><strong>Revenu familial :</strong> {st.session_state.revenu}</p>
    </div>
    """, unsafe_allow_html=True)

    # Scores des tests (bruts + étalonnés)
    st.markdown("**Tests psychotechniques — Scores bruts & étalonnés :**")
    cols5 = st.columns(5)
    for col, (code, brut) in zip(cols5, [
        ("D48", st.session_state.d48), ("KRX", st.session_state.krx),
        ("MECA", st.session_state.meca), ("BV11", st.session_state.bv11),
        ("PRC", st.session_state.prc)
    ]):
        etal = etalonner(brut)
        col.markdown(f"""
        <div class="score-card">
            <div style="font-weight:700;">{code}</div>
            <div class="score-raw">Brut : {brut:.1f}</div>
            <div class="score-etalon">{etal:.1f}/20</div>
        </div>
        """, unsafe_allow_html=True)

    st.write("")
    st.markdown("**Aptitudes calculées (étalonnées) & moyennes scolaires :**")
    m1, m2, m3, m4, m5 = st.columns(5)
    m1.metric("SA brute",             f"{SA_brut:.2f}/20")
    m2.metric("SA étalonnée",         f"{SA_etal:.2f}/20")
    m3.metric("LA étalonnée",         f"{LA_etal:.2f}/20")
    m4.metric(f"Moy. Sci. ({trim_label})", f"{moy_sci:.2f}/20")
    m5.metric(f"Moy. Lit. ({trim_label})", f"{moy_lit:.2f}/20")

    # Badge orientation + statut
    badge_color = "#10b981" if "C" in serie else ("#3b82f6" if "A" in serie else "#f59e0b")
    bg_color    = "#ecfdf5" if "C" in serie else ("#eff6ff" if "A" in serie else "#fffbeb")
    conf_color  = "#10b981" if score_conf >= 70 else ("#f59e0b" if score_conf >= 50 else "#ef4444")

    STATUT_LABELS = {
        "confirme":   ("✅ CONFIRMÉ",              "#10b981", "#ecfdf5"),
        "attente":    ("⏳ EN ATTENTE — T2 requis", "#0ea5e9", "#f0f9ff"),
        "probation":  ("⚠️ PROBATION",              "#f59e0b", "#fffbeb"),
        "revise":     ("🔄 RÉVISÉ — Définitif",     "#f43f5e", "#fef2f2"),
        "indetermine":("❓ INDÉTERMINÉ",             "#94a3b8", "#f8fafc"),
    }
    s_label, s_color, s_bg = STATUT_LABELS.get(
        st.session_state.statut, ("—", "#94a3b8", "#f8fafc"))

    st.write("")
    st.markdown(f"""
    <div style="text-align:center; padding:1.5rem; background:{bg_color};
                border-radius:16px; border:2px solid {badge_color}; margin:1rem 0;
                box-shadow: 0 4px 20px rgba(15,23,42,0.08);">
        <div style="font-size:0.8rem; color:#64748b; margin-bottom:0.5rem; text-transform:uppercase; letter-spacing:0.08em;">Orientation recommandée</div>
        <span style="background:{badge_color}; color:white; padding:0.5rem 2.5rem;
                     border-radius:24px; font-weight:700; font-size:1.5rem;
                     box-shadow:0 4px 12px rgba(0,0,0,0.15);">SÉRIE {serie}</span>
        <div style="margin-top:1rem; display:flex; justify-content:center; gap:16px; flex-wrap:wrap;">
            <span style="background:{s_bg}; color:{s_color}; padding:0.3rem 1rem;
                         border-radius:20px; font-size:0.82rem; font-weight:600;
                         border:1px solid {s_color};">{s_label}</span>
            <span style="background:#f8fafc; color:{conf_color}; padding:0.3rem 1rem;
                         border-radius:20px; font-size:0.82rem; font-weight:600;
                         border:1px solid {conf_color};">Confiance : {score_conf} %</span>
            <span style="background:#f8fafc; color:#64748b; padding:0.3rem 1rem;
                         border-radius:20px; font-size:0.82rem;
                         border:1px solid #e2e8f0;">Trimestre : {trim_label}</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if st.session_state.notes_conseiller:
        st.info(f"📝 Observations du conseiller : {st.session_state.notes_conseiller}")

    if st.session_state.chat_history:
        derniere_ia = [m["content"] for m in st.session_state.chat_history if m["role"] == "ia"]
        if derniere_ia:
            st.info(f"🤖 Synthèse entretien IA : {derniere_ia[-1]}")

    STATUT_TXT = {
        "confirme":   "CONFIRMÉ",
        "attente":    "EN ATTENTE — Notes T2 requises",
        "probation":  "PROBATION — Amélioration requise avant T3",
        "revise":     "RÉVISÉ DÉFINITIVEMENT",
        "indetermine":"INDÉTERMINÉ — Entretien requis",
    }
    statut_txt_export = STATUT_TXT.get(st.session_state.statut, "—")

    # Export enrichi
    fiche_txt = f"""FICHE D'ORIENTATION SCOLAIRE — CapAvenir CMR v2.1
Générée le {date_gen}

ÉLÈVE     : {nom} {prenom} | {st.session_state.age} ans | {st.session_state.sexe}
LYCÉE     : {st.session_state.lycee or 'non renseigné'}
CHOIX     : {st.session_state.choix_personnel}
PROJET    : {st.session_state.projet_pro or 'non renseigné'}
REVENU    : {st.session_state.revenu}

TESTS PSYCHOTECHNIQUES
  D48   : Brut={st.session_state.d48:.1f}/20  |  Étalonnée={etalonner(st.session_state.d48):.1f}/20
  KRX   : Brut={st.session_state.krx:.1f}/20  |  Étalonnée={etalonner(st.session_state.krx):.1f}/20
  MECA  : Brut={st.session_state.meca:.1f}/20  |  Étalonnée={etalonner(st.session_state.meca):.1f}/20
  BV11  : Brut={st.session_state.bv11:.1f}/20  |  Étalonnée={etalonner(st.session_state.bv11):.1f}/20
  PRC   : Brut={st.session_state.prc:.1f}/20  |  Étalonnée={etalonner(st.session_state.prc):.1f}/20

APTITUDES CALCULÉES
  SA brute    = (KRX + D48) / 2        = {SA_brut:.2f}/20
  SA étalonnée= (KRX_e + D48_e) / 2   = {SA_etal:.2f}/20
  LA étalonnée= (BV11_e + PRC_e) / 2  = {LA_etal:.2f}/20

NOTES SCOLAIRES ({trim_label})
  Maths={notes['maths']:.1f}  SciPhy={notes['sci_phy']:.1f}  SVT={notes['svt']:.1f}
  Français={notes['francais']:.1f}  HistGeo={notes['histgeo']:.1f}  Anglais={notes['anglais']:.1f}
  Moyenne Scientifique = {moy_sci:.2f}/20
  Moyenne Littéraire   = {moy_lit:.2f}/20

DÉCISION
  ORIENTATION : SÉRIE {serie}
  STATUT      : {statut_txt_export}
  CONFIANCE   : {score_conf} %

OBSERVATIONS DU CONSEILLER :
{st.session_state.notes_conseiller or 'Aucune observation saisie.'}
"""

    st.write("")    
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        if st.button("Retour au diagnostic ⬅️", use_container_width=True):
            st.session_state.step = 3
            rerun()
    with c2:
        if st.button("💾 Sauvegarder le dossier", use_container_width=True, type="primary"):
            res = db.sauvegarder_dossier(st.session_state, etalonner)
            if res["succes"]:
                st.success(res["message"])
            else:
                st.error(res["message"])
    with c3:
        st.download_button(
            "📥 Télécharger (.txt)",
            fiche_txt,
            file_name=f"orientation_{nom}_{prenom}_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with c4:
        derniere_ia = [m["content"] for m in st.session_state.chat_history if m["role"] == "ia"]
        ia_synthese  = derniere_ia[-1] if derniere_ia else ""
        pdf_data = {
            "nom": nom, "prenom": prenom,
            "age": st.session_state.age, "sexe": st.session_state.sexe,
            "lycee": st.session_state.lycee or "Non renseigné",
            "choix": st.session_state.choix_personnel,
            "projet_pro": st.session_state.projet_pro or "Non renseigné",
            "revenu": st.session_state.revenu,
            "d48": st.session_state.d48, "d48_e": etalonner(st.session_state.d48),
            "krx": st.session_state.krx, "krx_e": etalonner(st.session_state.krx),
            "meca": st.session_state.meca, "meca_e": etalonner(st.session_state.meca),
            "bv11": st.session_state.bv11, "bv11_e": etalonner(st.session_state.bv11),
            "prc": st.session_state.prc, "prc_e": etalonner(st.session_state.prc),
            "SA_brut": SA_brut, "SA_etal": SA_etal,
            "LA_brut": LA_brut, "LA_etal": LA_etal,
            "notes": notes, "trim_label": trim_label,
            "moy_sci": moy_sci, "moy_lit": moy_lit,
            "serie": serie, "probation": st.session_state.probation,
            "statut_label": statut_txt_export,
            "score_conf": score_conf,
            "notes_conseiller": st.session_state.notes_conseiller,
            "ia_synthese": ia_synthese,
            "date_gen": date_gen,
        }
        pdf_bytes = generate_pdf_fiche(pdf_data)
        st.download_button(
            "📄 Télécharger (.pdf)",
            pdf_bytes,
            file_name=f"orientation_{nom}_{prenom}_{datetime.now().strftime('%Y%m%d')}.pdf",
            mime="application/pdf",
            use_container_width=True,
            type="primary",
        )
    with c5:
        if st.button("Nouveau dossier 🔄", use_container_width=True):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            rerun()

    st.write("")
    st.markdown(
        "<center><small style='color:#94a3b8;'>CapAvenir CMR v2.1 — 2025 — Mémoire ENS Filière Informatique Niveau 5</small></center>",
        unsafe_allow_html=True
    )