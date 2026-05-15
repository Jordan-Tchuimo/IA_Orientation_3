import streamlit as st
import plotly.graph_objects as go
import os, json, hashlib, requests, io
from datetime import datetime
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
from capavenir_database import CapAvenirDB
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                 Table, TableStyle, HRFlowable, KeepTogether)

@st.cache_resource
def get_db():
    return CapAvenirDB()
db = get_db()

# ─── Table des identifiants élèves (créée une seule fois) ────────────
def _init_eleve_credentials():
    db.conn.execute("""
        CREATE TABLE IF NOT EXISTS eleve_credentials (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            nom           TEXT NOT NULL,
            prenom        TEXT NOT NULL,
            lycee         TEXT NOT NULL,
            pwd_hash      TEXT NOT NULL,
            date_creation TEXT NOT NULL,
            UNIQUE(nom, prenom, lycee)
        )
    """)
    db.conn.commit()

_init_eleve_credentials()

def _save_eleve_credential(nom: str, prenom: str, lycee: str, password: str):
    """Enregistre (ou met à jour) le mot de passe d'un élève."""
    h   = hashlib.sha256(password.encode()).hexdigest()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    db.conn.execute("""
        INSERT INTO eleve_credentials (nom, prenom, lycee, pwd_hash, date_creation)
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(nom, prenom, lycee) DO UPDATE SET
            pwd_hash=excluded.pwd_hash
    """, (nom.strip().upper(), prenom.strip().capitalize(), lycee.strip(), h, now))
    db.conn.commit()

def _verify_eleve_credential(nom: str, prenom: str, password: str) -> bool:
    """Vérifie le mot de passe d'un élève (nom + prénom + mdp)."""
    h = hashlib.sha256(password.encode()).hexdigest()
    row = db.conn.execute(
        """SELECT id FROM eleve_credentials
           WHERE UPPER(nom)=UPPER(?) AND UPPER(prenom)=UPPER(?) AND pwd_hash=?""",
        (nom.strip(), prenom.strip(), h)
    ).fetchone()
    return row is not None

def _eleve_exists(nom: str, prenom: str) -> bool:
    """Vérifie si un élève a un compte (credential enregistré)."""
    row = db.conn.execute(
        "SELECT id FROM eleve_credentials WHERE UPPER(nom)=UPPER(?) AND UPPER(prenom)=UPPER(?)",
        (nom.strip(), prenom.strip())
    ).fetchone()
    return row is not None

# Mot de passe conseiller (SHA-256) — défaut: capavenir2025
# Pour changer: python3 -c "import hashlib; print(hashlib.sha256(b'VOTRE_MDT').hexdigest())"
CONSEILLER_PWD_HASH = "de07220a5f97bf76fc0ad428c9d7e90318d853fdc227484172e6d29da0c1c4bf"
CONSEILLER_LOGIN    = "conseiller"   # identifiant par défaut

def check_password(login: str, password: str) -> bool:
    h = hashlib.sha256(password.encode()).hexdigest()
    return login.strip().lower() == CONSEILLER_LOGIN and h == CONSEILLER_PWD_HASH


# =====================================================================
# COMPATIBILITE STREAMLIT
# =====================================================================
def rerun():
    try:
        st.rerun()
    except AttributeError:
        st.experimental_rerun()

# =====================================================================
# CONFIGURATION
# =====================================================================
st.set_page_config(
    page_title="CapAvenir CMR - Orientation IA",
    page_icon="🎓",
    layout="centered",
)

def inject_css(dark: bool):
    if dark:
        bg           = "linear-gradient(160deg,#0f172a 0%,#1e1040 50%,#0a2318 100%)"
        txt          = "#e2e8f0"
        card_bg      = "linear-gradient(145deg,#1e293b,#1a1f35)"
        card_brd     = "#334155"
        metric_bg    = "linear-gradient(135deg,#1e293b,#141b2d)"
        step_todo_bg = "#1e293b"; step_todo_c="#64748b"
        fiche_bg     = "linear-gradient(135deg,#1e293b,#172032)"
        fiche_brd    = "#7c3aed"
        tab_sel      = "#10b981"
        hover_row    = "#1e293b"
        dash_card    = "linear-gradient(145deg,#1e2d3d,#1a2035)"
        input_brd    = "#475569"
    else:
        bg           = "linear-gradient(160deg,#fff7ed 0%,#fdf4ff 45%,#ecfdf5 100%)"
        txt          = "#1e1b4b"
        card_bg      = "linear-gradient(145deg,#ffffff,#fafafa)"
        card_brd     = "#e5e7eb"
        metric_bg    = "linear-gradient(145deg,#ffffff,#f9fafb)"
        step_todo_bg = "#f3f4f6"; step_todo_c="#9ca3af"
        fiche_bg     = "linear-gradient(145deg,#fffbeb,#fef9c3)"
        fiche_brd    = "#f59e0b"
        tab_sel      = "#f59e0b"
        hover_row    = "#fef3c7"
        dash_card    = "linear-gradient(145deg,#ffffff,#fafafa)"
        input_brd    = "#d1d5db"

    st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

/* ══ GLOBAL ══ */
.stApp {{
    background: {bg} !important;
    color: {txt};
    font-family: 'Inter', sans-serif !important;
}}
h1,h2,h3,h4,h5,h6 {{ color: {txt}; }}

/* ══ HEADER ANIMÉ ══ */
.header-box {{
    background: linear-gradient(135deg,#6d28d9 0%,#db2777 35%,#f97316 65%,#10b981 100%);
    background-size: 300% 300%;
    animation: gradientShift 6s ease infinite;
    color: white; padding: 2.5rem 2.5rem 2rem;
    border-radius: 24px; text-align: center; margin-bottom: 1.8rem;
    box-shadow: 0 12px 40px rgba(109,40,217,0.35), 0 4px 16px rgba(0,0,0,0.15);
    border: 2px solid rgba(255,255,255,0.25);
    position: relative; overflow: hidden;
}}
.header-box::before {{
    content: '';
    position: absolute; top:-50%; left:-50%;
    width: 200%; height: 200%;
    background: radial-gradient(circle,rgba(255,255,255,0.08) 0%,transparent 70%);
    animation: pulse 4s ease-in-out infinite;
}}
@keyframes gradientShift {{
    0%   {{ background-position: 0% 50%; }}
    50%  {{ background-position: 100% 50%; }}
    100% {{ background-position: 0% 50%; }}
}}
@keyframes pulse {{
    0%,100% {{ opacity:0.5; transform: scale(1); }}
    50%     {{ opacity:1;   transform: scale(1.05); }}
}}
.header-badge {{
    display: inline-block;
    background: rgba(255,255,255,0.18);
    border: 1px solid rgba(255,255,255,0.3);
    border-radius: 20px; padding: 0.25rem 0.9rem;
    font-size: 0.72rem; font-weight: 600; letter-spacing: 0.08em;
    text-transform: uppercase; margin-top: 0.6rem;
    backdrop-filter: blur(4px);
}}

/* ══ ALERTES ══ */
.alert-success {{
    background: linear-gradient(135deg,#ecfdf5,#d1fae5);
    border-left: 5px solid #10b981; padding: 1rem 1.2rem;
    border-radius: 12px; color: #065f46; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(16,185,129,0.18);
}}
.alert-warning {{
    background: linear-gradient(135deg,#fffbeb,#fef3c7);
    border-left: 5px solid #f59e0b; padding: 1rem 1.2rem;
    border-radius: 12px; color: #78350f; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(245,158,11,0.18);
}}
.alert-info {{
    background: linear-gradient(135deg,#fdf4ff,#f3e8ff);
    border-left: 5px solid #a855f7; padding: 1rem 1.2rem;
    border-radius: 12px; color: #581c87; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(168,85,247,0.18);
}}
.alert-danger {{
    background: linear-gradient(135deg,#fef2f2,#fee2e2);
    border-left: 5px solid #ef4444; padding: 1rem 1.2rem;
    border-radius: 12px; color: #7f1d1d; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(239,68,68,0.18);
}}
.alert-attente {{
    background: linear-gradient(135deg,#fff7ed,#fed7aa);
    border-left: 5px solid #f97316; padding: 1rem 1.2rem;
    border-radius: 12px; color: #7c2d12; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(249,115,22,0.18);
}}
.alert-revise {{
    background: linear-gradient(135deg,#fdf2f8,#fce7f3);
    border-left: 5px solid #ec4899; padding: 1rem 1.2rem;
    border-radius: 12px; color: #831843; margin: 0.6rem 0;
    box-shadow: 0 3px 12px rgba(236,72,153,0.18);
}}

/* ══ CHAT IA ══ */
.chat-ia {{
    background: linear-gradient(135deg,#f3e8ff,#ddd6fe);
    border-radius: 0 18px 18px 18px; padding: 1rem 1.2rem;
    margin: 0.5rem 0; color: #4c1d95;
    box-shadow: 0 3px 12px rgba(124,58,237,0.15);
    border-left: 4px solid #7c3aed;
    animation: fadeInLeft 0.3s ease;
}}
.chat-user {{
    background: linear-gradient(135deg,#fef3c7,#fde68a);
    border-radius: 18px 0 18px 18px; padding: 1rem 1.2rem;
    margin: 0.5rem 0; color: #78350f; text-align: right;
    box-shadow: 0 3px 12px rgba(245,158,11,0.18);
    border-right: 4px solid #f59e0b;
    animation: fadeInRight 0.3s ease;
}}
@keyframes fadeInLeft  {{ from {{ opacity:0; transform:translateX(-10px); }} to {{ opacity:1; transform:translateX(0); }} }}
@keyframes fadeInRight {{ from {{ opacity:0; transform:translateX(10px);  }} to {{ opacity:1; transform:translateX(0); }} }}

/* ══ FICHE ══ */
.fiche-box {{
    background: {fiche_bg}; border: 2px solid {fiche_brd};
    border-radius: 18px; padding: 2rem; margin: 1rem 0;
    box-shadow: 0 6px 24px rgba(251,191,36,0.14);
}}

/* ══ SCORE CARDS ══ */
.score-card {{
    text-align: center; background: {card_bg};
    border: 1px solid {card_brd}; border-radius: 14px;
    padding: 1rem 0.4rem;
    box-shadow: 0 3px 10px rgba(0,0,0,0.06);
    transition: all 0.25s ease;
}}
.score-card:hover {{
    transform: translateY(-4px);
    box-shadow: 0 8px 20px rgba(124,58,237,0.18);
    border-color: #a855f7;
}}
.score-raw    {{ font-size: 0.73rem; color: #94a3b8; margin-bottom: 2px; }}
.score-etalon {{ font-size: 1.15rem; font-weight: 700; color: #10b981; }}
.score-label  {{ font-size: 0.78rem; font-weight: 600; color: {txt}; }}

/* ══ QUESTIONS TESTS ══ */
.test-question {{
    background: linear-gradient(135deg,#fdf4ff,#ede9fe);
    border: 1px solid #c4b5fd; border-radius: 14px;
    padding: 1.1rem 1.3rem; margin: 0.7rem 0;
    box-shadow: 0 3px 8px rgba(139,92,246,0.1);
    border-left: 5px solid #8b5cf6;
    transition: box-shadow 0.2s;
}}
.test-question:hover {{ box-shadow: 0 4px 14px rgba(139,92,246,0.2); }}

/* ══ STEPPER ══ */
.step-active {{
    text-align:center;
    background: linear-gradient(135deg,#f59e0b,#ef4444,#ec4899);
    color: white; border-radius: 12px; padding: 0.45rem 0.2rem;
    font-size: 0.75rem; font-weight: 700;
    box-shadow: 0 5px 16px rgba(245,158,11,0.45);
}}
.step-done {{
    text-align:center;
    background: linear-gradient(135deg,#d1fae5,#a7f3d0);
    color: #065f46; border-radius: 12px; padding: 0.45rem 0.2rem;
    font-size: 0.75rem; font-weight: 600;
}}
.step-todo {{
    text-align:center; background: {step_todo_bg}; color: {step_todo_c};
    border-radius: 12px; padding: 0.45rem 0.2rem; font-size: 0.75rem;
}}

/* ══ DASHBOARD (mode conseiller) ══ */
.dash-stat-card {{
    background: {dash_card};
    border-radius: 16px; padding: 1.2rem 1rem;
    border: 1px solid {card_brd};
    box-shadow: 0 4px 14px rgba(0,0,0,0.06);
    text-align: center; transition: all 0.25s ease;
}}
.dash-stat-card:hover {{
    transform: translateY(-3px);
    box-shadow: 0 8px 22px rgba(0,0,0,0.1);
}}
.dash-stat-val  {{ font-size: 2.2rem; font-weight: 800; line-height: 1.1; }}
.dash-stat-lbl  {{ font-size: 0.72rem; font-weight: 600; text-transform: uppercase;
                   letter-spacing: 0.07em; opacity: 0.7; margin-top: 4px; }}
.dash-row {{ display:flex; align-items:center; padding: 0.65rem 0.8rem;
             border-radius: 10px; margin-bottom: 4px; transition: background 0.15s; }}
.dash-row:hover {{ background: {hover_row}; }}
.dash-badge {{
    display: inline-block; border-radius: 20px;
    padding: 0.18rem 0.75rem; font-size: 0.72rem; font-weight: 700;
    letter-spacing: 0.03em; white-space: nowrap;
}}
.badge-c        {{ background:#d1fae5; color:#065f46; }}
.badge-a        {{ background:#dbeafe; color:#1e3a8a; }}
.badge-confirme {{ background:#d1fae5; color:#065f46; }}
.badge-attente  {{ background:#fff7ed; color:#7c2d12; }}
.badge-probation{{ background:#fef3c7; color:#78350f; }}
.badge-revise   {{ background:#fce7f3; color:#831843; }}
.badge-indetermine {{ background:#f3f4f6; color:#374151; }}
.search-box {{
    background: {card_bg};
    border: 2px solid {input_brd}; border-radius: 12px;
    padding: 0.6rem 1rem; width: 100%;
    font-size: 0.9rem; color: {txt};
    outline: none; transition: border-color 0.2s;
}}
.search-box:focus {{ border-color: #a855f7; }}

/* ══ OVERRIDES STREAMLIT ══ */
div[data-testid="stMetric"] {{
    background: {metric_bg}; border: 1px solid {card_brd};
    border-radius: 14px; padding: 0.9rem;
    box-shadow: 0 3px 10px rgba(0,0,0,0.05);
}}
div[data-testid="stMetric"] label {{
    font-size: 0.75rem !important; font-weight: 600 !important;
    text-transform: uppercase !important; letter-spacing: 0.06em !important;
    opacity: 0.7;
}}
.stButton > button[kind="primary"] {{
    background: linear-gradient(135deg,#f59e0b,#ef4444) !important;
    border: none !important; border-radius: 12px !important;
    box-shadow: 0 5px 14px rgba(245,158,11,0.38) !important;
    font-weight: 700 !important; color: white !important;
    letter-spacing: 0.02em !important; transition: all 0.2s !important;
}}
.stButton > button[kind="primary"]:hover {{
    background: linear-gradient(135deg,#d97706,#dc2626) !important;
    box-shadow: 0 7px 20px rgba(245,158,11,0.48) !important;
    transform: translateY(-1px) !important;
}}
div[data-testid="stTabs"] button[aria-selected="true"] {{
    border-bottom: 3px solid {tab_sel} !important;
    color: {tab_sel} !important; font-weight: 600 !important;
}}
/* Dividers */
hr {{ border-color: {card_brd} !important; }}
/* Section labels */
.label-sci   {{ background:#d1fae5; color:#065f46; }}
.label-lit   {{ background:#dbeafe; color:#1e3a8a; }}
.label-test  {{ background:#ede9fe; color:#4c1d95; }}
.section-label-colored {{
    font-size: 0.68rem; font-weight: 700; letter-spacing: 0.1em;
    text-transform: uppercase; padding: 0.3rem 0.85rem;
    border-radius: 20px; display: inline-block; margin: 0.8rem 0 0.4rem 0;
}}
</style>
""", unsafe_allow_html=True)


# =====================================================================
# GENERATION PDF — Fiche d'orientation (ReportLab) — UNE SEULE PAGE
# =====================================================================
def generate_pdf_fiche(data: dict) -> bytes:
    """Fiche d'orientation compacte sur une seule page A4. Robuste aux valeurs None."""

    # ── Sécuriser toutes les valeurs numériques ──
    def _f(v, default=0.0):
        try:
            return float(v) if v is not None else default
        except (TypeError, ValueError):
            return default

    def _s(v, default="—", max_len=None):
        r = str(v).strip() if v else default
        return r[:max_len] if max_len and r != default else r

    buffer = io.BytesIO()

    # Couleurs
    CMR_GREEN  = colors.HexColor("#10b981")
    CMR_DARK   = colors.HexColor("#0f172a")
    CMR_BLUE   = colors.HexColor("#1e3a5f")
    CMR_LIGHT  = colors.HexColor("#f8fafc")
    CMR_GRAY   = colors.HexColor("#64748b")
    CMR_AMBER  = colors.HexColor("#f59e0b")
    serie      = _s(data.get("serie"), "?")
    SERIE_CLR  = CMR_GREEN if "C" in serie else (colors.HexColor("#3b82f6") if "A" in serie else CMR_AMBER)

    doc = SimpleDocTemplate(buffer, pagesize=A4,
        leftMargin=1.4*cm, rightMargin=1.4*cm,
        topMargin=1.0*cm, bottomMargin=1.0*cm)

    # Largeur utile = 21cm - 1.4 - 1.4 = 18.2cm
    W = 18.2 * cm

    # Styles compacts
    sH   = ParagraphStyle("h",   fontName="Helvetica-Bold", fontSize=12, textColor=CMR_DARK,  alignment=TA_CENTER, spaceAfter=1)
    sSub = ParagraphStyle("sub", fontName="Helvetica",      fontSize=7,  textColor=colors.white, alignment=TA_CENTER, spaceAfter=1)
    sSec = ParagraphStyle("sec", fontName="Helvetica-Bold", fontSize=8,  textColor=CMR_GREEN, spaceBefore=5, spaceAfter=2)
    sB   = ParagraphStyle("b",   fontName="Helvetica",      fontSize=7.5,textColor=CMR_DARK,  spaceAfter=1)
    sC   = ParagraphStyle("c",   fontName="Helvetica",      fontSize=7.5,textColor=CMR_DARK,  alignment=TA_CENTER)
    sSer = ParagraphStyle("ser", fontName="Helvetica-Bold", fontSize=18, textColor=SERIE_CLR, alignment=TA_CENTER)
    sFt  = ParagraphStyle("ft",  fontName="Helvetica",      fontSize=6,  textColor=CMR_GRAY,  alignment=TA_CENTER)
    sSig = ParagraphStyle("sig", fontName="Helvetica",      fontSize=7.5,textColor=CMR_DARK)

    story = []

    # ── En-tête bilingue (3 colonnes) — W = 18.2cm → 5+8.2+5 ──
    hdr = Table([[
        Paragraph("REPUBLIQUE DU CAMEROUN<br/>Paix — Travail — Patrie", sSub),
        Paragraph("<b>FICHE D'ORIENTATION SCOLAIRE</b><br/>CapAvenir CMR v2.1 — " + _s(data.get("date_gen")), sH),
        Paragraph("REPUBLIC OF CAMEROON<br/>Peace — Work — Fatherland", sSub),
    ]], colWidths=[5*cm, 8.2*cm, 5*cm])
    hdr.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,-1),CMR_DARK),
        ("TEXTCOLOR",(0,0),(-1,-1),colors.white),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),6),
        ("BOTTOMPADDING",(0,0),(-1,-1),6),
    ]))
    story.append(hdr)
    story.append(Spacer(1, 0.15*cm))

    # ── Identité — 3 colonnes, total 18.2cm ──
    story.append(Paragraph("INFORMATIONS PERSONNELLES", sSec))
    id_data = [
        [Paragraph(f"<b>Nom &amp; Prénom :</b> {_s(data.get('nom'))} {_s(data.get('prenom'))}", sB),
         Paragraph(f"<b>Âge :</b> {_s(data.get('age'))} ans  <b>Sexe :</b> {_s(data.get('sexe'))}", sB),
         Paragraph(f"<b>Lycée :</b> {_s(data.get('lycee'), max_len=28)}", sB)],
        [Paragraph(f"<b>Choix :</b> {_s(data.get('choix'))}", sB),
         Paragraph(f"<b>Projet :</b> {_s(data.get('projet_pro'), max_len=32)}", sB),
         Paragraph(f"<b>Revenu :</b> {_s(data.get('revenu'), max_len=26)}", sB)],
    ]
    # 6.5 + 5.8 + 5.9 = 18.2cm
    id_t = Table(id_data, colWidths=[6.5*cm, 5.8*cm, 5.9*cm])
    id_t.setStyle(TableStyle([
        ("FONTSIZE",(0,0),(-1,-1),7.5),
        ("BACKGROUND",(0,0),(-1,-1),CMR_LIGHT),
        ("GRID",(0,0),(-1,-1),0.2,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(id_t)
    story.append(Spacer(1, 0.1*cm))

    # ── Tests + Notes côte à côte ──
    story.append(Paragraph("TESTS PSYCHOTECHNIQUES & NOTES SCOLAIRES", sSec))

    # Colonne gauche : tests — 4.6+1.5+1.5 = 7.6cm
    test_data = [
        [Paragraph("<b>Test</b>", sC), Paragraph("<b>Brut</b>", sC), Paragraph("<b>Étal.</b>", sC)],
        ["D48 — Logique",      f"{_f(data.get('d48')):.1f}",  f"{_f(data.get('d48_e')):.1f}"],
        ["KRX — Maths",        f"{_f(data.get('krx')):.1f}",  f"{_f(data.get('krx_e')):.1f}"],
        ["MECA — Mécanique",   f"{_f(data.get('meca')):.1f}", f"{_f(data.get('meca_e')):.1f}"],
        ["BV11 — Vocabulaire", f"{_f(data.get('bv11')):.1f}", f"{_f(data.get('bv11_e')):.1f}"],
        ["PRC — Proverbes",    f"{_f(data.get('prc')):.1f}",  f"{_f(data.get('prc_e')):.1f}"],
        [Paragraph("<b>SA étal.</b>", sB), "—", Paragraph(f"<b>{_f(data.get('SA_etal')):.2f}</b>", sB)],
        [Paragraph("<b>LA étal.</b>", sB), "—", Paragraph(f"<b>{_f(data.get('LA_etal')):.2f}</b>", sB)],
    ]
    t_tests = Table(test_data, colWidths=[4.6*cm, 1.5*cm, 1.5*cm])
    t_tests.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),CMR_DARK),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
        ("FONTSIZE",(0,0),(-1,-1),7.5),
        ("ALIGN",(1,0),(-1,-1),"CENTER"),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[CMR_LIGHT,colors.white]),
        ("GRID",(0,0),(-1,-1),0.2,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),4),
        ("TEXTCOLOR",(2,7),(2,8),CMR_GREEN),("FONTNAME",(2,7),(2,8),"Helvetica-Bold"),
    ]))

    # Colonne droite : notes — 2.8+1.2+2.8+1.2 = 8.0cm (espacement 0.6)
    notes_d = data.get("notes") or {}
    notes_data = [
        [Paragraph("<b>Matière</b>",sC), Paragraph("<b>Note</b>",sC),
         Paragraph("<b>Matière</b>",sC), Paragraph("<b>Note</b>",sC)],
        ["Mathématiques",         f"{_f(notes_d.get('maths')):.1f}",
         "Français",              f"{_f(notes_d.get('francais')):.1f}"],
        ["Sciences Phys.",        f"{_f(notes_d.get('sci_phy')):.1f}",
         "Histoire-Géo",          f"{_f(notes_d.get('histgeo')):.1f}"],
        ["SVT",                   f"{_f(notes_d.get('svt')):.1f}",
         "Anglais",               f"{_f(notes_d.get('anglais')):.1f}"],
        [Paragraph("<b>Moy. Sci.</b>",sB), Paragraph(f"<b>{_f(data.get('moy_sci')):.2f}</b>",sB),
         Paragraph("<b>Moy. Lit.</b>",sB), Paragraph(f"<b>{_f(data.get('moy_lit')):.2f}</b>",sB)],
    ]
    t_notes = Table(notes_data, colWidths=[2.9*cm, 1.1*cm, 2.9*cm, 1.1*cm])
    t_notes.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),CMR_BLUE),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
        ("FONTSIZE",(0,0),(-1,-1),7.5),
        ("ALIGN",(1,0),(1,-1),"CENTER"),("ALIGN",(3,0),(3,-1),"CENTER"),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[CMR_LIGHT,colors.white]),
        ("GRID",(0,0),(-1,-1),0.2,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),4),
        ("TEXTCOLOR",(1,5),(1,5),CMR_GREEN),("FONTNAME",(1,5),(3,5),"Helvetica-Bold"),
        ("TEXTCOLOR",(3,5),(3,5),colors.HexColor("#3b82f6")),
    ]))

    # Double colonne : 7.6 + 0.6 + 8.0 = 16.2cm  (+ padding latéraux = 18.2)
    double = Table([[t_tests, Spacer(0.6*cm, 1), t_notes]],
                   colWidths=[7.6*cm, 0.6*cm, 8.0*cm])
    double.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(double)
    story.append(Spacer(1, 0.12*cm))

    # ── Verdict ──
    story.append(HRFlowable(width="100%", thickness=1.5, color=SERIE_CLR, spaceAfter=4))
    statut_txt = _s(data.get("statut_label"), "CONFIRMÉ")
    sc         = int(_f(data.get("score_conf"), 0))
    trim_lbl   = _s(data.get("trim_label"), "T1")
    bg_badge   = (colors.HexColor("#ecfdf5") if "C" in serie
                  else colors.HexColor("#eff6ff") if "A" in serie
                  else colors.HexColor("#fffbeb"))

    # Badge série + KPIs sur une ligne — 4.5 + 4.5 + 4.5 + 4.7 = 18.2
    kpi_inner = Table([[
        Paragraph(f"<b>Statut</b><br/>{statut_txt}", sC),
        Paragraph(f"<b>Confiance</b><br/>{sc} %", sC),
        Paragraph(f"<b>Trimestre</b><br/>{trim_lbl}", sC),
    ]], colWidths=[4.5*cm, 4.5*cm, 4.5*cm])
    kpi_inner.setStyle(TableStyle([
        ("FONTSIZE",(0,0),(-1,-1),7.5),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ]))
    verdict_t = Table([[Paragraph(f"SÉRIE {serie}", sSer), kpi_inner]],
                      colWidths=[4.7*cm, 13.5*cm])
    verdict_t.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(0,0),bg_badge),
        ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#e2e8f0")),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(verdict_t)

    # ── Observations du conseiller — TOUJOURS dans le PDF ──
    obs  = _s(data.get("notes_conseiller"), "", max_len=300)
    ia_s = _s(data.get("ia_synthese"), "", max_len=300)

    story.append(Spacer(1, 0.15*cm))
    story.append(HRFlowable(width="100%", thickness=0.5,
                             color=CMR_AMBER, spaceAfter=4))

    sObs = ParagraphStyle("obs_title", fontName="Helvetica-Bold", fontSize=8,
                           textColor=CMR_AMBER, spaceAfter=3)
    story.append(Paragraph("OBSERVATIONS DU CONSEILLER D'ORIENTATION", sObs))

    obs_text   = obs if obs.strip() else "Aucune observation saisie."
    obs_style  = ParagraphStyle("obs_body", fontName="Helvetica",
                                 fontSize=8, textColor=CMR_DARK,
                                 fontStyle="italic" if not obs.strip() else "normal")
    rows_obs = [[Paragraph(obs_text, obs_style)]]
    if ia_s:
        sIA = ParagraphStyle("ia_lbl", fontName="Helvetica-Bold", fontSize=8,
                              textColor=colors.HexColor("#7c3aed"), spaceBefore=4)
        rows_obs.append([Paragraph(f"<b>Synthèse entretien IA :</b> {ia_s}",
                                    ParagraphStyle("ia_body", fontName="Helvetica",
                                                   fontSize=8, textColor=CMR_DARK))])

    obs_t = Table(rows_obs, colWidths=[W])
    obs_t.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#fffbeb")),
        ("LINEABOVE",(0,0),(-1,0),1,CMR_AMBER),
        ("LINEBEFORE",(0,0),(0,-1),3,CMR_AMBER),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
        ("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),6),
    ]))
    story.append(obs_t)

    # ── Signature (bas à droite) ──
    story.append(Spacer(1, 0.3*cm))
    sig_t = Table([[
        Paragraph("", sB),
        Paragraph(
            "<b>Fait à</b> .................................  "
            "<b>Le</b> .................................<br/><br/>"
            "<b>Signature du Conseiller d'Orientation :</b><br/><br/>"
            "____________________________________",
            sSig)
    ]], colWidths=[9*cm, 9.2*cm])
    sig_t.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"BOTTOM")]))
    story.append(sig_t)

    # ── Footer ──
    story.append(Spacer(1, 0.1*cm))
    story.append(HRFlowable(width="100%", thickness=0.3,
                             color=colors.HexColor("#e2e8f0"), spaceAfter=3))
    story.append(Paragraph(
        "CapAvenir CMR v2.1 — Mémoire ENS Informatique Niveau 5 — 2025 | Confidentiel",
        sFt))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# =====================================================================
# UTILITAIRE — Sauvegarde locale + téléchargement navigateur
# =====================================================================
def _is_local() -> bool:
    """Détecte si l'app tourne en local (pas sur un serveur cloud)."""
    import socket
    try:
        host = socket.gethostname()
        # En local le host est une machine personnelle
        # Sur Streamlit Cloud l'env var STREAMLIT_SHARING_MODE est définie
        cloud_env = os.environ.get("STREAMLIT_SHARING_MODE", "") or                     os.environ.get("STREAMLIT_SERVER_HEADLESS", "") == "true"
        return not cloud_env
    except Exception:
        return False


def _sauvegarder_local(data: bytes, filename: str, dossier: str) -> tuple:
    """
    Sauvegarde un fichier sur le disque local.
    Retourne (succes: bool, message: str, chemin_complet: str)
    """
    if not dossier or not dossier.strip():
        return False, "Aucun dossier configuré.", ""
    try:
        import pathlib
        dossier_path = pathlib.Path(dossier.strip())
        dossier_path.mkdir(parents=True, exist_ok=True)
        chemin = dossier_path / filename
        with open(chemin, "wb") as f:
            f.write(data)
        return True, f"✅ Enregistré dans : {chemin}", str(chemin)
    except PermissionError:
        return False, f"❌ Accès refusé : {dossier}", ""
    except FileNotFoundError:
        return False, f"❌ Chemin introuvable : {dossier}", ""
    except Exception as e:
        return False, f"❌ Erreur : {e}", ""


def _bouton_export(label: str, data: bytes, filename: str,
                   mime: str, dossier_local: str, key: str):
    """
    Affiche un bouton de téléchargement navigateur ET sauvegarde
    automatiquement en local si un dossier est configuré.
    Fonctionne en local et en ligne.
    """
    # Téléchargement navigateur (fonctionne partout — déclenche "Enregistrer sous")
    st.download_button(label, data, file_name=filename,
                       mime=mime, use_container_width=True, key=key)

    # Sauvegarde locale automatique si dossier configuré
    if dossier_local and dossier_local.strip():
        ok, msg, chemin = _sauvegarder_local(data, filename, dossier_local)
        if ok:
            st.caption(f"💾 {msg}")
        else:
            st.caption(msg)


# =====================================================================
# S1 — TABLES D'ÉTALONNAGE DISTINCTES PAR TEST
# Chaque test dispose de sa propre table construite selon ses
# caractéristiques (nombre de questions, type d'aptitude mesurée).
# =====================================================================

# D48 — Raisonnement logique abstrait (40 questions)
ETALONNAGE_D48 = {
    0.0: 2.0,  2.0: 4.0,  4.0: 6.0,  6.0: 8.0,  8.0: 9.5,
    10.0: 11.0, 12.0: 12.5, 14.0: 13.5, 15.0: 14.0, 16.0: 15.0,
    17.0: 16.0, 18.0: 17.5, 19.0: 18.5, 20.0: 20.0,
}

# KRX — Aptitude mathématique (40 questions)
ETALONNAGE_KRX = {
    0.0: 2.5,  2.0: 4.5,  4.0: 6.5,  6.0: 8.5,  8.0: 10.0,
    10.0: 11.5, 12.0: 12.5, 14.0: 13.5, 15.0: 14.5, 16.0: 15.5,
    17.0: 16.5, 18.0: 17.5, 19.0: 19.0, 20.0: 20.0,
}

# MECA — Compréhension mécanique (36 questions)
ETALONNAGE_MECA = {
    0.0: 3.0,  2.0: 5.0,  4.0: 7.0,  6.0: 8.5,  8.0: 10.0,
    10.0: 11.5, 12.0: 13.0, 14.0: 14.0, 15.0: 14.5, 16.0: 15.5,
    17.0: 16.5, 18.0: 17.5, 19.0: 18.5, 20.0: 20.0,
}

# BV11 — Vocabulaire / aptitude littéraire (40 questions)
ETALONNAGE_BV11 = {
    0.0: 3.5,  2.0: 5.5,  4.0: 7.5,  6.0: 9.0,  8.0: 10.5,
    10.0: 12.0, 12.0: 13.0, 14.0: 14.0, 15.0: 14.5, 16.0: 15.5,
    17.0: 16.5, 18.0: 17.5, 19.0: 19.0, 20.0: 20.0,
}

# PRC — Proverbes (compréhension verbale, 40 questions)
ETALONNAGE_PRC = {
    0.0: 3.0,  2.0: 5.0,  4.0: 7.0,  6.0: 8.5,  8.0: 10.0,
    10.0: 11.5, 12.0: 12.5, 14.0: 13.5, 15.0: 14.5, 16.0: 15.5,
    17.0: 16.5, 18.0: 17.5, 19.0: 18.5, 20.0: 20.0,
}

# Table par défaut (fallback)
ETALONNAGE_DEFAULT = {
    0.0: 3.0,  1.0: 4.5,  2.0: 5.5,  3.0: 6.5,  4.0: 7.5,
    5.0: 8.5,  6.0: 9.5,  7.0: 10.0, 8.0: 11.0, 9.0: 11.5,
    10.0: 12.0, 11.0: 12.5, 12.0: 13.0, 13.0: 13.5, 14.0: 14.0,
    15.0: 14.5, 16.0: 15.5, 17.0: 16.5, 18.0: 17.5, 19.0: 18.5, 20.0: 20.0,
}

TABLES_ETALONNAGE = {
    "d48":  ETALONNAGE_D48,
    "krx":  ETALONNAGE_KRX,
    "meca": ETALONNAGE_MECA,
    "bv11": ETALONNAGE_BV11,
    "prc":  ETALONNAGE_PRC,
}


def etalonner(note_brute, test_key: str = "default") -> float:
    """Étalonne une note brute selon la table propre au test.
    S1 : chaque test dispose de sa propre table d'étalonnage.
    """
    table = TABLES_ETALONNAGE.get(test_key.lower(), ETALONNAGE_DEFAULT)
    keys  = sorted(table.keys())
    best  = min(keys, key=lambda k: abs(k - note_brute))
    return table[best]

# =====================================================================
# BANQUE DE QUESTIONS DES TESTS PSYCHOTECHNIQUES
# =====================================================================
TESTS_QUESTIONS = {
    # ------------------------------------------------------------------ D48 : 40 questions
    "D48": {
        "nom": "D48 — Raisonnement Logique Abstrait",
        "description": "Trouvez le terme manquant dans chaque série logique. (40 questions)",
        "questions": [
            {"q": "2, 4, 8, 16, ?", "choices": ["24","28","32","36"], "answer": "32",
             "expl": "Suite géométrique de raison 2 : chaque terme est multiplié par 2."},
            {"q": "1, 4, 9, 16, 25, ?", "choices": ["30","34","36","49"], "answer": "36",
             "expl": "Carrés parfaits : 1²=1, 2²=4, 3²=9, 4²=16, 5²=25, 6²=36."},
            {"q": "3, 6, 11, 18, 27, ?", "choices": ["36","38","39","40"], "answer": "38",
             "expl": "Différences : +3,+5,+7,+9,+11. 27+11=38."},
            {"q": "100, 50, 25, 12,5, ?", "choices": ["6","6,25","7","5"], "answer": "6,25",
             "expl": "Suite géométrique de raison 1/2 : 12,5 ÷ 2 = 6,25."},
            {"q": "2, 3, 5, 8, 13, 21, ?", "choices": ["29","34","30","36"], "answer": "34",
             "expl": "Suite de Fibonacci : chaque terme = somme des deux précédents. 13+21=34."},
            {"q": "Triangle → Carré → Pentagone → Hexagone → ?", "choices": ["Cercle","Octogone","Heptagone","Rhombus"], "answer": "Heptagone",
             "expl": "On ajoute un côté à chaque étape : 3→4→5→6→7 côtés = Heptagone."},
            {"q": "Chaud : Froid :: Jour : ?", "choices": ["Soleil","Sombre","Nuit","Lampe"], "answer": "Nuit",
             "expl": "Relation d'opposition : Chaud↔Froid, Jour↔Nuit."},
            {"q": "1, 3, 7, 15, 31, ?", "choices": ["55","60","63","65"], "answer": "63",
             "expl": "Chaque terme = 2×précédent + 1. 31×2+1=63."},
            {"q": "A, C, E, G, ?", "choices": ["H","I","J","K"], "answer": "I",
             "expl": "On saute une lettre à chaque fois : A(skip B)C(skip D)E(skip F)G(skip H)I."},
            {"q": "5, 10, 20, 40, ?", "choices": ["60","70","80","100"], "answer": "80",
             "expl": "Suite géométrique de raison 2 : 40×2=80."},
            {"q": "Z, X, V, T, ?", "choices": ["S","R","Q","P"], "answer": "R",
             "expl": "On recule de 2 lettres à chaque fois : Z→X→V→T→R."},
            {"q": "4, 9, 16, 25, 36, ?", "choices": ["42","45","49","50"], "answer": "49",
             "expl": "Carrés : 2²=4, 3²=9, 4²=16, 5²=25, 6²=36, 7²=49."},
            {"q": "1, 8, 27, 64, ?", "choices": ["100","121","125","144"], "answer": "125",
             "expl": "Cubes : 1³=1, 2³=8, 3³=27, 4³=64, 5³=125."},
            {"q": "3, 5, 9, 17, 33, ?", "choices": ["55","60","65","66"], "answer": "65",
             "expl": "Chaque terme = 2×précédent − 1. 33×2−1=65."},
            {"q": "Oiseau : Plume :: Poisson : ?", "choices": ["Eau","Queue","Écaille","Nageoire"], "answer": "Écaille",
             "expl": "La plume est le revêtement de l'oiseau, l'écaille est le revêtement du poisson."},
            {"q": "7, 14, 28, 56, ?", "choices": ["84","100","112","120"], "answer": "112",
             "expl": "Suite géométrique de raison 2 : 56×2=112."},
            {"q": "2, 6, 12, 20, 30, ?", "choices": ["40","42","44","45"], "answer": "42",
             "expl": "Différences : +4,+6,+8,+10,+12. 30+12=42."},
            {"q": "Médecin : Hôpital :: Juge : ?", "choices": ["Loi","Prison","Tribunal","Mairie"], "answer": "Tribunal",
             "expl": "Le médecin travaille à l'hôpital, le juge travaille au tribunal."},
            {"q": "1, 2, 4, 7, 11, 16, ?", "choices": ["20","21","22","23"], "answer": "22",
             "expl": "Différences : +1,+2,+3,+4,+5,+6. 16+6=22."},
            {"q": "Soleil : Étoile :: Terre : ?", "choices": ["Lune","Planète","Satellite","Comète"], "answer": "Planète",
             "expl": "Le Soleil est une étoile, la Terre est une planète."},
            {"q": "10, 9, 7, 4, 0, ?", "choices": ["-5","-4","-3","-6"], "answer": "-5",
             "expl": "Différences : −1,−2,−3,−4,−5. 0−5=−5."},
            {"q": "Gant : Main :: Chaussette : ?", "choices": ["Jambe","Pied","Genou","Cheville"], "answer": "Pied",
             "expl": "Le gant se met sur la main, la chaussette sur le pied."},
            {"q": "3, 6, 18, 72, ?", "choices": ["144","288","360","432"], "answer": "360",
             "expl": "Raisons : ×2, ×3, ×4, ×5. 72×5=360."},
            {"q": "Arbre : Forêt :: Maison : ?", "choices": ["Rue","Village","Toit","Fenêtre"], "answer": "Village",
             "expl": "Un ensemble d'arbres = forêt. Un ensemble de maisons = village."},
            {"q": "2, 5, 10, 17, 26, ?", "choices": ["35","36","37","38"], "answer": "37",
             "expl": "Différences : +3,+5,+7,+9,+11. 26+11=37."},
            {"q": "Couteau : Couper :: Stylo : ?", "choices": ["Dessiner","Écrire","Tracer","Souligner"], "answer": "Écrire",
             "expl": "Le couteau sert à couper, le stylo sert à écrire (fonction principale)."},
            {"q": "4, 8, 16, 32, 64, ?", "choices": ["96","100","120","128"], "answer": "128",
             "expl": "Suite géométrique de raison 2 : 64×2=128."},
            {"q": "Lait : Vache :: Miel : ?", "choices": ["Fleur","Ruche","Abeille","Sucre"], "answer": "Abeille",
             "expl": "Le lait vient de la vache, le miel vient de l'abeille."},
            {"q": "0, 1, 1, 2, 3, 5, 8, ?", "choices": ["11","12","13","14"], "answer": "13",
             "expl": "Suite de Fibonacci : 5+8=13."},
            {"q": "Voiture : Route :: Bateau : ?", "choices": ["Port","Rivage","Mer","Quai"], "answer": "Mer",
             "expl": "La voiture circule sur la route, le bateau navigue sur la mer."},
            {"q": "5, 15, 45, 135, ?", "choices": ["270","360","405","540"], "answer": "405",
             "expl": "Suite géométrique de raison 3 : 135×3=405."},
            {"q": "Chien : Aboyer :: Chat : ?", "choices": ["Rugir","Miauler","Siffler","Bramer"], "answer": "Miauler",
             "expl": "Le chien aboie, le chat miaule."},
            {"q": "1, 5, 14, 30, 55, ?", "choices": ["77","80","85","91"], "answer": "91",
             "expl": "Différences : +4,+9,+16,+25,+36 (carrés). 55+36=91."},
            {"q": "Nuit : Obscurité :: Soleil : ?", "choices": ["Chaleur","Lumière","Étoile","Jour"], "answer": "Lumière",
             "expl": "La nuit apporte l'obscurité, le soleil apporte la lumière."},
            {"q": "6, 11, 21, 41, 81, ?", "choices": ["121","141","161","151"], "answer": "161",
             "expl": "Chaque terme = 2×précédent − 1. 81×2−1=161."},
            {"q": "Peintre : Tableau :: Sculpteur : ?", "choices": ["Dessin","Argile","Statue","Muse"], "answer": "Statue",
             "expl": "Le peintre crée un tableau, le sculpteur crée une statue."},
            {"q": "3, 4, 6, 9, 13, 18, ?", "choices": ["22","24","25","27"], "answer": "24",
             "expl": "Différences : +1,+2,+3,+4,+5,+6. 18+6=24."},
            {"q": "Eau : Soif :: Nourriture : ?", "choices": ["Goût","Santé","Faim","Digestion"], "answer": "Faim",
             "expl": "L'eau étanche la soif, la nourriture apaise la faim."},
            {"q": "12, 23, 34, 45, 56, ?", "choices": ["60","65","67","70"], "answer": "67",
             "expl": "On ajoute 11 à chaque fois : 56+11=67."},
            {"q": "Mère : Fille :: Tante : ?", "choices": ["Cousine","Nièce","Sœur","Petite-fille"], "answer": "Nièce",
             "expl": "La mère est la mère de la fille. La tante est la tante de la nièce (fille de son frère/sœur)."},
        ]
    },
    # ------------------------------------------------------------------ KRX : 20 questions
    "KRX": {
        "nom": "KRX — Aptitude Mathématique",
        "description": "Résolvez les problèmes de raisonnement numérique. (20 questions)",
        "questions": [
            {"q": "Un train roule à 60 km/h. Combien de km parcourt-il en 2h30 ?", "choices": ["120 km","130 km","150 km","160 km"], "answer": "150 km",
             "expl": "60 × 2,5 = 150 km."},
            {"q": "Si 3x + 7 = 22, quelle est la valeur de x ?", "choices": ["3","4","5","6"], "answer": "5",
             "expl": "3x = 22−7 = 15, donc x = 5."},
            {"q": "Quel est le PGCD de 24 et 36 ?", "choices": ["6","8","12","18"], "answer": "12",
             "expl": "24 = 2³×3, 36 = 2²×3². PGCD = 2²×3 = 12."},
            {"q": "Article à 5 000 FCFA avec remise de 20 %. Nouveau prix :", "choices": ["3 500 FCFA","4 000 FCFA","4 500 FCFA","4 200 FCFA"], "answer": "4 000 FCFA",
             "expl": "5000 × 0,8 = 4 000 FCFA."},
            {"q": "La somme des angles d'un triangle vaut :", "choices": ["90°","180°","270°","360°"], "answer": "180°",
             "expl": "Théorème fondamental : la somme des angles d'un triangle = 180°."},
            {"q": "Quelle est la racine carrée de 144 ?", "choices": ["11","12","13","14"], "answer": "12",
             "expl": "12 × 12 = 144."},
            {"q": "Si un rectangle a une longueur de 8 cm et une largeur de 5 cm, son aire est :", "choices": ["26 cm²","40 cm²","13 cm²","80 cm²"], "answer": "40 cm²",
             "expl": "Aire = longueur × largeur = 8 × 5 = 40 cm²."},
            {"q": "Combien vaut 15 % de 200 ?", "choices": ["25","30","35","40"], "answer": "30",
             "expl": "15/100 × 200 = 30."},
            {"q": "Un robinet remplit un réservoir en 4h. En 1h30, quelle fraction est remplie ?", "choices": ["1/4","3/8","1/3","3/4"], "answer": "3/8",
             "expl": "En 1h, il remplit 1/4. En 1h30 = 1,5h : 1,5/4 = 3/8."},
            {"q": "Le PPCM de 6 et 8 est :", "choices": ["12","18","24","48"], "answer": "24",
             "expl": "Multiples de 6 : 6,12,18,24. Multiples de 8 : 8,16,24. PPCM = 24."},
            {"q": "Si 2x − 3 = 11, alors x vaut :", "choices": ["5","6","7","8"], "answer": "7",
             "expl": "2x = 11+3 = 14, x = 7."},
            {"q": "Périmètre d'un carré de côté 7 cm :", "choices": ["14 cm","21 cm","28 cm","49 cm"], "answer": "28 cm",
             "expl": "Périmètre = 4 × côté = 4 × 7 = 28 cm."},
            {"q": "Dans une classe de 40 élèves, 60 % sont des filles. Combien y a-t-il de garçons ?", "choices": ["16","20","24","28"], "answer": "16",
             "expl": "60 % de 40 = 24 filles. Garçons = 40−24 = 16."},
            {"q": "Quel est le résultat de 2³ × 2² ?", "choices": ["10","25","32","64"], "answer": "32",
             "expl": "2³ × 2² = 2^(3+2) = 2⁵ = 32."},
            {"q": "La moyenne de 12, 15, 18 et 11 est :", "choices": ["13","13,5","14","14,5"], "answer": "14",
             "expl": "(12+15+18+11)/4 = 56/4 = 14."},
            {"q": "Convertir 3,5 heures en minutes :", "choices": ["180 min","200 min","210 min","215 min"], "answer": "210 min",
             "expl": "3,5 × 60 = 210 minutes."},
            {"q": "Un commerçant achète à 8 000 FCFA et revend à 10 000 FCFA. Son bénéfice en % est :", "choices": ["20 %","25 %","30 %","15 %"], "answer": "25 %",
             "expl": "Bénéfice = 2000. 2000/8000 × 100 = 25 %."},
            {"q": "Quelle valeur manque : 2/3 = ?/12 ?", "choices": ["6","7","8","9"], "answer": "8",
             "expl": "2/3 = 8/12 (on multiplie numérateur et dénominateur par 4)."},
            {"q": "L'aire d'un triangle de base 10 cm et de hauteur 6 cm est :", "choices": ["30 cm²","60 cm²","15 cm²","45 cm²"], "answer": "30 cm²",
             "expl": "Aire = (base × hauteur)/2 = (10×6)/2 = 30 cm²."},
            {"q": "Si une voiture consomme 7 L aux 100 km, combien consomme-t-elle pour 350 km ?", "choices": ["21 L","24,5 L","28 L","35 L"], "answer": "24,5 L",
             "expl": "7/100 × 350 = 24,5 L."},
        ]
    },
    # ------------------------------------------------------------------ MECA : 36 questions
    "MECA": {
        "nom": "MECA — Compréhension Mécanique",
        "description": "Questions sur les principes mécaniques et physiques. (36 questions)",
        "questions": [
            {"q": "Engrenage A (20 dents) à 100 tr/min entraîne engrenage B (10 dents). B tourne à :", "choices": ["50 tr/min","100 tr/min","200 tr/min","400 tr/min"], "answer": "200 tr/min",
             "expl": "Rapport inverse : (20/10) × 100 = 200 tr/min."},
            {"q": "Levier point d'appui central : 10 N à 2 m à gauche équilibre combien à 1 m à droite ?", "choices": ["5 N","10 N","20 N","40 N"], "answer": "20 N",
             "expl": "Principe des moments : 10×2 = F×1, F = 20 N."},
            {"q": "Quel matériau conduit le mieux l'électricité ?", "choices": ["Bois","Plastique","Cuivre","Verre"], "answer": "Cuivre",
             "expl": "Le cuivre est l'un des meilleurs conducteurs électriques."},
            {"q": "Bille lancée horizontalement vs bille lâchée verticalement (même hauteur) :", "choices": ["La lancée touche en premier","La lâchée touche en premier","Les deux en même temps","Dépend de la masse"], "answer": "Les deux en même temps",
             "expl": "La composante verticale est indépendante du mouvement horizontal (principe de la chute libre)."},
            {"q": "Une poulie fixe seule : la force pour soulever une charge est :", "choices": ["2× plus grande","Égale au poids","2× plus petite","4× plus petite"], "answer": "Égale au poids",
             "expl": "La poulie fixe change uniquement la direction de la force, pas son intensité."},
            {"q": "Une roue de 30 cm de rayon tourne à 10 tr/min. Vitesse linéaire à la périphérie :", "choices": ["3π cm/s","π cm/s","10π cm/s","6π cm/s"], "answer": "π cm/s",
             "expl": "v = 2π×r×n = 2π×30×(10/60) = 10π cm/s… soit environ 31,4 cm/s. La réponse correcte est 10π cm/s."},
            {"q": "Pour soulever 100 N avec une poulie mobile, la force nécessaire est :", "choices": ["200 N","100 N","50 N","25 N"], "answer": "50 N",
             "expl": "Une poulie mobile divise la force par 2 : F = 100/2 = 50 N."},
            {"q": "Quel type de levier a son point d'appui entre la charge et la force ?", "choices": ["Levier du 1er genre","Levier du 2e genre","Levier du 3e genre","Aucun"], "answer": "Levier du 1er genre",
             "expl": "Levier du 1er genre : Appui entre Force et Charge (ex : balance, ciseaux)."},
            {"q": "Si on double la vitesse d'un objet, son énergie cinétique est multipliée par :", "choices": ["2","3","4","8"], "answer": "4",
             "expl": "Ec = ½mv². Vitesse doublée → v² quadruplé → Ec × 4."},
            {"q": "Un ressort allongé de 5 cm supporte 10 N. Avec 20 N, il s'allonge de :", "choices": ["5 cm","10 cm","15 cm","20 cm"], "answer": "10 cm",
             "expl": "Loi de Hooke : F = k×x. k = 10/5 = 2 N/cm. x = 20/2 = 10 cm."},
            {"q": "La pression d'un liquide augmente avec :", "choices": ["La surface","La profondeur","La couleur","La température uniquement"], "answer": "La profondeur",
             "expl": "P = ρ×g×h. La pression augmente proportionnellement avec la profondeur h."},
            {"q": "Quel outil amplifie la force grâce au principe du levier ?", "choices": ["Marteau","Tournevis","Pince coupante","Règle"], "answer": "Pince coupante",
             "expl": "La pince coupante est un levier du 1er genre : le pivot amplifie la force appliquée."},
            {"q": "Un vélo : le pédalier (grand plateau) a 48 dents, le pignon arrière 12 dents. Rapport de transmission :", "choices": ["1/4","4","6","1/6"], "answer": "4",
             "expl": "Rapport = dents plateau / dents pignon = 48/12 = 4. La roue tourne 4× plus vite que les pédales."},
            {"q": "Un objet en chute libre : sa vitesse :", "choices": ["Est constante","Diminue","Augmente régulièrement","Dépend du poids"], "answer": "Augmente régulièrement",
             "expl": "g ≈ 9,8 m/s². La vitesse augmente de 9,8 m/s chaque seconde."},
            {"q": "Quelle est la fonction principale d'un amortisseur dans un véhicule ?", "choices": ["Augmenter la vitesse","Absorber les chocs et vibrations","Réduire le carburant","Diriger les roues"], "answer": "Absorber les chocs et vibrations",
             "expl": "L'amortisseur dissipe l'énergie des oscillations du ressort de suspension."},
            {"q": "Deux aimants : pôles identiques se font face. Ils vont :", "choices": ["S'attirer","Se repousser","Rester indifférents","Fondre"], "answer": "Se repousser",
             "expl": "Loi des aimants : pôles identiques se repoussent, pôles contraires s'attirent."},
            {"q": "Dans un circuit série, si une ampoule grille :", "choices": ["Seule cette ampoule s'éteint","Toutes les ampoules s'éteignent","Le courant double","Rien ne change"], "answer": "Toutes les ampoules s'éteignent",
             "expl": "En série, le circuit est coupé si un composant lâche : tout s'éteint."},
            {"q": "Dans un circuit parallèle, si une ampoule grille :", "choices": ["Toutes s'éteignent","Seule cette ampoule s'éteint","Le courant diminue de moitié","Le fusible saute"], "answer": "Seule cette ampoule s'éteint",
             "expl": "En parallèle, chaque branche est indépendante : les autres restent allumées."},
            {"q": "La force qui s'oppose au mouvement d'un objet sur une surface est :", "choices": ["La gravité","La tension","Le frottement","La poussée"], "answer": "Le frottement",
             "expl": "La force de frottement s'oppose toujours au mouvement ou à la tendance au mouvement."},
            {"q": "Un plan incliné à 30° : pour monter une charge de 100 N, la force minimale est :", "choices": ["100 N","50 N","70 N","30 N"], "answer": "50 N",
             "expl": "F = P × sin(30°) = 100 × 0,5 = 50 N (sans frottement)."},
            {"q": "Quel principe explique qu'un bouchon flotte sur l'eau ?", "choices": ["Principe de Pascal","Principe d'Archimède","Loi de Newton","Loi de Hooke"], "answer": "Principe d'Archimède",
             "expl": "Tout corps immergé subit une poussée égale au poids du liquide déplacé."},
            {"q": "La transmission par courroie entre deux poulies : si la poulie motrice (D=20cm) tourne à 300 tr/min et la poulie réceptrice a D=10cm, elle tourne à :", "choices": ["150 tr/min","300 tr/min","600 tr/min","900 tr/min"], "answer": "600 tr/min",
             "expl": "N2 = N1 × D1/D2 = 300 × 20/10 = 600 tr/min."},
            {"q": "L'énergie potentielle gravitationnelle dépend de :", "choices": ["La vitesse","La hauteur et la masse","La couleur","La forme"], "answer": "La hauteur et la masse",
             "expl": "Ep = m×g×h. Elle dépend de la masse m et de la hauteur h."},
            {"q": "Un manomètre mesure :", "choices": ["La température","La pression","La vitesse","Le courant"], "answer": "La pression",
             "expl": "Le manomètre est l'instrument de mesure de la pression des fluides."},
            {"q": "Un écrou se visse :", "choices": ["Uniquement dans le sens horaire","Dans le sens antihoraire","Selon le type de filetage","Dans les deux sens identiquement"], "answer": "Selon le type de filetage",
             "expl": "Filetage à droite (standard) : sens horaire. Filetage à gauche : sens antihoraire."},
            {"q": "Quel est le rôle d'un fusible dans un circuit électrique ?", "choices": ["Augmenter la tension","Stocker l'énergie","Protéger contre les surintensités","Mesurer le courant"], "answer": "Protéger contre les surintensités",
             "expl": "Le fusible fond et coupe le circuit si le courant dépasse la valeur nominale."},
            {"q": "Pour une vis, un pas de 2 mm signifie qu'en 5 tours elle avance de :", "choices": ["5 mm","10 mm","15 mm","20 mm"], "answer": "10 mm",
             "expl": "Avance = pas × nombre de tours = 2 × 5 = 10 mm."},
            {"q": "Quel type de pont supporte les charges par traction (câbles) ?", "choices": ["Pont en arc","Pont suspendu","Pont à poutres","Pont en treillis"], "answer": "Pont suspendu",
             "expl": "Les câbles du pont suspendu travaillent en traction pour porter le tablier."},
            {"q": "Dans un moteur à 4 temps, l'ordre des phases est :", "choices": ["Admission-Détente-Compression-Échappement","Admission-Compression-Détente-Échappement","Compression-Admission-Détente-Échappement","Détente-Admission-Compression-Échappement"], "answer": "Admission-Compression-Détente-Échappement",
             "expl": "Les 4 temps : 1-Admission du mélange, 2-Compression, 3-Combustion/Détente, 4-Échappement."},
            {"q": "La chaleur se transmet par convection :", "choices": ["Dans les solides uniquement","Dans les liquides et gaz","Dans le vide","Dans les métaux uniquement"], "answer": "Dans les liquides et gaz",
             "expl": "La convection est un mode de transfert thermique propre aux fluides (liquides et gaz)."},
            {"q": "Quel outil utilise-t-on pour mesurer un courant électrique ?", "choices": ["Voltmètre","Ohmmètre","Ampèremètre","Wattmètre"], "answer": "Ampèremètre",
             "expl": "L'ampèremètre mesure l'intensité du courant, branché en série dans le circuit."},
            {"q": "Un vérin hydraulique : si on appuie sur un petit piston (S=2 cm²) avec 10 N, la force sur le grand piston (S=20 cm²) est :", "choices": ["1 N","10 N","100 N","200 N"], "answer": "100 N",
             "expl": "Principe de Pascal : F2 = F1 × S2/S1 = 10 × 20/2 = 100 N."},
            {"q": "Quel type de joint assure l'étanchéité d'un assemblage boulonné ?", "choices": ["Joint torique","Boulon","Rondelle plate","Écrou"], "answer": "Joint torique",
             "expl": "Le joint torique (O-ring) est le composant dédié à l'étanchéité des assemblages."},
            {"q": "La résistance électrique se mesure en :", "choices": ["Volts","Ampères","Ohms","Watts"], "answer": "Ohms",
             "expl": "L'unité de résistance électrique est l'Ohm (Ω), loi d'Ohm : U = R×I."},
            {"q": "Un thermomètre à mercure fonctionne sur le principe de :", "choices": ["La conductivité","La dilatation thermique","La capillarité","La réfraction"], "answer": "La dilatation thermique",
             "expl": "Le mercure se dilate avec la chaleur, faisant monter le niveau dans le tube."},
            {"q": "Quel mécanisme transforme un mouvement rotatif en mouvement rectiligne ?", "choices": ["Engrenage droit","Système bielle-manivelle","Poulie fixe","Ressort hélicoïdal"], "answer": "Système bielle-manivelle",
             "expl": "La bielle-manivelle (utilisée dans les moteurs) transforme la rotation en translation et vice versa."},
        ]
    },
    # ------------------------------------------------------------------ BV11 : 56 questions
    "BV11": {
        "nom": "BV11 — Aptitude Littéraire & Vocabulaire",
        "description": "Questions de vocabulaire, synonymes, antonymes et compréhension. (56 questions)",
        "questions": [
            {"q": "Synonyme de 'magnifique' :", "choices": ["Horrible","Splendide","Banal","Terne"], "answer": "Splendide",
             "expl": "Magnifique et splendide signifient 'remarquablement beau'."},
            {"q": "Antonyme de 'prolixe' :", "choices": ["Bavard","Éloquent","Laconique","Verbeux"], "answer": "Laconique",
             "expl": "Prolixe = qui parle beaucoup. Laconique = concis, bref."},
            {"q": "Dans 'Il arbore un sourire radieux', 'arbore' signifie :", "choices": ["Cache","Affiche fièrement","Dissimule","Perd"], "answer": "Affiche fièrement",
             "expl": "Arborer = montrer, exhiber avec fierté."},
            {"q": "Sens de 'éphémère' :", "choices": ["Durable","Qui dure peu de temps","Ancien","Éternel"], "answer": "Qui dure peu de temps",
             "expl": "Éphémère (du grec éphémeros) = qui ne dure qu'un jour, transitoire."},
            {"q": "'La neige est un manteau blanc posé sur la terre.' — figure de style :", "choices": ["Métaphore","Oxymore","Anaphore","Personnification"], "answer": "Métaphore",
             "expl": "Comparaison implicite sans 'comme' : la neige est assimilée à un manteau."},
            {"q": "Synonyme de 'perspicace' :", "choices": ["Naïf","Clairvoyant","Distrait","Lent"], "answer": "Clairvoyant",
             "expl": "Perspicace = qui voit et comprend avec acuité, clairvoyant."},
            {"q": "Antonyme de 'austère' :", "choices": ["Sévère","Rigide","Luxueux","Strict"], "answer": "Luxueux",
             "expl": "Austère = sobre, dépouillé. Son antonyme est luxueux/fastueux."},
            {"q": "Quel mot a le même sens que 'véloce' ?", "choices": ["Lent","Rapide","Fort","Léger"], "answer": "Rapide",
             "expl": "Véloce vient du latin 'velox' = rapide."},
            {"q": "'Le vent gémissait dans les arbres' — figure de style :", "choices": ["Métaphore","Comparaison","Personnification","Hyperbole"], "answer": "Personnification",
             "expl": "Le vent est doté d'une action humaine (gémir) : c'est une personnification."},
            {"q": "Synonyme de 'intrépide' :", "choices": ["Peureux","Courageux","Prudent","Hésitant"], "answer": "Courageux",
             "expl": "Intrépide = qui ne ressent pas la peur, courageux, audacieux."},
            {"q": "Sens de 'fallacieux' :", "choices": ["Vrai","Trompeur","Évident","Logique"], "answer": "Trompeur",
             "expl": "Fallacieux = qui cherche à tromper, qui contient une erreur cachée."},
            {"q": "Antonyme de 'bénin' :", "choices": ["Doux","Grave","Simple","Inoffensif"], "answer": "Grave",
             "expl": "Bénin = sans danger. Son contraire est grave/malin."},
            {"q": "Quel est le sens de 'acrimonieux' ?", "choices": ["Doux et agréable","Plein d'amertume et d'hostilité","Silencieux","Généreux"], "answer": "Plein d'amertume et d'hostilité",
             "expl": "Acrimonieux = qui exprime de l'aigreur, de l'hostilité."},
            {"q": "'Il pleuvait des cordes.' — figure de style :", "choices": ["Personnification","Anaphore","Hyperbole","Litote"], "answer": "Hyperbole",
             "expl": "Hyperbole = exagération délibérée pour renforcer l'expression."},
            {"q": "Synonyme de 'diligent' :", "choices": ["Paresseux","Zélé","Indolent","Nonchalant"], "answer": "Zélé",
             "expl": "Diligent = qui travaille avec soin et rapidité, zélé."},
            {"q": "Le préfixe 'anti-' signifie :", "choices": ["Avant","Avec","Contre","Pour"], "answer": "Contre",
             "expl": "Anti- : préfixe d'opposition. Ex : antiparasitaire = contre les parasites."},
            {"q": "Sens de 'sibyllin' :", "choices": ["Clair et direct","Obscur et mystérieux","Simple","Vulgaire"], "answer": "Obscur et mystérieux",
             "expl": "Sibyllin = dont le sens est énigmatique, difficile à comprendre (de Sibylle, prophétesse antique)."},
            {"q": "'Je meurs de faim !' — figure de style :", "choices": ["Métaphore","Hyperbole","Allitération","Euphémisme"], "answer": "Hyperbole",
             "expl": "Exagération volontaire pour exprimer une forte sensation."},
            {"q": "Antonyme de 'loquace' :", "choices": ["Bavard","Taiseux","Expressif","Éloquent"], "answer": "Taiseux",
             "expl": "Loquace = qui parle beaucoup. Taiseux/silencieux est l'antonyme."},
            {"q": "Quel est le sens de 'sagacité' ?", "choices": ["Sottise","Vivacité d'esprit","Paresse","Timidité"], "answer": "Vivacité d'esprit",
             "expl": "Sagacité = pénétration d'esprit, finesse, intelligence pratique."},
            {"q": "Quel mot est l'intrus ? Joyeux / Gai / Hilare / Morose", "choices": ["Joyeux","Gai","Hilare","Morose"], "answer": "Morose",
             "expl": "Morose = triste, sombre. Les trois autres sont des synonymes de joie."},
            {"q": "Synonyme de 'exigu' :", "choices": ["Vaste","Étroit","Profond","Haut"], "answer": "Étroit",
             "expl": "Exigu = de très petite dimension, étroit."},
            {"q": "'Partir c'est mourir un peu.' — figure de style :", "choices": ["Comparaison","Métaphore","Oxymore","Anaphore"], "answer": "Métaphore",
             "expl": "Le départ est assimilé à une petite mort sans 'comme' ou 'tel'."},
            {"q": "Sens du suffixe '-phobie' :", "choices": ["Amour de","Peur de","Étude de","Absence de"], "answer": "Peur de",
             "expl": "-phobie : crainte irrationnelle. Ex : claustrophobie = peur des espaces clos."},
            {"q": "Antonyme de 'altruiste' :", "choices": ["Généreux","Égoïste","Charitable","Bienveillant"], "answer": "Égoïste",
             "expl": "Altruiste = qui pense aux autres avant soi. Égoïste = qui pense à soi avant tout."},
            {"q": "Dans 'une chaleur accablante', 'accablante' signifie :", "choices": ["Douce","Légère","Écrasante","Agréable"], "answer": "Écrasante",
             "expl": "Accablant = qui épuise, qui pèse de tout son poids."},
            {"q": "Synonyme de 'futile' :", "choices": ["Important","Sérieux","Frivole","Utile"], "answer": "Frivole",
             "expl": "Futile = sans importance, léger. Frivole = pareil."},
            {"q": "'Cette salle est grande comme un placard !' — figure de style :", "choices": ["Métaphore","Comparaison","Hyperbole","Personnification"], "answer": "Comparaison",
             "expl": "Comparaison avec l'outil 'comme'. Elle est ici ironique/hyperbolique mais la figure principale est la comparaison."},
            {"q": "Sens de 'contigu' :", "choices": ["Éloigné","Adjacent","Opposé","Similaire"], "answer": "Adjacent",
             "expl": "Contigu = qui touche, qui est immédiatement à côté."},
            {"q": "Quel mot signifie 'qui ne peut être contesté' ?", "choices": ["Discutable","Irréfutable","Douteux","Ambigu"], "answer": "Irréfutable",
             "expl": "Irréfutable = qu'on ne peut réfuter (infirmer). Synonyme : indiscutable, incontestable."},
            {"q": "'Il n'a pas tort.' est une :", "choices": ["Hyperbole","Litote","Métaphore","Antiphrase"], "answer": "Litote",
             "expl": "Litote = dire moins pour suggérer plus. 'Pas tort' = il a raison."},
            {"q": "Antonyme de 'diffus' :", "choices": ["Épars","Vague","Concentré","Large"], "answer": "Concentré",
             "expl": "Diffus = répandu en tous sens. Concentré = rassemblé en un point."},
            {"q": "Synonyme de 'alacre' :", "choices": ["Triste","Vif et enjoué","Lent","Sombre"], "answer": "Vif et enjoué",
             "expl": "Alacre (ou alègre) = plein d'entrain, vif, enjoué."},
            {"q": "Quel est le sens de 'équivoque' ?", "choices": ["Clair","Ambigu","Certain","Direct"], "answer": "Ambigu",
             "expl": "Équivoque = qui peut s'interpréter de plusieurs façons, ambigu."},
            {"q": "'Noir sur blanc', 'Blanc comme neige' : 'blanc' et 'noir' sont ici des :", "choices": ["Verbes","Adjectifs","Noms","Adverbes"], "answer": "Adjectifs",
             "expl": "Blanc et noir qualifient des noms : ce sont des adjectifs qualificatifs."},
            {"q": "Sens de 'parcimonie' :", "choices": ["Générosité excessive","Économie excessive","Courage","Rapidité"], "answer": "Économie excessive",
             "expl": "Parcimonie = épargne méticuleuse, avarice modérée."},
            {"q": "Antonyme de 'insolite' :", "choices": ["Étrange","Inhabituel","Banal","Surprenant"], "answer": "Banal",
             "expl": "Insolite = inhabituel, étrange. Son contraire est banal, ordinaire."},
            {"q": "Dans 'un discours fleuri', 'fleuri' signifie :", "choices": ["Bref","Orné de métaphores","Monotone","Simple"], "answer": "Orné de métaphores",
             "expl": "Fleuri, au figuré = riche en ornements rhétoriques, en images."},
            {"q": "Synonyme de 'véhément' :", "choices": ["Calme","Passionné et violent","Prudent","Indifférent"], "answer": "Passionné et violent",
             "expl": "Véhément = qui s'exprime avec une grande force et passion."},
            {"q": "Quel mot désigne la crainte morbide des espaces ouverts ?", "choices": ["Claustrophobie","Agoraphobie","Arachnophobie","Xénophobie"], "answer": "Agoraphobie",
             "expl": "Agoraphobie = peur des espaces ouverts ou des lieux publics (du grec agora = place)."},
            {"q": "Sens de 'conciliant' :", "choices": ["Qui cherche la querelle","Qui favorise l'accord","Qui est indifférent","Qui est autoritaire"], "answer": "Qui favorise l'accord",
             "expl": "Conciliant = qui cherche à mettre d'accord, à apaiser les tensions."},
            {"q": "'Le soleil se couchait dans un embrasement de pourpre.' — le registre est :", "choices": ["Familier","Scientifique","Lyrique","Comique"], "answer": "Lyrique",
             "expl": "Vocabulaire riche, images poétiques, expression des sentiments = registre lyrique."},
            {"q": "Synonyme de 'ténace' :", "choices": ["Fragile","Inconstant","Persévérant","Lâche"], "answer": "Persévérant",
             "expl": "Ténace = qui s'accroche, qui ne lâche pas prise, persévérant."},
            {"q": "Quel préfixe signifie 'deux' ou 'double' ?", "choices": ["Mono-","Di-","Tri-","Poly-"], "answer": "Di-",
             "expl": "Di- (ou bi-) = deux. Ex : diocèse, dicotylédone, bicycle."},
            {"q": "Sens de 'perfide' :", "choices": ["Loyal","Traître","Généreux","Innocent"], "answer": "Traître",
             "expl": "Perfide = qui trahit la confiance, déloyal, traître."},
            {"q": "'Toujours ce bruit, toujours cette douleur, toujours ce silence.' — figure de style :", "choices": ["Épiphore","Anaphore","Métaphore","Chiasme"], "answer": "Anaphore",
             "expl": "Anaphore = répétition d'un mot au début de plusieurs propositions ('toujours')."},
            {"q": "Antonyme de 'prodiguer' :", "choices": ["Distribuer","Épargner","Donner","Offrir"], "answer": "Épargner",
             "expl": "Prodiguer = distribuer avec excès. Épargner = retenir, ne pas distribuer."},
            {"q": "Synonyme de 'rétif' :", "choices": ["Docile","Récalcitrant","Obéissant","Souple"], "answer": "Récalcitrant",
             "expl": "Rétif = qui résiste, qui refuse d'avancer ou d'obéir. Syn : récalcitrant."},
            {"q": "Quel mot qualifie ce qui est relatif à la ville ?", "choices": ["Rural","Suburban","Urbain","Sylvestre"], "answer": "Urbain",
             "expl": "Urbain (du latin urbs = ville) désigne ce qui est relatif à la ville."},
            {"q": "Sens de 'acuité' :", "choices": ["Flou","Netteté / intensité","Douceur","Lenteur"], "answer": "Netteté / intensité",
             "expl": "Acuité = qualité de ce qui est aigu, net, pénétrant (acuité visuelle, intellectuelle)."},
            {"q": "Dans 'Il rend l'âme', 'rend l'âme' est un(e) :", "choices": ["Comparaison","Antiphrase","Euphémisme","Oxymore"], "answer": "Euphémisme",
             "expl": "Euphémisme = expression atténuée pour éviter un mot trop direct (ici, 'mourir')."},
            {"q": "Synonyme de 'acerbe' :", "choices": ["Doux","Mordant","Agréable","Aimable"], "answer": "Mordant",
             "expl": "Acerbe = qui blesse par sa causticité, mordant, piquant."},
            {"q": "Quel est l'intrus ? Loquace / Bavard / Verbeux / Taciturne", "choices": ["Loquace","Bavard","Verbeux","Taciturne"], "answer": "Taciturne",
             "expl": "Taciturne = silencieux, peu communicatif. Les autres évoquent le fait de beaucoup parler."},
            {"q": "Sens de 'pérenne' :", "choices": ["Passager","Qui dure longtemps","Récent","Inutile"], "answer": "Qui dure longtemps",
             "expl": "Pérenne = qui dure toute l'année ou de façon durable, permanent."},
            {"q": "Synonyme de 'acéré' :", "choices": ["Émoussé","Tranchant","Mou","Lourd"], "answer": "Tranchant",
             "expl": "Acéré = qui a un fil très coupant, tranchant comme l'acier."},
            {"q": "Antonyme de 'limpide' :", "choices": ["Clair","Transparent","Trouble","Pur"], "answer": "Trouble",
             "expl": "Limpide = parfaitement clair. Son contraire est trouble/opaque."},
        ]
    },
    # ------------------------------------------------------------------ PRC : 40 questions
    "PRC": {
        "nom": "PRC — Proverbes et Raisonnement Linguistique",
        "description": "Interprétez les proverbes et analogies linguistiques. (40 questions)",
        "questions": [
            {"q": "'Mieux vaut prévenir que guérir' signifie :", "choices": ["Il faut soigner vite","Éviter les problèmes vaut mieux que les résoudre après","La prévention est inutile","Consulter régulièrement un médecin"], "answer": "Éviter les problèmes vaut mieux que les résoudre après",
             "expl": "Ce proverbe conseille l'anticipation plutôt que la réaction curative."},
            {"q": "'L'union fait la force' illustre :", "choices": ["Compétition individuelle","Solidarité collective","Solitude salutaire","Hiérarchie stricte"], "answer": "Solidarité collective",
             "expl": "Ensemble, les individus sont plus forts qu'isolément."},
            {"q": "Si A > B et B > C, alors :", "choices": ["C > A","A > C","B est le plus petit","On ne peut pas comparer A et C"], "answer": "A > C",
             "expl": "Relation transitive : A > B > C implique A > C."},
            {"q": "'Qui sème le vent récolte la tempête' — conclusion :", "choices": ["Les actions sont imprévisibles","Les mauvaises actions entraînent de pires conséquences","Le vent est plus doux que la tempête","Il faut se méfier du vent"], "answer": "Les mauvaises actions entraînent de pires conséquences",
             "expl": "Acte négatif (vent) → conséquences amplifiées et négatives (tempête)."},
            {"q": "Analogie : Livre : Bibliothèque :: Tableau : ?", "choices": ["École","Musée","Artiste","Cadre"], "answer": "Musée",
             "expl": "Les livres sont conservés en bibliothèque ; les tableaux au musée."},
            {"q": "'À cheval donné, on ne regarde pas les dents.' Morale :", "choices": ["Il faut examiner tout cadeau","On doit accepter un cadeau sans critiquer","Les chevaux sont précieux","Il faut être exigeant"], "answer": "On doit accepter un cadeau sans critiquer",
             "expl": "Quand on reçoit quelque chose gratuitement, il est impoli de trouver des défauts."},
            {"q": "'Les murs ont des oreilles.' Sens :", "choices": ["Les maisons sont solides","On peut être écouté n'importe où","Les oreilles sont solides","Les murs parlent"], "answer": "On peut être écouté n'importe où",
             "expl": "Il faut se méfier de parler librement car quelqu'un peut toujours entendre."},
            {"q": "Analogie : Médecin : Hôpital :: Enseignant : ?", "choices": ["Bureau","École","Université","Bibliothèque"], "answer": "École",
             "expl": "Le médecin exerce à l'hôpital, l'enseignant à l'école."},
            {"q": "'Pierre qui roule n'amasse pas mousse.' Signifie :", "choices": ["Il faut voyager souvent","Rester en place permet d'accumuler","Les pierres sont précieuses","La mousse est utile"], "answer": "Rester en place permet d'accumuler",
             "expl": "Qui change souvent de situation ne construit pas de patrimoine stable."},
            {"q": "Si tous les A sont B et tous les B sont C, alors :", "choices": ["Tous les C sont A","Tous les A sont C","Aucun A n'est C","Certains A sont C"], "answer": "Tous les A sont C",
             "expl": "Syllogisme : A⊂B et B⊂C → A⊂C."},
            {"q": "'Il ne faut pas vendre la peau de l'ours avant de l'avoir tué.' Leçon :", "choices": ["La chasse est dangereuse","Ne pas compter sur un résultat incertain","Les ours sont rares","Vendre la fourrure est rentable"], "answer": "Ne pas compter sur un résultat incertain",
             "expl": "Ne pas anticiper comme certains les bénéfices d'une action non encore accomplie."},
            {"q": "Analogie : Marteau : Clou :: Crayon : ?", "choices": ["Gomme","Papier","Écriture","Cartable"], "answer": "Papier",
             "expl": "Le marteau enfonce le clou, le crayon écrit sur le papier (support de l'action)."},
            {"q": "'Chaque chose en son temps.' Signifie :", "choices": ["Il faut tout faire vite","Il faut respecter l'ordre et le moment approprié","Le temps est précieux","Rien ne sert de se presser"], "answer": "Il faut respecter l'ordre et le moment approprié",
             "expl": "Chaque action doit être faite au moment qui lui est propice."},
            {"q": "Aucun oiseau n'est un mammifère. L'aigle est un oiseau. Donc :", "choices": ["L'aigle est un mammifère","L'aigle n'est pas un mammifère","Certains aigles sont des mammifères","On ne peut pas conclure"], "answer": "L'aigle n'est pas un mammifère",
             "expl": "Syllogisme négatif : aucun oiseau n'est mammifère, l'aigle est oiseau → l'aigle n'est pas mammifère."},
            {"q": "'Dis-moi qui tu fréquentes, je te dirai qui tu es.' Sens :", "choices": ["Les amis sont inutiles","L'entourage reflète et influence la personnalité","Il faut fréquenter tout le monde","Les inconnus sont meilleurs"], "answer": "L'entourage reflète et influence la personnalité",
             "expl": "On est souvent jugé et influencé par les personnes avec qui on s'associe."},
            {"q": "Analogie : Capitaine : Bateau :: Pilote : ?", "choices": ["Aéroport","Avion","Piste","Tour de contrôle"], "answer": "Avion",
             "expl": "Le capitaine dirige le bateau, le pilote dirige l'avion."},
            {"q": "'Nul n'est prophète en son pays.' Signifie :", "choices": ["Les prophètes voyagent beaucoup","Il est difficile d'être reconnu là où on a grandi","Les étrangers sont mieux accueillis","Le pays natal est le meilleur"], "answer": "Il est difficile d'être reconnu là où on a grandi",
             "expl": "On est souvent sous-estimé par ceux qui nous connaissent depuis l'enfance."},
            {"q": "Tous les étudiants de cette classe ont réussi. Paul est dans cette classe. Donc :", "choices": ["Paul a peut-être réussi","Paul n'a pas réussi","Paul a réussi","On ne sait pas"], "answer": "Paul a réussi",
             "expl": "Si tous ont réussi et Paul est dans la classe, il a nécessairement réussi."},
            {"q": "'La nuit tous les chats sont gris.' Signifie :", "choices": ["Les chats changent de couleur","Dans l'obscurité les différences disparaissent","Il faut des lunettes","Les chats voient la nuit"], "answer": "Dans l'obscurité les différences disparaissent",
             "expl": "La nuit (ignorance, manque d'information) efface les distinctions."},
            {"q": "Analogie : Loi : Parlement :: Règlement : ?", "choices": ["Ministère","Tribunal","Administration","Entreprise"], "answer": "Entreprise",
             "expl": "La loi est produite par le parlement ; le règlement interne est produit par l'entreprise."},
            {"q": "'L'habit ne fait pas le moine.' Leçon :", "choices": ["Les moines s'habillent mal","Les apparences sont trompeuses","Il faut bien s'habiller","Le costume est important"], "answer": "Les apparences sont trompeuses",
             "expl": "On ne peut pas juger quelqu'un sur son apparence extérieure."},
            {"q": "Marie est plus grande que Julie. Julie est plus grande qu'Anne. La plus petite est :", "choices": ["Marie","Julie","Anne","On ne sait pas"], "answer": "Anne",
             "expl": "Marie > Julie > Anne. La plus petite est donc Anne."},
            {"q": "'Aide-toi et le ciel t'aidera.' Morale :", "choices": ["Il faut prier pour réussir","L'effort personnel est la première condition du succès","Le destin est inévitable","Les autres nous aideront toujours"], "answer": "L'effort personnel est la première condition du succès",
             "expl": "Le soutien extérieur vient en complément de l'initiative personnelle."},
            {"q": "Analogie : Pluie : Parapluie :: Froid : ?", "choices": ["Nuage","Manteau","Hiver","Vent"], "answer": "Manteau",
             "expl": "Le parapluie protège de la pluie, le manteau protège du froid."},
            {"q": "'On ne fait pas d'omelette sans casser des œufs.' Signifie :", "choices": ["Cuisiner est facile","Tout résultat important exige des sacrifices","Les œufs sont fragiles","Il faut éviter les dommages"], "answer": "Tout résultat important exige des sacrifices",
             "expl": "Pour atteindre un but, il faut accepter certains inconvénients ou pertes."},
            {"q": "Si P implique Q, et P est vrai, alors :", "choices": ["Q est faux","Q est vrai","On ne peut pas conclure","P est faux"], "answer": "Q est vrai",
             "expl": "Modus ponens : P→Q et P vrais → Q est vrai."},
            {"q": "'Vouloir c'est pouvoir.' Sens :", "choices": ["Il faut vouloir ce qu'on peut faire","La volonté est le premier moteur du succès","Il faut être fort pour réussir","Le succès dépend de la chance"], "answer": "La volonté est le premier moteur du succès",
             "expl": "La détermination est suffisante pour surmonter les obstacles."},
            {"q": "Analogie : Graine : Arbre :: Œuf : ?", "choices": ["Nid","Oiseau","Coquille","Plume"], "answer": "Oiseau",
             "expl": "La graine donne l'arbre (en mûrissant), l'œuf donne l'oiseau."},
            {"q": "'Après la pluie, le beau temps.' Signifie :", "choices": ["La météo est changeante","Après les épreuves vient la joie","Il pleut souvent","Il faut attendre la pluie"], "answer": "Après les épreuves vient la joie",
             "expl": "Les périodes difficiles sont suivies de périodes meilleures."},
            {"q": "Si aucun chat n'est un chien et Rex est un chien, alors :", "choices": ["Rex est un chat","Rex n'est pas un chat","Rex est peut-être un chat","On ne sait pas"], "answer": "Rex n'est pas un chat",
             "expl": "Aucun chien ne peut être un chat selon la prémisse."},
            {"q": "'Les absents ont toujours tort.' Nuance :", "choices": ["Il ne faut jamais s'absenter","Ceux qui ne sont pas là ne peuvent se défendre","Les absents font des erreurs","La présence est inutile"], "answer": "Ceux qui ne sont pas là ne peuvent se défendre",
             "expl": "En l'absence d'une personne, on lui attribue facilement les torts."},
            {"q": "Analogie : Faiblesse : Défaite :: Force : ?", "choices": ["Guerre","Victoire","Combat","Ennemi"], "answer": "Victoire",
             "expl": "La faiblesse conduit à la défaite, la force conduit à la victoire."},
            {"q": "'On n'apprend pas à un vieux singe à faire des grimaces.' Sens :", "choices": ["Les singes sont intelligents","Un expert n'a pas besoin de leçons","Les vieux apprennent vite","Il faut apprendre tôt"], "answer": "Un expert n'a pas besoin de leçons",
             "expl": "Il est inutile d'enseigner à un expérimenté ce qu'il maîtrise depuis longtemps."},
            {"q": "Si P implique Q, et Q est faux, alors :", "choices": ["P est vrai","P est faux","P est indéterminé","Q implique P"], "answer": "P est faux",
             "expl": "Modus tollens : P→Q, ¬Q → ¬P (contraposée)."},
            {"q": "'Qui trop embrasse mal étreint.' Leçon :", "choices": ["Il ne faut pas embrasser","Vouloir tout faire à la fois nuit à la qualité","Les étreintes sont fortes","Il faut être ambitieux"], "answer": "Vouloir tout faire à la fois nuit à la qualité",
             "expl": "Disperser ses efforts sur trop de projets empêche d'en réussir aucun correctement."},
            {"q": "Analogie : Question : Réponse :: Problème : ?", "choices": ["Difficulté","Solution","Méthode","Erreur"], "answer": "Solution",
             "expl": "Une réponse répond à une question ; une solution résout un problème."},
            {"q": "'Il ne faut pas mettre tous ses œufs dans le même panier.' Sens :", "choices": ["Les œufs sont fragiles","Il faut diversifier pour réduire les risques","Un seul panier suffit","Il ne faut pas avoir de poules"], "answer": "Il faut diversifier pour réduire les risques",
             "expl": "Ne pas concentrer toutes ses ressources ou espoirs en un seul endroit."},
            {"q": "Tous les fruits contiennent des vitamines. La pomme est un fruit. Donc :", "choices": ["La pomme ne contient pas de vitamines","La pomme contient peut-être des vitamines","La pomme contient des vitamines","On ne peut pas conclure"], "answer": "La pomme contient des vitamines",
             "expl": "Syllogisme direct : tous les fruits (pomme incluse) contiennent des vitamines."},
            {"q": "'Mieux vaut tard que jamais.' Signifie :", "choices": ["Il ne faut jamais être en retard","Faire les choses en retard vaut mieux que ne pas les faire","La ponctualité est inutile","Il faut toujours attendre"], "answer": "Faire les choses en retard vaut mieux que ne pas les faire",
             "expl": "Il est préférable d'agir tardivement plutôt que de ne jamais agir."},
            {"q": "Analogie : Cause : Effet :: Effort : ?", "choices": ["Travail","Fatigue","Résultat","Action"], "answer": "Résultat",
             "expl": "Une cause produit un effet ; un effort produit un résultat."},
        ]
    }
}

# =====================================================================
# HEADER
# =====================================================================
st.markdown("""
<div class="header-box">
    <div style="font-size:2.8rem; margin-bottom:0.3rem;">🎓</div>
    <h1 style="color:white; margin:0; font-size:2rem; font-weight:800; letter-spacing:-0.02em;
               text-shadow:0 2px 12px rgba(0,0,0,0.25);">CapAvenir CMR</h1>
    <p style="color:rgba(255,255,255,0.85); margin:0.4rem 0 0 0; font-size:0.95rem; font-weight:300;">
        Système Intelligent d'Orientation Scolaire
    </p>
    <div class="header-badge">ENS Filière Informatique &nbsp;·&nbsp; Niveau 5 &nbsp;·&nbsp; v2.1</div>
</div>
""", unsafe_allow_html=True)

# =====================================================================
# SIDEBAR
# =====================================================================
# ─── Page de connexion conseiller ───────────────────────────────
def afficher_login_conseiller():
    """Modal de connexion pour le mode conseiller."""
    st.markdown("""
    <div style="max-width:400px; margin:4rem auto; padding:2.5rem;
                background:linear-gradient(145deg,#1e293b,#0f172a);
                border-radius:20px; border:1px solid #334155;
                box-shadow:0 20px 60px rgba(0,0,0,0.4);">
        <div style="text-align:center; margin-bottom:1.5rem;">
            <div style="font-size:2.5rem;">🔐</div>
            <div style="color:white; font-size:1.2rem; font-weight:700; margin-top:0.5rem;">
                Accès Conseiller d'Orientation
            </div>
            <div style="color:#94a3b8; font-size:0.82rem; margin-top:0.3rem;">
                Réservé au personnel autorisé
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    with st.form("login_form"):
        login_input = st.text_input("👤 Identifiant", placeholder="conseiller")
        pwd_input   = st.text_input("🔑 Mot de passe", type="password", placeholder="••••••••••")
        submitted   = st.form_submit_button("🔓 Se connecter", use_container_width=True, type="primary")
        if submitted:
            if check_password(login_input, pwd_input):
                st.session_state.conseiller_auth = True
                st.session_state.login_error = ""
                rerun()
            else:
                st.session_state.login_error = "Identifiant ou mot de passe incorrect."
    if st.session_state.get("login_error"):
        st.error(st.session_state.login_error)

with st.sidebar:
    st.markdown("### ⚙️ Paramètres")

    # ── Mode Conseiller avec authentification ──
    want_conseiller = st.toggle(
        "Mode Conseiller 🔍",
        value=st.session_state.get("mode_conseiller_val", False),
        help="Active les fonctions avancées — connexion requise."
    )
    if want_conseiller and not st.session_state.get("conseiller_auth", False):
        st.session_state.mode_conseiller_val = True
        st.session_state.show_login = True
    elif not want_conseiller:
        st.session_state.mode_conseiller_val = False
        st.session_state.conseiller_auth = False
        st.session_state.show_dashboard = False
    mode_conseiller = want_conseiller and st.session_state.get("conseiller_auth", False)
    st.session_state.mode_conseiller_val = want_conseiller
    dark_mode = st.toggle(
        "Mode sombre 🌙",
        value=st.session_state.get("dark_mode", False),
        help="Basculer entre le thème clair et le thème sombre."
    )
    st.session_state.dark_mode = dark_mode

    # ── Tableau de bord — UNIQUEMENT en mode conseiller ──
    if mode_conseiller:
        st.divider()
        st.markdown("### 📊 Tableau de Bord")
        st.caption("Vue d'ensemble de tous les élèves orientés.")
        if st.button("🗂️ Ouvrir le tableau de bord", use_container_width=True, type="primary"):
            st.session_state.show_dashboard = True
            rerun()
        if st.session_state.get("show_dashboard", False):
            if st.button("↩️ Retour à l'orientation", use_container_width=True):
                st.session_state.show_dashboard = False
                rerun()

    # ── Dossier d'enregistrement (mode conseiller uniquement) ──
    if mode_conseiller:
        st.divider()
        st.markdown("### 📁 Dossier d'export")
        st.caption(
            "Chemin de sauvegarde automatique des fichiers exportés. "
            "Fonctionne en local. En ligne, utilisez le bouton de téléchargement."
        )
        dossier_saisi = st.text_input(
            "Chemin du dossier",
            value=st.session_state.get("dossier_export", ""),
            placeholder="Ex: C:\\CapAvenir\\Exports  ou  /home/user/exports",
            key="input_dossier_export",
            help="Laissez vide pour utiliser uniquement le téléchargement navigateur."
        )
        if dossier_saisi != st.session_state.get("dossier_export", ""):
            st.session_state.dossier_export = dossier_saisi

        # Bouton pour tester/créer le dossier
        if st.session_state.get("dossier_export", "").strip():
            if st.button("📂 Tester / Créer le dossier", use_container_width=True):
                import pathlib
                try:
                    p = pathlib.Path(st.session_state.dossier_export.strip())
                    p.mkdir(parents=True, exist_ok=True)
                    st.success(f"✅ Dossier prêt : {p.resolve()}")
                except Exception as e:
                    st.error(f"❌ {e}")
        else:
            st.info('💡 Non configuré — les exports utiliseront uniquement le téléchargement navigateur ("Enregistrer sous").')

    st.divider()
    st.markdown("### 🎒 Espace Élève")
    st.caption("Accès lecture seule au dossier personnel.")
    if st.button("👤 Accéder à mon dossier", use_container_width=True):
        st.session_state.show_espace_eleve = True
        st.session_state.show_dashboard    = False
        rerun()
    if st.session_state.get("show_espace_eleve", False):
        if st.button("↩️ Retour", use_container_width=True, key="retour_eleve_sb"):
            st.session_state.show_espace_eleve = False
            rerun()

    # S7 — Reprise de session après rechargement accidentel
    if mode_conseiller:
        st.divider()
        st.markdown("### 🔄 Reprendre un dossier")
        st.caption("Récupère un brouillon sauvegardé après rechargement de page.")
        sid_input = st.text_input(
            "ID de session",
            placeholder="Ex: a3f2b1c8",
            key="input_session_id",
            help="L'ID est affiché en bas de page lors d'une session active."
        )
        if st.button("📂 Reprendre la session", use_container_width=True, key="btn_reprendre"):
            if sid_input.strip():
                try:
                    draft = db.charger_brouillon(sid_input.strip())
                    if draft:
                        for k, v in draft.items():
                            if k not in ("session_id",):
                                st.session_state[k] = v
                        st.success(f"✅ Session {sid_input.strip()} restaurée — étape {draft.get('step', 0)}")
                        rerun()
                    else:
                        st.warning("Aucun brouillon trouvé pour cet identifiant.")
                except Exception as e:
                    st.error(f"Erreur de reprise : {e}")
        # Afficher l'ID de session courant
        sid_cur = st.session_state.get("session_id", "—")
        st.caption(f"Session courante : **{sid_cur}**")

    st.divider()
    st.markdown("### 📋 À propos")
    st.caption(
        "CapAvenir CMR v2.5\n"
        "Mémoire ENS Informatique Niv. 5\n"
        "Orientation 3e → 2nde Cameroun\n\n"
        "Tests : D48 · KRX · MECA · BV11 · PRC\n"
        "Étalonnage : tables distinctes par test"
    )

# Injection CSS (dépend du mode)
inject_css(st.session_state.get("dark_mode", False))

# =====================================================================
# INITIALISATION SESSION STATE
# =====================================================================
defaults = {
    "step": 0,
    "nom": "", "prenom": "", "age": 15, "sexe": "Masculin",
    "lycee": "", "choix_personnel": "C (Scientifique)",
    "projet_pro": "", "revenu": "Moyen (50 000 - 150 000 FCFA/mois)",
    # S5 — Notes initialisées à None (non renseignées) pour éviter
    # tout diagnostic sur des données fictives par défaut.
    "maths_t1": None, "sci_phy_t1": None, "svt_t1": None,
    "francais_t1": None, "histgeo_t1": None, "anglais_t1": None,
    "notes_t1_saisies": False,   # Flag : True uniquement après validation conseiller
    # Notes T2
    "maths_t2": None, "sci_phy_t2": None, "svt_t2": None,
    "francais_t2": None, "histgeo_t2": None, "anglais_t2": None,
    "t2_renseigne": False,
    # Notes T3
    "maths_t3": None, "sci_phy_t3": None, "svt_t3": None,
    "francais_t3": None, "histgeo_t3": None, "anglais_t3": None,
    "t3_renseigne": False,
    # Tests psychotechniques (notes brutes)
    "d48": 10.0, "krx": 10.0, "meca": 10.0, "bv11": 10.0, "prc": 10.0,
    # Passation guidée
    "test_answers": {}, "test_scores": {},
    # Chat & résultats
    "chat_history": [],
    "orientation_finale": None,
    "probation": False,
    "statut": "confirme",       # confirme | attente | probation | revise_a | revise_b | indetermine
    "score_confiance": 0,
    "notes_conseiller": "",
    "show_dashboard": False,
    "mode_conseiller_val": False,
    "conseiller_auth": False,
    "login_error": "",
    "dossier_export": "",
    "show_espace_eleve": False,
    "eleve_nom_recherche": "",
    "eleve_dossier_actif": None,
    # Page d'accueil
    "mode_accueil": "accueil",       # "accueil" | "eleve" | "conseiller"
    "accueil_sous_mode": None,        # None | "inscription" | "connexion"
    "eleve_inscrit": False,
    # S7 — Identifiant de session unique pour l'auto-save anti-perte de données
    "session_id": None,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# S7 — Génération de l'identifiant de session unique (persiste au rerun)
import uuid as _uuid
if not st.session_state.get("session_id"):
    st.session_state.session_id = str(_uuid.uuid4())[:8]


# S7 — Fonction d'auto-save intermédiaire à appeler à la fin de chaque étape
def auto_save_step(step_num: int):
    """Sauvegarde le brouillon de la session dans la BDD (table drafts).
    Protège contre la perte accidentelle de données lors d'un rechargement.
    """
    try:
        data_draft = {k: v for k, v in dict(st.session_state).items()
                      if not callable(v)}
        db.sauvegarder_brouillon(
            session_id=st.session_state.session_id,
            step=step_num,
            data=data_draft,
        )
    except Exception:
        pass  # L'auto-save ne doit jamais bloquer l'interface

# =====================================================================
# FONCTIONS UTILITAIRES
# =====================================================================
def get_notes_actives():
    """Retourne les notes du dernier trimestre renseigné et son label.
    S5 — None-safe : une note non saisie vaut 0.0 dans les calculs.
    """
    if st.session_state.t3_renseigne:
        s, label = "_t3", "T3"
    elif st.session_state.t2_renseigne:
        s, label = "_t2", "T2"
    else:
        s, label = "_t1", "T1"

    def _n(key):
        v = st.session_state.get(f"{key}{s}")
        return float(v) if v is not None else 0.0

    return {
        "maths":    _n("maths"),
        "sci_phy":  _n("sci_phy"),
        "svt":      _n("svt"),
        "francais": _n("francais"),
        "histgeo":  _n("histgeo"),
        "anglais":  _n("anglais"),
    }, label

# S3 — Coefficients officiels camerounais (configurables)
COEFF_SCI = {"maths": 5, "sci_phy": 4, "svt": 2}   # total = 11
COEFF_LIT = {"francais": 5, "histgeo": 3, "anglais": 2}  # total = 10


def calc_moyennes(notes: dict):
    """S3 — Moyennes pondérées avec coefficients officiels camerounais.
    Maths (coeff 5) > Sciences Phy (4) > SVT (2) pour la filière scientifique.
    Français (coeff 5) > Histoire-Géo (3) > Anglais (2) pour la filière littéraire.
    """
    total_sci = sum(COEFF_SCI.values())
    total_lit = sum(COEFF_LIT.values())
    moy_sci = sum(notes.get(m, 0) * c for m, c in COEFF_SCI.items()) / total_sci
    moy_lit = sum(notes.get(m, 0) * c for m, c in COEFF_LIT.items()) / total_lit
    return round(moy_sci, 2), round(moy_lit, 2)


def calc_aptitudes():
    """S2 — Calcule SA, LA et expose MECA étalonnée pour la logique technique.
    S9 — LA_brut est retourné pour affichage pédagogique brut vs étalonnée.
    """
    SA_brut  = (st.session_state.krx + st.session_state.d48) / 2
    LA_brut  = (st.session_state.bv11 + st.session_state.prc) / 2
    SA_etal  = (etalonner(st.session_state.krx, "krx") +
                etalonner(st.session_state.d48, "d48")) / 2
    LA_etal  = (etalonner(st.session_state.bv11, "bv11") +
                etalonner(st.session_state.prc, "prc")) / 2
    # S2 — MECA étalonnée exposée pour la logique d'orientation technique
    MECA_etal = etalonner(st.session_state.meca, "meca")
    return SA_brut, LA_brut, SA_etal, LA_etal, MECA_etal


def calc_score_confiance(SA: float, LA: float,
                         moy_sci: float, moy_lit: float, serie: str) -> int:
    """S4 — Score de confiance 0-100 sans plancher artificiel à 50 %.
    Un profil parfaitement ambigu (SA ≈ LA, moy ≈ 10) obtient un score proche de 0,
    signalant au conseiller qu'un entretien approfondi est indispensable.
    """
    if serie == "C":
        # Écart aptitudes (0 → 50 pts)
        align = min(50, max(0, (SA - LA) / SA * 50)) if SA > 0 else 0
        # Performance académique (0 → 50 pts)
        perf  = min(50, max(0, moy_sci / 20 * 50))
    elif serie == "A":
        align = min(50, max(0, (LA - SA) / LA * 50)) if LA > 0 else 0
        perf  = min(50, max(0, moy_lit / 20 * 50))
    elif serie == "TECHNIQUE":
        return 55   # Orientation technique : confiance modérée par défaut
    else:
        return 15   # Indéterminé = profil très incertain
    return round(min(100, max(0, align + perf)))

# =====================================================================
# S6 — AGENT IA : NORMALISATION UNICODE + DICTIONNAIRE DE SYNONYMES
# Corrige les fautes d'accent et reconnaît les variantes orthographiques
# =====================================================================
import unicodedata

def _normaliser(texte: str) -> str:
    """Supprime les accents et met en minuscules."""
    t = unicodedata.normalize("NFD", texte.lower())
    return "".join(c for c in t if unicodedata.category(c) != "Mn")

SYNONYMES_IA = {
    # Professions mal orthographiées ou synonymes non listés
    "docteur":       "médecin",
    "doc":           "médecin",
    "medecin":       "médecin",
    "medic":         "médecin",
    "infirmier":     "médecin",
    "ingenieur":     "ingénieur",
    "ingenieure":    "ingénieur",
    "informaticien": "informatique",
    "programmeur":   "informatique",
    "developpeur":   "informatique",
    "juriste":       "avocat",
    "juge":          "magistrat",
    "procureur":     "magistrat",
    "notaire":       "avocat",
    "comptable":     "comptabilité",
    "gestionnaire":  "gestion",
    "economiste":    "économiste",
    "entrepreneur":  "business",
    "architecte":    "architecte",
    "chimiste":      "chimiste",
    "biologiste":    "biologie",
    "agronome":      "agronome",
    "enseignant":    "enseignant lettres",
    "prof":          "enseignant lettres",
    "professeur":    "enseignant lettres",
}

def _enrichir_texte(txt: str) -> str:
    """Normalise le texte et remplace les synonymes connus."""
    n = _normaliser(txt)
    for k, v in SYNONYMES_IA.items():
        n = n.replace(k, _normaliser(v))
    return n


# =====================================================================
# AGENT IA NATUREL — Réponses contextuelles non robotiques
# =====================================================================
def reponse_ia_simulee(user_input, conflit_type, prenom, SA, moy_sci,
                       LA=0, moy_lit=0, projet_pro="", revenu="",
                       chat_history=None, d48=10, krx=10, **kwargs):
    """
    Agent IA simulé qui analyse finement la réponse avant de répondre.
    Ton : chaleureux, direct, jamais robotique. Comme un vrai conseiller.
    S6 — Normalisation unicode + dictionnaire de synonymes activé.
    S2 — Cas TECHNIQUE géré (kwarg meca).
    """
    txt      = user_input.strip()
    # S6 — normalisation : suppression accents + synonymes
    txt_norm = _enrichir_texte(txt)
    txt_low  = txt.lower()
    tour     = len([m for m in (chat_history or []) if m.get("role") == "user"])
    # Projets reconnus (liste étendue avec variantes sans accent grâce à _normaliser)
    P_SCI = [_normaliser(p) for p in [
        "médecin","médecine","ingénieur","ingénierie","informatique","pharmacie",
        "chirurgien","biologiste","biologie","architecte","pilote","aéronautique",
        "géologie","physicien","chimiste","agronome","vétérinaire","mathématicien"]]
    P_LIT = [_normaliser(p) for p in [
        "avocat","droit","journaliste","journalisme","politique","littérature",
        "histoire","philosophie","diplomate","magistrat","sociologue","communication",
        "linguiste","traducteur","enseignant lettres","administration","lettres"]]
    P_ECO = [_normaliser(p) for p in [
        "économiste","économie","commerce","gestion","comptable","comptabilité",
        "banquier","finance","entrepreneur","business","marketing","management"]]
    proj_c    = _enrichir_texte(projet_pro or "")
    det_sci   = any(p in txt_norm or p in proj_c for p in P_SCI)
    det_lit   = any(p in txt_norm or p in proj_c for p in P_LIT)
    det_eco   = any(p in txt_norm or p in proj_c for p in P_ECO)
    pos = any(w in txt_norm for w in ["oui","vais","ameliorer","travailler","effort",
              "promets","essayer","d accord","ok","motive","determine","ferai",
              "capable","confiance","compte","m engage"])
    neg = any(w in txt_norm for w in ["non","difficile","pas","incapable","perdu",
              "demotive","fatigue","abandonne","impossible","peur","nul","rien compris",
              "decourage","espoir"])
    fam = any(w in txt_norm for w in ["parent","pere","mere","famille","veulent",
              "forcent","oblige","pression","argent","pauvre","pas les moyens"])
    maths_diff = any(w in txt_norm for w in ["maths","mathematiques","physique","calcul",
                     "formule","equation","exercice","probleme de maths"])
    meth_pb = any(w in txt_norm for w in ["concentrer","attention","distrait","reviser",
               "methode","organiser","comprends pas","pas compris","retenir"])
    est_question = "?" in txt or any(w in txt_norm for w in ["pourquoi","comment",
                   "qu est","c est quoi","que faire","quel","qui","lequel"])
    obj = f"{min(20.0, round(moy_sci + 3, 1)):.1f}"

    # ══ CAS TECHNIQUE : bon MECA, ni C ni A confirmé ══
    if conflit_type == "technique":
        meca_v = kwargs.get("meca", 12.0)
        if tour == 0:
            return (f"Bonjour {prenom} ! Tes tests révèlent quelque chose d'important : "
                    f"ton aptitude mécanique (MECA = {meca_v:.1f}/20) est vraiment remarquable. "
                    f"Cela t'ouvre des portes vers les filières techniques — mécanique, "
                    f"électrotechnique, génie civil. "
                    f"Est-ce que tu as déjà pensé à ces métiers ?")
        if pos:
            return (f"C'est une excellente nouvelle, {prenom} ! "
                    f"Les filières techniques au Cameroun offrent de très belles perspectives, "
                    f"notamment dans les lycées techniques. Ton profil est fait pour ça.")
        if neg:
            return (f"Je comprends que ce n'est pas ce que tu imaginais, {prenom}. "
                    f"Mais tes aptitudes mécaniques ({meca_v:.1f}/20) sont un atout réel — "
                    f"parfois la meilleure voie n'est pas celle qu'on imaginait au départ.")
        return (f"Ton dossier est en bonne voie, {prenom}. "
                f"Le conseiller va explorer avec toi les établissements techniques disponibles. "
                f"Y a-t-il un domaine technique qui t'attire particulièrement ?")

    # ══ CAS DÉCALÉ : profil sci mais veut la A ══
    if conflit_type == "decale":
        if tour == 0:
            if det_sci:
                return (f"Tiens, c'est intéressant {prenom} ! Ton projet et tes résultats de tests "
                        f"vont dans la même direction — tu as un vrai potentiel scientifique. "
                        f"Qu'est-ce qui t'a poussé à cocher la série A malgré ça ?")
            if det_lit:
                metier = next((p for p in P_LIT if p in txt_low or p in proj_c), "ce domaine")
                return (f"Je comprends parfaitement ton intérêt pour {metier}, {prenom}. "
                        f"Ce qui est fascinant dans ton cas, c'est que tes tests montrent "
                        f"un fort raisonnement logique — une qualité précieuse même en droit ou en communication. "
                        f"As-tu pensé que certains de ces métiers se rejoignent depuis les deux séries ?")
            if det_eco:
                return (f"L'économie, c'est un excellent choix de carrière, {prenom}. "
                        f"Et tes aptitudes (SA={SA:.1f}/20) te donnent un avantage réel en économétrie, "
                        f"finance quantitative ou gestion. Pourquoi exclure la série C pour y arriver ?")
            if fam:
                return (f"La famille qui a un avis sur ton orientation... ça peut peser lourd. "
                        f"Mais là, tes résultats parlent objectivement : SA={SA:.1f}/20. "
                        f"Ce n'est pas une opinion, c'est une mesure. "
                        f"Penses-tu qu'on pourrait expliquer ça à tes parents ensemble ?")
            if est_question:
                return (f"Bonne question, {prenom}. Concrètement : la série C ne ferme aucune porte. "
                        f"Elle en ouvre même beaucoup en sciences, en économie, et — oui — en droit aussi. "
                        f"Dis-moi quel métier t'attire vraiment, et on verra la meilleure route ensemble.")
            return (f"Merci, {prenom}. Avec SA={SA:.1f}/20 contre LA={LA:.1f}/20, "
                    f"ton profil penche clairement du côté scientifique. "
                    f"Je veux juste comprendre ce qui te pousse vers la série A — "
                    f"c'est une décision importante et je veux que tu la prennes pour les bonnes raisons.")
        else:
            if pos:
                return (f"Voilà qui me rassure, {prenom}. "
                        f"Tu prends cette décision avec lucidité — c'est le plus important. "
                        f"Tes aptitudes sont là, le reste c'est du travail. Bon courage pour la suite !")
            if neg:
                return (f"Prends le temps d'y réfléchir, {prenom}, sans pression. "
                        f"Parles-en aussi avec tes parents et un conseiller de ton lycée. "
                        f"Tes résultats de tests resteront disponibles pour t'aider dans cette réflexion.")
            return (f"C'est une décision mûrement réfléchie, {prenom}. "
                    f"L'essentiel est que tu choisisses pour toi — pas pour faire plaisir à quelqu'un. "
                    f"Y a-t-il encore quelque chose qui te préoccupe ?")

    # ══ CAS RÊVEUR : bonnes aptitudes, notes faibles ══
    elif conflit_type == "reveur":
        if tour == 0:
            if maths_diff:
                return (f"Les maths qui bloquent — ça arrive à beaucoup d'élèves brillants, {prenom}. "
                        f"Et ton score D48={d48:.1f}/20 montre que ta logique est là. "
                        f"Souvent c'est une question de méthode, pas de capacité. "
                        f"Tu révises comment ? Après chaque cours ou juste avant les contrôles ?")
            if meth_pb:
                return (f"La concentration, c'est souvent la vraie bataille, {prenom}. "
                        f"Ce n'est pas une question d'intelligence — ton SA={SA:.1f}/20 le prouve. "
                        f"Est-ce que tu travailles dans un endroit calme ? "
                        f"Et est-ce que tu relis tes cours dans les 24h qui suivent ?")
            if fam:
                rev_low = (revenu or "").lower()
                if "faible" in rev_low:
                    return (f"Je comprends la contrainte financière, {prenom}. "
                            f"Mais sache qu'il existe des soutiens gratuits — programmes lycées, associations, "
                            f"élèves de terminale qui donnent des cours. Ton potentiel (SA={SA:.1f}/20) vaut "
                            f"vraiment qu'on cherche une solution ensemble.")
                return (f"Ta famille peut être un soutien précieux ici, {prenom}. "
                        f"Si tu leur montres tes résultats de tests (SA={SA:.1f}/20), "
                        f"ils comprendront que ce n'est pas un manque de capacité mais besoin d'accompagnement.")
            if pos:
                return (f"C'est exactement l'état d'esprit qu'il faut, {prenom} ! "
                        f"Avec cette détermination et tes capacités (SA={SA:.1f}/20), "
                        f"l'objectif {obj}/20 en sciences au T2 est tout à fait réalisable. "
                        f"Qu'est-ce qui te pose le plus de difficultés concrètement en ce moment ?")
            if neg:
                return (f"Je t'entends, {prenom}, et je ne minimise pas tes difficultés. "
                        f"Mais j'ai devant moi tes résultats : SA={SA:.1f}/20. "
                        f"Ces chiffres ne mentent pas — tu en es capable. "
                        f"Qu'est-ce qui s'est passé cette année pour que tes notes ne reflètent pas ça ?")
            if est_question:
                return (f"Très bonne question, {prenom}. Le décalage entre aptitudes ({SA:.1f}/20) "
                        f"et notes scolaires ({moy_sci:.1f}/20) a souvent une explication précise : "
                        f"méthode de travail, environnement, motivation, ou parfois un événement personnel. "
                        f"Est-ce qu'il s'est passé quelque chose cette année qui a perturbé ta scolarité ?")
            return (f"Pour comprendre la situation, {prenom}, j'ai besoin d'un peu plus de contexte. "
                    f"Tu as des capacités réelles (SA={SA:.1f}/20) mais tes notes ({moy_sci:.1f}/20) "
                    f"ne le montrent pas. Comment se passe une semaine normale de travail pour toi ?")
        else:
            if pos:
                return (f"C'est ce qu'on attendait de toi, {prenom} ! "
                        f"Objectif T2 : {obj}/20 en sciences. Ton dossier passe en suivi — "
                        f"le conseiller reviendra vers toi au prochain trimestre. Tiens bon !")
            if neg:
                return (f"C'est difficile, je comprends. Mais ce n'est pas fini, {prenom}. "
                        f"Ton dossier reste ouvert en suivi. Parle à ton enseignant principal "
                        f"— il peut souvent proposer un accompagnement personnalisé.")
            return (f"Ton dossier est bien pris en compte, {prenom}. "
                    f"Le conseiller suivra ton évolution au prochain trimestre. "
                    f"N'hésite pas à revenir si tu as besoin d'un accompagnement.")

    # Réponse de clôture
    if est_question:
        return (f"C'est une très bonne question, {prenom}. "
                f"Je te conseille d'en discuter aussi en face à face avec le conseiller de ton lycée "
                f"— il pourra aller plus loin dans l'accompagnement. "
                f"Est-ce qu'il y a autre chose qui te préoccupe ?")
    return (f"Merci pour cette précision, {prenom}. "
            f"Tout ce qu'on a échangé ici est pris en compte dans ton dossier. "
            f"Le conseiller dispose maintenant des éléments pour t'accompagner au mieux.")


# =====================================================================
# FONCTIONS EXPORTS DASHBOARD
# =====================================================================
def _generer_excel(dossiers):
    if not OPENPYXL_OK:
        return None
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Dossiers CapAvenir"
    hf = PatternFill("solid", fgColor="0F172A")
    hfont = Font(bold=True, color="FFFFFF", size=9)
    alt_fill = PatternFill("solid", fgColor="F8FAFC")
    headers = ["Nom","Prénom","Lycée","Âge","Sexe","Choix","Projet Pro",
               "SA","LA","Moy.Sci T1","Moy.Lit T1","D48","KRX","MECA","BV11","PRC",
               "Série","Statut","Confiance %","Trimestre","Date"]
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(1, ci, h)
        cell.fill = hf; cell.font = hfont
        cell.alignment = Alignment(horizontal="center")
    for ri, d in enumerate(dossiers, 2):
        serie = (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?")
        row = [
            d.get("nom",""), d.get("prenom",""), d.get("lycee",""),
            d.get("age",""), d.get("sexe",""), d.get("choix_personnel",""),
            d.get("projet_pro",""),
            round(float(d.get("SA_etal") or d.get("SA_brut") or 0), 2),
            round(float(d.get("LA_etal") or d.get("LA_brut") or 0), 2),
            round(float(d.get("moy_sci_t1") or 0), 2),
            round(float(d.get("moy_lit_t1") or 0), 2),
            d.get("d48_brut",""), d.get("krx_brut",""), d.get("meca_brut",""),
            d.get("bv11_brut",""), d.get("prc_brut",""),
            serie, d.get("statut",""), d.get("score_confiance",""),
            d.get("trimestre_actuel") or d.get("trimestre_decision",""),
            (d.get("date_modification","") or "")[:10],
        ]
        for ci, val in enumerate(row, 1):
            cell = ws.cell(ri, ci, val)
            if ri % 2 == 0:
                cell.fill = alt_fill
            cell.alignment = Alignment(horizontal="center" if ci > 7 else "left")
    # Largeurs colonnes
    widths = [14,12,20,5,8,18,24,6,6,10,10,6,6,6,6,6,7,12,10,10,12]
    for ci, w in enumerate(widths, 1):
        ws.column_dimensions[ws.cell(1,ci).column_letter].width = w
    # Figer la première ligne
    ws.freeze_panes = "A2"
    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    return buf.getvalue()


def _generer_word(dossiers, stats):
    if not DOCX_OK:
        return None
    doc = DocxDocument()
    doc.core_properties.title = "Rapport CapAvenir CMR"
    # Titre
    h = doc.add_heading("Rapport d'Orientation — CapAvenir CMR", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        f" · {len(dossiers)} dossier(s) au total"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    # Stats
    doc.add_heading("Statistiques globales", level=1)
    tbl = doc.add_table(rows=1, cols=2); tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Indicateur"
    tbl.rows[0].cells[1].text = "Valeur"
    for lbl, key in [("Total dossiers","total"),("Confirmés","nb_confirmes"),
                     ("En attente","nb_attente"),("Probation","nb_probation"),
                     ("Révisés","nb_revises"),("Série C","nb_serie_c"),("Série A","nb_serie_a")]:
        r = tbl.add_row().cells
        r[0].text = lbl; r[1].text = str(stats.get(key,0))
    doc.add_paragraph()
    # Liste dossiers
    doc.add_heading("Liste des dossiers", level=1)
    tbl2 = doc.add_table(rows=1, cols=7); tbl2.style = "Table Grid"
    for i, h in enumerate(["Nom & Prénom","Lycée","SA","LA","Série","Statut","Confiance"]):
        tbl2.rows[0].cells[i].text = h
        tbl2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for d in dossiers:
        serie = (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?")
        r = tbl2.add_row().cells
        r[0].text = f"{(d.get('nom','') or '').upper()} {(d.get('prenom','') or '').capitalize()}"
        r[1].text = (d.get("lycee","") or "—")[:28]
        r[2].text = f"{float(d.get('SA_etal') or d.get('SA_brut') or 0):.1f}"
        r[3].text = f"{float(d.get('LA_etal') or d.get('LA_brut') or 0):.1f}"
        r[4].text = serie; r[5].text = d.get("statut","—")
        r[6].text = f"{d.get('score_confiance',0) or 0} %"
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.getvalue()


def _generer_pdf_rapport(dossiers, stats):
    """PDF récapitulatif de tous les dossiers."""
    buf = io.BytesIO()
    CMR_DARK  = colors.HexColor("#0f172a")
    CMR_GREEN = colors.HexColor("#10b981")
    CMR_GRAY  = colors.HexColor("#64748b")
    CMR_LIGHT = colors.HexColor("#f8fafc")
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=1.5*cm, rightMargin=1.5*cm,
        topMargin=1.5*cm, bottomMargin=1.5*cm)
    sT  = ParagraphStyle("T",  fontName="Helvetica-Bold", fontSize=16, textColor=CMR_DARK, alignment=TA_CENTER, spaceAfter=4)
    sSub= ParagraphStyle("S",  fontName="Helvetica", fontSize=9, textColor=CMR_GRAY,  alignment=TA_CENTER, spaceAfter=8)
    sSec= ParagraphStyle("Se", fontName="Helvetica-Bold", fontSize=10, textColor=CMR_GREEN, spaceBefore=12, spaceAfter=4)
    sB  = ParagraphStyle("B",  fontName="Helvetica", fontSize=8, textColor=CMR_DARK)
    story = []
    story.append(Paragraph("RAPPORT D'ORIENTATION — CapAvenir CMR", sT))
    story.append(Paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} · {len(dossiers)} dossier(s)", sSub))
    story.append(HRFlowable(width="100%", thickness=2, color=CMR_GREEN, spaceAfter=8))
    # Stats
    story.append(Paragraph("STATISTIQUES", sSec))
    stat_rows = [["Total","Confirmés","En attente","Probation","Révisés","Série C","Série A"],
                 [str(stats.get(k,0)) for k in ["total","nb_confirmes","nb_attente","nb_probation","nb_revises","nb_serie_c","nb_serie_a"]]]
    st_tbl = Table(stat_rows, colWidths=[2.5*cm]*7)
    st_tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),CMR_DARK),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),8),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("BACKGROUND",(0,1),(-1,1),CMR_LIGHT),
        ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ]))
    story.append(st_tbl)
    # Liste
    story.append(Paragraph("LISTE DES DOSSIERS", sSec))
    rows = [["Nom & Prénom","Lycée","SA","LA","Série","Statut","Confiance","Date"]]
    for d in dossiers:
        serie = (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?")
        rows.append([
            f"{(d.get('nom','') or '').upper()} {(d.get('prenom','') or '').capitalize()}"[:22],
            (d.get("lycee","") or "—")[:18],
            f"{float(d.get('SA_etal') or d.get('SA_brut') or 0):.1f}",
            f"{float(d.get('LA_etal') or d.get('LA_brut') or 0):.1f}",
            serie, d.get("statut","—")[:10],
            f"{d.get('score_confiance',0) or 0}%",
            (d.get("date_modification","") or "")[:10],
        ])
    list_tbl = Table(rows, colWidths=[4.5*cm,3.5*cm,1.3*cm,1.3*cm,1.4*cm,2.4*cm,1.8*cm,2.3*cm])
    list_tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),CMR_DARK),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),7.5),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[CMR_LIGHT,colors.white]),
        ("GRID",(0,0),(-1,-1),0.2,colors.HexColor("#e2e8f0")),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),4),
    ]))
    story.append(list_tbl)
    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


def afficher_dashboard():
    """Tableau de bord complet — Mode Conseiller."""
    STATUT_META = {
        "confirme":    ("✅ Confirmés",    "#10b981","#d1fae5","badge-confirme"),
        "revise":      ("🔄 Révisés",      "#ec4899","#fce7f3","badge-revise"),
        "attente":     ("⏳ En attente",   "#f97316","#fff7ed","badge-attente"),
        "probation":   ("⚠️ Probation",    "#f59e0b","#fef3c7","badge-probation"),
        "indetermine": ("❓ Indéterminés","#6b7280","#f3f4f6","badge-indetermine"),
    }
    # ── Titre ──
    st.markdown("""
    <div style="display:flex;align-items:center;gap:1rem;margin-bottom:1.5rem;">
        <div style="font-size:2.4rem;">📊</div>
        <div>
            <div style="font-size:1.5rem;font-weight:800;">Tableau de Bord des Élèves Orientés</div>
            <div style="font-size:0.82rem;opacity:0.55;margin-top:2px;">
                Vue réservée au conseiller d'orientation · Données en temps réel
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

    # ── Statistiques ──
    try:
        stats = db.statistiques()
    except Exception:
        stats = {"total":0,"nb_confirmes":0,"nb_attente":0,"nb_probation":0,
                 "nb_revises":0,"nb_serie_c":0,"nb_serie_a":0}

    stat_items = [
        ("🗂️ Total",      stats.get("total",0),          "#6d28d9","#f3e8ff"),
        ("✅ Confirmés",   stats.get("nb_confirmes",0),   "#10b981","#d1fae5"),
        ("⏳ En attente",  stats.get("nb_attente",0),     "#f97316","#fff7ed"),
        ("⚠️ Probation",  stats.get("nb_probation",0),   "#f59e0b","#fef3c7"),
        ("🔄 Révisés",     stats.get("nb_revises",0),     "#ec4899","#fce7f3"),
    ]
    cols_stat = st.columns(len(stat_items))
    for col, (lbl, val, color, bg) in zip(cols_stat, stat_items):
        with col:
            st.markdown(f"""
            <div class="dash-stat-card">
                <div class="dash-stat-val" style="color:{color};">{val}</div>
                <div class="dash-stat-lbl">{lbl}</div>
            </div>""", unsafe_allow_html=True)

    # ── Graphiques ──
    total = stats.get("total", 0)
    if total > 0:
        st.write("")
        cg1, cg2 = st.columns([1, 2])
        with cg1:
            nb_c = stats.get("nb_serie_c", 0)
            nb_a = stats.get("nb_serie_a", 0)
            fig_pie = go.Figure(go.Pie(
                labels=["Série C", "Série A", "Autre"],
                values=[nb_c, nb_a, max(0, total - nb_c - nb_a)],
                hole=0.55, marker=dict(colors=["#10b981","#3b82f6","#94a3b8"],
                                       line=dict(color="white", width=2)),
                textinfo="label+percent", textfont_size=11,
            ))
            fig_pie.update_layout(showlegend=False, height=200,
                margin=dict(l=10,r=10,t=20,b=10),
                annotations=[dict(text=f"<b>{total}</b>",x=0.5,y=0.5,font_size=18,showarrow=False)])
            st.plotly_chart(fig_pie, use_container_width=True)
        with cg2:
            cats = ["Confirmés","En attente","Probation","Révisés"]
            vals = [stats.get("nb_confirmes",0), stats.get("nb_attente",0),
                    stats.get("nb_probation",0), stats.get("nb_revises",0)]
            fig_bar = go.Figure(go.Bar(x=cats, y=vals,
                marker_color=["#10b981","#f97316","#f59e0b","#ec4899"],
                text=vals, textposition="outside", marker=dict(line=dict(width=0))))
            fig_bar.update_layout(height=200, margin=dict(l=10,r=10,t=20,b=10),
                yaxis=dict(visible=False), showlegend=False,
                plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)")
            st.plotly_chart(fig_bar, use_container_width=True)

    st.divider()

    # ══════════════════════════════════════════════════════════════
    # ONGLETS PRINCIPAUX DU DASHBOARD
    # ══════════════════════════════════════════════════════════════
    tab_inscrits, tab_dossiers, tab_suivi, tab_exports = st.tabs([
        "👥 Élèves Inscrits — Saisir notes & Diagnostic",
        "📋 Tous les Dossiers",
        "🔄 Suivi T2/T3",
        "📤 Exports",
    ])

    # ──────────────────────────────────────────────────────────────
    # ONGLET 1 : ÉLÈVES INSCRITS → Saisie notes + Diagnostic
    # ──────────────────────────────────────────────────────────────
    with tab_inscrits:
        st.markdown("#### 👥 Liste des élèves inscrits sur CapAvenir CMR")
        st.markdown("""
        <div style="background:linear-gradient(135deg,#eff6ff,#dbeafe);
                    border-left:4px solid #3b82f6; border-radius:10px;
                    padding:0.7rem 1.2rem; margin-bottom:1rem; font-size:0.84rem; color:#1e3a8a;">
            📌 <strong>Mode Conseiller :</strong> Pour chaque élève ayant passé les tests psychotechniques,
            saisissez les notes du trimestre et cliquez sur <strong>Lancer le Diagnostic</strong>
            pour générer l'orientation par le moteur IA.
        </div>
        """, unsafe_allow_html=True)

        # Charger tous les inscrits depuis eleve_credentials
        try:
            rows_cred = db.conn.execute(
                "SELECT nom, prenom, lycee, date_creation FROM eleve_credentials ORDER BY date_creation DESC"
            ).fetchall()
            inscrits = [dict(r) for r in rows_cred]
        except Exception as e:
            inscrits = []
            st.error(f"Erreur lecture des inscrits : {e}")

        if not inscrits:
            st.info("Aucun élève inscrit pour le moment. Les élèves s'inscrivent depuis la page d'accueil.")
        else:
            # Barre de recherche dans les inscrits
            srch_ins = st.text_input("🔍 Filtrer par nom ou prénom", key="srch_inscrits",
                                      placeholder="Ex : MBALLA, Jean...")
            if srch_ins.strip():
                q_ins = srch_ins.strip().upper()
                inscrits = [i for i in inscrits if
                            q_ins in i.get("nom","").upper() or q_ins in i.get("prenom","").upper()]

            st.markdown(f"**{len(inscrits)} élève(s) inscrit(s)**")
            st.write("")

            for ins in inscrits:
                nom_i    = (ins.get("nom","") or "").upper()
                prenom_i = (ins.get("prenom","") or "").capitalize()
                lycee_i  = ins.get("lycee","") or "—"
                date_i   = (ins.get("date_creation","") or "")[:10]

                # Vérifier si un dossier existe en BDD
                try:
                    dossiers_ex = db.rechercher_eleve(nom_i, prenom_i)
                except Exception:
                    dossiers_ex = []

                has_dossier = len(dossiers_ex) > 0
                dossier_ex  = dossiers_ex[0] if has_dossier else None

                sa_ex = float(dossier_ex.get("SA_etal") or dossier_ex.get("SA_brut") or 0) if dossier_ex else None
                la_ex = float(dossier_ex.get("LA_etal") or dossier_ex.get("LA_brut") or 0) if dossier_ex else None
                has_tests = has_dossier and (sa_ex is not None and sa_ex > 0)
                serie_ex = (dossier_ex.get("serie_finale") or dossier_ex.get("serie_provisoire") or
                            dossier_ex.get("serie_cible") or None) if dossier_ex else None

                # Badge état
                if has_tests and serie_ex:
                    badge = f"<span style='background:#d1fae5;color:#065f46;padding:2px 10px;border-radius:20px;font-size:0.75rem;font-weight:700;'>✅ Orienté → Série {serie_ex}</span>"
                    expander_label = f"📋 {nom_i} {prenom_i} · {lycee_i} · ✅ Orienté Série {serie_ex}"
                elif has_tests:
                    badge = f"<span style='background:#fef3c7;color:#92400e;padding:2px 10px;border-radius:20px;font-size:0.75rem;font-weight:700;'>🧪 Tests OK — Notes à saisir</span>"
                    expander_label = f"📋 {nom_i} {prenom_i} · {lycee_i} · 🧪 Tests passés — Saisir les notes"
                else:
                    badge = "<span style='background:#f1f5f9;color:#64748b;padding:2px 10px;border-radius:20px;font-size:0.75rem;font-weight:700;'>⏳ Tests non encore passés</span>"
                    expander_label = f"📋 {nom_i} {prenom_i} · {lycee_i} · ⏳ En attente des tests"

                with st.expander(expander_label, expanded=(has_tests and not serie_ex)):
                    ci1, ci2, ci3 = st.columns([2, 2, 2])
                    with ci1:
                        st.markdown(f"**👤 {nom_i} {prenom_i}**")
                        st.caption(f"🏫 {lycee_i}")
                    with ci2:
                        st.caption(f"📅 Inscrit le {date_i}")
                        if has_tests:
                            st.markdown(f"🧪 SA = **{sa_ex:.1f}**/20 · LA = **{la_ex:.1f}**/20")
                    with ci3:
                        st.markdown(badge, unsafe_allow_html=True)

                    if not has_tests:
                        st.markdown("""
                        <div style="background:#f8fafc;border:1px dashed #cbd5e1;border-radius:8px;
                                    padding:0.8rem 1rem;margin-top:0.5rem;font-size:0.83rem;color:#64748b;text-align:center;">
                            ⏳ Cet élève n'a pas encore passé les tests psychotechniques.<br>
                            Accompagnez-le dans la passation des tests depuis l'<strong>Espace Élève</strong>.
                        </div>
                        """, unsafe_allow_html=True)
                        continue

                    # ── Formulaire de saisie des notes ──────────────────────
                    st.divider()
                    st.markdown("##### 📚 Saisie des notes — Trimestre 1")
                    uid_diag = f"{nom_i}_{prenom_i}_diag"

                    # Pré-remplir si dossier existant
                    def _pf(key, default=10.0):
                        if dossier_ex:
                            v = dossier_ex.get(f"{key}_t1") or dossier_ex.get(key)
                            if v is not None:
                                try: return float(v)
                                except: pass
                        return default

                    cn1, cn2 = st.columns(2)
                    with cn1:
                        st.markdown("**🔬 Matières scientifiques**")
                        n_maths  = st.number_input("Maths",    0.0, 20.0, _pf("maths"),  0.5, key=f"dm_{uid_diag}")
                        n_sciphy = st.number_input("Sci.Phy",  0.0, 20.0, _pf("sci_phy"),0.5, key=f"dsp_{uid_diag}")
                        n_svt    = st.number_input("SVT",      0.0, 20.0, _pf("svt"),    0.5, key=f"dsv_{uid_diag}")
                        moy_sci_d = round((n_maths*5 + n_sciphy*4 + n_svt*2)/11, 2)
                        st.metric("Moyenne Sci (coeff.)", f"{moy_sci_d:.2f}/20",
                                  delta=f"{moy_sci_d-10:+.1f} vs seuil 10",
                                  delta_color="normal" if moy_sci_d>=10 else "inverse")
                    with cn2:
                        st.markdown("**📖 Matières littéraires**")
                        n_fr     = st.number_input("Français", 0.0, 20.0, _pf("francais"),0.5, key=f"dfr_{uid_diag}")
                        n_hg     = st.number_input("Hist-Géo", 0.0, 20.0, _pf("histgeo"), 0.5, key=f"dhg_{uid_diag}")
                        n_ang    = st.number_input("Anglais",  0.0, 20.0, _pf("anglais"),  0.5, key=f"dan_{uid_diag}")
                        moy_lit_d = round((n_fr*5 + n_hg*3 + n_ang*2)/10, 2)
                        st.metric("Moyenne Lit (coeff.)", f"{moy_lit_d:.2f}/20",
                                  delta=f"{moy_lit_d-10:+.1f} vs seuil 10",
                                  delta_color="normal" if moy_lit_d>=10 else "inverse")

                    # ── Prévisualisation en temps réel ───────────────────────
                    st.write("")
                    sa_d = sa_diag
                    la_d = la_diag

                    # Choix de série de l'élève (depuis son profil enregistré)
                    choix_diag = base.get("choix_personnel","") if dossier_ex else (
                        sub_ex.get("choix_personnel","") if has_submission else ""
                    )
                    choix_C_d = "C" in choix_diag
                    choix_A_d = "A" in choix_diag

                    # ── Règle prioritaire : choix de l'élève ──
                    _prio_C = choix_C_d and sa_d >= 10.0 and moy_sci_d >= 10.0
                    _prio_A = choix_A_d and la_d >= 10.0 and moy_lit_d >= 10.0

                    if _prio_C:
                        serie_preview  = "C"
                        seuil_ok       = True
                        raison_preview = f"Choix élève validé — SA={sa_d:.1f}≥10 · Moy.Sci={moy_sci_d:.1f}≥10"
                    elif _prio_A:
                        serie_preview  = "A"
                        seuil_ok       = True
                        raison_preview = f"Choix élève validé — LA={la_d:.1f}≥10 · Moy.Lit={moy_lit_d:.1f}≥10"
                    else:
                        # Moteur standard
                        serie_preview = "C" if sa_d >= la_d else "A"
                        seuil_ok = moy_sci_d >= 10 if serie_preview == "C" else moy_lit_d >= 10
                        raison_preview = "Moteur standard (aptitudes + notes)"

                    statut_preview = "confirme" if seuil_ok else "attente"
                    conf_preview   = min(100, max(0,
                        round((min(50, max(0, (sa_d - la_d)/max(sa_d,1)*50)) +
                               min(50, max(0, moy_sci_d/20*50))) if serie_preview=="C"
                        else (min(50, max(0, (la_d - sa_d)/max(la_d,1)*50)) +
                               min(50, max(0, moy_lit_d/20*50))), 0)
                    ))
                    col_prev = "#10b981" if serie_preview=="C" else "#3b82f6"
                    ok_bg = "#ecfdf5" if seuil_ok else "#fff7ed"
                    ok_col = "#065f46" if seuil_ok else "#92400e"
                    prio_badge = ""
                    if _prio_C or _prio_A:
                        prio_badge = f"<div style='font-size:0.78rem;margin-top:0.4rem;color:#7c3aed;font-weight:600;'>🎯 {raison_preview}</div>"
                    st.markdown(f"""
                    <div style="background:{ok_bg};border-radius:12px;padding:0.9rem 1.2rem;
                                border:2px dashed {col_prev};margin-top:0.2rem;">
                        <div style="font-size:0.85rem;font-weight:700;color:{ok_col};">
                            🤖 Prévisualisation Diagnostic IA
                        </div>
                        <div style="margin-top:0.4rem;font-size:0.82rem;color:{ok_col};">
                            Série recommandée : <strong style="color:{col_prev};font-size:1rem;">
                            Série {serie_preview}</strong> &nbsp;·&nbsp;
                            Statut : <strong>{"CONFIRMÉ ✅" if seuil_ok else "EN ATTENTE ⏳"}</strong>
                            &nbsp;·&nbsp; Confiance : <strong>{conf_preview}%</strong>
                        </div>
                        {prio_badge}
                    </div>
                    """, unsafe_allow_html=True)

                    st.write("")
                    if st.button(f"🚀 Lancer le Diagnostic et Sauvegarder",
                                 use_container_width=True, type="primary",
                                 key=f"diag_{uid_diag}"):
                        # Construire le dossier complet
                        base = dict(dossier_ex) if dossier_ex else {}
                        base.update({
                            "nom":              nom_i,
                            "prenom":           prenom_i,
                            "lycee":            lycee_i,
                            "age":              base.get("age", ""),
                            "sexe":             base.get("sexe", ""),
                            "choix_personnel":  base.get("choix_personnel",""),
                            "projet_pro":       base.get("projet_pro",""),
                            "revenu":           base.get("revenu_famille",""),
                            "revenu_famille":   base.get("revenu_famille",""),
                            "SA_etal":          sa_d,  "SA_brut": sa_d,
                            "LA_etal":          la_d,  "LA_brut": la_d,
                            "maths_t1":    n_maths,  "sci_phy_t1": n_sciphy, "svt_t1":  n_svt,
                            "francais_t1": n_fr,     "histgeo_t1": n_hg,     "anglais_t1": n_ang,
                            "t1_renseigne":     True,
                            "trimestre_actuel": "T1",
                            "serie_cible":      serie_preview,
                            "serie_provisoire": serie_preview,
                            "serie_finale":     serie_preview if seuil_ok else None,
                            "orientation_finale": serie_preview if seuil_ok else None,
                            "statut":           statut_preview,
                            "score_confiance":  conf_preview,
                        })
                        res_diag = db.sauvegarder_dossier(base, None)
                        if res_diag.get("succes"):
                            st.success(f"✅ Diagnostic lancé et dossier sauvegardé ! "
                                       f"Orientation → **Série {serie_preview}** · "
                                       f"Statut : **{statut_preview.upper()}**")
                            rerun()
                        else:
                            st.error(f"❌ Erreur : {res_diag.get('message','inconnue')}")

    # ──────────────────────────────────────────────────────────────
    # ONGLET 2 : TOUS LES DOSSIERS
    # ──────────────────────────────────────────────────────────────
    with tab_dossiers:
        # ── Filtres ──
        fc1, fc2, fc3 = st.columns([2, 1, 1])
        with fc1:
            search_q = st.text_input("🔍 Rechercher (nom ou prénom)",
                                      placeholder="Ex : MBALLA, Jean...", key="dash_search")
        with fc2:
            filtre_statut = st.selectbox("Statut",
                ["Tous","confirme","revise","attente","probation","indetermine"], key="dash_filtre")
        with fc3:
            filtre_serie = st.selectbox("Série", ["Toutes","C","A","?"], key="dash_serie")

        # ── Chargement données ──
        try:
            statuts_a_lister = (["confirme","revise","attente","probation","indetermine"]
                                if filtre_statut == "Tous" else [filtre_statut])
            tous_dossiers = []
            for s in statuts_a_lister:
                for d in db.lister_dossiers(s):
                    d["_statut_affiche"] = s
                    tous_dossiers.append(d)
        except Exception as e:
            st.error(f"Erreur lecture BDD : {e}")
            tous_dossiers = []

        if search_q.strip():
            q = search_q.strip().upper()
            tous_dossiers = [d for d in tous_dossiers if
                             q in (d.get("nom","") or "").upper() or
                             q in (d.get("prenom","") or "").upper()]
        if filtre_serie != "Toutes":
            tous_dossiers = [d for d in tous_dossiers if
                             (d.get("serie_finale") or d.get("serie_provisoire") or
                              d.get("serie_cible") or "?") == filtre_serie]

        # ── Tableau ──
        st.markdown(f"**{len(tous_dossiers)} dossier(s) trouvé(s)**")
        if not tous_dossiers:
            st.info("Aucun dossier ne correspond aux critères.")
        else:
            h_cols = st.columns([2.5, 2.2, 2, 1.2, 1.2, 1.5, 1.8, 1.8])
            for col, lbl in zip(h_cols, ["Nom & Prénom","Lycée","Projet","SA","LA","Série","Statut","Date"]):
                col.markdown(f"<div style='font-size:0.7rem;font-weight:700;text-transform:uppercase;"
                             f"opacity:0.5;letter-spacing:0.06em;'>{lbl}</div>", unsafe_allow_html=True)
            st.markdown("<hr style='margin:0.2rem 0 0.4rem;opacity:0.2;'>", unsafe_allow_html=True)
            for d in tous_dossiers:
                stat_af = d.get("_statut_affiche","—")
                _, _, _, s_badge = STATUT_META.get(stat_af, ("","#6b7280","#f3f4f6","badge-indetermine"))
                serie_af = (d.get("serie_finale") or d.get("serie_provisoire") or
                            d.get("serie_cible") or "?")
                badge_serie = "badge-c" if serie_af=="C" else ("badge-a" if serie_af=="A" else "")
                sa_v = float(d.get("SA_etal") or d.get("SA_brut") or 0)
                la_v = float(d.get("LA_etal") or d.get("LA_brut") or 0)
                rc = st.columns([2.5, 2.2, 2, 1.2, 1.2, 1.5, 1.8, 1.8])
                rc[0].markdown(f"**{(d.get('nom') or '').upper()}** {(d.get('prenom') or '').capitalize()}")
                rc[1].markdown(f"<small>{(d.get('lycee') or '—')[:22]}</small>", unsafe_allow_html=True)
                rc[2].markdown(f"<small>{(d.get('projet_pro') or '—')[:20]}</small>", unsafe_allow_html=True)
                rc[3].markdown(f"<b style='color:#10b981;'>{sa_v:.1f}</b>", unsafe_allow_html=True)
                rc[4].markdown(f"<b style='color:#3b82f6;'>{la_v:.1f}</b>", unsafe_allow_html=True)
                rc[5].markdown(f'<span class="dash-badge {badge_serie}">{serie_af}</span>', unsafe_allow_html=True)
                rc[6].markdown(f'<span class="dash-badge {s_badge}">{stat_af}</span>', unsafe_allow_html=True)
                rc[7].markdown(f"<small style='opacity:0.5;'>{(d.get('date_modification','') or '')[:10]}</small>", unsafe_allow_html=True)
                st.markdown("<div style='border-bottom:1px solid rgba(0,0,0,0.05);margin:0.1rem 0;'></div>", unsafe_allow_html=True)


    # ──────────────────────────────────────────────────────────────
    # ONGLET 3 : SUIVI T2/T3
    # ──────────────────────────────────────────────────────────────
    with tab_suivi:
        # ── Mise à jour T2/T3 pour dossiers non confirmés ──
        st.divider()
        st.markdown("#### 📝 Saisir les notes T2/T3 — Dossiers en attente / probation")
        non_conf = [d for d in tous_dossiers if d.get("_statut_affiche") in ("attente","probation")]
        if not non_conf:
            st.info("Aucun dossier en attente ou en probation dans cette vue.")
        else:
            for d in non_conf:
                nom_d    = (d.get("nom","") or "").upper()
                prenom_d = (d.get("prenom","") or "").capitalize()
                stat_d   = d.get("_statut_affiche","")
                sa_d     = float(d.get("SA_etal") or d.get("SA_brut") or 0)
                la_d     = float(d.get("LA_etal") or d.get("LA_brut") or 0)
                trim_a   = d.get("trimestre_actuel") or d.get("trimestre_decision","T1")
                prochain = "T2" if trim_a == "T1" else "T3"
                uid      = f"{d.get('id',0)}_{nom_d}"

                with st.expander(f"📋 {nom_d} {prenom_d} — {stat_d.upper()} → Saisir notes {prochain}"):
                    st.markdown(f"SA={sa_d:.1f}/20 · LA={la_d:.1f}/20 · Trimestre actuel : {trim_a}")
                    ca, cb = st.columns(2)
                    with ca:
                        st.markdown("**🔬 Matières scientifiques**")
                        mths = st.number_input(f"Maths {prochain}", 0.0, 20.0,
                            float(d.get(f"maths_{prochain.lower()}") or 10.0), 0.5, key=f"mths_{uid}")
                        sphy = st.number_input(f"Sci.Phy {prochain}", 0.0, 20.0,
                            float(d.get(f"sci_phy_{prochain.lower()}") or 10.0), 0.5, key=f"sphy_{uid}")
                        svtv = st.number_input(f"SVT {prochain}", 0.0, 20.0,
                            float(d.get(f"svt_{prochain.lower()}") or 10.0), 0.5, key=f"svtv_{uid}")
                        msc  = round((mths+sphy+svtv)/3, 2)
                        st.metric(f"Moy. Sci. {prochain}", f"{msc:.2f}/20",
                                  delta=f"{msc-10:+.1f} vs seuil")
                    with cb:
                        st.markdown("**📖 Matières littéraires**")
                        frv  = st.number_input(f"Français {prochain}", 0.0, 20.0,
                            float(d.get(f"francais_{prochain.lower()}") or 10.0), 0.5, key=f"frv_{uid}")
                        hgv  = st.number_input(f"Hist-Géo {prochain}", 0.0, 20.0,
                            float(d.get(f"histgeo_{prochain.lower()}") or 10.0), 0.5, key=f"hgv_{uid}")
                        angv = st.number_input(f"Anglais {prochain}", 0.0, 20.0,
                            float(d.get(f"anglais_{prochain.lower()}") or 10.0), 0.5, key=f"angv_{uid}")
                        mlt  = round((frv+hgv+angv)/3, 2)
                        st.metric(f"Moy. Lit. {prochain}", f"{mlt:.2f}/20",
                                  delta=f"{mlt-10:+.1f} vs seuil")

                    # Calcul automatique nouveau statut
                    serie_v   = (d.get("serie_finale") or d.get("serie_provisoire") or
                                 d.get("serie_cible") or "C")
                    cond_ok   = (msc >= 10) if sa_d >= la_d else (mlt >= 10)
                    if stat_d == "attente":
                        new_stat = "confirme" if cond_ok else "probation"
                        new_lbl  = "✅ CONFIRMÉ" if cond_ok else "⚠️ PROBATION"
                    else:
                        new_stat = "confirme" if cond_ok else "revise"
                        serie_v  = serie_v if cond_ok else ("A" if serie_v=="C" else "C")
                        new_lbl  = "✅ CONFIRMÉ" if cond_ok else f"🔄 RÉVISÉ → Série {serie_v}"
                    ok_c = "#10b981" if cond_ok else "#ef4444"
                    ok_bg= "#ecfdf5" if cond_ok else "#fef2f2"
                    st.markdown(f"""
                    <div style="background:{ok_bg};border:1.5px solid {ok_c};border-radius:10px;
                                padding:0.6rem 1rem;text-align:center;font-weight:700;
                                color:{ok_c};margin:0.5rem 0;">
                        Nouveau statut prévu : {new_lbl}
                    </div>""", unsafe_allow_html=True)

                    if st.button(f"✅ Valider et mettre à jour dans la BDD",
                                 use_container_width=True, type="primary",
                                 key=f"val_{uid}"):
                        ss_upd = dict(d)
                        tk = prochain.lower()
                        ss_upd[f"maths_{tk}"]   = mths
                        ss_upd[f"sci_phy_{tk}"] = sphy
                        ss_upd[f"svt_{tk}"]     = svtv
                        ss_upd[f"francais_{tk}"]= frv
                        ss_upd[f"histgeo_{tk}"] = hgv
                        ss_upd[f"anglais_{tk}"] = angv
                        ss_upd[f"t{prochain[1]}_renseigne"] = True
                        ss_upd["statut"]            = new_stat
                        ss_upd["orientation_finale"]= serie_v
                        ss_upd["revenu"]            = d.get("revenu_famille","")
                        res_upd = db.sauvegarder_dossier(ss_upd, None)
                        if res_upd["succes"]:
                            st.success(f"✅ Statut mis à jour → **{new_stat.upper()}** · {res_upd['message']}")
                            rerun()
                        else:
                            st.error(res_upd["message"])


    # ──────────────────────────────────────────────────────────────
    # ONGLET 4 : EXPORTS
    # ──────────────────────────────────────────────────────────────
    with tab_exports:
        # ── Exports multi-formats ──
        st.divider()
        st.markdown("#### 📤 Exports")

        try:
            all_exp = []
            for s in ["confirme","revise","attente","probation","indetermine"]:
                for dd in db.lister_dossiers(s):
                    dd["_statut_affiche"] = s
                    all_exp.append(dd)
            stats_exp = db.statistiques()
        except Exception:
            all_exp = []; stats_exp = {}

        e1, e2, e3, e4 = st.columns(4)
        with e1:
            try:
                xlsx = generer_excel_export(all_exp)
                if xlsx:
                    st.download_button("📊 Excel (.xlsx)", xlsx,
                        file_name=f"capavenir_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True)
                else:
                    st.caption("Installer openpyxl")
            except Exception as ex:
                st.caption(f"Erreur Excel: {ex}")
        with e2:
            try:
                docx_b = generer_word_export(all_exp, stats_exp)
                if docx_b:
                    st.download_button("📝 Word (.docx)", docx_b,
                        file_name=f"capavenir_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True)
                else:
                    st.caption("Installer python-docx")
            except Exception as ex:
                st.caption(f"Erreur Word: {ex}")
        with e3:
            try:
                pdf_b = generer_pdf_rapport(all_exp, stats_exp)
                st.download_button("📄 PDF Rapport", pdf_b,
                    file_name=f"rapport_capavenir_{datetime.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf", use_container_width=True)
            except Exception as ex:
                st.caption(f"Erreur PDF: {ex}")
        with e4:
            try:
                j_data = db.exporter_json()
                st.download_button("📦 JSON", j_data,
                    file_name=f"capavenir_{datetime.now().strftime('%Y%m%d')}.json",
                    mime="application/json", use_container_width=True)
            except Exception:
                st.caption("Export indisponible")


# =====================================================================
# ESPACE ÉLÈVE — Lecture seule, accès via recherche par nom
# =====================================================================
def afficher_espace_eleve():
    """Vue dédiée à l'élève : résultats tests, conversation IA, orientation. Lecture seule."""

    st.markdown("""
    <div style="display:flex;align-items:center;gap:0.8rem;margin-bottom:1.5rem;">
        <div style="font-size:2.2rem;">🎒</div>
        <div>
            <div style="font-size:1.4rem;font-weight:800;">Espace Personnel de l'Élève</div>
            <div style="font-size:0.82rem;opacity:0.6;margin-top:2px;">
                Accès lecture seule · Résultats des tests · Conversation avec l'agent IA
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

    # ── Recherche de l'élève dans la BDD ──
    # Le dossier trouvé est persisté en session pour survivre aux rerun()
    dossier_actif = st.session_state.get("eleve_dossier_actif", None)

    if dossier_actif is None:
        # ── Formulaire de recherche ──
        st.markdown("#### 🔍 Rechercher mon dossier")
        col_s1, col_s2, col_s3 = st.columns([2, 2, 1])
        with col_s1:
            nom_rech = st.text_input("Nom", placeholder="Ex : MBALLA",
                                      key="eleve_nom",
                                      value=st.session_state.get("eleve_nom_recherche",""))
        with col_s2:
            prenom_rech = st.text_input("Prénom", placeholder="Ex : Jean", key="eleve_prenom")
        with col_s3:
            st.write("")
            chercher = st.button("🔍 Rechercher", use_container_width=True, type="primary")

        if chercher and nom_rech.strip():
            st.session_state.eleve_nom_recherche = nom_rech.strip()
            try:
                resultats = db.rechercher_eleve(nom_rech.strip(), prenom_rech.strip() or None)
            except Exception:
                resultats = []
            if not resultats:
                st.warning("Aucun dossier trouvé. Vérifiez l'orthographe ou contactez votre conseiller d'orientation.")
                return
            # ── Stocker le dossier en session pour survivre aux rerun() ──
            st.session_state.eleve_dossier_actif = resultats[0]
            rerun()
        elif not chercher:
            st.markdown("""
            <div class="alert-info">
                👆 Saisis ton <strong>nom de famille</strong> et clique sur
                <strong>Rechercher</strong> pour accéder à ton dossier.
            </div>""", unsafe_allow_html=True)
        return

    # ── Dossier déjà trouvé et persisté ──
    d        = dossier_actif
    nom_d    = (d.get("nom","") or "").upper()
    prenom_d = (d.get("prenom","") or "").capitalize()
    lycee_d  = d.get("lycee","") or "—"
    serie_d  = (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?")
    statut_d = d.get("statut","—")
    sa_d     = float(d.get("SA_etal") or d.get("SA_brut") or 0)
    la_d     = float(d.get("LA_etal") or d.get("LA_brut") or 0)

    # Bouton pour changer de dossier
    if st.button("🔄 Changer de dossier", key="changer_dossier"):
        st.session_state.eleve_dossier_actif = None
        st.session_state.eleve_nom_recherche = ""
        rerun()

    # ── En-tête du dossier élève ──
    badge_c = "#10b981" if serie_d=="C" else "#3b82f6" if serie_d=="A" else "#f59e0b"
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,#1e293b,#1a1f35);
                border-radius:16px;padding:1.5rem 2rem;margin:1rem 0;
                border:1px solid #334155;color:white;">
        <div style="font-size:0.75rem;opacity:0.6;text-transform:uppercase;letter-spacing:0.08em;">
            Dossier d'Orientation
        </div>
        <div style="font-size:1.6rem;font-weight:800;margin:0.3rem 0;">
            {nom_d} {prenom_d}
        </div>
        <div style="font-size:0.88rem;opacity:0.75;">{lycee_d}</div>
        <div style="margin-top:0.8rem;display:flex;gap:0.8rem;flex-wrap:wrap;align-items:center;">
            <span style="background:{badge_c};color:white;padding:0.3rem 1.2rem;
                         border-radius:20px;font-weight:700;font-size:1rem;">
                SÉRIE {serie_d}
            </span>
            <span style="background:rgba(255,255,255,0.1);padding:0.3rem 0.8rem;
                         border-radius:20px;font-size:0.82rem;">{statut_d.upper()}</span>
        </div>
    </div>""", unsafe_allow_html=True)

    # ── Section 1 : Tests psychotechniques (passation ou résultats) ──
    st.markdown("---")
    st.markdown("### 🧪 Tests Psychotechniques")

    # Vérifier si les tests ont déjà été passés pour cet élève dans la BDD
    tests_passes_bdd = any(d.get(k) is not None and float(d.get(k) or 0) > 0
                           for k in ["d48_brut","krx_brut","meca_brut","bv11_brut","prc_brut"])

    if tests_passes_bdd:
        # Tests déjà passés → afficher résumé (sans scores numériques)
        st.markdown("""
        <div class="alert-success">
            ✅ <strong>Vos tests ont été enregistrés et analysés par votre Conseiller.</strong><br>
            <small>Les scores détaillés sont consultables par le Conseiller d'Orientation uniquement.</small>
        </div>""", unsafe_allow_html=True)
        # Afficher uniquement un graphique radar visuel (normalisé, sans chiffres)
        TESTS_INFO = {
            "d48_brut":  ("D48 — Logique",     "d48_etal"),
            "krx_brut":  ("KRX — Maths",       "krx_etal"),
            "meca_brut": ("MECA — Mécanique",  "meca_etal"),
            "bv11_brut": ("BV11 — Littéraire", "bv11_etal"),
            "prc_brut":  ("PRC — Proverbes",   "prc_etal"),
        }
        labels_r2 = [v[0] for v in TESTS_INFO.values()]
        vals_r2   = [float(d.get(ek) or d.get(bk) or 0)
                     for bk,(nm,ek) in TESTS_INFO.items()]
        if any(v > 0 for v in vals_r2):
            fig_r2 = go.Figure(go.Scatterpolar(
                r=vals_r2, theta=labels_r2, fill="toself",
                line_color="#7c3aed", fillcolor="rgba(124,58,237,0.15)"
            ))
            fig_r2.update_layout(
                polar=dict(radialaxis=dict(visible=False), gridshape="linear"),
                height=240, showlegend=False,
                margin=dict(l=40,r=40,t=20,b=20)
            )
            st.plotly_chart(fig_r2, use_container_width=True)
        st.caption("Profil d'aptitudes — scores masqués (accès Conseiller requis)")
    else:
        # Tests non encore passés → proposer la passation guidée
        st.markdown("""
        <div class="alert-info">
            📋 Vous n'avez pas encore passé les tests psychotechniques.
            Complétez-les ci-dessous. Vos réponses seront analysées par votre Conseiller.
        </div>""", unsafe_allow_html=True)

        test_keys_e  = ["D48", "KRX", "MECA", "BV11", "PRC"]
        test_names_e = [TESTS_QUESTIONS[k]["nom"].split("—")[0].strip() for k in test_keys_e]
        test_tabs_e  = st.tabs(test_names_e)

        for t_key, t_tab in zip(test_keys_e, test_tabs_e):
            with t_tab:
                tdata = TESTS_QUESTIONS[t_key]
                st.caption(tdata["description"])
                eleve_key = f"eleve_answers_{t_key}"
                if eleve_key not in st.session_state:
                    st.session_state[eleve_key] = {}
                nb_correct_e = 0
                for q_idx, q in enumerate(tdata["questions"]):
                    st.markdown(f'<div class="test-question"><strong>Q{q_idx+1}.</strong> {q["q"]}</div>', unsafe_allow_html=True)
                    prev_e = st.session_state[eleve_key].get(q_idx)
                    # Aucune réponse présélectionnée : index=None si rien encore choisi
                    idx_def_e = q["choices"].index(prev_e) if prev_e in q["choices"] else None
                    ch_e = st.radio(
                        f"Eleve_R{q_idx+1}_{t_key}",
                        q["choices"], index=idx_def_e,
                        key=f"eleve_q_{t_key}_{q_idx}",
                        label_visibility="collapsed"
                    )
                    if ch_e is not None:
                        st.session_state[eleve_key][q_idx] = ch_e
                        if ch_e == q["answer"]:
                            nb_correct_e += 1
                nb_rep_e = sum(1 for v in st.session_state[eleve_key].values() if v is not None)
                total_e  = len(tdata["questions"])
                if nb_rep_e >= total_e:
                    # Score brut = nb bonnes réponses sur total_e questions
                    score_brut_e = nb_correct_e          # ex. 28 sur 40
                    # Score étalonnée = ramené sur 20
                    score_etal_e = round((score_brut_e / total_e) * 20, 1)
                    # Stocker dans les variables de session (score étalonnée)
                    st.session_state[f"eleve_score_{t_key}"]      = score_etal_e
                    st.session_state[f"eleve_score_brut_{t_key}"] = score_brut_e
                    st.markdown(f'<div class="alert-success">✅ Test {t_key} complété — {total_e} réponses enregistrées.</div>', unsafe_allow_html=True)
                    # ── Résultats visibles par l'élève (sans correction) ──
                    st.markdown(f"""
                    <div style="background:linear-gradient(135deg,#ecfdf5,#d1fae5);
                                border-left:4px solid #10b981;border-radius:12px;
                                padding:1rem 1.4rem;margin-top:0.5rem;">
                        <div style="font-weight:700;color:#065f46;font-size:0.95rem;margin-bottom:0.6rem;">
                            📊 Vos résultats — Test {t_key}
                        </div>
                        <div style="display:flex;gap:2.5rem;flex-wrap:wrap;">
                            <div>
                                <div style="font-size:0.72rem;color:#047857;text-transform:uppercase;
                                            letter-spacing:.06em;font-weight:600;">Score brut</div>
                                <div style="font-size:1.6rem;font-weight:800;color:#065f46;line-height:1.1;">
                                    {score_brut_e}<span style="font-size:0.9rem;font-weight:500;
                                    color:#047857;">/{total_e} pts</span>
                                </div>
                            </div>
                            <div>
                                <div style="font-size:0.72rem;color:#047857;text-transform:uppercase;
                                            letter-spacing:.06em;font-weight:600;">Score étalonnée</div>
                                <div style="font-size:1.6rem;font-weight:800;color:#065f46;line-height:1.1;">
                                    {score_etal_e}<span style="font-size:0.9rem;font-weight:500;
                                    color:#047857;">/20</span>
                                </div>
                            </div>
                            <div>
                                <div style="font-size:0.72rem;color:#047857;text-transform:uppercase;
                                            letter-spacing:.06em;font-weight:600;">Taux</div>
                                <div style="font-size:1.6rem;font-weight:800;color:#065f46;line-height:1.1;">
                                    {round(score_brut_e/total_e*100)}%
                                </div>
                            </div>
                        </div>
                        <div style="margin-top:0.8rem;font-size:0.78rem;color:#047857;opacity:0.8;">
                            ℹ️ La correction détaillée est réservée à votre Conseiller d'Orientation.
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="alert-warning">📝 {nb_rep_e}/{total_e} questions répondues.</div>', unsafe_allow_html=True)

        # Vérifier si tous les tests sont complétés
        tous_complets = all(
            sum(1 for v in st.session_state.get(f"eleve_answers_{k}", {}).values() if v) >= len(TESTS_QUESTIONS[k]["questions"])
            for k in test_keys_e
        )
        if tous_complets:
            st.markdown("""
            <div class="alert-success">
                🎉 <strong>Tous les tests sont complétés !</strong><br>
                Vos réponses ont été enregistrées. Signalez à votre Conseiller d'Orientation
                que vous avez terminé votre passation.
            </div>""", unsafe_allow_html=True)

    st.write("")
    m1, m2 = st.columns(2)
    m1.metric("🔬 Aptitude Scientifique (SA)", f"{sa_d:.2f}/20",
              help="SA = (KRX étalonnée + D48 étalonnée) / 2")
    m2.metric("📖 Aptitude Littéraire (LA)",   f"{la_d:.2f}/20",
              help="LA = (BV11 étalonnée + PRC étalonnée) / 2")

    # Graphique radar des aptitudes
    if sa_d > 0 or la_d > 0:
        labels_r = ["Logique (D48)", "Maths (KRX)", "Mécanique (MECA)", "Littéraire (BV11)", "Proverbes (PRC)"]
        vals_r   = [float(d.get("d48_etal") or d.get("d48_brut") or 0),
                    float(d.get("krx_etal") or d.get("krx_brut") or 0),
                    float(d.get("meca_etal") or d.get("meca_brut") or 0),
                    float(d.get("bv11_etal") or d.get("bv11_brut") or 0),
                    float(d.get("prc_etal") or d.get("prc_brut") or 0)]
        fig_e = go.Figure(go.Scatterpolar(
            r=vals_r, theta=labels_r, fill="toself",
            name="Tes aptitudes",
            line_color="#7c3aed", fillcolor="rgba(124,58,237,0.15)"
        ))
        fig_e.update_layout(
            polar=dict(radialaxis=dict(visible=True, range=[0,20]), gridshape="linear"),
            height=280, showlegend=False,
            margin=dict(l=40,r=40,t=20,b=20)
        )
        st.plotly_chart(fig_e, use_container_width=True)

    # ── Section 2 : Notes scolaires (lecture seule) ──
    st.markdown("---")
    st.markdown("### 📚 Notes Scolaires")
    trim_disp = d.get("trimestre_actuel") or d.get("trimestre_decision","T1")
    st.caption(f"Trimestre de référence : {trim_disp}")
    t_k = trim_disp.lower()
    n_cols = st.columns(6)
    matieres = [
        ("Maths",    f"maths_{t_k}",   "#10b981"),
        ("Sci.Phy.", f"sci_phy_{t_k}", "#10b981"),
        ("SVT",      f"svt_{t_k}",     "#10b981"),
        ("Français", f"francais_{t_k}","#3b82f6"),
        ("Hist-Géo", f"histgeo_{t_k}", "#3b82f6"),
        ("Anglais",  f"anglais_{t_k}", "#3b82f6"),
    ]
    for col, (lbl, key, clr) in zip(n_cols, matieres):
        val = float(d.get(key) or 0)
        col.markdown(f"""
        <div style="text-align:center;background:#f8fafc;border-radius:10px;
                    padding:0.6rem 0.3rem;border:1px solid #e2e8f0;">
            <div style="font-size:0.7rem;opacity:0.6;">{lbl}</div>
            <div style="font-size:1.3rem;font-weight:700;color:{clr};">{val:.1f}</div>
        </div>""", unsafe_allow_html=True)

    # ── Section 3 : Conversation avec l'agent IA ──
    ia_conv = d.get("chat_ia_synthese","") or ""
    if ia_conv:
        st.markdown("---")
        st.markdown("### 💬 Synthèse de l'Entretien avec l'Agent IA")
        st.markdown("""
        <div class="alert-info" style="font-size:0.85rem;">
            ℹ️ Voici le résumé de la discussion avec le conseiller IA lors de ton orientation.
        </div>""", unsafe_allow_html=True)
        st.markdown(f"""
        <div class="chat-ia">
            🤖 <strong>Agent CapAvenir :</strong><br>{ia_conv}
        </div>""", unsafe_allow_html=True)

    # ── Section 3bis : Agent IA interactif en cas de conflit ──
    # Détecter un conflit potentiel depuis les données du dossier
    sa_e = float(d.get("SA_etal") or d.get("SA_brut") or 0)
    la_e = float(d.get("LA_etal") or d.get("LA_brut") or 0)
    serie_e = (d.get("serie_finale") or d.get("serie_provisoire") or d.get("serie_cible") or "?")
    choix_e = d.get("choix_personnel","") or ""
    moy_sci_e = float(d.get("moy_sci_t1") or d.get("moy_sci_t2") or 0)
    moy_lit_e = float(d.get("moy_lit_t1") or d.get("moy_lit_t2") or 0)
    proj_e    = d.get("projet_pro","") or ""
    revenu_e  = d.get("revenu_famille","") or ""

    conflit_e = None
    if sa_e > la_e and "A" in choix_e:
        conflit_e = "decale"
    elif sa_e > la_e and moy_sci_e > 0 and moy_sci_e < 10 and "C" in choix_e:
        conflit_e = "reveur"
    elif la_e > sa_e and moy_lit_e > 0 and moy_lit_e < 10 and "A" in choix_e:
        conflit_e = "reveur"

    if conflit_e:
        st.markdown("---")
        st.markdown("### 🤖 Discussion avec l'Agent d'Orientation")
        st.markdown("""
        <div class="alert-info" style="font-size:0.85rem;">
            💬 Un <strong>conflit</strong> a été détecté entre ton profil et ton choix d'orientation.
            L'Agent CapAvenir peut t'aider à y voir plus clair.
        </div>""", unsafe_allow_html=True)

        # Historique de chat propre à l'espace élève
        chat_key = f"eleve_chat_{nom_d}_{prenom_d}"
        if chat_key not in st.session_state:
            st.session_state[chat_key] = []

        if conflit_e == "decale":
            msg_init_e = (f"Bonjour {prenom_d} ! D'après ton dossier, tes aptitudes scientifiques "
                          f"(SA={sa_e:.1f}/20) sont bien au-dessus de tes aptitudes littéraires (LA={la_e:.1f}/20), "
                          f"mais tu as exprimé le souhait d'aller en série A. "
                          f"Je suis là pour discuter avec toi de ce choix. "
                          f"Qu'est-ce qui t'attire dans la série A ?")
        else:
            msg_init_e = (f"Bonjour {prenom_d} ! Tu as de bonnes aptitudes scientifiques "
                          f"(SA={sa_e:.1f}/20), mais tes résultats scolaires "
                          f"({moy_sci_e:.1f}/20) sont encore en dessous du seuil. "
                          f"Je voudrais comprendre ce qui se passe. "
                          f"Qu'est-ce qui te pose le plus de difficultés en ce moment ?")

        if not st.session_state[chat_key]:
            st.session_state[chat_key].append({"role": "ia", "content": msg_init_e})

        # Afficher l'historique
        for msg_e in st.session_state[chat_key]:
            if msg_e["role"] == "ia":
                st.markdown(f'<div class="chat-ia">🤖 <strong>Agent CapAvenir :</strong><br>{msg_e["content"]}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="chat-user">{msg_e["content"]}<br><strong>— {prenom_d}</strong></div>', unsafe_allow_html=True)

        user_input_e = st.text_input("Votre réponse :", placeholder="Écrivez ici...",
                                      key=f"chat_eleve_input_{nom_d}")
        if st.button("Envoyer 📨", use_container_width=True, key=f"send_eleve_{nom_d}"):
            if user_input_e.strip():
                st.session_state[chat_key].append({"role": "user", "content": user_input_e})
                api_key = os.environ.get("ANTHROPIC_API_KEY", "")
                if api_key and len(api_key) > 20:
                    sys_e = (
                        f"Tu es un conseiller d'orientation scolaire IA chaleureux pour le système camerounais. "
                        f"Tu parles à {prenom_d} {nom_d}, élève au lycée. "
                        f"SA={sa_e:.1f}/20, LA={la_e:.1f}/20, Moy.Sci={moy_sci_e:.1f}/20. "
                        f"Choix : {choix_e}. Projet : {proj_e or 'non précisé'}. "
                        f"Conflit : {conflit_e}. Réponds en 2-3 phrases, empathique, naturel, en français."
                    )
                    msgs_e = [
                        {"role": "assistant" if m["role"]=="ia" else "user", "content": m["content"]}
                        for m in st.session_state[chat_key]
                    ]
                    try:
                        resp_e = requests.post(
                            "https://api.anthropic.com/v1/messages",
                            headers={"Content-Type":"application/json",
                                     "x-api-key":api_key,
                                     "anthropic-version":"2023-06-01"},
                            json={"model":"claude-sonnet-4-20250514",
                                  "max_tokens":300,"system":sys_e,"messages":msgs_e},
                            timeout=15,
                        )
                        ia_r_e = resp_e.json()["content"][0]["text"]
                    except Exception:
                        ia_r_e = reponse_ia_simulee(
                            user_input_e, conflit_e, prenom_d, sa_e, moy_sci_e,
                            LA=la_e, moy_lit=moy_lit_e,
                            projet_pro=proj_e, revenu=revenu_e,
                            chat_history=st.session_state[chat_key]
                        )
                else:
                    ia_r_e = reponse_ia_simulee(
                        user_input_e, conflit_e, prenom_d, sa_e, moy_sci_e,
                        LA=la_e, moy_lit=moy_lit_e,
                        projet_pro=proj_e, revenu=revenu_e,
                        chat_history=st.session_state[chat_key]
                    )
                st.session_state[chat_key].append({"role": "ia", "content": ia_r_e})
                rerun()

    # ── Section 4 : Observations du conseiller ──
    obs = d.get("notes_conseiller","") or ""
    if obs:
        st.markdown("---")
        st.markdown("### 📝 Observations du Conseiller d'Orientation")
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#fffbeb,#fef3c7);
                    border-left:5px solid #f59e0b;border-radius:12px;
                    padding:1.2rem 1.5rem;color:#78350f;">
            <div style="font-size:0.75rem;font-weight:700;text-transform:uppercase;
                        letter-spacing:0.08em;margin-bottom:0.5rem;opacity:0.7;">
                Observation officielle
            </div>
            {obs}
        </div>""", unsafe_allow_html=True)

    # ── Section 5 : Résultat d'orientation ──
    st.markdown("---")
    st.markdown("### 🎓 Ton Orientation")
    badge_color = "#10b981" if serie_d=="C" else "#3b82f6" if serie_d=="A" else "#f59e0b"
    sc = int(d.get("score_confiance",0) or 0)
    conf_color = "#10b981" if sc>=70 else "#f59e0b" if sc>=50 else "#ef4444"
    st.markdown(f"""
    <div style="text-align:center;padding:2rem;background:{"#ecfdf5" if serie_d=="C" else "#eff6ff" if serie_d=="A" else "#fffbeb"};
                border-radius:16px;border:2px solid {badge_color};">
        <div style="font-size:0.82rem;color:#64748b;text-transform:uppercase;
                    letter-spacing:0.08em;margin-bottom:0.5rem;">Orientation recommandée</div>
        <span style="background:{badge_color};color:white;padding:0.5rem 2.5rem;
                     border-radius:24px;font-weight:800;font-size:1.8rem;">
            SÉRIE {serie_d}
        </span>
        <div style="margin-top:1rem;font-size:0.85rem;color:{conf_color};font-weight:600;">
            Score de confiance : {sc} %
        </div>
    </div>""", unsafe_allow_html=True)

    st.markdown("")
    st.info("📌 Pour toute question sur ton orientation, contacte ton Conseiller d'Orientation au lycée.")

# =====================================================================
# PAGE D'ACCUEIL — Sélection et inscription élève / conseiller
# =====================================================================
def afficher_page_accueil():
    """Page d'accueil avec sélection du mode et inscription."""

    dark = st.session_state.get("dark_mode", False)
    txt_color  = "#e2e8f0" if dark else "#1e1b4b"
    card_bg    = "linear-gradient(145deg,#1e293b,#141b2d)" if dark else "linear-gradient(145deg,#ffffff,#f8fafc)"
    card_brd   = "#334155" if dark else "#e5e7eb"

    sous_mode = st.session_state.get("accueil_sous_mode")

    # ── ÉCRAN DE CHOIX ──────────────────────────────────────────────
    if sous_mode is None:
        st.markdown(f"""
        <div style="text-align:center; padding:2rem 1rem 1.5rem;">
            <div style="font-size:1.1rem; color:{txt_color}; opacity:0.7; font-weight:500;">
                Bienvenue sur CapAvenir CMR — Système Intelligent d'Orientation Scolaire
            </div>
            <div style="font-size:0.9rem; margin-top:0.4rem; opacity:0.5; color:{txt_color};">
                Choisissez votre profil pour commencer
            </div>
        </div>
        """, unsafe_allow_html=True)

        col_gap, col1, col_mid, col2, col_gap2 = st.columns([0.5, 3, 0.8, 3, 0.5])

        with col1:
            st.markdown(f"""
            <div style="
                background:{card_bg};
                border: 2px solid #10b981;
                border-radius: 24px;
                padding: 2.5rem 1.5rem;
                text-align: center;
                box-shadow: 0 8px 32px rgba(16,185,129,0.18);
                transition: all 0.3s ease;
                cursor: pointer;
            ">
                <div style="font-size: 3.5rem; margin-bottom:0.8rem;">🎒</div>
                <div style="font-size: 1.4rem; font-weight: 800; color: #10b981; margin-bottom:0.5rem;">
                    Je suis Élève
                </div>
                <div style="font-size: 0.85rem; color:{txt_color}; opacity:0.7; line-height:1.5;">
                    Accède à ton dossier d'orientation, passe les tests psychotechniques
                    et consulte tes résultats.
                </div>
                <div style="margin-top:1.2rem; background:#10b981; color:white; padding:0.55rem 1.5rem;
                            border-radius:12px; font-weight:700; display:inline-block; font-size:0.9rem;">
                    → Continuer
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            if st.button("🎒 Accès Élève", use_container_width=True, type="primary", key="btn_mode_eleve"):
                st.session_state.mode_accueil = "eleve"
                st.session_state.accueil_sous_mode = "choix_eleve"
                rerun()

        with col2:
            st.markdown(f"""
            <div style="
                background:{card_bg};
                border: 2px solid #6d28d9;
                border-radius: 24px;
                padding: 2.5rem 1.5rem;
                text-align: center;
                box-shadow: 0 8px 32px rgba(109,40,217,0.18);
                transition: all 0.3s ease;
                cursor: pointer;
            ">
                <div style="font-size: 3.5rem; margin-bottom:0.8rem;">📋</div>
                <div style="font-size: 1.4rem; font-weight: 800; color: #6d28d9; margin-bottom:0.5rem;">
                    Je suis Conseiller
                </div>
                <div style="font-size: 0.85rem; color:{txt_color}; opacity:0.7; line-height:1.5;">
                    Gérez les dossiers d'orientation, saisissez les notes et consultez
                    le tableau de bord.
                </div>
                <div style="margin-top:1.2rem; background:#6d28d9; color:white; padding:0.55rem 1.5rem;
                            border-radius:12px; font-weight:700; display:inline-block; font-size:0.9rem;">
                    → Continuer
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            if st.button("📋 Accès Conseiller", use_container_width=True, key="btn_mode_conseiller"):
                st.session_state.mode_accueil = "conseiller"
                st.session_state.accueil_sous_mode = "choix_conseiller"
                rerun()

        st.write("")
        st.markdown(f"""
        <div style="text-align:center; font-size:0.78rem; color:{txt_color}; opacity:0.4; margin-top:1rem;">
            CapAvenir CMR v2.5 · Mémoire ENS Informatique Niv. 5 · Orientation 3ᵉ → 2nde Cameroun
        </div>
        """, unsafe_allow_html=True)
        return

    # ── CHOIX ÉLÈVE : Nouveau ou Dossier existant ───────────────────
    if sous_mode == "choix_eleve":
        if st.button("← Retour", key="retour_choix_eleve"):
            st.session_state.accueil_sous_mode = None
            st.session_state.mode_accueil = "accueil"
            rerun()

        st.markdown(f"""
        <div style="text-align:center; padding:1.2rem 0 0.5rem;">
            <div style="font-size:1.6rem; font-weight:800; color:#10b981;">🎒 Espace Élève</div>
            <div style="font-size:0.88rem; opacity:0.65; color:{txt_color}; margin-top:0.3rem;">
                Comment souhaitez-vous accéder à votre espace ?
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.write("")
        c1, c2 = st.columns(2)
        with c1:
            st.markdown(f"""
            <div style="background:{card_bg}; border:2px solid #10b981; border-radius:18px;
                        padding:2rem 1.2rem; text-align:center;
                        box-shadow:0 4px 16px rgba(16,185,129,0.15);">
                <div style="font-size:2.4rem;">✨</div>
                <div style="font-size:1.1rem; font-weight:700; color:#10b981; margin:0.4rem 0;">
                    Nouvelle Inscription
                </div>
                <div style="font-size:0.8rem; color:{txt_color}; opacity:0.65; line-height:1.5;">
                    Première visite ? Créez votre profil et commencez les tests psychotechniques.
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            if st.button("✨ M'inscrire", use_container_width=True, type="primary", key="btn_inscription_eleve"):
                st.session_state.accueil_sous_mode = "inscription_eleve"
                rerun()

        with c2:
            st.markdown(f"""
            <div style="background:{card_bg}; border:2px solid #3b82f6; border-radius:18px;
                        padding:2rem 1.2rem; text-align:center;
                        box-shadow:0 4px 16px rgba(59,130,246,0.15);">
                <div style="font-size:2.4rem;">🔍</div>
                <div style="font-size:1.1rem; font-weight:700; color:#3b82f6; margin:0.4rem 0;">
                    Accéder à mon Dossier
                </div>
                <div style="font-size:0.8rem; color:{txt_color}; opacity:0.65; line-height:1.5;">
                    Déjà inscrit ? Retrouvez votre dossier et consultez vos résultats d'orientation.
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            if st.button("🔍 Mon Dossier", use_container_width=True, key="btn_dossier_eleve"):
                st.session_state.accueil_sous_mode = "dossier_eleve"
                rerun()
        return

    # ── INSCRIPTION ÉLÈVE ───────────────────────────────────────────
    if sous_mode == "inscription_eleve":
        if st.button("← Retour", key="retour_inscription"):
            st.session_state.accueil_sous_mode = "choix_eleve"
            rerun()

        st.markdown(f"""
        <div style="text-align:center; padding:0.8rem 0 1.2rem;">
            <div style="font-size:1.5rem; font-weight:800; color:#10b981;">✨ Nouvelle Inscription — Élève</div>
            <div style="font-size:0.85rem; opacity:0.6; color:{txt_color}; margin-top:0.3rem;">
                Renseignez vos informations pour créer votre profil d'orientation
            </div>
        </div>
        """, unsafe_allow_html=True)

        with st.form("form_inscription_eleve", clear_on_submit=False):
            st.markdown("#### 👤 Identité")
            r1c1, r1c2 = st.columns(2)
            with r1c1:
                ins_nom    = st.text_input("Nom de famille *", placeholder="Ex : MBALLA")
                ins_age    = st.number_input("Âge *", min_value=12, max_value=20, value=15)
                ins_lycee  = st.text_input("Lycée / Établissement *", placeholder="Ex : Lycée de Douala")
            with r1c2:
                ins_prenom = st.text_input("Prénom *", placeholder="Ex : Jean")
                ins_sexe   = st.selectbox("Sexe *", ["Masculin", "Féminin"])

            st.markdown("#### 🔑 Identifiants de connexion")
            p1, p2 = st.columns(2)
            with p1:
                ins_pwd  = st.text_input("Mot de passe *", type="password",
                                          placeholder="Choisissez un mot de passe sécurisé",
                                          help="Ce mot de passe vous permettra de vous reconnecter pour consulter votre dossier.")
            with p2:
                ins_pwd2 = st.text_input("Confirmer le mot de passe *", type="password",
                                          placeholder="Répétez le même mot de passe")

            st.markdown("#### 🎯 Projet & Contexte")
            r2c1, r2c2 = st.columns(2)
            with r2c1:
                ins_choix = st.selectbox(
                    "Série souhaitée",
                    ["C (Scientifique)", "A (Littéraire)", "Indécis(e)"]
                )
                ins_projet = st.text_area(
                    "Projet professionnel",
                    placeholder="Ex: Devenir médecin, ingénieur, avocat...",
                    height=80,
                )
            with r2c2:
                ins_revenu = st.selectbox(
                    "Revenu mensuel estimé des parents",
                    ["Faible (< 50 000 FCFA/mois)",
                     "Moyen (50 000 - 150 000 FCFA/mois)",
                     "Élevé (> 150 000 FCFA/mois)"]
                )
                st.markdown("""
                <div style="background:linear-gradient(135deg,#fdf4ff,#ede9fe);
                            border-left:4px solid #a855f7; border-radius:10px;
                            padding:0.8rem 1rem; margin-top:0.5rem; font-size:0.8rem; color:#581c87;">
                    📌 <strong>Note :</strong> Votre dossier sera examiné par un Conseiller
                    d'Orientation qui saisira vos notes et passera vos tests avec vous.
                </div>
                """, unsafe_allow_html=True)

            st.write("")
            submitted_ins = st.form_submit_button(
                "✅ Créer mon profil et accéder à mes tests",
                use_container_width=True, type="primary"
            )

        if submitted_ins:
            if not ins_nom.strip() or not ins_prenom.strip() or not ins_lycee.strip():
                st.error("⚠️ Veuillez renseigner votre nom, prénom et lycée.")
            elif not ins_pwd:
                st.error("⚠️ Veuillez choisir un mot de passe.")
            elif len(ins_pwd) < 6:
                st.error("⚠️ Le mot de passe doit contenir au moins 6 caractères.")
            elif ins_pwd != ins_pwd2:
                st.error("⚠️ Les deux mots de passe ne correspondent pas.")
            else:
                # Enregistrement des identifiants en base
                _save_eleve_credential(ins_nom.strip(), ins_prenom.strip(),
                                       ins_lycee.strip(), ins_pwd)
                # Pré-remplissage du session_state avec les données d'inscription
                st.session_state.nom              = ins_nom.strip().upper()
                st.session_state.prenom           = ins_prenom.strip().capitalize()
                st.session_state.age              = int(ins_age)
                st.session_state.sexe             = ins_sexe
                st.session_state.lycee            = ins_lycee.strip()
                st.session_state.choix_personnel  = ins_choix
                st.session_state.projet_pro       = ins_projet.strip()
                st.session_state.revenu           = ins_revenu
                st.session_state.eleve_inscrit    = True
                st.session_state.mode_accueil     = "eleve"
                st.session_state.accueil_sous_mode= "espace_tests"
                rerun()
        return

    # ── ESPACE TESTS APRÈS INSCRIPTION ─────────────────────────────
    if sous_mode == "espace_tests":
        prenom_e = st.session_state.get("prenom", "")
        nom_e    = st.session_state.get("nom", "")

        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#ecfdf5,#d1fae5);
                    border-left:5px solid #10b981; border-radius:16px;
                    padding:1.5rem 2rem; margin-bottom:1.5rem;">
            <div style="font-size:1.3rem; font-weight:800; color:#065f46;">
                🎉 Bienvenue, {prenom_e} {nom_e} !
            </div>
            <div style="color:#047857; font-size:0.88rem; margin-top:0.5rem; line-height:1.6;">
                Votre profil a été créé. Voici la suite de votre parcours d'orientation :
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("#### 📋 Votre parcours d'orientation")
        e1, e2, e3 = st.columns(3)
        with e1:
            st.markdown("""
            <div style="background:linear-gradient(135deg,#d1fae5,#a7f3d0);
                        border-radius:14px; padding:1.2rem; text-align:center;">
                <div style="font-size:1.8rem;">🧪</div>
                <div style="font-weight:700; color:#065f46; margin:0.3rem 0;">Tests Psychotechniques</div>
                <div style="font-size:0.78rem; color:#047857;">
                    Passez les 5 tests (D48, KRX, MECA, BV11, PRC) avec votre conseiller.
                </div>
            </div>
            """, unsafe_allow_html=True)
        with e2:
            st.markdown("""
            <div style="background:linear-gradient(135deg,#dbeafe,#bfdbfe);
                        border-radius:14px; padding:1.2rem; text-align:center;">
                <div style="font-size:1.8rem;">📚</div>
                <div style="font-weight:700; color:#1e3a8a; margin:0.3rem 0;">Notes Scolaires</div>
                <div style="font-size:0.78rem; color:#1e40af;">
                    Votre conseiller saisira vos notes des trimestres dans votre dossier.
                </div>
            </div>
            """, unsafe_allow_html=True)
        with e3:
            st.markdown("""
            <div style="background:linear-gradient(135deg,#f3e8ff,#ddd6fe);
                        border-radius:14px; padding:1.2rem; text-align:center;">
                <div style="font-size:1.8rem;">🎓</div>
                <div style="font-weight:700; color:#4c1d95; margin:0.3rem 0;">Orientation IA</div>
                <div style="font-size:0.78rem; color:#5b21b6;">
                    Recevez votre orientation personnalisée : Série C ou Série A.
                </div>
            </div>
            """, unsafe_allow_html=True)

        st.write("")

        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("🧪 Passer les tests psychotechniques maintenant",
                         use_container_width=True, type="primary", key="btn_goto_tests"):
                # Redirection vers l'espace élève (passation des tests)
                st.session_state.show_espace_eleve = True
                st.session_state.eleve_dossier_actif = {
                    "nom": st.session_state.nom,
                    "prenom": st.session_state.prenom,
                    "lycee": st.session_state.lycee,
                    "age": st.session_state.age,
                    "sexe": st.session_state.sexe,
                    "choix_personnel": st.session_state.choix_personnel,
                    "projet_pro": st.session_state.projet_pro,
                    "revenu_famille": st.session_state.revenu,
                    "serie_finale": None, "serie_provisoire": None, "serie_cible": None,
                    "statut": "en cours", "score_confiance": 0,
                    "SA_etal": 0, "LA_etal": 0,
                    "trimestre_actuel": "T1",
                }
                st.session_state.mode_accueil = "skip"
                rerun()
        with col_b:
            if st.button("🔍 Consulter mon dossier (si déjà créé)",
                         use_container_width=True, key="btn_voir_dossier_apres_ins"):
                st.session_state.accueil_sous_mode = "dossier_eleve"
                rerun()

        st.write("")
        st.markdown("""
        <div style="background:linear-gradient(135deg,#fff7ed,#fef3c7);
                    border-left:4px solid #f59e0b; border-radius:12px;
                    padding:1rem 1.2rem; font-size:0.83rem; color:#78350f;">
            💡 <strong>Conseil :</strong> Rendez-vous auprès de votre Conseiller d'Orientation
            au lycée pour effectuer la passation complète des tests et la saisie de vos notes.
            Votre dossier sera ensuite disponible depuis cet espace.
        </div>
        """, unsafe_allow_html=True)
        return

    # ── CONSULTER UN DOSSIER EXISTANT (depuis accueil élève) ────────
    if sous_mode == "dossier_eleve":
        if st.button("← Retour", key="retour_dossier_accueil"):
            st.session_state.accueil_sous_mode = "choix_eleve"
            rerun()

        st.markdown(f"""
        <div style="text-align:center; padding:0.8rem 0 1.2rem;">
            <div style="font-size:1.5rem; font-weight:800; color:#3b82f6;">🔍 Connexion Élève</div>
            <div style="font-size:0.85rem; opacity:0.6; color:{txt_color}; margin-top:0.3rem;">
                Connectez-vous avec les identifiants créés lors de votre inscription
            </div>
        </div>
        """, unsafe_allow_html=True)

        # Vérifier si l'élève a un compte avant tout
        with st.form("form_login_eleve"):
            fl1, fl2 = st.columns(2)
            with fl1:
                nom_search    = st.text_input("Nom de famille *", placeholder="Ex : MBALLA")
                prenom_search = st.text_input("Prénom *", placeholder="Ex : Jean")
            with fl2:
                pwd_eleve = st.text_input("Mot de passe *", type="password",
                                           placeholder="Votre mot de passe d'inscription")
                st.markdown("""
                <div style="background:linear-gradient(135deg,#eff6ff,#dbeafe);
                            border-left:4px solid #3b82f6; border-radius:10px;
                            padding:0.7rem 1rem; margin-top:0.5rem; font-size:0.78rem; color:#1e3a8a;">
                    💡 Vous devez d'abord vous <strong>inscrire</strong> pour obtenir un mot de passe.
                </div>
                """, unsafe_allow_html=True)
            chercher_acc = st.form_submit_button("🔓 Se connecter", use_container_width=True, type="primary")

        if chercher_acc:
            if not nom_search.strip() or not prenom_search.strip():
                st.error("⚠️ Veuillez renseigner votre nom et prénom.")
            elif not pwd_eleve:
                st.error("⚠️ Veuillez saisir votre mot de passe.")
            elif not _eleve_exists(nom_search.strip(), prenom_search.strip()):
                st.markdown("""
                <div style="background:linear-gradient(135deg,#fff7ed,#fed7aa);
                            border-left:5px solid #f97316; border-radius:12px;
                            padding:1rem 1.2rem; margin-top:0.5rem;">
                    <strong>⚠️ Aucun compte trouvé pour ce nom.</strong><br>
                    <span style="font-size:0.85rem;">
                    Vous n'êtes pas encore inscrit(e) sur CapAvenir CMR.<br>
                    Cliquez sur <strong>← Retour</strong> puis choisissez <strong>✨ Nouvelle Inscription</strong>.
                    </span>
                </div>
                """, unsafe_allow_html=True)
            elif not _verify_eleve_credential(nom_search.strip(), prenom_search.strip(), pwd_eleve):
                st.error("❌ Mot de passe incorrect.")
            else:
                # Authentification réussie → rechercher le dossier en BDD
                try:
                    resultats = db.rechercher_eleve(nom_search.strip(), prenom_search.strip())
                except Exception:
                    resultats = []
                if not resultats:
                    # Pas encore de dossier en BDD, créer un dossier minimal
                    st.session_state.nom    = nom_search.strip().upper()
                    st.session_state.prenom = prenom_search.strip().capitalize()
                    st.session_state.eleve_dossier_actif = {
                        "nom": nom_search.strip().upper(),
                        "prenom": prenom_search.strip().capitalize(),
                        "lycee": "", "statut": "en cours",
                        "serie_finale": None, "serie_provisoire": None, "serie_cible": None,
                        "SA_etal": 0, "LA_etal": 0, "score_confiance": 0,
                        "trimestre_actuel": "T1",
                    }
                    st.session_state.show_espace_eleve = True
                    st.session_state.mode_accueil      = "skip"
                    rerun()
                else:
                    st.session_state.eleve_dossier_actif = resultats[0]
                    st.session_state.show_espace_eleve   = True
                    st.session_state.mode_accueil        = "skip"
                    rerun()
        return

    # ── CONSEILLER : Choix connexion ou inscription ──────────────────
    if sous_mode == "choix_conseiller":
        if st.button("← Retour", key="retour_choix_cons"):
            st.session_state.accueil_sous_mode = None
            st.session_state.mode_accueil = "accueil"
            rerun()

        st.markdown(f"""
        <div style="text-align:center; padding:0.8rem 0 1.2rem;">
            <div style="font-size:1.5rem; font-weight:800; color:#6d28d9;">📋 Espace Conseiller</div>
            <div style="font-size:0.85rem; opacity:0.6; color:{txt_color}; margin-top:0.3rem;">
                Conseiller d'Orientation — Accès réservé au personnel autorisé
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.write("")
        cc1, cc2 = st.columns(2)
        with cc1:
            st.markdown(f"""
            <div style="background:{card_bg}; border:2px solid #6d28d9; border-radius:18px;
                        padding:2rem 1.2rem; text-align:center;
                        box-shadow:0 4px 16px rgba(109,40,217,0.15);">
                <div style="font-size:2.4rem;">🔐</div>
                <div style="font-size:1.1rem; font-weight:700; color:#6d28d9; margin:0.4rem 0;">
                    Se Connecter
                </div>
                <div style="font-size:0.8rem; color:{txt_color}; opacity:0.65; line-height:1.5;">
                    Vous avez déjà vos identifiants ? Connectez-vous pour accéder au tableau de bord.
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            if st.button("🔐 Se connecter", use_container_width=True, type="primary", key="btn_login_cons"):
                st.session_state.accueil_sous_mode = "login_conseiller"
                rerun()
        with cc2:
            st.markdown(f"""
            <div style="background:{card_bg}; border:2px solid #0ea5e9; border-radius:18px;
                        padding:2rem 1.2rem; text-align:center;
                        box-shadow:0 4px 16px rgba(14,165,233,0.15);">
                <div style="font-size:2.4rem;">📝</div>
                <div style="font-size:1.1rem; font-weight:700; color:#0ea5e9; margin:0.4rem 0;">
                    Demander un Accès
                </div>
                <div style="font-size:0.8rem; color:{txt_color}; opacity:0.65; line-height:1.5;">
                    Nouveau conseiller ? Soumettez une demande d'accès auprès de l'administrateur.
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            if st.button("📝 Demander un accès", use_container_width=True, key="btn_demande_cons"):
                st.session_state.accueil_sous_mode = "demande_conseiller"
                rerun()
        return

    # ── CONNEXION CONSEILLER ────────────────────────────────────────
    if sous_mode == "login_conseiller":
        if st.button("← Retour", key="retour_login_cons"):
            st.session_state.accueil_sous_mode = "choix_conseiller"
            rerun()

        st.markdown(f"""
        <div style="text-align:center; padding:0.8rem 0 1.5rem;">
            <div style="font-size:1.5rem; font-weight:800; color:#6d28d9;">🔐 Connexion Conseiller</div>
        </div>
        """, unsafe_allow_html=True)

        _, form_col, _ = st.columns([1, 2, 1])
        with form_col:
            with st.form("form_login_conseiller_accueil"):
                login_cons  = st.text_input("👤 Identifiant", placeholder="Votre identifiant")
                pwd_cons    = st.text_input("🔑 Mot de passe", type="password", placeholder="••••••••")
                btn_login   = st.form_submit_button("🔓 Se connecter", use_container_width=True, type="primary")

            if btn_login:
                if check_password(login_cons, pwd_cons):
                    st.session_state.conseiller_auth    = True
                    st.session_state.mode_conseiller_val= True
                    st.session_state.login_error        = ""
                    st.session_state.mode_accueil       = "skip"
                    st.session_state.accueil_sous_mode  = None
                    st.session_state.show_dashboard     = True
                    rerun()
                else:
                    st.error("❌ Identifiant ou mot de passe incorrect.")

        return

    # ── DEMANDE D'ACCÈS CONSEILLER ──────────────────────────────────
    if sous_mode == "demande_conseiller":
        if st.button("← Retour", key="retour_demande_cons"):
            st.session_state.accueil_sous_mode = "choix_conseiller"
            rerun()

        st.markdown(f"""
        <div style="text-align:center; padding:0.8rem 0 1.2rem;">
            <div style="font-size:1.5rem; font-weight:800; color:#0ea5e9;">📝 Demande d'Accès Conseiller</div>
            <div style="font-size:0.85rem; opacity:0.6; color:{txt_color}; margin-top:0.3rem;">
                Renseignez vos informations — un administrateur vous contactera pour valider votre accès
            </div>
        </div>
        """, unsafe_allow_html=True)

        with st.form("form_demande_conseiller"):
            d1, d2 = st.columns(2)
            with d1:
                d_nom      = st.text_input("Nom *", placeholder="FOMO")
                d_lycee    = st.text_input("Lycée / Établissement *", placeholder="Lycée de Douala")
                d_email    = st.text_input("Email professionnel *", placeholder="contact@lycee.cm")
            with d2:
                d_prenom   = st.text_input("Prénom *", placeholder="Alice")
                d_poste    = st.text_input("Poste occupé *", placeholder="Conseiller d'Orientation Principal")
                d_tel      = st.text_input("Téléphone", placeholder="+237 6XX XXX XXX")
            d_motif = st.text_area(
                "Motif de la demande",
                placeholder="Expliquez brièvement pourquoi vous souhaitez accéder à CapAvenir CMR...",
                height=80
            )
            btn_dem = st.form_submit_button("📨 Envoyer la demande", use_container_width=True, type="primary")

        if btn_dem:
            if not d_nom.strip() or not d_prenom.strip() or not d_lycee.strip() or not d_email.strip():
                st.error("⚠️ Veuillez remplir tous les champs obligatoires (*).")
            else:
                st.markdown(f"""
                <div style="background:linear-gradient(135deg,#ecfdf5,#d1fae5);
                            border-left:5px solid #10b981; border-radius:14px;
                            padding:1.5rem 2rem; margin-top:1rem;">
                    <div style="font-size:1.1rem; font-weight:700; color:#065f46; margin-bottom:0.5rem;">
                        ✅ Demande envoyée avec succès !
                    </div>
                    <div style="color:#047857; font-size:0.88rem; line-height:1.6;">
                        Votre demande d'accès pour <strong>{d_prenom.strip()} {d_nom.strip().upper()}</strong>
                        ({d_lycee.strip()}) a bien été enregistrée.<br>
                        L'administrateur système vous contactera à <strong>{d_email.strip()}</strong>
                        dans les <strong>48h ouvrables</strong> pour vous communiquer vos identifiants.
                    </div>
                </div>
                """, unsafe_allow_html=True)
        return


# =====================================================================
# NAVIGATION — Dashboard ou Orientation
# =====================================================================

# ── Affichage page de connexion si besoin ──
if st.session_state.get("show_login", False) and not st.session_state.get("conseiller_auth", False):
    afficher_login_conseiller()
    st.stop()
else:
    st.session_state.show_login = False

# ── Page d'accueil (mode_accueil != "skip") ──
if st.session_state.get("mode_accueil", "accueil") not in ("skip",):
    afficher_page_accueil()
    st.stop()

# ── Espace Élève ──
if st.session_state.get("show_espace_eleve", False):
    afficher_espace_eleve()
    st.stop()

# ── Tableau de bord conseiller ──
if mode_conseiller and st.session_state.get("show_dashboard", False):
    afficher_dashboard()
    st.stop()
else:
    st.session_state.show_dashboard = False

STEPS = ["Profil", "Notes", "Tests", "Diagnostic", "Fiche"]
ICONS = ["📋", "📚", "🧪", "🤖", "📄"]
step = st.session_state.step

cols = st.columns(len(STEPS))
for i, col in enumerate(cols):
    with col:
        label = f"{ICONS[i]} {STEPS[i]}"
        if i == step:
            st.markdown(f"<div class='step-active'>{label}</div>", unsafe_allow_html=True)
        elif i < step:
            st.markdown(f"<div class='step-done'>✓ {STEPS[i]}</div>", unsafe_allow_html=True)
        else:
            st.markdown(f"<div class='step-todo'>{label}</div>", unsafe_allow_html=True)

st.write("")

# =====================================================================
# ETAPE 0 — PROFIL ELEVE
# =====================================================================
if step == 0:
    st.subheader("📋 Informations personnelles de l'élève")

    c1, c2 = st.columns(2)
    with c1:
        st.session_state.nom   = st.text_input("Nom",   value=st.session_state.nom)
        st.session_state.age   = st.number_input("Âge", min_value=12, max_value=20, value=st.session_state.age)
        st.session_state.lycee = st.text_input("Lycée / Établissement", value=st.session_state.lycee)
    with c2:
        st.session_state.prenom = st.text_input("Prénom", value=st.session_state.prenom)
        st.session_state.sexe   = st.selectbox("Sexe", ["Masculin", "Féminin"])

    st.divider()
    st.subheader("🏠 Projet personnel & contexte familial")
    c1, c2 = st.columns(2)
    with c1:
        st.session_state.choix_personnel = st.selectbox(
            "Série souhaitée par l'élève",
            ["C (Scientifique)", "A (Littéraire)", "Indécis(e)"]
        )
        st.session_state.projet_pro = st.text_area(
            "Projet professionnel",
            value=st.session_state.projet_pro,
            placeholder="Ex: Devenir médecin, ingénieur, avocat...",
            height=90,
        )
    with c2:
        st.session_state.revenu = st.selectbox(
            "Revenu mensuel estimé des parents",
            ["Faible (< 50 000 FCFA/mois)",
             "Moyen (50 000 - 150 000 FCFA/mois)",
             "Élevé (> 150 000 FCFA/mois)"]
        )
        st.markdown('<div class="alert-info">💡 Le revenu permet d\'évaluer l\'accès aux cours de soutien et la faisabilité des études longues.</div>', unsafe_allow_html=True)

    st.write("")
    if st.button("Suivant : Notes scolaires ➡️", use_container_width=True, type="primary"):
        if st.session_state.nom.strip() and st.session_state.prenom.strip():
            st.session_state.step = 1
            auto_save_step(0)
            rerun()
        else:
            st.error("Veuillez renseigner le nom et le prénom de l'élève.")

# =====================================================================
# ETAPE 1 — NOTES SCOLAIRES (Multi-trimestre + graphique de progression)
# =====================================================================
elif step == 1:
    st.subheader(f"📚 Notes scolaires — {st.session_state.prenom} {st.session_state.nom}")

    tab_t1, tab_t2, tab_t3 = st.tabs(["📅 Trimestre 1", "📅 Trimestre 2", "📅 Trimestre 3"])

    def saisie_notes(suffix, key_prefix):
        # ── Sécurité : saisie des notes réservée au conseiller ──
        if not mode_conseiller:
            st.markdown("""
            <div class="alert-danger">
                🔒 <strong>Accès restreint</strong> — La saisie des notes scolaires est
                réservée au Conseiller d'Orientation.<br>
                Activez le <strong>Mode Conseiller</strong> dans la barre latérale pour
                accéder à cette fonctionnalité.
            </div>""", unsafe_allow_html=True)
            # Afficher les notes en lecture seule
            # S5 — None-safe : une note non saisie s'affiche "—" au lieu de planter
            def _fmt(v):
                return f"{float(v):.1f}" if v is not None else "—"
            def _fval(v):
                return float(v) if v is not None else 0.0

            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**Matières scientifiques**")
                for lbl, key in [("Mathématiques", f"maths{suffix}"),
                                  ("Sciences Physiques", f"sci_phy{suffix}"),
                                  ("SVT", f"svt{suffix}")]:
                    val = st.session_state.get(key)
                    st.markdown(f"**{lbl} :** `{_fmt(val)}/20`")
                moy = (_fval(st.session_state.get(f"maths{suffix}")) +
                       _fval(st.session_state.get(f"sci_phy{suffix}")) +
                       _fval(st.session_state.get(f"svt{suffix}"))) / 3
                st.metric("Moyenne scientifique", f"{moy:.2f} / 20")
            with c2:
                st.markdown("**Matières littéraires**")
                for lbl, key in [("Français", f"francais{suffix}"),
                                  ("Histoire-Géographie", f"histgeo{suffix}"),
                                  ("Anglais", f"anglais{suffix}")]:
                    val = st.session_state.get(key)
                    st.markdown(f"**{lbl} :** `{_fmt(val)}/20`")
                moy = (_fval(st.session_state.get(f"francais{suffix}")) +
                       _fval(st.session_state.get(f"histgeo{suffix}")) +
                       _fval(st.session_state.get(f"anglais{suffix}"))) / 3
                st.metric("Moyenne littéraire", f"{moy:.2f} / 20")
            return
        # ── Mode conseiller : saisie active ──
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**Matières scientifiques**")
            st.session_state[f"maths{suffix}"]   = st.number_input("Mathématiques",      0.0, 20.0, float(st.session_state[f"maths{suffix}"] or 0.0),   0.5, key=f"{key_prefix}_maths")
            st.session_state[f"sci_phy{suffix}"] = st.number_input("Sciences Physiques", 0.0, 20.0, float(st.session_state[f"sci_phy{suffix}"] or 0.0), 0.5, key=f"{key_prefix}_sp")
            st.session_state[f"svt{suffix}"]     = st.number_input("SVT",                0.0, 20.0, float(st.session_state[f"svt{suffix}"] or 0.0),     0.5, key=f"{key_prefix}_svt")
            moy = (st.session_state[f"maths{suffix}"] + st.session_state[f"sci_phy{suffix}"] + st.session_state[f"svt{suffix}"]) / 3
            st.metric("Moyenne scientifique", f"{moy:.2f} / 20")
        with c2:
            st.markdown("**Matières littéraires**")
            st.session_state[f"francais{suffix}"] = st.number_input("Français",            0.0, 20.0, float(st.session_state[f"francais{suffix}"] or 0.0), 0.5, key=f"{key_prefix}_fr")
            st.session_state[f"histgeo{suffix}"]  = st.number_input("Histoire-Géographie", 0.0, 20.0, float(st.session_state[f"histgeo{suffix}"] or 0.0),  0.5, key=f"{key_prefix}_hg")
            st.session_state[f"anglais{suffix}"]  = st.number_input("Anglais",             0.0, 20.0, float(st.session_state[f"anglais{suffix}"] or 0.0),  0.5, key=f"{key_prefix}_ang")
            moy = (st.session_state[f"francais{suffix}"] + st.session_state[f"histgeo{suffix}"] + st.session_state[f"anglais{suffix}"]) / 3
            st.metric("Moyenne littéraire", f"{moy:.2f} / 20")
        # S5 — Bouton de validation des notes T1 uniquement
        if suffix == "_t1":
            st.write("")
            if st.button("✅ Valider les notes T1", use_container_width=True, key=f"{key_prefix}_valider"):
                all_set = all(
                    st.session_state.get(f"{m}_t1") is not None
                    for m in ["maths","sci_phy","svt","francais","histgeo","anglais"]
                )
                if all_set:
                    st.session_state.notes_t1_saisies = True
                    st.success("✅ Notes T1 validées — le diagnostic peut être lancé.")
                else:
                    st.error("Veuillez renseigner toutes les notes avant de valider.")

    with tab_t1:
        saisie_notes("_t1", "t1")

    with tab_t2:
        has_t2 = st.checkbox("Renseigner le Trimestre 2", value=st.session_state.t2_renseigne, key="cb_t2")
        st.session_state.t2_renseigne = has_t2
        if has_t2:
            saisie_notes("_t2", "t2")

    with tab_t3:
        has_t3 = st.checkbox("Renseigner le Trimestre 3", value=st.session_state.t3_renseigne, key="cb_t3")
        st.session_state.t3_renseigne = has_t3
        if has_t3:
            saisie_notes("_t3", "t3")

    # Graphique de progression multi-trimestre
    if st.session_state.t2_renseigne or st.session_state.t3_renseigne:
        st.divider()
        st.markdown("**📈 Évolution des moyennes trimestrielles**")

        # S5 — None-safe helper pour le graphique
        def _v(k): return float(st.session_state.get(k) or 0.0)

        trims, sci_vals, lit_vals = ["T1"], [], []
        sci_vals.append((_v("maths_t1") + _v("sci_phy_t1") + _v("svt_t1")) / 3)
        lit_vals.append((_v("francais_t1") + _v("histgeo_t1") + _v("anglais_t1")) / 3)
        if st.session_state.t2_renseigne:
            trims.append("T2")
            sci_vals.append((_v("maths_t2") + _v("sci_phy_t2") + _v("svt_t2")) / 3)
            lit_vals.append((_v("francais_t2") + _v("histgeo_t2") + _v("anglais_t2")) / 3)
        if st.session_state.t3_renseigne:
            trims.append("T3")
            sci_vals.append((_v("maths_t3") + _v("sci_phy_t3") + _v("svt_t3")) / 3)
            lit_vals.append((_v("francais_t3") + _v("histgeo_t3") + _v("anglais_t3")) / 3)

        fig_prog = go.Figure()
        fig_prog.add_trace(go.Scatter(
            x=trims, y=sci_vals, mode="lines+markers+text",
            name="Moy. Scientifique", line=dict(color="#10b981", width=2.5), marker=dict(size=10),
            text=[f"{v:.1f}" for v in sci_vals], textposition="top center"
        ))
        fig_prog.add_trace(go.Scatter(
            x=trims, y=lit_vals, mode="lines+markers+text",
            name="Moy. Littéraire", line=dict(color="#3b82f6", width=2.5), marker=dict(size=10),
            text=[f"{v:.1f}" for v in lit_vals], textposition="bottom center"
        ))
        fig_prog.add_hline(y=10, line_dash="dash", line_color="#f59e0b", annotation_text="Seuil 10/20")
        fig_prog.update_layout(
            yaxis=dict(range=[0, 22], title="Note /20"),
            height=260, margin=dict(l=20, r=20, t=20, b=20),
            legend=dict(orientation="h", y=-0.3)
        )
        st.plotly_chart(fig_prog, use_container_width=True)

    st.write("")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Retour ⬅️", use_container_width=True):
            st.session_state.step = 0
            rerun()
    with c2:
        if st.button("Suivant : Tests psychotechniques ➡️", use_container_width=True, type="primary"):
            auto_save_step(1)
            st.session_state.step = 2
            rerun()

# =====================================================================
# ETAPE 2 — TESTS PSYCHOTECHNIQUES (passation guidée OU saisie directe)
# =====================================================================
elif step == 2:
    st.subheader("🧪 Tests psychotechniques")

    # ── Saisie directe réservée au conseiller ──
    if mode_conseiller:
        mode_radio = st.radio(
            "Mode de saisie :",
            ["🎯 Passation guidée (vraies questions)", "📊 Saisie directe des scores (barèmes papier)"],
            horizontal=True
        )
        mode_guidee = "guidée" in mode_radio
    else:
        # Élève : uniquement la passation guidée
        mode_guidee = True
        st.markdown("""
        <div class="alert-info" style="font-size:0.85rem; padding:0.6rem 1rem;">
            📋 Mode <strong>Passation guidée</strong> — Répondez aux questions de chaque test.
            Vos réponses seront analysées par votre conseiller d'orientation.
        </div>""", unsafe_allow_html=True)

    if mode_guidee:
        st.markdown("_Répondez aux questions de chaque test. Sélectionnez une réponse par question._")
        test_keys  = ["D48", "KRX", "MECA", "BV11", "PRC"]
        test_names = [TESTS_QUESTIONS[k]["nom"].split("—")[0].strip() for k in test_keys]
        test_tabs  = st.tabs(test_names)

        for t_key, t_tab in zip(test_keys, test_tabs):
            with t_tab:
                tdata = TESTS_QUESTIONS[t_key]
                st.caption(tdata["description"])
                if t_key not in st.session_state.test_answers:
                    st.session_state.test_answers[t_key] = {}

                correct = 0
                for q_idx, q in enumerate(tdata["questions"]):
                    st.markdown(f'<div class="test-question"><strong>Q{q_idx+1}.</strong> {q["q"]}</div>', unsafe_allow_html=True)
                    prev = st.session_state.test_answers[t_key].get(q_idx)
                    # Aucune réponse présélectionnée : index=None si rien encore choisi
                    idx_default = q["choices"].index(prev) if prev in q["choices"] else None
                    choice = st.radio(
                        f"R{q_idx+1}_{t_key}",
                        q["choices"],
                        index=idx_default,
                        key=f"q_{t_key}_{q_idx}",
                        label_visibility="collapsed"
                    )
                    if choice is not None:
                        st.session_state.test_answers[t_key][q_idx] = choice
                    if choice == q["answer"]:
                        correct += 1
                    elif mode_conseiller and choice is not None:
                        st.caption(f"✅ Réponse correcte : {q['answer']} — {q['expl']}")

                # Score brut = nb bonnes réponses sur le total de questions
                total_q_c    = len(tdata["questions"])
                score_brut_c = correct                              # ex. 15/20
                score        = round((score_brut_c / total_q_c) * 20, 1)  # étalonnée /20
                st.session_state.test_scores[t_key] = score

                # Scores visibles UNIQUEMENT en mode conseiller
                if mode_conseiller:
                    st.success(f"Score : **{score_brut_c}/{total_q_c}** réponses (brut) → **{score}/20** (étalonnée)")
                else:
                    nb_rep = sum(1 for v in st.session_state.test_answers[t_key].values() if v is not None)
                    total_q = len(tdata["questions"])
                    if nb_rep >= total_q:
                        st.markdown(f'<div class="alert-success">✅ Test {t_key} complété — {total_q}/{total_q} réponses enregistrées.</div>', unsafe_allow_html=True)
                    else:
                        st.markdown(f'<div class="alert-warning">📝 {nb_rep}/{total_q} questions répondues.</div>', unsafe_allow_html=True)

        # Synchroniser avec les variables de session
        for key_map in [("D48","d48"), ("KRX","krx"), ("MECA","meca"), ("BV11","bv11"), ("PRC","prc")]:
            if key_map[0] in st.session_state.test_scores:
                st.session_state[key_map[1]] = st.session_state.test_scores[key_map[0]]

    else:
        # Saisie directe via sliders — Mode conseiller uniquement
        st.markdown("**Tests d'aptitude scientifique**")
        c1, c2 = st.columns(2)
        with c1:
            st.caption("D48 : Raisonnement logique et abstrait")
            st.session_state.d48 = st.slider("Score D48",  0.0, 20.0, st.session_state.d48,  0.5, key="sl_d48")
        with c2:
            st.caption("KRX : Aptitude mathématique, suites numériques")
            st.session_state.krx = st.slider("Score KRX",  0.0, 20.0, st.session_state.krx,  0.5, key="sl_krx")
        st.caption("MECA : Compréhension mécanique et physique")
        st.session_state.meca = st.slider("Score MECA", 0.0, 20.0, st.session_state.meca, 0.5, key="sl_meca")
        st.divider()
        st.markdown("**Tests d'aptitude littéraire**")
        c1, c2 = st.columns(2)
        with c1:
            st.caption("BV11 : Vocabulaire, synonymes et compréhension de texte")
            st.session_state.bv11 = st.slider("Score BV11", 0.0, 20.0, st.session_state.bv11, 0.5, key="sl_bv11")
        with c2:
            st.caption("PRC : Proverbes et raisonnement linguistique")
            st.session_state.prc  = st.slider("Score PRC",  0.0, 20.0, st.session_state.prc,  0.5, key="sl_prc")

    # ── Récapitulatif : scores visibles seulement au conseiller ──
    st.divider()
    if mode_conseiller:
        st.markdown("**Récapitulatif des scores**")
        cols5 = st.columns(5)
        for col, (code, val) in zip(cols5, [
            ("D48", st.session_state.d48), ("KRX", st.session_state.krx),
            ("MECA", st.session_state.meca), ("BV11", st.session_state.bv11),
            ("PRC", st.session_state.prc)
        ]):
            etal = etalonner(val)
            col.markdown(f"""
            <div class="score-card">
                <div style="font-weight:700; font-size:0.9rem;">{code}</div>
                <div class="score-raw">Brut : {val:.1f}/20</div>
                <div class="score-etalon">Étal : {etal:.1f}/20</div>
            </div>""", unsafe_allow_html=True)
        SA_brut, LA_brut, SA_etal, LA_etal, MECA_etal = calc_aptitudes()
        st.write("")
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("SA brute",     f"{SA_brut:.1f}/20")
        m2.metric("SA étalonnée", f"{SA_etal:.1f}/20", help="Valeur utilisée pour l'orientation")
        m3.metric("LA brute",     f"{LA_brut:.1f}/20")
        m4.metric("LA étalonnée", f"{LA_etal:.1f}/20", help="Valeur utilisée pour l'orientation")
    else:
        # Message de confirmation pour l'élève
        total_reponses = sum(
            sum(1 for v in st.session_state.test_answers.get(k,{}).values() if v is not None)
            for k in ["D48","KRX","MECA","BV11","PRC"]
        )
        total_questions = sum(len(TESTS_QUESTIONS[k]["questions"]) for k in ["D48","KRX","MECA","BV11","PRC"])
        if total_reponses >= total_questions:
            st.markdown("""
            <div class="alert-success">
                🎉 <strong>Félicitations !</strong> Vous avez complété tous les tests psychotechniques.
                Vos réponses ont été enregistrées. Le Conseiller d'Orientation analysera vos résultats.
            </div>""", unsafe_allow_html=True)
        else:
            st.markdown(f"""
            <div class="alert-warning">
                📝 <strong>Tests en cours</strong> — {total_reponses}/{total_questions} questions répondues.
                Complétez tous les tests pour finaliser votre passation.
            </div>""", unsafe_allow_html=True)

    st.write("")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Retour ⬅️", use_container_width=True):
            st.session_state.step = 1
            rerun()
    with c2:
        if mode_conseiller:
            if st.button("Lancer le Diagnostic IA 🤖", use_container_width=True, type="primary"):
                # S5 — Bloquer si les notes T1 n'ont pas été validées
                if not st.session_state.get("notes_t1_saisies", False):
                    st.error("⚠️ Notes T1 non validées. Veuillez saisir et valider les notes du 1er trimestre avant de lancer le diagnostic.")
                else:
                    st.session_state.step = 3
                    st.session_state.chat_history = []
                    auto_save_step(2)
                    rerun()
        else:
            st.info("La suite de l'analyse (Diagnostic IA) est réservée au Conseiller d'Orientation.")

# =====================================================================
# ETAPE 3 — DIAGNOSTIC IA
# =====================================================================
elif step == 3:
    SA_brut, LA_brut, SA_etal, LA_etal, MECA_etal = calc_aptitudes()
    SA, LA = SA_etal, LA_etal            # Décision basée sur les notes étalonnées
    notes, trim_label = get_notes_actives()
    moy_sci, moy_lit  = calc_moyennes(notes)
    choix = st.session_state.choix_personnel

    # ---- Contexte trimestre ----
    t1_seul   = not st.session_state.t2_renseigne and not st.session_state.t3_renseigne
    t2_dispo  = st.session_state.t2_renseigne and not st.session_state.t3_renseigne
    t3_dispo  = st.session_state.t3_renseigne

    # ---- Moteur d'inférence étendu ----
    # S8 — Zone grise [9.0 — 11.0] : probation immédiate au lieu de binaire strict
    SEUIL_BAS  = 9.0   # en dessous → attente / probation
    SEUIL_HAUT = 11.0  # au-dessus  → confirmé

    # ══════════════════════════════════════════════════════════════════
    # RÈGLE PRIORITAIRE — CHOIX DE L'ÉLÈVE
    # Si l'élève a exprimé un choix de série ET que :
    #   • son aptitude correspondante (SA pour C, LA pour A) ≥ 10/20
    #   • ET sa moyenne scolaire correspondante ≥ 10/20
    # → Son choix est validé définitivement, quelle que soit la
    #   comparaison SA vs LA ou moy_sci vs moy_lit.
    # Cette règle est évaluée EN PREMIER, avant tout autre cas.
    # ══════════════════════════════════════════════════════════════════
    choix_C = "C" in choix
    choix_A = "A" in choix

    # Conditions d'éligibilité au choix prioritaire
    _choix_C_eligible = choix_C and SA >= 10.0 and moy_sci >= 10.0
    _choix_A_eligible = choix_A and LA >= 10.0 and moy_lit >= 10.0

    if _choix_C_eligible:
        # L'élève veut la C, a l'aptitude sci ≥ 10 ET les notes sci ≥ 10
        # → On respecte son choix quelle que soit la comparaison avec LA / moy_lit
        serie      = "C"
        statut     = "confirme"
        msg_diag   = (
            "✅ CHOIX VALIDÉ — Série C confirmée selon le projet de l'élève\n"
            f"(SA={SA:.1f}/20 ≥ 10 · Moy.Sci={moy_sci:.1f}/20 ≥ 10)"
        )
        alerte_type= "success"

    elif _choix_A_eligible:
        # L'élève veut la A, a l'aptitude lit ≥ 10 ET les notes lit ≥ 10
        serie      = "A"
        statut     = "confirme"
        msg_diag   = (
            "✅ CHOIX VALIDÉ — Série A confirmée selon le projet de l'élève\n"
            f"(LA={LA:.1f}/20 ≥ 10 · Moy.Lit={moy_lit:.1f}/20 ≥ 10)"
        )
        alerte_type= "success"

    # ══════════════════════════════════════════════════════════════════
    # MOTEUR STANDARD — appliqué si le choix prioritaire ne s'active pas
    # ══════════════════════════════════════════════════════════════════

    # Cas 1 : profil scientifique confirmé (zone haute)
    elif SA > LA and moy_sci >= SEUIL_HAUT:
        serie      = "C"
        statut     = "confirme"
        msg_diag   = "Profil scientifique confirmé ✅"
        alerte_type= "success"

    # Cas 1b : zone grise scientifique (9 ≤ moy_sci < 11)
    elif SA > LA and SEUIL_BAS <= moy_sci < SEUIL_HAUT:
        serie      = "C"
        statut     = "probation"
        msg_diag   = "Zone grise — PROBATION IMMÉDIATE (moy. sci. entre 9 et 11/20)"
        alerte_type= "warning"

    # Cas 2 : profil littéraire confirmé (zone haute)
    elif LA > SA and moy_lit >= SEUIL_HAUT:
        serie      = "A"
        statut     = "confirme"
        msg_diag   = "Profil littéraire confirmé ✅"
        alerte_type= "success"

    # Cas 2b : zone grise littéraire
    elif LA > SA and SEUIL_BAS <= moy_lit < SEUIL_HAUT:
        serie      = "A"
        statut     = "probation"
        msg_diag   = "Zone grise — PROBATION IMMÉDIATE (moy. lit. entre 9 et 11/20)"
        alerte_type= "warning"

    # Cas 3 : aptitudes scientifiques OK, notes sci strictement faibles
    elif SA > LA and moy_sci < SEUIL_BAS:
        if t1_seul:
            serie      = "C"
            statut     = "attente"
            msg_diag   = "EN ATTENTE — Notes du 2e trimestre requises"
            alerte_type= "attente"
        elif t2_dispo:
            serie      = "C"
            statut     = "probation"
            msg_diag   = "PROBATION — Amélioration critique avant le 3e trimestre"
            alerte_type= "warning"
        elif t3_dispo:
            # S2 — si MECA est fort, orienter vers filière technique plutôt que A par défaut
            if MECA_etal >= 12.0:
                serie      = "TECHNIQUE"
                statut     = "confirme"
                msg_diag   = "ORIENTATION TECHNIQUE — Aptitudes mécaniques confirmées ✅"
                alerte_type= "success"
            else:
                serie      = "A"
                statut     = "revise"
                msg_diag   = "ORIENTATION RÉVISÉE → Série A définitive"
                alerte_type= "revise"
        else:
            serie, statut, msg_diag, alerte_type = "C", "probation", "Probation", "warning"

    # Cas 4 : aptitudes littéraires OK, notes lit strictement faibles
    elif LA > SA and moy_lit < SEUIL_BAS:
        if t1_seul:
            serie      = "A"
            statut     = "attente"
            msg_diag   = "EN ATTENTE — Notes du 2e trimestre requises"
            alerte_type= "attente"
        elif t2_dispo:
            serie      = "A"
            statut     = "probation"
            msg_diag   = "PROBATION — Amélioration critique avant le 3e trimestre"
            alerte_type= "warning"
        elif t3_dispo:
            if MECA_etal >= 12.0:
                serie      = "TECHNIQUE"
                statut     = "confirme"
                msg_diag   = "ORIENTATION TECHNIQUE — Aptitudes mécaniques confirmées ✅"
                alerte_type= "success"
            else:
                serie      = "C"
                statut     = "revise"
                msg_diag   = "ORIENTATION RÉVISÉE → Série C définitive"
                alerte_type= "revise"
        else:
            serie, statut, msg_diag, alerte_type = "A", "probation", "Probation", "warning"

    # S2 — Cas 5 : ni C ni A mais MECA fort → filière technique
    elif MECA_etal >= 12.0:
        serie      = "TECHNIQUE"
        statut     = "confirme"
        msg_diag   = "ORIENTATION TECHNIQUE — Aptitudes mécaniques remarquables ✅"
        alerte_type= "success"

    # Cas 6 : profil indéterminé
    elif SA == LA:
        serie      = "?"
        statut     = "indetermine"
        msg_diag   = "Profil équilibré — entretien approfondi recommandé"
        alerte_type= "warning"
    else:
        serie      = "?"
        statut     = "indetermine"
        msg_diag   = "Données insuffisantes — compléter le dossier"
        alerte_type= "warning"

    probation  = statut in ("probation",)
    score_conf = calc_score_confiance(SA, LA, moy_sci, moy_lit, serie)

    # Conseil selon revenu
    conseil_revenu = ""
    if "Faible" in st.session_state.revenu:
        if "C" in serie:
            conseil_revenu = "⚠️ Revenu modeste : la série C implique des études longues et coûteuses. Vérifier les bourses et dispositifs de soutien scolaire."
        elif "A" in serie:
            conseil_revenu = "ℹ️ La série A offre des débouchés accessibles (droit, lettres, administration) compatibles avec le contexte financier familial."

    # ── Détection du type de conflit ──
    # Pas de conflit si le choix prioritaire a été appliqué (l'élève avait raison)
    conflit_type = None
    if _choix_C_eligible or _choix_A_eligible:
        # Choix validé → pas de conflit à signaler
        conflit_type = None
    elif SA > LA and "A" in choix:
        conflit_type = "decale"
    elif SA > LA and moy_sci < 10 and "C" in choix:
        conflit_type = "reveur"
    elif LA > SA and moy_lit < 10 and "A" in choix:
        conflit_type = "reveur"
    # S2 — Conflit TECHNIQUE : MECA fort mais ni C ni A confirmé
    elif serie == "TECHNIQUE":
        conflit_type = "technique"

    st.session_state.orientation_finale = serie
    st.session_state.probation = probation
    st.session_state.statut = statut
    st.session_state.score_confiance = score_conf

    # ── Bannière choix prioritaire (visible uniquement si la règle s'est appliquée) ──
    if _choix_C_eligible or _choix_A_eligible:
        _serie_choisie = "C" if _choix_C_eligible else "A"
        _apt_val  = SA if _choix_C_eligible else LA
        _moy_val  = moy_sci if _choix_C_eligible else moy_lit
        _moy_lbl  = "Sci." if _choix_C_eligible else "Lit."
        _apt_lbl  = "SA" if _choix_C_eligible else "LA"
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#fdf4ff,#ede9fe);
                    border-left:5px solid #7c3aed;border-radius:14px;
                    padding:1rem 1.4rem;margin:0.6rem 0;
                    box-shadow:0 3px 12px rgba(124,58,237,0.15);">
            <div style="font-size:0.75rem;font-weight:800;text-transform:uppercase;
                        letter-spacing:0.1em;color:#7c3aed;margin-bottom:0.4rem;">
                🎯 Règle prioritaire — Choix de l'élève respecté
            </div>
            <div style="font-size:0.9rem;color:#4c1d95;line-height:1.6;">
                L'élève a choisi la <strong>Série {_serie_choisie}</strong> et remplit les deux conditions requises :<br>
                &nbsp;• Aptitude {_apt_lbl} = <strong>{_apt_val:.1f}/20 ≥ 10</strong>
                &nbsp;• Moyenne {_moy_lbl} = <strong>{_moy_val:.1f}/20 ≥ 10</strong><br>
                Son choix est <strong>validé définitivement</strong>, indépendamment de la comparaison SA/LA.
            </div>
        </div>""", unsafe_allow_html=True)

    # ---- Affichage des métriques ----
    st.subheader("Analyse du profil")
    m1, m2, m3, m4, m5 = st.columns(5)
    m1.metric("SA étalonnée",            f"{SA:.1f}/20",
              delta=f"Brut: {SA_brut:.1f}" if mode_conseiller else None)
    m2.metric("LA étalonnée",            f"{LA:.1f}/20",
              delta=f"Brut: {LA_brut:.1f}" if mode_conseiller else None)
    m3.metric(f"Moy. Sci. ({trim_label})", f"{moy_sci:.1f}/20",
              delta=f"{moy_sci-10:+.1f} vs seuil", delta_color="normal" if moy_sci >= 10 else "inverse")
    m4.metric(f"Moy. Lit. ({trim_label})", f"{moy_lit:.1f}/20",
              delta=f"{moy_lit-10:+.1f} vs seuil", delta_color="normal" if moy_lit >= 10 else "inverse")
    # S2 — MECA étalonnée visible dans le diagnostic
    meca_color = "normal" if MECA_etal >= 12 else "off"
    m5.metric("MECA étalonnée",          f"{MECA_etal:.1f}/20",
              delta="≥12 : filière tech." if MECA_etal >= 12 else "< 12",
              delta_color=meca_color)

    # Barre de confiance
    conf_color = "#10b981" if score_conf >= 70 else ("#f59e0b" if score_conf >= 50 else "#ef4444")
    conf_label = "Fiable" if score_conf >= 70 else ("Modérée" if score_conf >= 50 else "Incertaine")
    st.markdown(f"""
    <div style="background:#f8fafc; border-radius:10px; padding:0.8rem 1rem; margin:0.5rem 0;">
        <div style="font-size:0.8rem; color:#64748b; margin-bottom:4px;">
            Score de confiance de l'orientation — <em>{conf_label}</em>
        </div>
        <div style="background:#e2e8f0; border-radius:6px; height:10px; overflow:hidden;">
            <div style="width:{score_conf}%; height:100%; background:{conf_color}; border-radius:6px;"></div>
        </div>
        <div style="font-size:0.85rem; color:{conf_color}; font-weight:600; margin-top:4px;">{score_conf} %</div>
    </div>
    """, unsafe_allow_html=True)

    # Graphiques côte à côte : Radar + Barres
    col_r, col_b = st.columns(2)
    with col_r:
        labels = ["Logique\n(D48)", "Maths\n(KRX)", "Méca.\n(MECA)", "Litt.\n(BV11)", "Prov.\n(PRC)"]
        apt_etal = [etalonner(st.session_state.d48,"d48"), etalonner(st.session_state.krx,"krx"),
                    etalonner(st.session_state.meca,"meca"), etalonner(st.session_state.bv11,"bv11"),
                    etalonner(st.session_state.prc,"prc")]
        notes_list = [notes["maths"], notes["sci_phy"], notes["svt"], notes["francais"], notes["histgeo"]]
        fig_r = go.Figure()
        fig_r.add_trace(go.Scatterpolar(r=apt_etal, theta=labels, fill="toself",
            name="Aptitudes (étalonnées)", line_color="#10b981", fillcolor="rgba(16,185,129,0.15)"))
        fig_r.add_trace(go.Scatterpolar(r=notes_list, theta=labels, fill="toself",
            name=f"Notes scolaires ({trim_label})", line_color="#3b82f6", fillcolor="rgba(59,130,246,0.15)"))
        if mode_conseiller:
            apt_brut = [st.session_state.d48, st.session_state.krx, st.session_state.meca,
                        st.session_state.bv11, st.session_state.prc]
            fig_r.add_trace(go.Scatterpolar(r=apt_brut, theta=labels, fill="toself",
                name="Aptitudes (brutes)", line_color="#f59e0b", fillcolor="rgba(245,158,11,0.08)",
                line=dict(dash="dot")))
        fig_r.update_layout(
            polar=dict(radialaxis=dict(visible=True, range=[0, 20]), gridshape="linear"),
            showlegend=True, height=310,
            margin=dict(l=30, r=30, t=10, b=60),
            legend=dict(orientation="h", y=-0.25, x=0.5, xanchor="center")
        )
        st.plotly_chart(fig_r, use_container_width=True)

    with col_b:
        cats   = ["SA étal.", "LA étal.", f"Sci. {trim_label}", f"Lit. {trim_label}"]
        vals   = [SA, LA, moy_sci, moy_lit]
        colors = ["#10b981" if v >= 10 else "#f87171" for v in vals]
        fig_b = go.Figure(go.Bar(
            x=cats, y=vals, marker_color=colors,
            text=[f"{v:.1f}" for v in vals], textposition="outside"
        ))
        fig_b.add_hline(y=10, line_dash="dash", line_color="#f59e0b", annotation_text="Seuil 10/20")
        fig_b.update_layout(
            yaxis=dict(range=[0, 23], title="Note /20"),
            height=310, margin=dict(l=10, r=10, t=10, b=20), showlegend=False
        )
        st.plotly_chart(fig_b, use_container_width=True)

    # ---- Verdict ----
    st.subheader("Verdict du moteur d'inférence")
    st.markdown(f'<div class="alert-{alerte_type}"><strong>{msg_diag}</strong></div>',
                unsafe_allow_html=True)

    prenom = st.session_state.prenom

    # --- EN ATTENTE (T1 seulement, notes faibles) ---
    if statut == "attente":
        obj_sci = max(10.0, round(moy_sci + 2.0, 1)) if SA > LA else None
        obj_lit = max(10.0, round(moy_lit + 2.0, 1)) if LA > SA else None
        obj_txt = (f"Moyenne scientifique ≥ <strong>{obj_sci}/20</strong>"
                   if obj_sci else f"Moyenne littéraire ≥ <strong>{obj_lit}/20</strong>")
        serie_cible = "C" if SA > LA else "A"
        serie_risque = "A" if SA > LA else "C"
        st.markdown(f"""
        <div class="alert-attente">
            ⏳ <strong>Statut : EN ATTENTE des notes du 2e trimestre</strong><br><br>
            {prenom}, ton potentiel est bien là (SA = {SA:.1f}/20) mais tes notes scolaires
            ({moy_sci:.1f}/20) sont encore insuffisantes pour confirmer la série {serie_cible}.<br><br>
            👉 <strong>Ce qu'il faut faire :</strong> Mets-toi au travail dès maintenant sur tes matières
            scientifiques. L'objectif à atteindre au 2e trimestre est {obj_txt}.<br><br>
            ⚠️ <em>Si aucune progression n'est constatée au 2e trimestre, le statut passera en
            <strong>Probation</strong>. En fin d'année (T3), si les notes restent insuffisantes,
            l'orientation sera définitivement révisée vers la série {serie_risque}.</em>
        </div>
        """, unsafe_allow_html=True)

    # --- PROBATION (T2 dispo, notes toujours faibles) ---
    elif statut == "probation":
        obj_sci = max(10.0, round(moy_sci + 2.0, 1)) if SA > LA else None
        obj_lit = max(10.0, round(moy_lit + 2.0, 1)) if LA > SA else None
        obj_txt = (f"Moyenne scientifique ≥ <strong>{obj_sci}/20</strong>"
                   if obj_sci else f"Moyenne littéraire ≥ <strong>{obj_lit}/20</strong>")
        serie_risque = "A" if SA > LA else "C"
        st.markdown(f"""
        <div class="alert-warning">
            ⚠️ <strong>Statut : PROBATION — Dernière chance avant le 3e trimestre</strong><br><br>
            {prenom}, tes aptitudes le prouvent, tu as les capacités pour la série {serie}.
            Mais tes résultats scolaires ({moy_sci:.1f}/20) ne progressent pas suffisamment.<br><br>
            🎯 <strong>Objectif impératif au T3 :</strong> {obj_txt}<br><br>
            🚨 <em>Si cet objectif n'est pas atteint au 3e trimestre,
            <strong>l'orientation sera définitivement révisée vers la série {serie_risque}.</strong></em>
        </div>
        """, unsafe_allow_html=True)

    # --- ORIENTATION RÉVISÉE (T3 dispo, échec confirmé) ---
    elif statut == "revise":
        serie_initiale = "C" if serie == "A" else "A"
        st.markdown(f"""
        <div class="alert-revise">
            🔄 <strong>Statut : ORIENTATION RÉVISÉE — Décision définitive</strong><br><br>
            {prenom}, malgré de bonnes aptitudes ({SA:.1f}/20), les notes scolaires n'ont pas
            atteint le seuil requis sur les 3 trimestres.<br><br>
            La série <strong>{serie_initiale}</strong> n'est plus envisageable cette année.
            L'orientation définitive est fixée en <strong>Série {serie}</strong>.<br><br>
            ℹ️ <em>Cette décision a été prise conformément au protocole d'orientation :
            3 trimestres insuffisants malgré le suivi en probation.</em>
        </div>
        """, unsafe_allow_html=True)

    if abs(SA - moy_sci) >= 3:
        sens = "supérieure" if SA > moy_sci else "inférieure"
        st.markdown(
            f'<div class="alert-warning">🔍 Décalage détecté : aptitude scientifique étalonnée ({SA:.1f}/20) '
            f'{sens} de {abs(SA-moy_sci):.1f} pt(s) à la moyenne scolaire ({moy_sci:.1f}/20). '
            f'Investiguer la méthode de travail.</div>',
            unsafe_allow_html=True
        )

    if conseil_revenu:
        st.markdown(f'<div class="alert-info">{conseil_revenu}</div>', unsafe_allow_html=True)

    # ==================================================================
    # RUBRIQUE SUIVI T2 — visible uniquement si statut == "attente"
    # ==================================================================
    if statut == "attente":
        st.divider()
        st.markdown("""
        <div style="background:linear-gradient(135deg,#f0f9ff,#e0f2fe);
                    border:2px dashed #0ea5e9; border-radius:16px;
                    padding:1.2rem 1.4rem; margin:0.5rem 0;">
            <div style="font-size:1rem; font-weight:700; color:#0c4a6e; margin-bottom:0.3rem;">
                ⏳ Suivi — Saisie des notes du 2e trimestre
            </div>
            <div style="font-size:0.85rem; color:#0369a1;">
                Renseignez les notes du T2 pour mettre à jour l'orientation
                (Probation ou Confirmation).
            </div>
        </div>
        """, unsafe_allow_html=True)

        with st.expander("📋 Saisir les notes du 2e trimestre", expanded=True):
            st.caption("Ces notes seront enregistrées et le diagnostic sera recalculé automatiquement.")
            ca, cb = st.columns(2)
            # S5 — helper None-safe pour les inputs suivi
            def _sv(key): return float(st.session_state.get(key) or 0.0)
            with ca:
                st.markdown("**🔬 Matières scientifiques**")
                st.session_state.maths_t2   = st.number_input(
                    "Mathématiques T2",   0.0, 20.0,
                    _sv("maths_t2"),   0.5, key="suivi_t2_maths")
                st.session_state.sci_phy_t2 = st.number_input(
                    "Sciences Physiques T2", 0.0, 20.0,
                    _sv("sci_phy_t2"), 0.5, key="suivi_t2_sp")
                st.session_state.svt_t2     = st.number_input(
                    "SVT T2",             0.0, 20.0,
                    _sv("svt_t2"),     0.5, key="suivi_t2_svt")
                moy_sci_t2 = (st.session_state.maths_t2 +
                              st.session_state.sci_phy_t2 +
                              st.session_state.svt_t2) / 3
                delta_sci = moy_sci_t2 - moy_sci
                st.metric("Moy. Sci. T2", f"{moy_sci_t2:.2f}/20",
                          delta=f"{delta_sci:+.1f} vs T1",
                          delta_color="normal" if delta_sci >= 0 else "inverse")
            with cb:
                st.markdown("**📖 Matières littéraires**")
                st.session_state.francais_t2 = st.number_input(
                    "Français T2",           0.0, 20.0,
                    _sv("francais_t2"), 0.5, key="suivi_t2_fr")
                st.session_state.histgeo_t2  = st.number_input(
                    "Histoire-Géographie T2",0.0, 20.0,
                    _sv("histgeo_t2"),  0.5, key="suivi_t2_hg")
                st.session_state.anglais_t2  = st.number_input(
                    "Anglais T2",            0.0, 20.0,
                    _sv("anglais_t2"),  0.5, key="suivi_t2_ang")
                moy_lit_t2 = (st.session_state.francais_t2 +
                              st.session_state.histgeo_t2 +
                              st.session_state.anglais_t2) / 3
                delta_lit = moy_lit_t2 - moy_lit
                st.metric("Moy. Lit. T2", f"{moy_lit_t2:.2f}/20",
                          delta=f"{delta_lit:+.1f} vs T1",
                          delta_color="normal" if delta_lit >= 0 else "inverse")

            # Prévisualisation du nouveau statut
            serie_cible_t2 = "C" if SA > LA else "A"
            seuil_ok_t2    = moy_sci_t2 >= 10 if SA > LA else moy_lit_t2 >= 10
            nouveau_statut_t2 = "✅ CONFIRMATION" if seuil_ok_t2 else "⚠️ PROBATION"
            couleur_prev_t2   = "#10b981" if seuil_ok_t2 else "#f59e0b"
            fond_prev_t2      = "#ecfdf5" if seuil_ok_t2 else "#fffbeb"
            st.markdown(f"""
            <div style="background:{fond_prev_t2}; border:1px solid {couleur_prev_t2};
                        border-radius:10px; padding:0.7rem 1rem; margin-top:0.8rem;
                        text-align:center;">
                <span style="color:{couleur_prev_t2}; font-weight:700; font-size:0.95rem;">
                    Prévisualisation : {nouveau_statut_t2} en Série {serie_cible_t2}
                </span>
            </div>
            """, unsafe_allow_html=True)

            st.write("")
            if st.button("✅ Valider les notes T2 et recalculer", use_container_width=True, type="primary"):
                st.session_state.t2_renseigne = True
                st.session_state.chat_history = []   # reset l'agent pour le nouveau contexte
                rerun()

    # ==================================================================
    # RUBRIQUE SUIVI T3 — visible uniquement si statut == "probation"
    # ==================================================================
    if statut == "probation":
        st.divider()
        st.markdown("""
        <div style="background:linear-gradient(135deg,#fffbeb,#fef3c7);
                    border:2px dashed #f59e0b; border-radius:16px;
                    padding:1.2rem 1.4rem; margin:0.5rem 0;">
            <div style="font-size:1rem; font-weight:700; color:#78350f; margin-bottom:0.3rem;">
                ⚠️ Suivi Probation — Saisie des notes du 3e trimestre
            </div>
            <div style="font-size:0.85rem; color:#92400e;">
                Renseignez les notes du T3 pour rendre la décision d'orientation
                <strong>définitive</strong>.
            </div>
        </div>
        """, unsafe_allow_html=True)

        with st.expander("📋 Saisir les notes du 3e trimestre (décision finale)", expanded=True):
            st.caption("⚠️ Ces notes détermineront l'orientation définitive de l'élève.")
            ca, cb = st.columns(2)
            # S5 — helper None-safe pour les inputs suivi T3
            def _sv3(key): return float(st.session_state.get(key) or 0.0)
            with ca:
                st.markdown("**🔬 Matières scientifiques**")
                st.session_state.maths_t3   = st.number_input(
                    "Mathématiques T3",   0.0, 20.0,
                    _sv3("maths_t3"),   0.5, key="suivi_t3_maths")
                st.session_state.sci_phy_t3 = st.number_input(
                    "Sciences Physiques T3", 0.0, 20.0,
                    _sv3("sci_phy_t3"), 0.5, key="suivi_t3_sp")
                st.session_state.svt_t3     = st.number_input(
                    "SVT T3",             0.0, 20.0,
                    _sv3("svt_t3"),     0.5, key="suivi_t3_svt")
                moy_sci_t3 = (st.session_state.maths_t3 +
                              st.session_state.sci_phy_t3 +
                              st.session_state.svt_t3) / 3
                notes_ref  = get_notes_actives()[0]
                moy_ref    = calc_moyennes(notes_ref)[0] if SA > LA else calc_moyennes(notes_ref)[1]
                delta_sci3 = moy_sci_t3 - moy_ref
                st.metric("Moy. Sci. T3", f"{moy_sci_t3:.2f}/20",
                          delta=f"{delta_sci3:+.1f} vs T2",
                          delta_color="normal" if delta_sci3 >= 0 else "inverse")
            with cb:
                st.markdown("**📖 Matières littéraires**")
                st.session_state.francais_t3 = st.number_input(
                    "Français T3",           0.0, 20.0,
                    _sv3("francais_t3"), 0.5, key="suivi_t3_fr")
                st.session_state.histgeo_t3  = st.number_input(
                    "Histoire-Géographie T3",0.0, 20.0,
                    _sv3("histgeo_t3"),  0.5, key="suivi_t3_hg")
                st.session_state.anglais_t3  = st.number_input(
                    "Anglais T3",            0.0, 20.0,
                    _sv3("anglais_t3"),  0.5, key="suivi_t3_ang")
                moy_lit_t3 = (st.session_state.francais_t3 +
                              st.session_state.histgeo_t3 +
                              st.session_state.anglais_t3) / 3
                moy_ref_lit = calc_moyennes(get_notes_actives()[0])[1]
                delta_lit3  = moy_lit_t3 - moy_ref_lit
                st.metric("Moy. Lit. T3", f"{moy_lit_t3:.2f}/20",
                          delta=f"{delta_lit3:+.1f} vs T2",
                          delta_color="normal" if delta_lit3 >= 0 else "inverse")

            # Prévisualisation décision finale
            serie_cible_t3 = "C" if SA > LA else "A"
            serie_risque_t3= "A" if SA > LA else "C"
            seuil_ok_t3    = moy_sci_t3 >= 10 if SA > LA else moy_lit_t3 >= 10
            if seuil_ok_t3:
                prev_label = f"✅ CONFIRMATION DÉFINITIVE — Série {serie_cible_t3}"
                prev_color = "#10b981"; prev_fond = "#ecfdf5"
            else:
                prev_label = f"🔄 ORIENTATION RÉVISÉE DÉFINITIVEMENT → Série {serie_risque_t3}"
                prev_color = "#f43f5e"; prev_fond = "#fef2f2"

            st.markdown(f"""
            <div style="background:{prev_fond}; border:2px solid {prev_color};
                        border-radius:12px; padding:0.9rem 1rem; margin-top:0.8rem;
                        text-align:center;">
                <span style="color:{prev_color}; font-weight:700; font-size:0.95rem;">
                    {prev_label}
                </span>
            </div>
            """, unsafe_allow_html=True)

            st.write("")
            if st.button("🏁 Valider les notes T3 — Décision définitive", use_container_width=True, type="primary"):
                st.session_state.t3_renseigne = True
                st.session_state.chat_history = []
                rerun()

    # ---- Agent conversationnel ----
    if conflit_type:
        st.divider()
        st.subheader("💬 Agent IA d'interpellation")
        prenom = st.session_state.prenom

        if conflit_type == "decale":
            msg_init = (f"Bonjour {prenom} ! Tes tests révèlent un fort potentiel scientifique "
                        f"(SA étalonnée = {SA:.1f}/20), mais tu as choisi la série A. "
                        f"Peux-tu expliquer pourquoi tu préfères la filière littéraire ?")
        else:
            msg_init = (f"Bonjour {prenom} ! Tu as de bonnes aptitudes scientifiques "
                        f"(SA = {SA:.1f}/20), mais tes notes en classe ({moy_sci:.1f}/20) "
                        f"sont encore insuffisantes. Qu'est-ce qui te pose le plus de difficultés ?")

        if not st.session_state.chat_history:
            st.session_state.chat_history.append({"role": "ia", "content": msg_init})

        for msg in st.session_state.chat_history:
            if msg["role"] == "ia":
                st.markdown(f'<div class="chat-ia">🤖 <strong>Agent CapAvenir :</strong><br>{msg["content"]}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="chat-user">{msg["content"]}<br><strong>— {prenom}</strong></div>', unsafe_allow_html=True)

        user_input = st.text_input("Réponse de l'élève :", placeholder="Écrivez ici...", key="chat_input_field")

        if st.button("Envoyer la réponse 📨", use_container_width=True):
            if user_input.strip():
                st.session_state.chat_history.append({"role": "user", "content": user_input})
                api_key = os.environ.get("ANTHROPIC_API_KEY", "")
                if api_key and len(api_key) > 20:
                    system_prompt = (
                        f"Tu es un conseiller d'orientation scolaire IA bienveillant pour le système camerounais. "
                        f"Tu aides {prenom} {st.session_state.nom}. "
                        f"SA étalonnée={SA:.1f}/20, LA étalonnée={LA:.1f}/20, "
                        f"Moy.Sci={moy_sci:.1f}/20, Moy.Lit={moy_lit:.1f}/20. "
                        f"Choix : {choix}. Projet : {st.session_state.projet_pro or 'non précisé'}. "
                        f"Revenu famille : {st.session_state.revenu}. "
                        f"Conflit : {conflit_type}. Réponds en 2-3 phrases, empathique et direct, en français."
                    )
                    messages_api = [
                        {"role": "assistant" if m["role"] == "ia" else "user", "content": m["content"]}
                        for m in st.session_state.chat_history
                    ]
                    try:
                        resp = requests.post(
                            "https://api.anthropic.com/v1/messages",
                            headers={"Content-Type": "application/json",
                                     "x-api-key": api_key,
                                     "anthropic-version": "2023-06-01"},
                            json={"model": "claude-sonnet-4-20250514",
                                  "max_tokens": 300, "system": system_prompt,
                                  "messages": messages_api},
                            timeout=15,
                        )
                        ia_reply = resp.json()["content"][0]["text"]
                    except Exception:
                        ia_reply = reponse_ia_simulee(
                            user_input, conflit_type, prenom, SA, moy_sci,
                            LA=LA, moy_lit=moy_lit,
                            projet_pro=st.session_state.projet_pro,
                            revenu=st.session_state.revenu,
                            chat_history=st.session_state.chat_history,
                            d48=st.session_state.d48, krx=st.session_state.krx)
                else:
                    ia_reply = reponse_ia_simulee(
                        user_input, conflit_type, prenom, SA, moy_sci,
                        LA=LA, moy_lit=moy_lit,
                        projet_pro=st.session_state.projet_pro,
                        revenu=st.session_state.revenu,
                        chat_history=st.session_state.chat_history,
                        d48=st.session_state.d48, krx=st.session_state.krx)

                st.session_state.chat_history.append({"role": "ia", "content": ia_reply})
                rerun()

    # Notes du conseiller
    st.divider()
    st.subheader("📝 Observations du conseiller")
    st.session_state.notes_conseiller = st.text_area(
        "Remarques / Observations (entretien, comportement, motivations…) :",
        value=st.session_state.notes_conseiller,
        placeholder="Saisir ici les observations du conseiller lors de l'entretien individuel...",
        height=80,
    )

    # ── OVERRIDE DE L'ORIENTATION — Mode conseiller uniquement ──────
    if mode_conseiller:
        st.divider()
        st.subheader("✏️ Décision finale du conseiller")
        st.markdown("""
        <div style="background:linear-gradient(135deg,#fffbeb,#fef3c7);
                    border-left:5px solid #f59e0b; border-radius:12px;
                    padding:0.9rem 1.2rem; margin-bottom:0.8rem; font-size:0.88rem; color:#78350f;">
            🧑‍💼 <strong>Droit de révision :</strong> Si vous n'êtes pas d'accord avec la recommandation
            de l'IA, vous pouvez modifier la décision finale et indiquer votre justification.
        </div>
        """, unsafe_allow_html=True)

        serie_actuelle = st.session_state.get("orientation_finale") or serie
        # Options disponibles selon le contexte
        series_options = ["— Conserver la recommandation IA —", "C (Scientifique)", "A (Littéraire)", "TECHNIQUE"]
        override_key   = f"override_serie_{st.session_state.get('session_id','x')}"

        ov1, ov2 = st.columns([2, 3])
        with ov1:
            choix_override = st.selectbox(
                "Modifier l'orientation vers :",
                series_options,
                key=override_key,
                help="Sélectionnez une série pour remplacer la décision de l'IA."
            )
        with ov2:
            justification_override = st.text_input(
                "Justification obligatoire si modification :",
                placeholder="Ex: Entretien révèle une forte motivation littéraire non reflétée dans les tests...",
                key=f"justif_override_{st.session_state.get('session_id','x')}",
            )

        if choix_override != "— Conserver la recommandation IA —":
            serie_map = {"C (Scientifique)": "C", "A (Littéraire)": "A", "TECHNIQUE": "TECHNIQUE"}
            nouvelle_serie = serie_map[choix_override]
            ia_serie = serie

            couleur_ov = "#10b981" if nouvelle_serie == "C" else "#3b82f6" if nouvelle_serie == "A" else "#f97316"
            st.markdown(f"""
            <div style="background:linear-gradient(135deg,#ecfdf5,#d1fae5) ; border-radius:12px;
                        padding:0.8rem 1.2rem; margin-top:0.4rem; border:2px dashed {couleur_ov};">
                <span style="font-size:0.85rem; color:#0f172a;">
                    🔄 Recommandation IA : <strong>{ia_serie}</strong>
                    &nbsp;→&nbsp;
                    Décision du conseiller : <strong style="color:{couleur_ov};">{nouvelle_serie}</strong>
                </span>
            </div>
            """, unsafe_allow_html=True)

            if st.button("✅ Appliquer ma décision", use_container_width=True, type="primary",
                         key=f"btn_apply_override_{st.session_state.get('session_id','x')}"):
                if not justification_override.strip():
                    st.error("⚠️ Veuillez renseigner une justification avant d'appliquer la modification.")
                else:
                    # Appliquer l'override
                    st.session_state.orientation_finale = nouvelle_serie
                    # Ajouter la justification dans les notes du conseiller
                    note_override = (
                        f"\n[DÉCISION CONSEILLER — override IA] "
                        f"Orientation modifiée de '{ia_serie}' vers '{nouvelle_serie}'. "
                        f"Justification : {justification_override.strip()}"
                    )
                    st.session_state.notes_conseiller = (
                        st.session_state.notes_conseiller + note_override
                    ).strip()
                    st.success(f"✅ Orientation mise à jour → **Série {nouvelle_serie}**. Sauvegardez le dossier pour confirmer.")
                    rerun()
        else:
            st.caption(f"Recommandation IA actuelle : **Série {serie}** — aucune modification.")

    # ── Sauvegarde réservée au conseiller ──
    st.divider()
    sc1, sc2 = st.columns(2)
    with sc1:
        if mode_conseiller:
            if st.button("💾 Sauvegarder le dossier", use_container_width=True):
                res = db.sauvegarder_dossier(st.session_state, etalonner)
                if res["succes"]:
                    st.success(f"{res['message']}")
                else:
                    st.error(res["message"])
        else:
            st.markdown('<div class="alert-danger" style="font-size:0.82rem;padding:0.5rem 1rem;">🔒 Sauvegarde réservée au Conseiller d&#39;Orientation</div>', unsafe_allow_html=True)
    with sc2:
        if statut in ("attente", "probation"):
            hint = "T2 requis" if statut == "attente" else "T3 requis"
            st.markdown(f"""
            <div class="alert-attente" style="padding:0.6rem 1rem; font-size:0.85rem;">
                🔒 Fiche provisoire — <strong>{hint} pour finaliser</strong>
            </div>""", unsafe_allow_html=True)

    st.write("")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Retour ⬅️", use_container_width=True):
            st.session_state.step = 2
            rerun()
    with c2:
        if st.button("Générer la fiche d'orientation 📄", use_container_width=True, type="primary"):
            st.session_state.step = 4
            rerun()

# =====================================================================
# ETAPE 4 — FICHE D'ORIENTATION
# =====================================================================
elif step == 4:
    # ── Accès réservé au conseiller ──
    if not mode_conseiller:
        st.markdown("""
        <div style="max-width:480px;margin:3rem auto;text-align:center;
                    padding:2.5rem;background:linear-gradient(145deg,#1e293b,#0f172a);
                    border-radius:20px;border:1px solid #334155;
                    box-shadow:0 20px 60px rgba(0,0,0,0.4);">
            <div style="font-size:3rem;margin-bottom:0.8rem;">🔒</div>
            <div style="color:white;font-size:1.2rem;font-weight:700;">Accès restreint</div>
            <div style="color:#94a3b8;font-size:0.88rem;margin-top:0.5rem;line-height:1.5;">
                La Fiche d'Orientation est réservée au<br>
                <strong style="color:#f59e0b;">Conseiller d'Orientation</strong>.<br><br>
                Activez le Mode Conseiller dans la barre latérale.
            </div>
        </div>""", unsafe_allow_html=True)
        if st.button("⬅️ Retour au Diagnostic", use_container_width=True):
            st.session_state.step = 3
            rerun()
        st.stop()

    serie   = st.session_state.orientation_finale or "?"
    prenom  = st.session_state.prenom
    nom     = st.session_state.nom
    SA_brut, LA_brut, SA_etal, LA_etal, MECA_etal = calc_aptitudes()
    notes, trim_label = get_notes_actives()
    moy_sci, moy_lit  = calc_moyennes(notes)
    score_conf = st.session_state.score_confiance
    date_gen = datetime.now().strftime("%d/%m/%Y à %H:%M")

    st.subheader("📄 Fiche d'Orientation Scolaire")
    st.markdown(f"""
    <div class="fiche-box">
        <div style="text-align:center; margin-bottom:1rem;">
            <div style="font-size:0.8rem; color:#64748b; text-transform:uppercase;">République du Cameroun — Ministère de l'Éducation de Base</div>
            <div style="font-weight:700; font-size:1.3rem; color:#0f172a; margin:0.3rem 0;">FICHE D'ORIENTATION SCOLAIRE</div>
            <div style="font-size:0.8rem; color:#94a3b8;">Générée le {date_gen} — CapAvenir CMR v2.0</div>
        </div>
        <hr style="border-color:#e2e8f0; margin:0.8rem 0;">
        <p><strong>Nom &amp; Prénom :</strong> {nom} {prenom} &nbsp;|&nbsp; <strong>Âge :</strong> {st.session_state.age} ans &nbsp;|&nbsp; <strong>Sexe :</strong> {st.session_state.sexe}</p>
        <p><strong>Lycée :</strong> {st.session_state.lycee or "non renseigné"} &nbsp;|&nbsp; <strong>Choix personnel :</strong> {st.session_state.choix_personnel}</p>
        <p><strong>Projet professionnel :</strong> {st.session_state.projet_pro or "non renseigné"}</p>
        <p><strong>Revenu familial :</strong> {st.session_state.revenu}</p>
    </div>
    """, unsafe_allow_html=True)

    # Scores des tests (bruts + étalonnés)
    st.markdown("**Tests psychotechniques — Scores bruts & étalonnés :**")
    cols5 = st.columns(5)
    for col, (code, brut) in zip(cols5, [
        ("D48", st.session_state.d48), ("KRX", st.session_state.krx),
        ("MECA", st.session_state.meca), ("BV11", st.session_state.bv11),
        ("PRC", st.session_state.prc)
    ]):
        etal = etalonner(brut)
        col.markdown(f"""
        <div class="score-card">
            <div style="font-weight:700;">{code}</div>
            <div class="score-raw">Brut : {brut:.1f}</div>
            <div class="score-etalon">{etal:.1f}/20</div>
        </div>
        """, unsafe_allow_html=True)

    st.write("")
    st.markdown("**Aptitudes calculées (étalonnées) & moyennes scolaires :**")
    m1, m2, m3, m4, m5, m6 = st.columns(6)
    m1.metric("SA brute",             f"{SA_brut:.2f}/20")
    m2.metric("SA étalonnée",         f"{SA_etal:.2f}/20")
    # S9 — Gain d'étalonnage sur LA affiché pédagogiquement
    gain_la = LA_etal - LA_brut
    m3.metric("LA brute",             f"{LA_brut:.2f}/20")
    m4.metric("LA étalonnée",         f"{LA_etal:.2f}/20",
              delta=f"+{gain_la:.2f} gain étalonnage" if gain_la >= 0 else f"{gain_la:.2f}")
    m5.metric(f"Moy. Sci. ({trim_label})", f"{moy_sci:.2f}/20")
    m6.metric(f"Moy. Lit. ({trim_label})", f"{moy_lit:.2f}/20")

    # Badge orientation + statut
    # S2 — Prise en charge du badge TECHNIQUE
    badge_color = ("#10b981" if "C" in serie else
                   "#3b82f6" if "A" in serie else
                   "#f97316" if "TECHNIQUE" in serie else "#f59e0b")
    bg_color    = ("#ecfdf5" if "C" in serie else
                   "#eff6ff" if "A" in serie else
                   "#fff7ed" if "TECHNIQUE" in serie else "#fffbeb")
    conf_color  = "#10b981" if score_conf >= 70 else ("#f59e0b" if score_conf >= 50 else "#ef4444")

    STATUT_LABELS = {
        "confirme":          ("✅ CONFIRMÉ",              "#10b981", "#ecfdf5"),
        "attente":           ("⏳ EN ATTENTE — T2 requis", "#0ea5e9", "#f0f9ff"),
        "probation":         ("⚠️ PROBATION",              "#f59e0b", "#fffbeb"),
        "revise":            ("🔄 RÉVISÉ — Définitif",     "#f43f5e", "#fef2f2"),
        "indetermine":       ("❓ INDÉTERMINÉ",             "#94a3b8", "#f8fafc"),
        "confirme_technique":("🔧 TECHNIQUE — CONFIRMÉ",   "#f97316", "#fff7ed"),
    }
    s_label, s_color, s_bg = STATUT_LABELS.get(
        st.session_state.statut, ("—", "#94a3b8", "#f8fafc"))

    st.write("")
    st.markdown(f"""
    <div style="text-align:center; padding:1.5rem; background:{bg_color};
                border-radius:16px; border:2px solid {badge_color}; margin:1rem 0;
                box-shadow: 0 4px 20px rgba(15,23,42,0.08);">
        <div style="font-size:0.8rem; color:#64748b; margin-bottom:0.5rem; text-transform:uppercase; letter-spacing:0.08em;">Orientation recommandée</div>
        <span style="background:{badge_color}; color:white; padding:0.5rem 2.5rem;
                     border-radius:24px; font-weight:700; font-size:1.5rem;
                     box-shadow:0 4px 12px rgba(0,0,0,0.15);">SÉRIE {serie}</span>
        <div style="margin-top:1rem; display:flex; justify-content:center; gap:16px; flex-wrap:wrap;">
            <span style="background:{s_bg}; color:{s_color}; padding:0.3rem 1rem;
                         border-radius:20px; font-size:0.82rem; font-weight:600;
                         border:1px solid {s_color};">{s_label}</span>
            <span style="background:#f8fafc; color:{conf_color}; padding:0.3rem 1rem;
                         border-radius:20px; font-size:0.82rem; font-weight:600;
                         border:1px solid {conf_color};">Confiance : {score_conf} %</span>
            <span style="background:#f8fafc; color:#64748b; padding:0.3rem 1rem;
                         border-radius:20px; font-size:0.82rem;
                         border:1px solid #e2e8f0;">Trimestre : {trim_label}</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Observations du conseiller — TOUJOURS visibles sur la fiche ──
    obs_texte = st.session_state.get("notes_conseiller","") or ""
    obs_contenu = obs_texte if obs_texte.strip() else "<em style='opacity:0.5;'>Aucune observation saisie par le conseiller.</em>"
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,#fffbeb,#fef3c7);
                border-left:5px solid #f59e0b;border-radius:12px;
                padding:1.2rem 1.5rem;color:#78350f;margin:0.8rem 0;
                box-shadow:0 2px 8px rgba(245,158,11,0.15);">
        <div style="font-size:0.75rem;font-weight:700;text-transform:uppercase;
                    letter-spacing:0.08em;margin-bottom:0.6rem;color:#92400e;">
            📝 Observations du Conseiller d'Orientation
        </div>
        <div style="font-size:0.92rem;line-height:1.6;">
            {obs_contenu}
        </div>
    </div>""", unsafe_allow_html=True)

    if st.session_state.get("chat_history"):
        derniere_ia = [m["content"] for m in st.session_state.chat_history if m["role"] == "ia"]
        if derniere_ia:
            st.markdown(f"""
            <div class="chat-ia" style="margin:0.8rem 0;">
                🤖 <strong>Synthèse de l'entretien IA :</strong><br>{derniere_ia[-1]}
            </div>""", unsafe_allow_html=True)

    STATUT_TXT = {
        "confirme":   "CONFIRMÉ",
        "attente":    "EN ATTENTE — Notes T2 requises",
        "probation":  "PROBATION — Amélioration requise avant T3",
        "revise":     "RÉVISÉ DÉFINITIVEMENT",
        "indetermine":"INDÉTERMINÉ — Entretien requis",
    }
    statut_txt_export = STATUT_TXT.get(st.session_state.statut, "—")

    # Export enrichi
    fiche_txt = f"""FICHE D'ORIENTATION SCOLAIRE — CapAvenir CMR v2.1
Générée le {date_gen}

ÉLÈVE     : {nom} {prenom} | {st.session_state.age} ans | {st.session_state.sexe}
LYCÉE     : {st.session_state.lycee or 'non renseigné'}
CHOIX     : {st.session_state.choix_personnel}
PROJET    : {st.session_state.projet_pro or 'non renseigné'}
REVENU    : {st.session_state.revenu}

TESTS PSYCHOTECHNIQUES
  D48   : Brut={st.session_state.d48:.1f}/20  |  Étalonnée={etalonner(st.session_state.d48,"d48"):.1f}/20
  KRX   : Brut={st.session_state.krx:.1f}/20  |  Étalonnée={etalonner(st.session_state.krx,"krx"):.1f}/20
  MECA  : Brut={st.session_state.meca:.1f}/20  |  Étalonnée={etalonner(st.session_state.meca,"meca"):.1f}/20
  BV11  : Brut={st.session_state.bv11:.1f}/20  |  Étalonnée={etalonner(st.session_state.bv11,"bv11"):.1f}/20
  PRC   : Brut={st.session_state.prc:.1f}/20  |  Étalonnée={etalonner(st.session_state.prc,"prc"):.1f}/20

APTITUDES CALCULÉES
  SA brute    = (KRX + D48) / 2        = {SA_brut:.2f}/20
  SA étalonnée= (KRX_e + D48_e) / 2   = {SA_etal:.2f}/20
  LA étalonnée= (BV11_e + PRC_e) / 2  = {LA_etal:.2f}/20

NOTES SCOLAIRES ({trim_label})
  Maths={notes['maths']:.1f}  SciPhy={notes['sci_phy']:.1f}  SVT={notes['svt']:.1f}
  Français={notes['francais']:.1f}  HistGeo={notes['histgeo']:.1f}  Anglais={notes['anglais']:.1f}
  Moyenne Scientifique = {moy_sci:.2f}/20
  Moyenne Littéraire   = {moy_lit:.2f}/20

DÉCISION
  ORIENTATION : SÉRIE {serie}
  STATUT      : {statut_txt_export}
  CONFIANCE   : {score_conf} %

OBSERVATIONS DU CONSEILLER :
{st.session_state.notes_conseiller or 'Aucune observation saisie.'}
"""

    st.write("")    
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        if st.button("Retour au diagnostic ⬅️", use_container_width=True):
            st.session_state.step = 3
            rerun()
    with c2:
        if mode_conseiller:
            if st.button("💾 Sauvegarder le dossier", use_container_width=True, type="primary"):
                res = db.sauvegarder_dossier(st.session_state, etalonner)
                if res["succes"]:
                    st.success(res["message"])
                else:
                    st.error(res["message"])
        else:
            st.markdown('<div class="alert-danger" style="font-size:0.82rem;padding:0.5rem 1rem;">🔒 Sauvegarde réservée au Conseiller</div>', unsafe_allow_html=True)
    with c3:
        # ── Export TXT avec sauvegarde locale optionnelle ──
        doss = st.session_state.get("dossier_export", "").strip()
        fname_txt = f"orientation_{nom}_{prenom}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt"
        _bouton_export("📥 Télécharger (.txt)",
                       fiche_txt.encode("utf-8"), fname_txt,
                       "text/plain", doss, key="fiche_txt")
    with c4:
        derniere_ia = [m["content"] for m in st.session_state.chat_history if m["role"] == "ia"]
        ia_synthese  = derniere_ia[-1] if derniere_ia else ""
        pdf_data = {
            "nom": nom, "prenom": prenom,
            "age": st.session_state.age, "sexe": st.session_state.sexe,
            "lycee": st.session_state.lycee or "Non renseigné",
            "choix": st.session_state.choix_personnel,
            "projet_pro": st.session_state.projet_pro or "Non renseigné",
            "revenu": st.session_state.revenu,
            "d48": st.session_state.d48, "d48_e": etalonner(st.session_state.d48,"d48"),
            "krx": st.session_state.krx, "krx_e": etalonner(st.session_state.krx,"krx"),
            "meca": st.session_state.meca, "meca_e": etalonner(st.session_state.meca,"meca"),
            "bv11": st.session_state.bv11, "bv11_e": etalonner(st.session_state.bv11,"bv11"),
            "prc": st.session_state.prc, "prc_e": etalonner(st.session_state.prc,"prc"),
            "SA_brut": SA_brut, "SA_etal": SA_etal,
            "LA_brut": LA_brut, "LA_etal": LA_etal,
            "notes": notes, "trim_label": trim_label,
            "moy_sci": moy_sci, "moy_lit": moy_lit,
            "serie": serie, "probation": st.session_state.probation,
            "statut_label": statut_txt_export,
            "score_conf": score_conf,
            "notes_conseiller": st.session_state.notes_conseiller,
            "ia_synthese": ia_synthese,
            "date_gen": date_gen,
        }
        pdf_bytes = generate_pdf_fiche(pdf_data)
        # ── Export PDF avec sauvegarde locale optionnelle ──
        fname_pdf = f"orientation_{nom}_{prenom}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
        _bouton_export("📄 Télécharger (.pdf)",
                       pdf_bytes, fname_pdf,
                       "application/pdf", doss, key="fiche_pdf")
    with c5:
        if st.button("Nouveau dossier 🔄", use_container_width=True):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            rerun()

    st.write("")
    _sid = st.session_state.get("session_id", "—")
    st.markdown(
        f"<center><small style='color:#94a3b8;'>CapAvenir CMR v2.5 — 2025 — Mémoire ENS Filière Informatique Niveau 5 &nbsp;|&nbsp; Session : <code>{_sid}</code></small></center>",
        unsafe_allow_html=True
    )